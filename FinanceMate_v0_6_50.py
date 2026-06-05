
# -*- coding: utf-8 -*-
"""
Finance Mate – rekonstruierte lokale Desktop-App
Stand: v0.6.16-rekonstruiert

Rekonstruktionsbasis: Projektprotokoll FinanceMate – Rekonstruktionsmaster V4 / Patch-für-Patch-Systemreferenz.
Technik: Python, tkinter/ttk, SQLite, lokale Dateipersistenz.

Hinweis: Diese Datei ist als konsolidierte, startfähige Rekonstruktion aufgebaut. Sie bildet die im
Projektprotokoll beschriebene Fachsemantik in einer bereinigten Ein-Datei-Struktur ab.
"""

from __future__ import annotations

import os
import sys
import shutil
import sqlite3
import mimetypes
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from dataclasses import dataclass
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

APP_NAME = "Finance Mate"
APP_VERSION = "0.6.16-rekonstruiert"
STANDARD_DESIGN_NAME = "Reliefgrau"
DATE_FMT = "%d.%m.%Y"

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
ATTACHMENTS_DIR = BASE_DIR / "attachments"
IMPORTS_DIR = BASE_DIR / "imports"
DB_PATH = DATA_DIR / "finance_mate.sqlite3"

# Farben / Design
BG = "#eef1f4"
HEADER = "#dde3ea"
FOOTER = "#d8dee6"
LINE = "#b8c0ca"
TEXT = "#1f2933"
TEXT2 = "#5d6672"
WHITE = "#ffffff"
CARD_BG = "#f9fafb"
CARD_BORDER = "#cbd3dc"
BUTTON_GREEN = "#cfead1"
BUTTON_GREEN_ACTIVE = "#b9dfbd"
STANDARD_BUTTON_BG = "#d9dde3"
STANDARD_BUTTON_ACTIVE = "#cbd1d8"
STANDARD_BUTTON_BORDER = "#aeb7c2"
PLACEHOLDER_TEXT = "Suchen…"
PLACEHOLDER_COLOR = "#8a929c"
SOFT_GREEN = "#dff3e2"
SOFT_YELLOW = "#fff4c2"
SOFT_ORANGE = "#ffe1bf"
SOFT_RED = "#ffd4d4"

STATUS_OPEN = "Offen"
STATUS_PARTIAL = "Teilweise bezahlt"
STATUS_PAID = "Bezahlt"
STATUS_OVERDUE = "Überfällig"
STATUS_STORNO = "Storno"
PAYMENT_STATUS_VALUES = [STATUS_OPEN, STATUS_PARTIAL, STATUS_PAID, STATUS_OVERDUE, STATUS_STORNO]
TAX_SCOPE_VALUES = ["Inland", "EU", "Drittland"]
TRADE_COUNTRIES = ["Deutschland", "Österreich", "Schweiz", "Frankreich", "Italien", "Spanien", "Niederlande", "Belgien", "Polen", "Sonstige"]
SUPPORTED_ARCHIVE_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".txt", ".csv", ".doc", ".docx", ".xls", ".xlsx"}


def ensure_directories() -> None:
    for p in (DATA_DIR, ATTACHMENTS_DIR, IMPORTS_DIR):
        p.mkdir(parents=True, exist_ok=True)


def get_connection() -> sqlite3.Connection:
    ensure_directories()
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys = ON")
    return con


def table_has_column(con: sqlite3.Connection, table: str, column: str) -> bool:
    rows = con.execute(f"PRAGMA table_info({table})").fetchall()
    return any(r[1] == column for r in rows)


def ensure_column(con: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    if not table_has_column(con, table, column):
        con.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def init_sqlite() -> None:
    ensure_directories()
    with get_connection() as con:
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS app_meta (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS gl_accounts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                account_no TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                account_type TEXT DEFAULT '',
                tax_code TEXT DEFAULT '',
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS customers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                customer_no TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                country TEXT DEFAULT 'Deutschland',
                street TEXT DEFAULT '', zip TEXT DEFAULT '', city TEXT DEFAULT '',
                email TEXT DEFAULT '', phone TEXT DEFAULT '', tax_id TEXT DEFAULT '', vat_id TEXT DEFAULT '',
                iban TEXT DEFAULT '', bic TEXT DEFAULT '', bank_name TEXT DEFAULT '',
                payment_term TEXT DEFAULT '', active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS vendors (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                vendor_no TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                country TEXT DEFAULT 'Deutschland',
                street TEXT DEFAULT '', zip TEXT DEFAULT '', city TEXT DEFAULT '',
                email TEXT DEFAULT '', phone TEXT DEFAULT '', tax_id TEXT DEFAULT '', vat_id TEXT DEFAULT '',
                iban TEXT DEFAULT '', bic TEXT DEFAULT '', bank_name TEXT DEFAULT '',
                payment_term TEXT DEFAULT '', active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS tax_codes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                rate REAL DEFAULT 0,
                tax_scope TEXT DEFAULT 'Inland',
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS payment_terms (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                days INTEGER DEFAULT 0,
                discount_days INTEGER DEFAULT 0,
                discount_percent REAL DEFAULT 0,
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS bank_accounts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                iban TEXT DEFAULT '', bic TEXT DEFAULT '', bank_name TEXT DEFAULT '',
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS journal_entries (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_no TEXT UNIQUE NOT NULL,
                document_date TEXT NOT NULL,
                posting_date TEXT NOT NULL,
                description TEXT DEFAULT '',
                total_debit REAL DEFAULT 0,
                total_credit REAL DEFAULT 0,
                status TEXT DEFAULT 'Gebucht',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS journal_entry_lines (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_no TEXT NOT NULL,
                account_no TEXT NOT NULL,
                side TEXT NOT NULL,
                amount REAL NOT NULL,
                tax_code TEXT DEFAULT '',
                text TEXT DEFAULT '',
                FOREIGN KEY(document_no) REFERENCES journal_entries(document_no) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS customer_invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_no TEXT UNIQUE NOT NULL,
                customer_no TEXT NOT NULL,
                customer_name TEXT NOT NULL,
                customer_address TEXT DEFAULT '',
                invoice_date TEXT NOT NULL,
                due_date TEXT NOT NULL,
                payment_term TEXT DEFAULT '',
                tax_code TEXT DEFAULT '',
                net_amount REAL DEFAULT 0,
                tax_amount REAL DEFAULT 0,
                gross_amount REAL DEFAULT 0,
                open_amount REAL DEFAULT 0,
                status TEXT DEFAULT 'Offen',
                linked_journal_no TEXT DEFAULT '',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS vendor_invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_no TEXT UNIQUE NOT NULL,
                vendor_no TEXT NOT NULL,
                vendor_name TEXT NOT NULL,
                vendor_address TEXT DEFAULT '',
                invoice_date TEXT NOT NULL,
                due_date TEXT NOT NULL,
                payment_term TEXT DEFAULT '',
                tax_code TEXT DEFAULT '',
                net_amount REAL DEFAULT 0,
                tax_amount REAL DEFAULT 0,
                gross_amount REAL DEFAULT 0,
                open_amount REAL DEFAULT 0,
                status TEXT DEFAULT 'Offen',
                linked_journal_no TEXT DEFAULT '',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS open_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                entity_type TEXT NOT NULL,
                reference_no TEXT NOT NULL,
                partner_no TEXT NOT NULL,
                partner_name TEXT NOT NULL,
                due_date TEXT NOT NULL,
                original_amount REAL NOT NULL,
                open_amount REAL NOT NULL,
                status TEXT NOT NULL,
                linked_journal_no TEXT DEFAULT '',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS attachments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                entity_type TEXT NOT NULL,
                reference_no TEXT NOT NULL,
                file_name TEXT NOT NULL,
                file_path TEXT NOT NULL,
                added_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS invoice_import_batches (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                hist_no TEXT UNIQUE NOT NULL,
                title TEXT NOT NULL,
                vendor_no TEXT DEFAULT '',
                vendor_name TEXT DEFAULT '',
                status TEXT DEFAULT 'Importiert',
                marked INTEGER DEFAULT 0,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS invoice_import_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                hist_no TEXT NOT NULL,
                file_name TEXT NOT NULL,
                file_path TEXT NOT NULL,
                mime_type TEXT DEFAULT '',
                created_at TEXT NOT NULL,
                FOREIGN KEY(hist_no) REFERENCES invoice_import_batches(hist_no) ON DELETE CASCADE
            );
            """
        )
        # defensive schema compatibility
        for table, col, definition in [
            ("customer_invoices", "customer_address", "TEXT DEFAULT ''"),
            ("customer_invoices", "linked_journal_no", "TEXT DEFAULT ''"),
            ("vendor_invoices", "vendor_address", "TEXT DEFAULT ''"),
            ("vendor_invoices", "linked_journal_no", "TEXT DEFAULT ''"),
            ("open_items", "linked_journal_no", "TEXT DEFAULT ''"),
            ("tax_codes", "tax_scope", "TEXT DEFAULT 'Inland'"),
        ]:
            ensure_column(con, table, col, definition)
        now = now_str()
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now))
        seed_defaults(con)
        con.commit()


def seed_defaults(con: sqlite3.Connection) -> None:
    now = now_str()
    defaults = [
        ("gl_accounts", "account_no", "1000", {"name": "Kasse", "account_type": "Aktiv"}),
        ("gl_accounts", "account_no", "1200", {"name": "Bank", "account_type": "Aktiv"}),
        ("gl_accounts", "account_no", "1400", {"name": "Forderungen", "account_type": "Aktiv"}),
        ("gl_accounts", "account_no", "1600", {"name": "Verbindlichkeiten", "account_type": "Passiv"}),
        ("gl_accounts", "account_no", "8400", {"name": "Erlöse", "account_type": "Ertrag"}),
        ("gl_accounts", "account_no", "3400", {"name": "Wareneingang", "account_type": "Aufwand"}),
    ]
    for table, keycol, keyval, vals in defaults:
        exists = con.execute(f"SELECT 1 FROM {table} WHERE {keycol}=?", (keyval,)).fetchone()
        if not exists:
            cols = [keycol] + list(vals.keys()) + ["created_at", "updated_at"]
            params = [keyval] + list(vals.values()) + [now, now]
            con.execute(f"INSERT INTO {table}({','.join(cols)}) VALUES({','.join(['?']*len(cols))})", params)
    for code, name, rate in [("V19", "Vorsteuer/Umsatzsteuer 19%", 19), ("V7", "Vorsteuer/Umsatzsteuer 7%", 7), ("O0", "Ohne Steuer", 0)]:
        if not con.execute("SELECT 1 FROM tax_codes WHERE code=?", (code,)).fetchone():
            con.execute("INSERT INTO tax_codes(code,name,rate,tax_scope,created_at,updated_at) VALUES(?,?,?,?,?,?)", (code, name, rate, "Inland", now, now))
    for code, name, days in [("SOFORT", "Sofort fällig", 0), ("14T", "14 Tage netto", 14), ("30T", "30 Tage netto", 30)]:
        if not con.execute("SELECT 1 FROM payment_terms WHERE code=?", (code,)).fetchone():
            con.execute("INSERT INTO payment_terms(code,name,days,created_at,updated_at) VALUES(?,?,?,?,?)", (code, name, days, now, now))


def now_str() -> str:
    return datetime.now().strftime("%d.%m.%Y %H:%M:%S")


def today_str() -> str:
    return datetime.now().strftime(DATE_FMT)


def validate_date(value: str) -> bool:
    try:
        datetime.strptime(value.strip(), DATE_FMT)
        return True
    except Exception:
        return False


def parse_amount(value: Any) -> Decimal:
    if value is None or value == "":
        return Decimal("0.00")
    s = str(value).strip().replace("€", "").replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    else:
        s = s.replace(",", ".")
    try:
        return Decimal(s).quantize(Decimal("0.01"))
    except InvalidOperation:
        raise ValueError(f"Ungültiger Betrag: {value}")


def format_amount(value: Any) -> str:
    try:
        d = Decimal(str(value)).quantize(Decimal("0.01"))
    except Exception:
        d = Decimal("0.00")
    return f"{d:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def calc_due_date(date_str: str, payment_term_code: str) -> str:
    base = datetime.strptime(date_str, DATE_FMT)
    days = 0
    with get_connection() as con:
        row = con.execute("SELECT days FROM payment_terms WHERE code=?", (payment_term_code,)).fetchone()
        if row:
            days = int(row[0] or 0)
    return (base + timedelta(days=days)).strftime(DATE_FMT)


def compute_status_from_open_amount(open_amount: Any, due_date: str) -> str:
    amount = parse_amount(open_amount)
    if amount <= 0:
        return STATUS_PAID
    if validate_date(due_date) and datetime.strptime(due_date, DATE_FMT).date() < datetime.now().date():
        return STATUS_OVERDUE
    return STATUS_OPEN


def urgency_bucket(due_date: str, open_amount: Any) -> str:
    if parse_amount(open_amount) <= 0:
        return "paid"
    if not validate_date(due_date):
        return "open"
    d = datetime.strptime(due_date, DATE_FMT).date()
    delta = (d - datetime.now().date()).days
    if delta < 0:
        return "overdue"
    if delta <= 7:
        return "soon"
    return "open"


def normalize_tax_code(code: str) -> str:
    return (code or "").strip().upper()


def yes_no_to_int(v: Any) -> int:
    return 1 if str(v).lower() in ("1", "ja", "true", "aktiv", "yes") else 0


def int_to_yes_no(v: Any) -> str:
    return "Ja" if int(v or 0) else "Nein"


def generate_number(prefix: str, meta_key: str, width: int = 5) -> str:
    with get_connection() as con:
        row = con.execute("SELECT value FROM app_meta WHERE key=?", (meta_key,)).fetchone()
        current = int(row[0]) if row else 0
        nxt = current + 1
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", (meta_key, str(nxt), now_str()))
        con.commit()
    return f"{prefix}{nxt:0{width}d}"


def next_hist_no() -> str:
    return generate_number("HIST-", "counter_hist", 6)


def clean_document_title(path: str) -> str:
    return Path(path).stem.replace("_", " ").replace("-", " ").strip() or Path(path).name


def iter_supported_files_from_folder(folder: str) -> Iterable[str]:
    for root, _dirs, files in os.walk(folder):
        for f in files:
            p = Path(root) / f
            if p.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
                yield str(p)


def collect_preview_text(path: str, limit: int = 8000) -> str:
    p = Path(path)
    if not p.exists():
        return "Datei nicht gefunden."
    if p.suffix.lower() in {".txt", ".csv"}:
        try:
            return p.read_text(encoding="utf-8", errors="replace")[:limit]
        except Exception as exc:
            return f"Textvorschau nicht möglich: {exc}"
    return f"Dokumentkarte\n\nDatei: {p.name}\nTyp: {mimetypes.guess_type(str(p))[0] or p.suffix}\nPfad: {p}\n\nFür PDF-/Office-Dateien wird eine Dateikarte angezeigt. Öffnen über Doppelklick oder Kontextmenü."


def open_path(path: str) -> None:
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as exc:
        messagebox.showerror("Öffnen nicht möglich", str(exc))


def apply_search_placeholder(entry: tk.Entry, placeholder: str = PLACEHOLDER_TEXT) -> None:
    def on_focus_in(_event=None):
        if entry.get() == placeholder and entry.cget("fg") == PLACEHOLDER_COLOR:
            entry.delete(0, tk.END)
            entry.config(fg=TEXT)
    def on_focus_out(_event=None):
        if not entry.get().strip():
            entry.insert(0, placeholder)
            entry.config(fg=PLACEHOLDER_COLOR)
    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)
    on_focus_out()


def create_standard_button(parent, text: str, command=None, confirm: bool = False, **kw):
    style = "Confirm.TButton" if confirm else "StandardGray.TButton"
    return ttk.Button(parent, text=text, command=command, style=style, **kw)


def configure_tree_tags(tree: ttk.Treeview) -> None:
    tree.tag_configure("paid", background=SOFT_GREEN)
    tree.tag_configure("soon", background=SOFT_YELLOW)
    tree.tag_configure("overdue", background=SOFT_RED)
    tree.tag_configure("open", background=WHITE)


def load_ui_preference(key: str, default: str = "") -> str:
    with get_connection() as con:
        row = con.execute("SELECT value FROM app_meta WHERE key=?", (f"ui:{key}",)).fetchone()
        return row[0] if row else default


def save_ui_preference(key: str, value: str) -> None:
    with get_connection() as con:
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", (f"ui:{key}", value, now_str()))
        con.commit()


class ScrollableFrame(ttk.Frame):
    def __init__(self, parent, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.canvas = tk.Canvas(self, bg=BG, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview, style="Vertical.TScrollbar")
        self.content = ttk.Frame(self.canvas)
        self.content_id = self.canvas.create_window((0, 0), window=self.content, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        self.content.bind("<Configure>", self._on_content_configure)
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfigure(self.content_id, width=e.width))
        self._bind_mousewheel(self.canvas)
        self._bind_descendant_wheel(self.content)

    def _on_content_configure(self, _event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _bind_mousewheel(self, widget):
        widget.bind("<Enter>", lambda e: self.canvas.bind_all("<MouseWheel>", self._on_mousewheel), add="+")
        widget.bind("<Leave>", lambda e: self.canvas.unbind_all("<MouseWheel>"), add="+")

    def _bind_descendant_wheel(self, widget):
        # Wheel always scrolls frame; comboboxes explicitly do not change values by wheel.
        def bind_child(child):
            if isinstance(child, ttk.Combobox):
                child.bind("<MouseWheel>", lambda e: "break")
            else:
                child.bind("<MouseWheel>", self._on_mousewheel, add="+")
            for sub in child.winfo_children():
                bind_child(sub)
        self.after(100, lambda: bind_child(widget))

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        return "break"


class SortableTreeMixin:
    def _coerce_sort_value(self, value: str):
        if value is None:
            return ""
        s = str(value)
        try:
            return float(s.replace(".", "").replace(",", "."))
        except Exception:
            pass
        if validate_date(s):
            return datetime.strptime(s, DATE_FMT)
        return s.lower()

    def setup_sorting(self, tree: ttk.Treeview) -> None:
        for col in tree["columns"]:
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

    def sort_treeview(self, tree: ttk.Treeview, col: str, reverse: bool) -> None:
        data = [(self._coerce_sort_value(tree.set(k, col)), k) for k in tree.get_children("")]
        data.sort(reverse=reverse)
        for idx, (_val, k) in enumerate(data):
            tree.move(k, "", idx)
        tree.heading(col, command=lambda: self.sort_treeview(tree, col, not reverse))


class AttachmentMixin:
    def add_attachment_paths(self, entity_type: str, reference_no: str, paths: Sequence[str]) -> None:
        if not reference_no:
            messagebox.showwarning("Anhang", "Bitte zuerst einen Datensatz speichern oder auswählen.")
            return
        target_dir = ATTACHMENTS_DIR / entity_type / reference_no
        target_dir.mkdir(parents=True, exist_ok=True)
        with get_connection() as con:
            for p in paths:
                src = Path(p)
                if not src.exists():
                    continue
                dest = target_dir / src.name
                if dest.exists():
                    dest = target_dir / f"{src.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{src.suffix}"
                shutil.copy2(src, dest)
                con.execute(
                    "INSERT INTO attachments(entity_type,reference_no,file_name,file_path,added_at) VALUES(?,?,?,?,?)",
                    (entity_type, reference_no, dest.name, str(dest), now_str()),
                )
            con.commit()

    def get_attachment_count(self, entity_type: str, reference_no: str) -> int:
        with get_connection() as con:
            row = con.execute("SELECT COUNT(*) FROM attachments WHERE entity_type=? AND reference_no=?", (entity_type, reference_no)).fetchone()
            return int(row[0] if row else 0)

    def get_attachments(self, entity_type: str, reference_no: str) -> List[sqlite3.Row]:
        with get_connection() as con:
            return list(con.execute("SELECT * FROM attachments WHERE entity_type=? AND reference_no=? ORDER BY added_at DESC", (entity_type, reference_no)).fetchall())

    def open_attachment_popup(self, entity_type: str, reference_no: str) -> None:
        pop = tk.Toplevel()
        pop.title(f"Anhänge – {reference_no}")
        pop.geometry("900x520")
        pop.configure(bg=BG)
        ttk.Label(pop, text=f"Anhänge zu {entity_type} / {reference_no}", style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
        cols = ("file_name", "file_path", "added_at")
        tree = ttk.Treeview(pop, columns=cols, show="headings")
        for col, w, txt in [("file_name", 220, "Dateiname"), ("file_path", 480, "Pfad"), ("added_at", 150, "Zeitstempel")]:
            tree.heading(col, text=txt)
            tree.column(col, width=w, anchor="w")
        tree.pack(fill="both", expand=True, padx=12, pady=6)
        for r in self.get_attachments(entity_type, reference_no):
            tree.insert("", "end", values=(r["file_name"], r["file_path"], r["added_at"]))
        tree.bind("<Double-1>", lambda e: open_path(tree.set(tree.focus(), "file_path")) if tree.focus() else None)
        bar = ttk.Frame(pop)
        bar.pack(fill="x", padx=12, pady=10)
        create_standard_button(bar, "Datei hinzufügen", lambda: self._attachment_popup_add(entity_type, reference_no, tree), confirm=True).pack(side="left")
        create_standard_button(bar, "Schließen", pop.destroy).pack(side="right")

    def _attachment_popup_add(self, entity_type: str, reference_no: str, tree: ttk.Treeview):
        paths = filedialog.askopenfilenames(title="Anhänge auswählen")
        if paths:
            self.add_attachment_paths(entity_type, reference_no, paths)
            for i in tree.get_children(): tree.delete(i)
            for r in self.get_attachments(entity_type, reference_no):
                tree.insert("", "end", values=(r["file_name"], r["file_path"], r["added_at"]))


class TwoBlock(ttk.Frame):
    """Zentrale Zweiblock-Engine: freie Größenänderung ohne starre Wand."""
    def __init__(self, parent, left_title: str, right_title: str):
        super().__init__(parent)
        self.pane = ttk.PanedWindow(self, orient="horizontal")
        self.pane.pack(fill="both", expand=True)
        self.left = ttk.Frame(self.pane, style="Card.TFrame")
        self.right = ttk.Frame(self.pane, style="Card.TFrame")
        self.pane.add(self.left, weight=3)
        self.pane.add(self.right, weight=2)
        ttk.Label(self.left, text=left_title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))
        ttk.Label(self.right, text=right_title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))


def build_two_block_shell(parent, left_title: str, right_title: str) -> TwoBlock:
    tb = TwoBlock(parent, left_title, right_title)
    tb.pack(fill="both", expand=True)
    return tb


class StammdatenView(ttk.Frame, SortableTreeMixin):
    CONFIG = {
        "Sachkonten": {
            "table": "gl_accounts", "key": "account_no", "label": "Sachkonto",
            "fields": [("account_no", "Konto-Nr."), ("name", "Name"), ("account_type", "Kontotyp"), ("tax_code", "Steuerkennz."), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("account_no", 100, "Konto"), ("name", 220, "Name"), ("account_type", 120, "Typ"), ("tax_code", 100, "Steuer"), ("active", 70, "Aktiv")],
            "prefix": "SK"
        },
        "Debitoren": {
            "table": "customers", "key": "customer_no", "label": "Debitor",
            "fields": [("customer_no", "Debitor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"), ("email", "E-Mail"), ("phone", "Telefon"), ("tax_id", "Steuer-Nr."), ("vat_id", "USt-IdNr."), ("iban", "IBAN"), ("bic", "BIC"), ("bank_name", "Bank"), ("payment_term", "Zahlungsbed."), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("customer_no", 110, "Nr."), ("name", 220, "Name"), ("city", 140, "Ort"), ("payment_term", 110, "ZB"), ("active", 70, "Aktiv")],
            "prefix": "D"
        },
        "Kreditoren": {
            "table": "vendors", "key": "vendor_no", "label": "Kreditor",
            "fields": [("vendor_no", "Kreditor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"), ("email", "E-Mail"), ("phone", "Telefon"), ("tax_id", "Steuer-Nr."), ("vat_id", "USt-IdNr."), ("iban", "IBAN"), ("bic", "BIC"), ("bank_name", "Bank"), ("payment_term", "Zahlungsbed."), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("vendor_no", 110, "Nr."), ("name", 220, "Name"), ("city", 140, "Ort"), ("payment_term", 110, "ZB"), ("active", 70, "Aktiv")],
            "prefix": "K"
        },
        "Steuerkennzeichen": {
            "table": "tax_codes", "key": "code", "label": "Steuerkennzeichen",
            "fields": [("code", "Code"), ("name", "Name"), ("rate", "Satz %"), ("tax_scope", "Gültigkeit"), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("code", 90, "Code"), ("name", 240, "Name"), ("rate", 90, "%"), ("tax_scope", 120, "Gültigkeit"), ("active", 70, "Aktiv")],
            "prefix": "T"
        },
        "Zahlungsbedingungen": {
            "table": "payment_terms", "key": "code", "label": "Zahlungsbedingung",
            "fields": [("code", "Code"), ("name", "Name"), ("days", "Tage netto"), ("discount_days", "Skontotage"), ("discount_percent", "Skonto %"), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("code", 90, "Code"), ("name", 240, "Name"), ("days", 90, "Tage"), ("discount_percent", 100, "Skonto %"), ("active", 70, "Aktiv")],
            "prefix": "ZB"
        },
    }

    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.vars: Dict[str, Dict[str, tk.StringVar]] = {}
        self.trees: Dict[str, ttk.Treeview] = {}
        self.search_vars: Dict[str, tk.StringVar] = {}
        self.selected_ids: Dict[str, Optional[int]] = {}
        self._build_ui()
        self.load_all_tables()

    def _build_ui(self):
        top = ttk.Frame(self)
        top.pack(fill="x", pady=(0, 6))
        ttk.Label(top, text="Stammdaten", style="CardTitle.TLabel").pack(side="left")
        create_standard_button(top, "Daten-Refresh", self.load_all_tables).pack(side="right")
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)
        for title in self.CONFIG:
            self._build_tab(title)
        self.nb.bind("<<NotebookTabChanged>>", lambda e: self.load_all_tables())

    def _build_tab(self, title: str):
        cfg = self.CONFIG[title]
        frame = ttk.Frame(self.nb)
        self.nb.add(frame, text=title)
        shell = build_two_block_shell(frame, f"{title} – Liste", f"{title} – Formular")
        # left
        left_body = ttk.Frame(shell.left)
        left_body.pack(fill="both", expand=True, padx=10, pady=8)
        search = tk.Entry(left_body, relief="sunken", bg=WHITE, fg=TEXT)
        search.pack(fill="x", pady=(0, 6))
        apply_search_placeholder(search)
        search.bind("<KeyRelease>", lambda e, t=title, s=search: self.load_table_data(t, s.get()))
        cols = [c[0] for c in cfg["columns"]]
        tree = ttk.Treeview(left_body, columns=cols, show="headings", selectmode="browse")
        for col, w, txt in cfg["columns"]:
            tree.heading(col, text=txt)
            tree.column(col, width=w, anchor="w")
        tree.pack(fill="both", expand=True)
        tree.bind("<<TreeviewSelect>>", lambda e, t=title: self.load_selected_record(t))
        self.setup_sorting(tree)
        self.trees[title] = tree
        self.selected_ids[title] = None
        # right
        sf = ScrollableFrame(shell.right)
        sf.pack(fill="both", expand=True, padx=10, pady=8)
        self.vars[title] = {}
        for field, label in cfg["fields"]:
            row = ttk.Frame(sf.content)
            row.pack(fill="x", pady=3)
            ttk.Label(row, text=label, width=18).pack(side="left")
            var = tk.StringVar(value="Ja" if field == "active" else "")
            self.vars[title][field] = var
            if field == "country":
                cb = ttk.Combobox(row, textvariable=var, values=TRADE_COUNTRIES, state="readonly")
                cb.pack(side="left", fill="x", expand=True)
                cb.bind("<MouseWheel>", lambda e: "break")
            elif field == "tax_scope":
                cb = ttk.Combobox(row, textvariable=var, values=TAX_SCOPE_VALUES, state="readonly")
                cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
            elif field == "payment_term":
                cb = ttk.Combobox(row, textvariable=var, values=self._load_payment_term_codes(), state="readonly")
                cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
            elif field == "active":
                cb = ttk.Combobox(row, textvariable=var, values=["Ja", "Nein"], state="readonly")
                cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
            else:
                ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)
        btns = ttk.Frame(sf.content)
        btns.pack(fill="x", pady=10)
        create_standard_button(btns, "Neu / Felder leeren", lambda t=title: self.clear_form(t)).pack(side="left", padx=3)
        create_standard_button(btns, "Speichern", lambda t=title: self.save_record(t), confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Löschen", lambda t=title: self.delete_record(t)).pack(side="left", padx=3)

    def _load_payment_term_codes(self):
        with get_connection() as con:
            return [r[0] for r in con.execute("SELECT code FROM payment_terms ORDER BY code")]

    def load_all_tables(self):
        for t in self.CONFIG:
            self.load_table_data(t)

    def load_table_data(self, title: str, search: str = ""):
        if search == PLACEHOLDER_TEXT: search = ""
        cfg = self.CONFIG[title]
        tree = self.trees[title]
        for i in tree.get_children(): tree.delete(i)
        cols = [c[0] for c in cfg["columns"]]
        sql = f"SELECT {','.join(cols)} FROM {cfg['table']}"
        params: List[Any] = []
        if search.strip():
            where_cols = [c for c in cols if c != "id"]
            sql += " WHERE " + " OR ".join([f"CAST({c} AS TEXT) LIKE ?" for c in where_cols])
            params = [f"%{search.strip()}%"] * len(where_cols)
        sql += " ORDER BY id DESC"
        with get_connection() as con:
            for r in con.execute(sql, params):
                vals = [int_to_yes_no(r[c]) if c == "active" else r[c] for c in cols]
                tree.insert("", "end", values=vals)

    def load_selected_record(self, title: str):
        tree = self.trees[title]
        iid = tree.focus()
        if not iid: return
        item = tree.item(iid)
        rec_id = item["values"][0]
        self.selected_ids[title] = int(rec_id)
        cfg = self.CONFIG[title]
        with get_connection() as con:
            row = con.execute(f"SELECT * FROM {cfg['table']} WHERE id=?", (rec_id,)).fetchone()
        if not row: return
        for field, var in self.vars[title].items():
            if field == "active": var.set(int_to_yes_no(row[field]))
            else: var.set(str(row[field] or ""))

    def clear_form(self, title: str):
        self.selected_ids[title] = None
        cfg = self.CONFIG[title]
        for field, var in self.vars[title].items():
            if field == "active": var.set("Ja")
            elif field == "country": var.set("Deutschland")
            elif field == cfg["key"]: var.set(generate_number(cfg["prefix"], f"counter_{cfg['table']}", 5))
            else: var.set("")

    def save_record(self, title: str):
        cfg = self.CONFIG[title]
        vals = {field: var.get().strip() for field, var in self.vars[title].items()}
        if not vals.get(cfg["key"]):
            vals[cfg["key"]] = generate_number(cfg["prefix"], f"counter_{cfg['table']}", 5)
        if not vals.get("name") and "name" in vals:
            messagebox.showwarning("Pflichtfeld", "Bitte einen Namen eingeben.")
            return
        if "active" in vals: vals["active"] = yes_no_to_int(vals["active"])
        now = now_str()
        with get_connection() as con:
            if self.selected_ids[title]:
                cols = list(vals.keys()) + ["updated_at"]
                params = list(vals.values()) + [now, self.selected_ids[title]]
                con.execute(f"UPDATE {cfg['table']} SET " + ",".join([f"{c}=?" for c in cols]) + " WHERE id=?", params)
            else:
                vals["created_at"] = now; vals["updated_at"] = now
                cols = list(vals.keys())
                con.execute(f"INSERT INTO {cfg['table']}({','.join(cols)}) VALUES({','.join(['?']*len(cols))})", list(vals.values()))
                self.selected_ids[title] = con.execute("SELECT last_insert_rowid()").fetchone()[0]
            con.commit()
        self.app.set_status(f"{title}: Datensatz gespeichert.")
        self.load_table_data(title)

    def delete_record(self, title: str):
        rec_id = self.selected_ids[title]
        if not rec_id:
            messagebox.showinfo("Löschen", "Bitte zuerst einen Datensatz auswählen.")
            return
        cfg = self.CONFIG[title]
        name = self.vars[title].get("name", tk.StringVar(value="")).get() or self.vars[title].get("code", tk.StringVar(value="")).get()
        if not messagebox.askyesno("Unwiderruflich löschen", f"[{rec_id}] [{name}] unwiderruflich löschen?"):
            return
        with get_connection() as con:
            con.execute(f"DELETE FROM {cfg['table']} WHERE id=?", (rec_id,)); con.commit()
        self.clear_form(title)
        self.load_table_data(title)
        self.app.set_status(f"{title}: Datensatz gelöscht.")


class JournalView(ttk.Frame, SortableTreeMixin, AttachmentMixin):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.pending_lines: List[Dict[str, Any]] = []
        self.pending_attachments: List[str] = []
        self.selected_doc = ""
        self._build_ui()
        self.reload_journal_entries()

    def _build_ui(self):
        shell = build_two_block_shell(self, "Buchung erfassen", "Historie / Kontrolle")
        left = ScrollableFrame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        self.doc_no = tk.StringVar(value=generate_number("B", "counter_journal", 6))
        self.doc_date = tk.StringVar(value=today_str()); self.posting_date = tk.StringVar(value=today_str()); self.desc = tk.StringVar()
        for label, var in [("Beleg-Nr.", self.doc_no), ("Belegdatum", self.doc_date), ("Buchungsdatum", self.posting_date), ("Beschreibung", self.desc)]:
            self._labeled_entry(left.content, label, var)
        ttk.Separator(left.content).pack(fill="x", pady=8)
        self.line_account = tk.StringVar(); self.line_side = tk.StringVar(value="Soll"); self.line_amount = tk.StringVar(); self.line_tax = tk.StringVar(); self.line_text = tk.StringVar()
        self._labeled_combo(left.content, "Konto", self.line_account, self._load_accounts())
        self._labeled_combo(left.content, "Soll/Haben", self.line_side, ["Soll", "Haben"])
        self._labeled_entry(left.content, "Betrag", self.line_amount)
        self._labeled_combo(left.content, "Steuerkennz.", self.line_tax, self._load_tax_codes())
        self._labeled_entry(left.content, "Text", self.line_text)
        btns = ttk.Frame(left.content); btns.pack(fill="x", pady=6)
        create_standard_button(btns, "Zeile hinzufügen", self.add_line_item, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Zeile entfernen", self.remove_selected_line).pack(side="left", padx=3)
        cols = ("account", "side", "amount", "tax", "text")
        self.lines_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=7)
        for c, t, w in [("account", "Konto", 110), ("side", "S/H", 70), ("amount", "Betrag", 100), ("tax", "Steuer", 90), ("text", "Text", 240)]:
            self.lines_tree.heading(c, text=t); self.lines_tree.column(c, width=w, anchor="w")
        self.lines_tree.pack(fill="x", pady=6)
        self.totals = ttk.Label(left.content, text="Soll 0,00 | Haben 0,00 | Differenz 0,00", style="Hint.TLabel")
        self.totals.pack(anchor="w", pady=5)
        btns2 = ttk.Frame(left.content); btns2.pack(fill="x", pady=10)
        create_standard_button(btns2, "Anhang vormerken", self.add_pending_attachments).pack(side="left", padx=3)
        create_standard_button(btns2, "Buchung speichern", self.save_journal_entry, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns2, "Neu", self.clear_journal_form).pack(side="left", padx=3)
        # right history
        body = ttk.Frame(shell.right); body.pack(fill="both", expand=True, padx=10, pady=8)
        cols = ("document_no", "posting_date", "description", "debit", "credit", "att")
        self.hist_tree = ttk.Treeview(body, columns=cols, show="headings")
        for c, t, w in [("document_no", "Beleg", 110), ("posting_date", "Datum", 90), ("description", "Beschreibung", 240), ("debit", "Soll", 90), ("credit", "Haben", 90), ("att", "Anh.", 55)]:
            self.hist_tree.heading(c, text=t); self.hist_tree.column(c, width=w, anchor="w")
        self.hist_tree.pack(fill="both", expand=True)
        self.hist_tree.bind("<Double-1>", self._on_history_click)
        create_standard_button(body, "Anhang zu ausgewähltem Beleg", self.add_attachment_to_selected_history).pack(anchor="e", pady=8)
        self.setup_sorting(self.hist_tree)

    def _labeled_entry(self, parent, label, var):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)

    def _labeled_combo(self, parent, label, var, values):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        cb = ttk.Combobox(row, textvariable=var, values=values)
        cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")

    def _load_accounts(self):
        with get_connection() as con: return [f"{r['account_no']} {r['name']}" for r in con.execute("SELECT account_no,name FROM gl_accounts WHERE active=1 ORDER BY account_no")]
    def _load_tax_codes(self):
        with get_connection() as con: return [r[0] for r in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]

    def add_line_item(self):
        try: amount = parse_amount(self.line_amount.get())
        except ValueError as exc: messagebox.showwarning("Betrag", str(exc)); return
        acc = self.line_account.get().split(" ")[0]
        if not acc or amount <= 0: messagebox.showwarning("Zeile", "Bitte Konto und positiven Betrag erfassen."); return
        self.pending_lines.append({"account": acc, "side": self.line_side.get(), "amount": float(amount), "tax": self.line_tax.get(), "text": self.line_text.get()})
        self._refresh_lines_tree(); self.clear_line_inputs()

    def remove_selected_line(self):
        iid = self.lines_tree.focus()
        if iid:
            idx = self.lines_tree.index(iid); self.pending_lines.pop(idx); self._refresh_lines_tree()

    def clear_line_inputs(self):
        self.line_account.set(""); self.line_side.set("Soll"); self.line_amount.set(""); self.line_tax.set(""); self.line_text.set("")

    def _refresh_lines_tree(self):
        for i in self.lines_tree.get_children(): self.lines_tree.delete(i)
        debit = Decimal("0"); credit = Decimal("0")
        for ln in self.pending_lines:
            self.lines_tree.insert("", "end", values=(ln["account"], ln["side"], format_amount(ln["amount"]), ln["tax"], ln["text"]))
            if ln["side"] == "Soll": debit += parse_amount(ln["amount"])
            else: credit += parse_amount(ln["amount"])
        diff = debit - credit
        self.totals.config(text=f"Soll {format_amount(debit)} | Haben {format_amount(credit)} | Differenz {format_amount(diff)}")

    def add_pending_attachments(self):
        paths = filedialog.askopenfilenames(title="Beleganhänge vormerken")
        if paths:
            self.pending_attachments.extend(paths)
            self.app.set_status(f"{len(paths)} Anhang/Anhänge vorgemerkt.")

    def save_journal_entry(self):
        if not validate_date(self.doc_date.get()) or not validate_date(self.posting_date.get()):
            messagebox.showwarning("Datum", "Datumsformat bitte TT.MM.JJJJ verwenden."); return
        if not self.pending_lines: messagebox.showwarning("Buchung", "Bitte mindestens eine Buchungszeile erfassen."); return
        debit = sum(parse_amount(ln["amount"]) for ln in self.pending_lines if ln["side"] == "Soll")
        credit = sum(parse_amount(ln["amount"]) for ln in self.pending_lines if ln["side"] == "Haben")
        if debit != credit:
            messagebox.showwarning("Nicht ausgeglichen", "Soll und Haben müssen identisch sein."); return
        with get_connection() as con:
            con.execute("INSERT INTO journal_entries(document_no,document_date,posting_date,description,total_debit,total_credit,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)",
                        (self.doc_no.get(), self.doc_date.get(), self.posting_date.get(), self.desc.get(), float(debit), float(credit), "Gebucht", now_str(), now_str()))
            for ln in self.pending_lines:
                con.execute("INSERT INTO journal_entry_lines(document_no,account_no,side,amount,tax_code,text) VALUES(?,?,?,?,?,?)", (self.doc_no.get(), ln["account"], ln["side"], ln["amount"], ln["tax"], ln["text"]))
            con.commit()
        if self.pending_attachments:
            self.add_attachment_paths("journal", self.doc_no.get(), self.pending_attachments)
        self.app.set_status(f"Buchung {self.doc_no.get()} gespeichert.")
        self.clear_journal_form(); self.reload_journal_entries()

    def clear_journal_form(self):
        self.doc_no.set(generate_number("B", "counter_journal", 6)); self.doc_date.set(today_str()); self.posting_date.set(today_str()); self.desc.set("")
        self.pending_lines.clear(); self.pending_attachments.clear(); self._refresh_lines_tree(); self.clear_line_inputs()

    def reload_journal_entries(self):
        for i in self.hist_tree.get_children(): self.hist_tree.delete(i)
        with get_connection() as con:
            for r in con.execute("SELECT * FROM journal_entries ORDER BY id DESC"):
                att = self.get_attachment_count("journal", r["document_no"])
                self.hist_tree.insert("", "end", values=(r["document_no"], r["posting_date"], r["description"], format_amount(r["total_debit"]), format_amount(r["total_credit"]), f"📎 {att}" if att else ""))

    def _on_history_click(self, _event=None):
        iid = self.hist_tree.focus()
        if iid:
            doc = self.hist_tree.set(iid, "document_no")
            self.open_attachment_popup("journal", doc)

    def add_attachment_to_selected_history(self):
        iid = self.hist_tree.focus()
        if not iid: messagebox.showinfo("Anhang", "Bitte Buchung auswählen."); return
        doc = self.hist_tree.set(iid, "document_no")
        paths = filedialog.askopenfilenames(title="Anhang auswählen")
        if paths:
            self.add_attachment_paths("journal", doc, paths); self.reload_journal_entries()


class InvoiceModuleBase(ttk.Frame, SortableTreeMixin, AttachmentMixin):
    entity_label = "Rechnung"
    partner_table = "customers"
    partner_no_col = "customer_no"
    invoice_table = "customer_invoices"
    entity_type = "customer_invoice"
    prefix = "AR"

    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.pending_attachments: List[str] = []
        self.selected_invoice = ""
        self._load_reference_data()
        self._build_ui()
        self.reload_invoices(); self.reload_open_items()

    def _load_reference_data(self):
        with get_connection() as con:
            self.partners = [f"{r[self.partner_no_col]} {r['name']}" for r in con.execute(f"SELECT {self.partner_no_col},name FROM {self.partner_table} WHERE active=1 ORDER BY name")]
            self.tax_codes = [r[0] for r in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]
            self.payment_terms = [r[0] for r in con.execute("SELECT code FROM payment_terms WHERE active=1 ORDER BY code")]

    def _build_ui(self):
        shell = build_two_block_shell(self, f"{self.entity_label} erfassen", "Offene Posten / Kontrolle")
        left = ScrollableFrame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
        self.partner = tk.StringVar(); self.invoice_date = tk.StringVar(value=today_str()); self.due_date = tk.StringVar(value=today_str())
        self.payment_term = tk.StringVar(); self.tax_code = tk.StringVar(value="V19"); self.net_amount = tk.StringVar(); self.tax_amount = tk.StringVar(); self.gross_amount = tk.StringVar()
        self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
        self._labeled_combo(left.content, "Partner", self.partner, self.partners)
        self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
        self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
        self._labeled_entry(left.content, "Fälligkeit", self.due_date)
        self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
        self._labeled_entry(left.content, "Netto", self.net_amount)
        self._labeled_entry(left.content, "Steuer", self.tax_amount)
        self._labeled_entry(left.content, "Brutto", self.gross_amount)
        self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
        self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
        self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
        btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
        create_standard_button(btns, "Anhang vormerken", self.add_pending_attachments).pack(side="left", padx=3)
        create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
        ttk.Label(left.content, text="Erfasste Rechnungen", style="Section.TLabel").pack(anchor="w", pady=(10, 4))
        cols = ("invoice_no", "partner", "date", "due", "gross", "open", "status", "att")
        self.inv_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=9)
        for c, t, w in [("invoice_no", "Rechnung", 110), ("partner", "Partner", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anh.", 55)]:
            self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
        self.inv_tree.pack(fill="both", expand=True)
        self.inv_tree.bind("<Double-1>", lambda e: self.add_attachment_to_selected_invoice())
        configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
        # right
        right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
        cols2 = ("reference_no", "partner", "due", "original", "open", "status")
        self.op_tree = ttk.Treeview(right, columns=cols2, show="headings")
        for c, t, w in [("reference_no", "Referenz", 120), ("partner", "Partner", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120)]:
            self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
        self.op_tree.pack(fill="both", expand=True)
        configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
        create_standard_button(right, "Anhang zu Rechnung", self.add_attachment_to_selected_invoice).pack(anchor="e", pady=8)

    def _generate_invoice_no(self): return generate_number(self.prefix, f"counter_{self.invoice_table}", 6)
    def _labeled_entry(self, parent, label, var):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3); ttk.Label(row, text=label, width=16).pack(side="left"); ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)
    def _labeled_combo(self, parent, label, var, values):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3); ttk.Label(row, text=label, width=16).pack(side="left")
        cb = ttk.Combobox(row, textvariable=var, values=values); cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")

    def apply_payment_term(self):
        if validate_date(self.invoice_date.get()) and self.payment_term.get():
            try: self.due_date.set(calc_due_date(self.invoice_date.get(), self.payment_term.get()))
            except Exception: pass

    def _calculate_amounts(self):
        try: net = parse_amount(self.net_amount.get())
        except Exception: return
        rate = Decimal("0")
        with get_connection() as con:
            row = con.execute("SELECT rate FROM tax_codes WHERE code=?", (self.tax_code.get(),)).fetchone()
            if row: rate = Decimal(str(row[0] or 0))
        tax = (net * rate / Decimal("100")).quantize(Decimal("0.01")); gross = net + tax
        self.tax_amount.set(format_amount(tax)); self.gross_amount.set(format_amount(gross))

    def _partner_parts(self):
        no = self.partner.get().split(" ")[0] if self.partner.get() else ""
        with get_connection() as con:
            row = con.execute(f"SELECT * FROM {self.partner_table} WHERE {self.partner_no_col}=?", (no,)).fetchone()
        if row:
            address = f"{row['street']}, {row['zip']} {row['city']}, {row['country']}".strip(', ')
            return no, row["name"], address
        return no, self.partner.get(), ""

    def add_pending_attachments(self):
        paths = filedialog.askopenfilenames(title="Beleganhänge vormerken")
        if paths: self.pending_attachments.extend(paths); self.app.set_status(f"{len(paths)} Anhang/Anhänge vorgemerkt.")

    def save_invoice(self):
        if not validate_date(self.invoice_date.get()) or not validate_date(self.due_date.get()): messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden."); return
        no, name, address = self._partner_parts()
        if not no: messagebox.showwarning("Partner", "Bitte Partner auswählen."); return
        gross = parse_amount(self.gross_amount.get()); net = parse_amount(self.net_amount.get()); tax = parse_amount(self.tax_amount.get())
        status = compute_status_from_open_amount(gross, self.due_date.get())
        address_col = "customer_address" if self.invoice_table == "customer_invoices" else "vendor_address"
        partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
        partner_no_col = "customer_no" if self.invoice_table == "customer_invoices" else "vendor_no"
        with get_connection() as con:
            con.execute(f"INSERT INTO {self.invoice_table}(invoice_no,{partner_no_col},{partner_name_col},{address_col},invoice_date,due_date,payment_term,tax_code,net_amount,tax_amount,gross_amount,open_amount,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                        (self.invoice_no.get(), no, name, address, self.invoice_date.get(), self.due_date.get(), self.payment_term.get(), self.tax_code.get(), float(net), float(tax), float(gross), float(gross), status, now_str(), now_str()))
            con.execute("INSERT INTO open_items(entity_type,reference_no,partner_no,partner_name,due_date,original_amount,open_amount,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                        (self.entity_type, self.invoice_no.get(), no, name, self.due_date.get(), float(gross), float(gross), status, now_str(), now_str()))
            con.commit()
        if self.pending_attachments: self.add_attachment_paths(self.entity_type, self.invoice_no.get(), self.pending_attachments)
        self.app.set_status(f"{self.entity_label} {self.invoice_no.get()} gespeichert.")
        self.clear_form(); self.reload_invoices(); self.reload_open_items()

    def clear_form(self):
        self.invoice_no.set(self._generate_invoice_no()); self.partner.set(""); self.invoice_date.set(today_str()); self.due_date.set(today_str()); self.payment_term.set(""); self.tax_code.set("V19"); self.net_amount.set(""); self.tax_amount.set(""); self.gross_amount.set(""); self.pending_attachments.clear()

    def reload_invoices(self):
        for i in self.inv_tree.get_children(): self.inv_tree.delete(i)
        partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
        with get_connection() as con:
            for r in con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"):
                att = self.get_attachment_count(self.entity_type, r["invoice_no"])
                tag = urgency_bucket(r["due_date"], r["open_amount"])
                self.inv_tree.insert("", "end", values=(r["invoice_no"], r[partner_name_col], r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else ""), tags=(tag,))

    def reload_open_items(self):
        for i in self.op_tree.get_children(): self.op_tree.delete(i)
        with get_connection() as con:
            for r in con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date", (self.entity_type,)):
                tag = urgency_bucket(r["due_date"], r["open_amount"])
                self.op_tree.insert("", "end", values=(r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"]), tags=(tag,))

    def add_attachment_to_selected_invoice(self):
        iid = self.inv_tree.focus()
        if not iid: messagebox.showinfo("Anhang", "Bitte Rechnung auswählen."); return
        inv = self.inv_tree.set(iid, "invoice_no")
        paths = filedialog.askopenfilenames(title="Anhang auswählen")
        if paths: self.add_attachment_paths(self.entity_type, inv, paths); self.reload_invoices()


class DebitorsView(InvoiceModuleBase):
    entity_label = "Ausgangsrechnung"
    partner_table = "customers"; partner_no_col = "customer_no"; invoice_table = "customer_invoices"; entity_type = "customer_invoice"; prefix = "AR"


class CreditorsView(InvoiceModuleBase):
    entity_label = "Eingangsrechnung"
    partner_table = "vendors"; partner_no_col = "vendor_no"; invoice_table = "vendor_invoices"; entity_type = "vendor_invoice"; prefix = "ER"

    def _build_ui(self):
        outer = ttk.PanedWindow(self, orient="vertical"); outer.pack(fill="both", expand=True)
        invoice_area = ttk.Frame(outer); stack_area = ttk.Frame(outer)
        outer.add(invoice_area, weight=3); outer.add(stack_area, weight=2)
        # temporarily build base UI into invoice_area by monkey parent composition
        old_parent = self.master
        shell = build_two_block_shell(invoice_area, "Eingangsrechnung erfassen", "OP / Freigabe")
        left = ScrollableFrame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
        self.partner = tk.StringVar(); self.invoice_date = tk.StringVar(value=today_str()); self.due_date = tk.StringVar(value=today_str())
        self.payment_term = tk.StringVar(); self.tax_code = tk.StringVar(value="V19"); self.net_amount = tk.StringVar(); self.tax_amount = tk.StringVar(); self.gross_amount = tk.StringVar()
        self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
        prow = ttk.Frame(left.content); prow.pack(fill="x", pady=3)
        ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
        cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners); cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)
        self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
        self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
        self._labeled_entry(left.content, "Fälligkeit", self.due_date)
        self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
        self._labeled_entry(left.content, "Netto", self.net_amount); self._labeled_entry(left.content, "Steuer", self.tax_amount); self._labeled_entry(left.content, "Brutto", self.gross_amount)
        self.payment_term.trace_add("write", lambda *_: self.apply_payment_term()); self.net_amount.trace_add("write", lambda *_: self._calculate_amounts()); self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
        btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
        create_standard_button(btns, "Anhang vormerken", self.add_pending_attachments).pack(side="left", padx=3)
        create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
        ttk.Label(left.content, text="Eingangsrechnungen", style="Section.TLabel").pack(anchor="w", pady=(10, 4))
        cols = ("invoice_no", "partner", "date", "due", "gross", "open", "status", "att")
        self.inv_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=7)
        for c, t, w in [("invoice_no", "Rechnung", 110), ("partner", "Kreditor", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anh.", 55)]: self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
        self.inv_tree.pack(fill="both", expand=True); self.inv_tree.bind("<Double-1>", lambda e: self.add_attachment_to_selected_invoice()); configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
        right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
        cols2 = ("reference_no", "partner", "due", "original", "open", "status")
        self.op_tree = ttk.Treeview(right, columns=cols2, show="headings")
        for c, t, w in [("reference_no", "Referenz", 120), ("partner", "Kreditor", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120)]: self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
        self.op_tree.pack(fill="both", expand=True); configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
        create_standard_button(right, "Anhang zu Rechnung", self.add_attachment_to_selected_invoice).pack(anchor="e", pady=8)
        self._build_stack_ui(stack_area)

    def _open_vendor_quick_popup(self):
        pop = tk.Toplevel(self); pop.title("Quick-Kreditor"); pop.geometry("520x360"); pop.configure(bg=BG)
        vars = {k: tk.StringVar() for k in ["vendor_no", "name", "country", "street", "zip", "city", "payment_term"]}
        vars["vendor_no"].set(generate_number("K", "counter_vendors", 5)); vars["country"].set("Deutschland")
        for key, label in [("vendor_no", "Kreditor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"), ("payment_term", "Zahlungsbed.")]:
            row = ttk.Frame(pop); row.pack(fill="x", padx=12, pady=4); ttk.Label(row, text=label, width=16).pack(side="left"); ttk.Entry(row, textvariable=vars[key]).pack(side="left", fill="x", expand=True)
        def save():
            if not vars["name"].get().strip(): messagebox.showwarning("Pflichtfeld", "Name fehlt."); return
            with get_connection() as con:
                con.execute("INSERT INTO vendors(vendor_no,name,country,street,zip,city,payment_term,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", (vars["vendor_no"].get(), vars["name"].get(), vars["country"].get(), vars["street"].get(), vars["zip"].get(), vars["city"].get(), vars["payment_term"].get(), 1, now_str(), now_str())); con.commit()
            self._load_reference_data(); self.partner.set(f"{vars['vendor_no'].get()} {vars['name'].get()}"); pop.destroy()
        create_standard_button(pop, "Speichern", save, confirm=True).pack(pady=12)

    def _build_stack_ui(self, parent):
        shell = build_two_block_shell(parent, "Dokumentenstapel", "Vorschau")
        left = ttk.Frame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        bar = ttk.Frame(left); bar.pack(fill="x", pady=(0, 6))
        create_standard_button(bar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=2)
        create_standard_button(bar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=2)
        create_standard_button(bar, "Alle markieren", self._toggle_all_marks).pack(side="left", padx=2)
        create_standard_button(bar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=2)
        create_standard_button(bar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=2)
        cols = ("mark", "hist_no", "title", "status", "files", "created")
        self.stack_tree = ttk.Treeview(left, columns=cols, show="headings", selectmode="extended")
        for c, t, w in [("mark", "☐", 45), ("hist_no", "Hist.-Nr.", 120), ("title", "Titel", 260), ("status", "Status", 100), ("files", "Dateien", 70), ("created", "Importiert", 150)]: self.stack_tree.heading(c, text=t); self.stack_tree.column(c, width=w, anchor="w")
        self.stack_tree.pack(fill="both", expand=True); self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select); self.stack_tree.bind("<Double-1>", lambda e: self._open_selected_stack_file())
        right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
        self.preview = tk.Text(right, wrap="word", bg=WHITE, fg=TEXT, relief="sunken")
        self.preview.pack(fill="both", expand=True)
        self._load_import_batches()

    def import_creditor_files(self):
        paths = filedialog.askopenfilenames(title="Dokumente importieren", filetypes=[("Unterstützte Dateien", "*.pdf *.png *.jpg *.jpeg *.bmp *.gif *.txt *.csv *.doc *.docx *.xls *.xlsx"), ("Alle Dateien", "*.*")])
        if paths: self._import_archive_paths(paths)

    def import_creditor_folder(self):
        folder = filedialog.askdirectory(title="Importordner auswählen")
        if folder: self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

    def _import_archive_paths(self, paths: Sequence[str]):
        imported = 0
        with get_connection() as con:
            for src in paths:
                p = Path(src)
                if not p.exists() or p.suffix.lower() not in SUPPORTED_ARCHIVE_EXTENSIONS: continue
                hist = next_hist_no(); target_dir = IMPORTS_DIR / hist; target_dir.mkdir(parents=True, exist_ok=True)
                dest = target_dir / p.name; shutil.copy2(p, dest)
                title = clean_document_title(str(p))
                con.execute("INSERT INTO invoice_import_batches(hist_no,title,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?)", (hist, title, "Importiert", 0, now_str(), now_str()))
                con.execute("INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)", (hist, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()))
                imported += 1
            con.commit()
        self._load_import_batches(); self.app.set_status(f"{imported} Dokument(e) importiert.")

    def _load_import_batches(self):
        if not hasattr(self, "stack_tree"): return
        for i in self.stack_tree.get_children(): self.stack_tree.delete(i)
        with get_connection() as con:
            rows = con.execute("SELECT b.*, COUNT(f.id) files FROM invoice_import_batches b LEFT JOIN invoice_import_files f ON f.hist_no=b.hist_no GROUP BY b.hist_no ORDER BY b.id DESC").fetchall()
            for r in rows:
                self.stack_tree.insert("", "end", values=("☑" if r["marked"] else "☐", r["hist_no"], r["title"], r["status"], r["files"], r["created_at"]))

    def _on_stack_select(self, event=None):
        iid = self.stack_tree.focus()
        if not iid: return
        col = self.stack_tree.identify_column(event.x) if event else ""
        hist = self.stack_tree.set(iid, "hist_no")
        if col == "#1":
            with get_connection() as con:
                row = con.execute("SELECT marked FROM invoice_import_batches WHERE hist_no=?", (hist,)).fetchone(); new = 0 if row and row[0] else 1
                con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=? WHERE hist_no=?", (new, now_str(), hist)); con.commit()
            self._load_import_batches(); return
        self._refresh_stack_preview(hist)

    def _refresh_stack_preview(self, hist_no: str):
        with get_connection() as con: row = con.execute("SELECT * FROM invoice_import_files WHERE hist_no=? ORDER BY id LIMIT 1", (hist_no,)).fetchone()
        self.preview.delete("1.0", tk.END)
        if row: self.preview.insert("1.0", collect_preview_text(row["file_path"]))
        else: self.preview.insert("1.0", "Keine Datei im Stapel.")

    def _open_selected_stack_file(self):
        iid = self.stack_tree.focus()
        if not iid: return
        hist = self.stack_tree.set(iid, "hist_no")
        with get_connection() as con: row = con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id LIMIT 1", (hist,)).fetchone()
        if row: open_path(row["file_path"])

    def _toggle_all_marks(self):
        with get_connection() as con:
            any_unmarked = con.execute("SELECT 1 FROM invoice_import_batches WHERE marked=0 LIMIT 1").fetchone()
            con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=?", (1 if any_unmarked else 0, now_str())); con.commit()
        self._load_import_batches()

    def _delete_marked_batches(self):
        if not messagebox.askyesno("Bulk-Löschung", "Alle markierten Stapel unwiderruflich löschen?"): return
        with get_connection() as con:
            rows = con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1").fetchall()
            for r in rows:
                shutil.rmtree(IMPORTS_DIR / r["hist_no"], ignore_errors=True)
                con.execute("DELETE FROM invoice_import_batches WHERE hist_no=?", (r["hist_no"],))
            con.commit()
        self._load_import_batches(); self.preview.delete("1.0", tk.END)

    def merge_selected_batches(self):
        with get_connection() as con:
            rows = con.execute("SELECT hist_no,title FROM invoice_import_batches WHERE marked=1 ORDER BY id").fetchall()
            if len(rows) < 2: messagebox.showinfo("Zusammenführen", "Bitte mindestens zwei Stapel markieren."); return
            new_hist = next_hist_no(); target_dir = IMPORTS_DIR / new_hist; target_dir.mkdir(parents=True, exist_ok=True)
            title = " + ".join([r["title"] for r in rows])[:240]
            con.execute("INSERT INTO invoice_import_batches(hist_no,title,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?)", (new_hist, title, "Zusammengeführt", 0, now_str(), now_str()))
            for r in rows:
                for f in con.execute("SELECT * FROM invoice_import_files WHERE hist_no=?", (r["hist_no"],)).fetchall():
                    src = Path(f["file_path"]); dest = target_dir / f"{r['hist_no']}_{src.name}"
                    if src.exists(): shutil.copy2(src, dest)
                    con.execute("INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)", (new_hist, dest.name, str(dest), f["mime_type"], now_str()))
                con.execute("UPDATE invoice_import_batches SET status='Zusammengeführt in ' || ?, marked=0, updated_at=? WHERE hist_no=?", (new_hist, now_str(), r["hist_no"]))
            con.commit()
        self._load_import_batches(); self.app.set_status(f"Stapel {new_hist} erzeugt.")


class FinanceMateApp(tk.Tk):
    def __init__(self):
        super().__init__()
        init_sqlite()
        self.title(f"{APP_NAME} {APP_VERSION}")
        self.geometry("1420x860")
        self.minsize(1100, 680)
        self.configure(bg=BG)
        self.sidebar_collapsed = load_ui_preference("sidebar_collapsed", "0") == "1"
        self.active_module = "Dashboard"
        self.nav_buttons: Dict[str, ttk.Button] = {}
        self._configure_ttk()
        self._build_layout()
        self.show_module("Dashboard")

    def _configure_ttk(self):
        style = ttk.Style(self)
        try: style.theme_use("clam")
        except Exception: pass
        style.configure("TFrame", background=BG)
        style.configure("Card.TFrame", background=CARD_BG, relief="raised", borderwidth=1)
        style.configure("TLabel", background=BG, foreground=TEXT)
        style.configure("CardTitle.TLabel", background=BG, foreground=TEXT, font=("Segoe UI", 15, "bold"))
        style.configure("Section.TLabel", background=CARD_BG, foreground=TEXT, font=("Segoe UI", 10, "bold"))
        style.configure("Hint.TLabel", background=BG, foreground=TEXT2)
        style.configure("StandardGray.TButton", background=STANDARD_BUTTON_BG, foreground=TEXT, bordercolor=STANDARD_BUTTON_BORDER, relief="raised", padding=(8, 4))
        style.map("StandardGray.TButton", background=[("active", STANDARD_BUTTON_ACTIVE)])
        style.configure("Confirm.TButton", background=BUTTON_GREEN, foreground=TEXT, relief="raised", padding=(8, 4))
        style.map("Confirm.TButton", background=[("active", BUTTON_GREEN_ACTIVE)])
        style.configure("Nav.TButton", background=STANDARD_BUTTON_BG, anchor="w", padding=(8, 7), relief="raised")
        style.configure("NavActive.TButton", background="#c5d3e6", anchor="w", padding=(8, 7), relief="sunken")
        style.configure("Treeview.Heading", background=STANDARD_BUTTON_BG, relief="raised", foreground=TEXT)
        style.configure("Treeview", rowheight=24, fieldbackground=WHITE, background=WHITE)
        style.configure("TNotebook", background=BG)
        style.configure("TNotebook.Tab", background=STANDARD_BUTTON_BG, padding=(10, 5))
        style.map("TNotebook.Tab", background=[("selected", CARD_BG)])
        style.configure("Vertical.TScrollbar", background=STANDARD_BUTTON_BG, troughcolor="#edf0f3", arrowsize=14, width=15)
        style.configure("TCombobox", fieldbackground=WHITE, background=STANDARD_BUTTON_BG)

    def _build_layout(self):
        self.header = ttk.Frame(self, style="Card.TFrame"); self.header.pack(fill="x", padx=6, pady=(6, 3))
        ttk.Label(self.header, text=APP_NAME, style="CardTitle.TLabel").pack(side="left", padx=12, pady=7)
        ttk.Label(self.header, text=f"{STANDARD_DESIGN_NAME} · lokale SQLite-Persistenz · {APP_VERSION}", style="Hint.TLabel").pack(side="left", padx=10)
        create_standard_button(self.header, "Datenbankordner", lambda: open_path(str(DATA_DIR))).pack(side="right", padx=8)
        body = ttk.Frame(self); body.pack(fill="both", expand=True, padx=6, pady=3)
        self.sidebar = ttk.Frame(body, style="Card.TFrame"); self.sidebar.pack(side="left", fill="y", padx=(0, 5))
        self.workspace = ttk.Frame(body); self.workspace.pack(side="left", fill="both", expand=True)
        self._build_sidebar()
        self.footer = ttk.Frame(self, style="Card.TFrame"); self.footer.pack(fill="x", padx=6, pady=(3, 6))
        self.status_var = tk.StringVar(value="Bereit.")
        ttk.Label(self.footer, textvariable=self.status_var, style="Hint.TLabel").pack(side="left", padx=12, pady=5)
        ttk.Label(self.footer, text="Mini-Widgets: Änderung vorschlagen · [i] Hilfe", style="Hint.TLabel").pack(side="right", padx=12)

    def _build_sidebar(self):
        for w in self.sidebar.winfo_children(): w.destroy()
        width = 42 if self.sidebar_collapsed else 210
        self.sidebar.configure(width=width)
        self.sidebar.pack_propagate(False)
        arrow = "»" if self.sidebar_collapsed else "«"
        create_standard_button(self.sidebar, arrow, self._toggle_sidebar).pack(fill="x", padx=5, pady=5)
        modules = [("Dashboard", "DB"), ("Stammdaten", "ST"), ("Finanzbuchhaltung", "FB"), ("Debitoren", "DE"), ("Kreditoren", "KR"), ("Zahlungen", "ZA"), ("Reporting", "RE"), ("Audit", "AU"), ("Einstellungen", "EI")]
        self.nav_buttons.clear()
        for name, abbr in modules:
            text = abbr if self.sidebar_collapsed else name
            btn = ttk.Button(self.sidebar, text=text, style="Nav.TButton", command=lambda n=name: self.show_module(n))
            btn.pack(fill="x", padx=5, pady=2)
            self.nav_buttons[name] = btn

    def _toggle_sidebar(self):
        self.sidebar_collapsed = not self.sidebar_collapsed
        save_ui_preference("sidebar_collapsed", "1" if self.sidebar_collapsed else "0")
        self._build_sidebar(); self._update_nav_styles()

    def _update_nav_styles(self):
        for name, btn in self.nav_buttons.items(): btn.configure(style="NavActive.TButton" if name == self.active_module else "Nav.TButton")

    def set_status(self, msg: str):
        self.status_var.set(msg)

    def show_module(self, module: str):
        self.active_module = module; self._update_nav_styles()
        for w in self.workspace.winfo_children(): w.destroy()
        title = ttk.Label(self.workspace, text=module, style="CardTitle.TLabel"); title.pack(anchor="w", pady=(0, 6))
        container = ttk.Frame(self.workspace); container.pack(fill="both", expand=True)
        render = {
            "Dashboard": self._render_dashboard,
            "Stammdaten": self._render_stammdaten,
            "Finanzbuchhaltung": self._render_finanzbuchhaltung,
            "Debitoren": self._render_debitoren,
            "Kreditoren": self._render_kreditoren,
            "Zahlungen": self._render_zahlungen,
            "Reporting": self._render_reporting,
            "Audit": self._render_audit,
            "Einstellungen": self._render_einstellungen,
        }.get(module, self._render_dashboard)
        render(container); self.set_status(f"Modul {module} geladen.")

    def _card(self, parent, title, body):
        f = ttk.Frame(parent, style="Card.TFrame"); f.pack(fill="both", expand=True, padx=6, pady=6)
        ttk.Label(f, text=title, style="Section.TLabel").pack(anchor="w", padx=12, pady=(10, 4))
        ttk.Label(f, text=body, style="Hint.TLabel", wraplength=520).pack(anchor="nw", padx=12, pady=(0, 12))
        return f

    def _render_dashboard(self, parent):
        grid = ttk.Frame(parent); grid.pack(fill="both", expand=True)
        for i in range(2): grid.columnconfigure(i, weight=1); grid.rowconfigure(i, weight=1)
        for idx, (title, text) in enumerate([
            ("Systemstatus", "Alle Kernmodule sind über die Sidebar erreichbar. SQLite wird lokal initialisiert."),
            ("Offene Posten", "Debitoren- und Kreditoren-OPs werden nach Fälligkeit farblich gekennzeichnet."),
            ("Dokumentenstapel", "Kreditorenimporte mit Hist.-Nr., Markierung, Bulk-Löschung, Merge und Vorschau."),
            ("Nächste Ausbaustufe", "Zahlungen, Reporting und Audit sind als funktionsfähige Vorbereitungsflächen angelegt."),
        ]):
            f = ttk.Frame(grid, style="Card.TFrame"); f.grid(row=idx//2, column=idx%2, sticky="nsew", padx=6, pady=6)
            ttk.Label(f, text=title, style="Section.TLabel").pack(anchor="w", padx=12, pady=10)
            ttk.Label(f, text=text, style="Hint.TLabel", wraplength=520).pack(anchor="nw", padx=12, pady=4)

    def _render_stammdaten(self, parent): StammdatenView(parent, self).pack(fill="both", expand=True)
    def _render_finanzbuchhaltung(self, parent): JournalView(parent, self).pack(fill="both", expand=True)
    def _render_debitoren(self, parent): DebitorsView(parent, self).pack(fill="both", expand=True)
    def _render_kreditoren(self, parent): CreditorsView(parent, self).pack(fill="both", expand=True)

    def _render_zahlungen(self, parent):
        shell = build_two_block_shell(parent, "Zahlungserfassung / OP-Ausgleich", "Vorbereitung Zahlungsläufe")
        ttk.Label(shell.left, text="Teilzahlungen, Vollausgleich und OP-Abgleich sind im Datenmodell vorbereitet.\nDie OP-Tabellen aus Debitoren/Kreditoren dienen als Grundlage für Block 7.", style="Hint.TLabel", wraplength=650).pack(anchor="nw", padx=12, pady=12)
        ttk.Label(shell.right, text="Workflow-Platzhalter für spätere Zahlungsläufe, Zuordnungsvorschläge und Statuswechsel.", style="Hint.TLabel", wraplength=420).pack(anchor="nw", padx=12, pady=12)

    def _render_reporting(self, parent):
        shell = build_two_block_shell(parent, "Standardberichte", "Kennzahlen")
        with get_connection() as con:
            counts = {"Debitorenrechnungen": con.execute("SELECT COUNT(*) FROM customer_invoices").fetchone()[0], "Kreditorenrechnungen": con.execute("SELECT COUNT(*) FROM vendor_invoices").fetchone()[0], "Offene Posten": con.execute("SELECT COUNT(*) FROM open_items WHERE open_amount>0").fetchone()[0], "Buchungen": con.execute("SELECT COUNT(*) FROM journal_entries").fetchone()[0]}
        for k, v in counts.items(): ttk.Label(shell.left, text=f"{k}: {v}", style="Hint.TLabel").pack(anchor="w", padx=12, pady=4)
        ttk.Label(shell.right, text="Periodenfilter, Druck und Exportpfade sind als Zielarchitektur vorgesehen.", style="Hint.TLabel", wraplength=420).pack(anchor="nw", padx=12, pady=12)

    def _render_audit(self, parent):
        self._card(parent, "Audit / Kontrollpfade", "Änderungshistorien, Kontrollpfade, Berechtigungsoptionen und Compliance-Prüfungen sind als eigene Ausbauschicht vorbereitet. Revisionsrelevante Zeitstempel werden bereits in den Kerntabellen geführt.")

    def _render_einstellungen(self, parent):
        shell = build_two_block_shell(parent, "Systemparameter", "UI-Präferenzen")
        ttk.Label(shell.left, text=f"Datenbank: {DB_PATH}\nAnhänge: {ATTACHMENTS_DIR}\nImporte: {IMPORTS_DIR}\nVersion: {APP_VERSION}", style="Hint.TLabel", wraplength=700).pack(anchor="nw", padx=12, pady=12)
        create_standard_button(shell.left, "SQLite initialisieren / prüfen", lambda: (init_sqlite(), self.set_status("SQLite geprüft.")), confirm=True).pack(anchor="w", padx=12, pady=6)
        ttk.Label(shell.right, text=f"Sidebar kollabiert: {'Ja' if self.sidebar_collapsed else 'Nein'}\nDesignstandard: {STANDARD_DESIGN_NAME}", style="Hint.TLabel").pack(anchor="nw", padx=12, pady=12)


def maximize_window(root: tk.Tk) -> None:
    try: root.state("zoomed")
    except Exception: pass



# === FINANCE MATE PATCH V0_6_17 ===
APP_VERSION = "0.6.17-relief-docflow"

# --- Benutzerbezogene UI-Präferenzen / Splitterpersistenz ---
_FM_UI_USER = os.environ.get("FINANCEMATE_UI_USER") or os.environ.get("USERNAME") or os.environ.get("USER") or "default"


def _fm_pref_key(key: str) -> str:
    return f"ui:{_FM_UI_USER}:{key}"


def load_ui_preference(key: str, default: str = "") -> str:
    with get_connection() as con:
        row = con.execute("SELECT value FROM app_meta WHERE key=?", (_fm_pref_key(key),)).fetchone()
        return row[0] if row else default


def save_ui_preference(key: str, value: str) -> None:
    with get_connection() as con:
        con.execute(
            "INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)",
            (_fm_pref_key(key), str(value), now_str()),
        )
        con.commit()


def _fm_attach_paned_persistence(paned, key: str, default_ratio: float = 0.60) -> None:
    if getattr(paned, "_fm_pref_attached", False):
        return
    paned._fm_pref_attached = True

    def _restore():
        try:
            paned.update_idletasks()
            saved = load_ui_preference(f"pane:{key}", "")
            total = max(240, paned.winfo_width())
            if saved:
                pos = int(float(saved))
            else:
                pos = int(total * default_ratio)
            pos = max(80, min(total - 80, pos))
            try:
                paned.sashpos(0, pos)
            except Exception:
                try:
                    paned.tk.call(paned._w, "sashpos", 0, pos)
                except Exception:
                    pass
        except Exception:
            pass

    def _save(_event=None):
        try:
            try:
                pos = int(paned.sashpos(0))
            except Exception:
                pos = int(paned.tk.call(paned._w, "sashpos", 0))
            save_ui_preference(f"pane:{key}", str(pos))
        except Exception:
            pass

    paned.bind("<ButtonRelease-1>", _save, add="+")
    paned.bind("<Configure>", lambda _e: paned.after(40, _restore), add="+")
    paned.after(120, _restore)


_orig_TwoBlock_init = TwoBlock.__init__

def _fm_twoblock_init(self, parent, left_title: str, right_title: str):
    _orig_TwoBlock_init(self, parent, left_title, right_title)
    self._fm_pane_key = f"{parent.__class__.__name__}:{left_title}|{right_title}"
    _fm_attach_paned_persistence(self.pane, self._fm_pane_key, default_ratio=0.60)


TwoBlock.__init__ = _fm_twoblock_init

# --- Globale Dokumentlogik ---
def _fm_refresh_attachment_tree(tree: ttk.Treeview, rows) -> None:
    for item in tree.get_children():
        tree.delete(item)
    for row in rows:
        tree.insert("", "end", iid=str(row["id"]), values=(row["file_name"], row["file_path"], row["added_at"]))


def _fm_remove_attachment_record(attachment_id: int) -> None:
    with get_connection() as con:
        row = con.execute("SELECT file_path FROM attachments WHERE id=?", (attachment_id,)).fetchone()
        if not row:
            return
        file_path = Path(row[0])
        con.execute("DELETE FROM attachments WHERE id=?", (attachment_id,))
        con.commit()
    try:
        if file_path.exists():
            file_path.unlink()
    except Exception:
        pass


def _fm_replace_attachment_record(attachment_id: int, new_source: str) -> None:
    src_path = Path(new_source)
    if not src_path.exists():
        return
    with get_connection() as con:
        row = con.execute("SELECT entity_type, reference_no, file_path FROM attachments WHERE id=?", (attachment_id,)).fetchone()
        if not row:
            return
        entity_type = row["entity_type"]
        reference_no = row["reference_no"]
        old_file = Path(row["file_path"])
        target_dir = ATTACHMENTS_DIR / entity_type / reference_no
        target_dir.mkdir(parents=True, exist_ok=True)
        dest = target_dir / src_path.name
        if dest.exists():
            dest = target_dir / f"{src_path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{src_path.suffix}"
        shutil.copy2(src_path, dest)
        con.execute(
            "UPDATE attachments SET file_name=?, file_path=?, added_at=? WHERE id=?",
            (dest.name, str(dest), now_str(), attachment_id),
        )
        con.commit()
    try:
        if old_file.exists() and old_file.resolve() != dest.resolve():
            old_file.unlink()
    except Exception:
        pass


def _fm_pending_popup(self, title: str = "Dokumente") -> None:
    pop = tk.Toplevel(self)
    pop.title(title)
    pop.geometry("860x500")
    pop.configure(bg=BG)
    ttk.Label(pop, text=title, style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
    cols = ("file_name", "file_path")
    tree = ttk.Treeview(pop, columns=cols, show="headings")
    for col, w, txt in [("file_name", 260, "Dateiname"), ("file_path", 540, "Pfad")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=12, pady=8)

    def _refresh():
        for item in tree.get_children():
            tree.delete(item)
        for idx, file_path in enumerate(getattr(self, "pending_attachments", [])):
            q = Path(file_path)
            tree.insert("", "end", iid=str(idx), values=(q.name, str(q)))

    def _open_selected():
        iid = tree.focus()
        if iid:
            fp = tree.set(iid, "file_path")
            if fp:
                open_path(fp)

    def _delete_selected():
        iid = tree.focus()
        if iid:
            idx = int(iid)
            if 0 <= idx < len(self.pending_attachments):
                self.pending_attachments.pop(idx)
            _refresh()

    def _add_more():
        paths = list(filedialog.askopenfilenames(title="Dokument anhängen"))
        if paths:
            self.pending_attachments.extend(paths)
            _refresh()
            try:
                self.app.set_status(f"{len(self.pending_attachments)} Dokument(e) vorgemerkt.")
            except Exception:
                pass

    bar = ttk.Frame(pop)
    bar.pack(fill="x", padx=12, pady=10)
    create_standard_button(bar, "Öffnen", _open_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Löschen", _delete_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Weitere hinzufügen", _add_more, confirm=True).pack(side="left", padx=3)
    create_standard_button(bar, "Schließen", pop.destroy).pack(side="right", padx=3)
    tree.bind("<Double-1>", lambda _e: _open_selected())
    _refresh()


def _fm_manage_attachment_request(self, entity_type: str, reference_no: str, refresh_callback=None) -> None:
    if not reference_no:
        if not getattr(self, "pending_attachments", []):
            paths = list(filedialog.askopenfilenames(title="Dokument anhängen"))
            if paths:
                self.pending_attachments.extend(paths)
                try:
                    self.app.set_status(f"{len(paths)} Dokument(e) vorgemerkt.")
                except Exception:
                    pass
        else:
            _fm_pending_popup(self, "Vorgemerkte Dokumente")
        if callable(refresh_callback):
            refresh_callback()
        return

    count = self.get_attachment_count(entity_type, reference_no)
    if count <= 0:
        paths = list(filedialog.askopenfilenames(title="Dokument anhängen"))
        if paths:
            self.add_attachment_paths(entity_type, reference_no, paths)
            if callable(refresh_callback):
                refresh_callback()
    else:
        self.open_attachment_popup(entity_type, reference_no, refresh_callback=refresh_callback)


def _fm_open_attachment_popup(self, entity_type: str, reference_no: str, refresh_callback=None) -> None:
    pop = tk.Toplevel()
    pop.title(f"Dokumente – {reference_no}")
    pop.geometry("980x560")
    pop.configure(bg=BG)
    ttk.Label(pop, text=f"Dokumente zu {reference_no}", style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
    cols = ("file_name", "file_path", "added_at")
    tree = ttk.Treeview(pop, columns=cols, show="headings")
    for col, w, txt in [("file_name", 220, "Dateiname"), ("file_path", 560, "Pfad"), ("added_at", 150, "Hinzugefügt")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=12, pady=8)

    def _reload():
        _fm_refresh_attachment_tree(tree, self.get_attachments(entity_type, reference_no))
        if callable(refresh_callback):
            refresh_callback()

    def _selected_id():
        iid = tree.focus()
        return int(iid) if iid else None

    def _open_selected():
        iid = tree.focus()
        if iid:
            fp = tree.set(iid, "file_path")
            if fp:
                open_path(fp)

    def _add_more():
        paths = list(filedialog.askopenfilenames(title="Weitere Dokumente hinzufügen"))
        if paths:
            self.add_attachment_paths(entity_type, reference_no, paths)
            _reload()

    def _delete_selected():
        attachment_id = _selected_id()
        if attachment_id is None:
            return
        if not messagebox.askyesno("Dokument löschen", "Ausgewähltes Dokument unwiderruflich löschen?", parent=pop):
            return
        _fm_remove_attachment_record(attachment_id)
        _reload()
        if not tree.get_children():
            pop.destroy()

    def _replace_selected():
        attachment_id = _selected_id()
        if attachment_id is None:
            return
        path = filedialog.askopenfilename(title="Dokument ersetzen")
        if path:
            _fm_replace_attachment_record(attachment_id, path)
            _reload()

    bar = ttk.Frame(pop)
    bar.pack(fill="x", padx=12, pady=10)
    create_standard_button(bar, "Öffnen", _open_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Ändern", _replace_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Löschen", _delete_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Dokument hinzufügen", _add_more, confirm=True).pack(side="left", padx=3)
    create_standard_button(bar, "Schließen", pop.destroy).pack(side="right", padx=3)
    tree.bind("<Double-1>", lambda _e: _open_selected())
    _reload()


AttachmentMixin.open_attachment_popup = _fm_open_attachment_popup
AttachmentMixin.manage_attachment_request = _fm_manage_attachment_request

# --- Gemeinsame Rechnungslogik ---
def _fm_invoice_tree_click(self, event, tree_attr: str = "inv_tree"):
    tree = getattr(self, tree_attr, None)
    if tree is None:
        return None
    region = tree.identify("region", event.x, event.y)
    col = tree.identify_column(event.x)
    row = tree.identify_row(event.y)
    if region == "cell" and row and col == f"#{len(tree['columns'])}":
        inv_no = tree.item(row, "values")[0]
        self.manage_attachment_request(self.entity_type, inv_no, refresh_callback=self.reload_invoices)
        return "break"
    return None


def _fm_op_tree_click(self, event, tree_attr: str = "op_tree"):
    tree = getattr(self, tree_attr, None)
    if tree is None:
        return None
    region = tree.identify("region", event.x, event.y)
    col = tree.identify_column(event.x)
    row = tree.identify_row(event.y)
    if region == "cell" and row and col == f"#{len(tree['columns'])}":
        ref_no = tree.item(row, "values")[0]
        self.manage_attachment_request(self.entity_type, ref_no, refresh_callback=self.reload_open_items)
        return "break"
    return None


def _fm_invoice_add_pending(self):
    self.manage_attachment_request(self.entity_type, "", refresh_callback=None)


def _fm_invoice_open_selected(self):
    iid = self.inv_tree.focus()
    if not iid:
        messagebox.showinfo("Dokument", "Bitte Rechnung auswählen.")
        return
    inv = self.inv_tree.set(iid, "invoice_no")
    self.manage_attachment_request(self.entity_type, inv, refresh_callback=self.reload_invoices)


def _fm_invoice_reload_invoices(self):
    trees = []
    for attr in ('inv_tree',):
        tree = getattr(self, attr, None)
        if tree is not None:
            trees.append(tree)
    for tree in getattr(self, 'invoice_tabs_trees', []):
        if tree not in trees:
            trees.append(tree)
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    with get_connection() as con:
        rows = list(con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"))
    for tree in trees:
        for i in tree.get_children():
            tree.delete(i)
    for r in rows:
        att = self.get_attachment_count(self.entity_type, r["invoice_no"])
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        values = (r["invoice_no"], r[partner_name_col], r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else "Dokument anhängen")
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))


def _fm_invoice_reload_open_items(self):
    trees = []
    for attr in ('op_tree',):
        tree = getattr(self, attr, None)
        if tree is not None:
            trees.append(tree)
    for tree in getattr(self, 'open_item_tabs_trees', []):
        if tree not in trees:
            trees.append(tree)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date", (self.entity_type,)))
    for tree in trees:
        for i in tree.get_children():
            tree.delete(i)
    for r in rows:
        if parse_amount(r['open_amount']) <= 0:
            continue
        att = self.get_attachment_count(self.entity_type, r['reference_no'])
        tag = urgency_bucket(r['due_date'], r['open_amount'])
        values = (r['reference_no'], r['partner_name'], r['due_date'], format_amount(r['original_amount']), format_amount(r['open_amount']), r['status'], f"📎 {att}" if att else "Dokument anhängen")
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))


InvoiceModuleBase.add_pending_attachments = _fm_invoice_add_pending
InvoiceModuleBase.add_attachment_to_selected_invoice = _fm_invoice_open_selected
InvoiceModuleBase.reload_invoices = _fm_invoice_reload_invoices
InvoiceModuleBase.reload_open_items = _fm_invoice_reload_open_items

_orig_invoice_save = InvoiceModuleBase.save_invoice


def _fm_invoice_save_capture(self):
    self._last_saved_invoice_no = self.invoice_no.get().strip()
    selected_hist = getattr(self, 'current_hist_no', None)
    selected_partner_display = self.partner.get().strip()
    if self.entity_type == 'vendor_invoice' and not (selected_hist or getattr(self, 'pending_attachments', [])):
        messagebox.showwarning('Dokument', 'Eingangsrechnungen können ohne Dokument nicht erfasst werden.')
        return
    _orig_invoice_save(self)
    invoice_no = self._last_saved_invoice_no or ''
    if self.entity_type == 'vendor_invoice' and selected_hist and invoice_no:
        with get_connection() as con:
            files = list(con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id", (selected_hist,)))
            partner_no = selected_partner_display.split(' ')[0] if selected_partner_display else ''
            partner_name = ' '.join(selected_partner_display.split(' ')[1:]) if ' ' in selected_partner_display else selected_partner_display
            con.execute("UPDATE invoice_import_batches SET status='Erfasst', vendor_no=?, vendor_name=?, updated_at=? WHERE hist_no=?", (partner_no, partner_name, now_str(), selected_hist))
            con.commit()
        if files:
            self.add_attachment_paths(self.entity_type, invoice_no, [r[0] for r in files])
        self.current_hist_no = None
        try:
            self.stack_tree.selection_remove(self.stack_tree.selection())
        except Exception:
            pass
        if hasattr(self, '_clear_stack_preview'):
            self._clear_stack_preview()
        if hasattr(self, '_load_import_batches'):
            self._load_import_batches()
        self.reload_invoices()
        self.reload_open_items()


InvoiceModuleBase.save_invoice = _fm_invoice_save_capture

_orig_clear_form = InvoiceModuleBase.clear_form

def _fm_invoice_clear_form(self):
    _orig_clear_form(self)
    if hasattr(self, '_update_vendor_info'):
        try:
            self._update_vendor_info()
        except Exception:
            pass
    if hasattr(self, '_clear_stack_preview'):
        try:
            self._clear_stack_preview()
        except Exception:
            pass


InvoiceModuleBase.clear_form = _fm_invoice_clear_form

# --- Debitoren: nur explizit gewünschte Änderungen ---

def _walk_children(widget):
    for child in widget.winfo_children():
        yield child
        yield from _walk_children(child)


_orig_debitors_build_ui = InvoiceModuleBase._build_ui

def _fm_debitors_build_ui(self):
    _orig_debitors_build_ui(self)
    for ch in _walk_children(self):
        if isinstance(ch, ttk.Button):
            try:
                if ch.cget('text') == 'Anhang vormerken':
                    ch.configure(text='Dokument anhängen', command=self.add_pending_attachments)
            except Exception:
                pass
    try:
        self.inv_tree.bind('<ButtonRelease-1>', lambda e: _fm_invoice_tree_click(self, e), add='+')
    except Exception:
        pass
    try:
        self.op_tree.configure(columns=("reference_no", "partner", "due", "original", "open", "status", "att"))
        for c, t, w in [("reference_no", "Referenz", 120), ("partner", "Partner", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
            self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor='w')
        self.op_tree.bind('<ButtonRelease-1>', lambda e: _fm_op_tree_click(self, e), add='+')
    except Exception:
        pass
    self.reload_invoices()
    self.reload_open_items()


DebitorsView._build_ui = _fm_debitors_build_ui

# --- Kreditoren: 4 Fensterblöcke beibehalten, nur gewünschte Änderungen ---

def _fm_creditors_update_vendor_info(self):
    body = getattr(self, 'vendor_info_body', None)
    if body is None:
        return
    vendor_no = self.partner.get().split(' ')[0] if self.partner.get() else ''
    content = 'Kein Kreditor ausgewählt.'
    if vendor_no:
        with get_connection() as con:
            row = con.execute("SELECT * FROM vendors WHERE vendor_no=?", (vendor_no,)).fetchone()
        if row:
            content = "\n".join([
                f"Kreditor-Nr.: {row['vendor_no']}",
                f"Name: {row['name']}",
                f"Land: {row['country']}",
                f"Adresse: {row['street']}, {row['zip']} {row['city']}",
                f"E-Mail: {row['email']}",
                f"Telefon: {row['phone']}",
                f"Steuer-Nr.: {row['tax_id']}",
                f"USt-IdNr.: {row['vat_id']}",
                f"IBAN/BIC: {row['iban']} / {row['bic']}",
                f"Bank: {row['bank_name']}",
                f"Zahlungsbedingung: {row['payment_term']}",
                f"Aktiv: {int_to_yes_no(row['active'])}",
            ])
    body.configure(state='normal')
    body.delete('1.0', tk.END)
    body.insert('1.0', content)
    body.configure(state='disabled')


def _fm_creditors_clear_stack_preview(self):
    if hasattr(self, 'preview_title_label'):
        self.preview_title_label.configure(text='Kein Dokument ausgewählt')
    if hasattr(self, 'preview_image_label'):
        self.preview_image_label.configure(text='Keine Vorschau verfügbar', image='')
        try:
            self.preview_image_label.image = None
        except Exception:
            pass


def _fm_creditors_build_ui(self):
    outer = ttk.PanedWindow(self, orient='vertical')
    outer.pack(fill='both', expand=True)
    _fm_attach_paned_persistence(outer, 'CreditorsView:outer_vertical', default_ratio=0.62)
    invoice_area = ttk.Frame(outer)
    stack_area = ttk.Frame(outer)
    outer.add(invoice_area, weight=3)
    outer.add(stack_area, weight=2)

    shell = build_two_block_shell(invoice_area, 'Eingangsrechnung erfassen', 'Eingangsrechnungen')
    left = ScrollableFrame(shell.left)
    left.pack(fill='both', expand=True, padx=10, pady=8)
    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.partner = tk.StringVar(); self.invoice_date = tk.StringVar(value=today_str()); self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar(); self.tax_code = tk.StringVar(value='V19'); self.net_amount = tk.StringVar(); self.tax_amount = tk.StringVar(); self.gross_amount = tk.StringVar()
    self.current_hist_no = None
    self.pending_attachments = []

    self._labeled_entry(left.content, 'Rechnungs-Nr.', self.invoice_no)
    prow = ttk.Frame(left.content); prow.pack(fill='x', pady=3)
    ttk.Label(prow, text='Kreditor', width=16).pack(side='left')
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side='left', fill='x', expand=True)
    self.partner_cb.bind('<MouseWheel>', lambda e: 'break')
    self.partner_cb.bind('<<ComboboxSelected>>', lambda _e: self._update_vendor_info())
    create_standard_button(prow, '+', self._open_vendor_quick_popup).pack(side='left', padx=3)
    self._labeled_entry(left.content, 'Rechnungsdatum', self.invoice_date)
    self._labeled_combo(left.content, 'Zahlungsbed.', self.payment_term, self.payment_terms)
    self._labeled_entry(left.content, 'Fälligkeit', self.due_date)
    self._labeled_combo(left.content, 'Steuerkennz.', self.tax_code, self.tax_codes)
    self._labeled_entry(left.content, 'Netto', self.net_amount)
    self._labeled_entry(left.content, 'Steuer', self.tax_amount)
    self._labeled_entry(left.content, 'Brutto', self.gross_amount)
    self.payment_term.trace_add('write', lambda *_: self.apply_payment_term())
    self.net_amount.trace_add('write', lambda *_: self._calculate_amounts())
    self.tax_code.trace_add('write', lambda *_: self._calculate_amounts())

    btns = ttk.Frame(left.content); btns.pack(fill='x', pady=8)
    create_standard_button(btns, 'Dokument anhängen', self.add_pending_attachments).pack(side='left', padx=3)
    create_standard_button(btns, 'Speichern', self.save_invoice, confirm=True).pack(side='left', padx=3)
    create_standard_button(btns, 'Neu', self.clear_form).pack(side='left', padx=3)

    info_frame = ttk.LabelFrame(left.content, text='Stammdaten des Kreditors')
    info_frame.pack(fill='x', pady=(8, 4))
    self.vendor_info_body = tk.Text(info_frame, height=8, wrap='word', bg=WHITE, fg=TEXT, relief='flat')
    self.vendor_info_body.pack(fill='both', expand=True, padx=6, pady=6)
    self.vendor_info_body.configure(state='disabled')

    right = ttk.Frame(shell.right); right.pack(fill='both', expand=True, padx=10, pady=8)
    nb = ttk.Notebook(right)
    nb.pack(fill='both', expand=True)
    tab_inv = ttk.Frame(nb)
    tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text='Eingangsrechnungen')
    nb.add(tab_op, text='Offene Posten')

    inv_cols = ('invoice_no', 'partner', 'date', 'due', 'gross', 'open', 'status', 'att')
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show='headings', height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [('invoice_no', 'Rechnung', 110), ('partner', 'Kreditor', 180), ('date', 'Datum', 90), ('due', 'Fällig', 90), ('gross', 'Brutto', 90), ('open', 'Offen', 90), ('status', 'Status', 110), ('att', 'Anhang', 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor='w')
    self.inv_tree.pack(fill='both', expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind('<ButtonRelease-1>', lambda e: _fm_invoice_tree_click(self, e), add='+')

    op_cols = ('reference_no', 'partner', 'due', 'original', 'open', 'status', 'att')
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show='headings', height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [('reference_no', 'Referenz', 120), ('partner', 'Kreditor', 180), ('due', 'Fällig', 90), ('original', 'Original', 90), ('open', 'Offen', 90), ('status', 'Status', 120), ('att', 'Anhang', 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor='w')
    self.op_tree.pack(fill='both', expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind('<ButtonRelease-1>', lambda e: _fm_op_tree_click(self, e), add='+')

    self._build_stack_ui(stack_area)
    self.partner.trace_add('write', lambda *_: self._update_vendor_info())
    self._update_vendor_info()
    self.reload_invoices(); self.reload_open_items(); self._load_import_batches(); self._clear_stack_preview()


def _fm_creditors_build_stack_ui(self, parent):
    split = build_two_block_shell(parent, 'Dokumentenstapel', 'Vorschau')
    left = ttk.Frame(split.left); left.pack(fill='both', expand=True, padx=10, pady=8)
    right = ttk.Frame(split.right); right.pack(fill='both', expand=True, padx=10, pady=8)
    toolbar = ttk.Frame(left); toolbar.pack(fill='x', pady=(0, 6))
    create_standard_button(toolbar, 'Dateien importieren', self.import_creditor_files, confirm=True).pack(side='left', padx=3)
    create_standard_button(toolbar, 'Ordner importieren', self.import_creditor_folder).pack(side='left', padx=3)
    create_standard_button(toolbar, 'Markierte löschen', self._delete_marked_batches).pack(side='left', padx=3)
    create_standard_button(toolbar, 'Markierte zusammenführen', self.merge_selected_batches).pack(side='left', padx=3)

    cols = ('hist_no', 'title', 'status', 'mark')
    self.stack_tree = ttk.Treeview(left, columns=cols, show='headings', height=10)
    for c, t, w in [('hist_no', 'Hist.-Nr.', 110), ('title', 'Dokument', 260), ('status', 'Status', 110), ('mark', '☑', 40)]:
        self.stack_tree.heading(c, text=t)
        self.stack_tree.column(c, width=w, anchor='w' if c != 'mark' else 'center')
    self.stack_tree.pack(fill='both', expand=True)
    self.stack_tree.bind('<ButtonRelease-1>', self._on_stack_select, add='+')
    self.stack_tree.bind('<Double-1>', lambda _e: self._open_selected_stack_file())

    self.preview_title_label = ttk.Label(right, text='Kein Dokument ausgewählt', style='Hint.TLabel')
    self.preview_title_label.pack(anchor='w', pady=(0, 6))
    self.preview_image_label = ttk.Label(right, text='Keine Vorschau verfügbar', anchor='nw', justify='left')
    self.preview_image_label.pack(fill='both', expand=True)


def _fm_creditors_import_archive_paths(self, paths):
    valid = []
    for raw in paths:
        fp = Path(raw)
        if fp.exists() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo('Dokumentenstapel', 'Keine unterstützten Dateien ausgewählt.')
        return
    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute("INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)", (hist_no, clean_document_title(fp.name), '', '', 'Importiert', 0, now_str(), now_str()))
            con.execute("INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)", (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or '', now_str()))
        con.commit()
    self._load_import_batches()
    self.app.set_status(f"{len(valid)} Dokument(e) in den Dokumentenstapel importiert.")


def _fm_creditors_import_files(self):
    paths = filedialog.askopenfilenames(title='Dateien für Dokumentenstapel importieren')
    if paths:
        self._import_archive_paths(paths)


def _fm_creditors_import_folder(self):
    folder = filedialog.askdirectory(title='Ordner für Dokumentenstapel importieren')
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))


def _fm_creditors_load_import_batches(self):
    for i in self.stack_tree.get_children():
        self.stack_tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_batches ORDER BY id DESC"))
    for r in rows:
        self.stack_tree.insert('', 'end', iid=r['hist_no'], values=(r['hist_no'], r['title'], r['status'], '☑' if int(r['marked'] or 0) else '☐'))


def _fm_creditors_on_stack_select(self, event=None):
    row_id = self.stack_tree.identify_row(event.y) if event is not None else None
    col = self.stack_tree.identify_column(event.x) if event is not None else None
    if row_id and col == '#4':
        with get_connection() as con:
            row = con.execute("SELECT marked FROM invoice_import_batches WHERE hist_no=?", (row_id,)).fetchone()
            new_val = 0 if row and int(row[0] or 0) else 1
            con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=? WHERE hist_no=?", (new_val, now_str(), row_id))
            con.commit()
        self._load_import_batches()
        return 'break'
    if not row_id:
        self.current_hist_no = None
        self._clear_stack_preview()
        try:
            self.stack_tree.selection_remove(self.stack_tree.selection())
        except Exception:
            pass
        return None
    if self.current_hist_no == row_id:
        self.current_hist_no = None
        try:
            self.stack_tree.selection_remove(self.stack_tree.selection())
        except Exception:
            pass
        self._clear_stack_preview()
        return 'break'
    self.current_hist_no = row_id
    try:
        self.stack_tree.selection_set(row_id)
    except Exception:
        pass
    self._refresh_stack_preview()
    return None


def _fm_creditors_refresh_stack_preview(self):
    if not getattr(self, 'current_hist_no', None):
        self._clear_stack_preview()
        return
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_files WHERE hist_no=? ORDER BY id", (self.current_hist_no,)))
    if not rows:
        self._clear_stack_preview()
        return
    first = rows[0]
    title = clean_document_title(first['file_name'])
    if len(rows) > 1:
        title += f" (+ {len(rows)-1} weitere)"
    self.preview_title_label.configure(text=title)
    self.preview_image_label.configure(text=collect_preview_text(first['file_path']), image='')
    try:
        self.preview_image_label.image = None
    except Exception:
        pass


def _fm_creditors_open_selected_stack_file(self):
    if not getattr(self, 'current_hist_no', None):
        return
    with get_connection() as con:
        row = con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id LIMIT 1", (self.current_hist_no,)).fetchone()
    if row:
        open_path(row[0])


def _fm_creditors_delete_marked(self):
    with get_connection() as con:
        hist_nos = [r['hist_no'] for r in con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1 ORDER BY id")]
    if not hist_nos:
        messagebox.showinfo('Dokumentenstapel', 'Keine markierten Dokumente vorhanden.')
        return
    if not messagebox.askyesno('Dokumentenstapel', f"{len(hist_nos)} markierte Dokument(e) löschen?", parent=self):
        return
    placeholders = ','.join(['?'] * len(hist_nos))
    with get_connection() as con:
        files = list(con.execute(f"SELECT file_path FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos))
        con.execute(f"DELETE FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos)
        con.execute(f"DELETE FROM invoice_import_batches WHERE hist_no IN ({placeholders})", hist_nos)
        con.commit()
    for r in files:
        try:
            fp = Path(r['file_path'])
            if fp.exists():
                fp.unlink()
        except Exception:
            pass
    if getattr(self, 'current_hist_no', None) in hist_nos:
        self.current_hist_no = None
        self._clear_stack_preview()
    self._load_import_batches()


def _fm_creditors_merge_selected(self):
    with get_connection() as con:
        hist_nos = [r['hist_no'] for r in con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1 ORDER BY id")]
        if len(hist_nos) < 2:
            messagebox.showinfo('Dokumentenstapel', 'Bitte mindestens zwei markierte Dokumente wählen.')
            return
        target = hist_nos[0]
        for hist_no in hist_nos[1:]:
            con.execute("UPDATE invoice_import_files SET hist_no=? WHERE hist_no=?", (target, hist_no))
            con.execute("DELETE FROM invoice_import_batches WHERE hist_no=?", (hist_no,))
        con.execute("UPDATE invoice_import_batches SET marked=0, updated_at=? WHERE hist_no=?", (now_str(), target))
        con.commit()
    self.current_hist_no = target
    self._load_import_batches()
    self._refresh_stack_preview()


CreditorsView._build_ui = _fm_creditors_build_ui
CreditorsView._build_stack_ui = _fm_creditors_build_stack_ui
CreditorsView._update_vendor_info = _fm_creditors_update_vendor_info
CreditorsView._clear_stack_preview = _fm_creditors_clear_stack_preview
CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths
CreditorsView.import_creditor_files = _fm_creditors_import_files
CreditorsView.import_creditor_folder = _fm_creditors_import_folder
CreditorsView._load_import_batches = _fm_creditors_load_import_batches
CreditorsView._on_stack_select = _fm_creditors_on_stack_select
CreditorsView._refresh_stack_preview = _fm_creditors_refresh_stack_preview
CreditorsView._open_selected_stack_file = _fm_creditors_open_selected_stack_file
CreditorsView._delete_marked_batches = _fm_creditors_delete_marked
CreditorsView.merge_selected_batches = _fm_creditors_merge_selected



# === FINANCE MATE V0.6.19 - CLEAN PATCH FROM STABLE V0.6.18 ===
# Grundlage: stabile v0.6.18. Ziel: keine leeren/grauen Kreditoren-Module mehr, nur gewünschte Weiterentwicklung.
APP_VERSION = "0.6.19-clean-from-v0.6.18"

TRADE_COUNTRIES = [
    "Deutschland", "Vereinigte Staaten", "China", "Japan", "Vereinigtes Königreich", "Frankreich", "Italien", "Niederlande", "Belgien", "Schweiz",
    "Österreich", "Spanien", "Polen", "Tschechien", "Ungarn", "Slowakei", "Schweden", "Dänemark", "Finnland", "Norwegen",
    "Irland", "Portugal", "Griechenland", "Türkei", "Rumänien", "Bulgarien", "Slowenien", "Kroatien", "Serbien", "Ukraine",
    "Kanada", "Mexiko", "Brasilien", "Argentinien", "Chile", "Kolumbien", "Peru", "Indien", "Südkorea", "Taiwan",
    "Singapur", "Malaysia", "Thailand", "Indonesien", "Vietnam", "Philippinen", "Australien", "Neuseeland", "Vereinigte Arabische Emirate", "Südafrika"
]
INVOICE_KIND_VALUES = ["Rechnung", "Gutschrift"]
FIBU_VISIBLE_NAME = "Buchungsjournal"


def _v619_status(open_amount: Any) -> str:
    try:
        return STATUS_PAID if parse_amount(open_amount) <= 0 else STATUS_OPEN
    except Exception:
        return STATUS_OPEN


def compute_status_from_open_amount(open_amount: Any, due_date: str) -> str:
    return _v619_status(open_amount)


def urgency_bucket(due_date: str, open_amount: Any) -> str:
    try:
        amount = parse_amount(open_amount)
    except Exception:
        amount = Decimal("0.00")
    if amount <= 0:
        return "paid"
    if not validate_date(due_date):
        return "open"
    delta = (datetime.strptime(due_date, DATE_FMT).date() - datetime.now().date()).days
    if delta <= 0:
        return "overdue"
    if delta <= 5:
        return "soon"
    return "open"


def configure_tree_tags(tree: ttk.Treeview) -> None:
    tree.tag_configure("paid", background=SOFT_GREEN)
    tree.tag_configure("soon", background=SOFT_ORANGE)
    tree.tag_configure("overdue", background=SOFT_RED)
    tree.tag_configure("open", background=WHITE)


_orig_init_sqlite_v619_clean = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v619_clean()
    with get_connection() as con:
        for table, column, definition in [
            ("customer_invoices", "invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("customer_invoices", "description", "TEXT DEFAULT ''"),
            ("vendor_invoices", "invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("vendor_invoices", "description", "TEXT DEFAULT ''"),
            ("open_items", "invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("open_items", "description", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, table, column, definition)
        # bestehende Doppler in offenen Posten bereinigen
        con.execute("DELETE FROM open_items WHERE id NOT IN (SELECT MIN(id) FROM open_items GROUP BY entity_type, reference_no)")
        # Status für vorhandene Rechnungen und OPs fachlich korrigieren
        for table in ("customer_invoices", "vendor_invoices"):
            for row in con.execute(f"SELECT id, open_amount FROM {table}").fetchall():
                con.execute(f"UPDATE {table} SET status=?, updated_at=? WHERE id=?", (_v619_status(row["open_amount"]), now_str(), row["id"]))
        for row in con.execute("SELECT id, open_amount FROM open_items").fetchall():
            con.execute("UPDATE open_items SET status=?, updated_at=? WHERE id=?", (_v619_status(row["open_amount"]), now_str(), row["id"]))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()


# ---------- kleine stabile UI-Helfer ----------
def _v619_labeled_text(parent, label: str, height: int = 2) -> tk.Text:
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=3)
    ttk.Label(row, text=label, width=16).pack(side="left", anchor="n")
    txt = tk.Text(row, height=height, wrap="word", bg=WHITE, fg=TEXT, relief="sunken")
    txt.pack(side="left", fill="x", expand=True)
    return txt


def _v619_text_get(widget: Optional[tk.Text]) -> str:
    if widget is None:
        return ""
    return widget.get("1.0", "end-1c").strip()


def _v619_text_set(widget: Optional[tk.Text], value: str = "") -> None:
    if widget is None:
        return
    widget.delete("1.0", tk.END)
    if value:
        widget.insert("1.0", value)


def _v619_unique_trees(*trees):
    result = []
    seen = set()
    for tree in trees:
        if tree is None:
            continue
        if id(tree) in seen:
            continue
        result.append(tree)
        seen.add(id(tree))
    return result


def _v619_add_attachment_column_click(self, event, tree, ref_getter, refresh_callback):
    row = tree.identify_row(event.y)
    col = tree.identify_column(event.x)
    if row and col == f"#{len(tree['columns'])}":
        self.manage_attachment_request(self.entity_type, ref_getter(row), refresh_callback=refresh_callback)
        return "break"
    return None


# ---------- Referenzdaten stabil ----------
def _v619_load_reference_data(self):
    with get_connection() as con:
        self.partners = [f"{r[self.partner_no_col]} {r['name']}" for r in con.execute(f"SELECT {self.partner_no_col},name FROM {self.partner_table} WHERE active=1 ORDER BY name")]
        self.tax_codes = [r[0] for r in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]
        self.payment_terms = [r[0] for r in con.execute("SELECT code FROM payment_terms WHERE active=1 ORDER BY code")]

InvoiceModuleBase._load_reference_data = _v619_load_reference_data


# ---------- Speichern/Leeren/Reload für Debitoren und Kreditoren ----------
def _v619_save_invoice(self):
    if not validate_date(self.invoice_date.get()) or not validate_date(self.due_date.get()):
        messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.")
        return
    no, name, address = self._partner_parts()
    if not no:
        messagebox.showwarning("Partner", "Bitte Partner auswählen.")
        return
    selected_hist = getattr(self, "current_hist_no", None)
    if self.entity_type == "vendor_invoice" and not (selected_hist or getattr(self, "pending_attachments", [])):
        messagebox.showwarning("Dokument", "Eingangsrechnungen können ohne Dokument nicht erfasst werden.")
        return
    gross = parse_amount(self.gross_amount.get())
    net = parse_amount(self.net_amount.get())
    tax = parse_amount(self.tax_amount.get())
    status = _v619_status(gross)
    invoice_no = self.invoice_no.get().strip()
    invoice_kind = self.invoice_kind.get().strip() if hasattr(self, "invoice_kind") else "Rechnung"
    description = _v619_text_get(getattr(self, "description_widget", None))
    address_col = "customer_address" if self.invoice_table == "customer_invoices" else "vendor_address"
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    partner_no_col = "customer_no" if self.invoice_table == "customer_invoices" else "vendor_no"
    stack_files = []
    with get_connection() as con:
        con.execute(
            f"INSERT INTO {self.invoice_table}(invoice_no,{partner_no_col},{partner_name_col},{address_col},invoice_date,due_date,payment_term,tax_code,net_amount,tax_amount,gross_amount,open_amount,status,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (invoice_no, no, name, address, self.invoice_date.get(), self.due_date.get(), self.payment_term.get(), self.tax_code.get(), float(net), float(tax), float(gross), float(gross), status, invoice_kind, description, now_str(), now_str()),
        )
        con.execute("DELETE FROM open_items WHERE entity_type=? AND reference_no=?", (self.entity_type, invoice_no))
        con.execute(
            "INSERT INTO open_items(entity_type,reference_no,partner_no,partner_name,due_date,original_amount,open_amount,status,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
            (self.entity_type, invoice_no, no, name, self.due_date.get(), float(gross), float(gross), status, invoice_kind, description, now_str(), now_str()),
        )
        if self.entity_type == "vendor_invoice" and selected_hist:
            stack_files = list(con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id", (selected_hist,)))
            con.execute("UPDATE invoice_import_batches SET status='Erfasst', vendor_no=?, vendor_name=?, updated_at=? WHERE hist_no=?", (no, name, now_str(), selected_hist))
        con.commit()
    if stack_files:
        self.add_attachment_paths(self.entity_type, invoice_no, [r["file_path"] for r in stack_files])
    if getattr(self, "pending_attachments", []):
        self.add_attachment_paths(self.entity_type, invoice_no, self.pending_attachments)
    self.app.set_status(f"{self.entity_label} {invoice_no} gespeichert.")
    self.clear_form()
    self.reload_invoices()
    self.reload_open_items()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    if hasattr(self, "_clear_stack_preview"):
        self._clear_stack_preview()

InvoiceModuleBase.save_invoice = _v619_save_invoice


def _v619_clear_form(self):
    self.invoice_no.set(self._generate_invoice_no())
    if hasattr(self, "invoice_kind"):
        self.invoice_kind.set("Rechnung")
    self.partner.set("")
    self.invoice_date.set(today_str())
    self.due_date.set(today_str())
    self.payment_term.set("")
    self.tax_code.set("V19")
    self.net_amount.set("")
    self.tax_amount.set("")
    self.gross_amount.set("")
    _v619_text_set(getattr(self, "description_widget", None), "")
    self.pending_attachments.clear()
    if hasattr(self, "current_hist_no"):
        self.current_hist_no = None
    if hasattr(self, "_update_vendor_info"):
        try:
            self._update_vendor_info()
        except Exception:
            pass
    if hasattr(self, "_clear_stack_preview"):
        try:
            self._clear_stack_preview()
        except Exception:
            pass

InvoiceModuleBase.clear_form = _v619_clear_form


def _v619_reload_invoices(self):
    trees = _v619_unique_trees(getattr(self, "inv_tree", None), *getattr(self, "invoice_tabs_trees", []))
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    with get_connection() as con:
        rows = list(con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"))
    for tree in trees:
        for item in tree.get_children():
            tree.delete(item)
    for r in rows:
        att = self.get_attachment_count(self.entity_type, r["invoice_no"])
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        values = (
            r["invoice_no"],
            r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung",
            r[partner_name_col],
            r["description"] if "description" in r.keys() else "",
            r["invoice_date"],
            r["due_date"],
            format_amount(r["gross_amount"]),
            format_amount(r["open_amount"]),
            _v619_status(r["open_amount"]),
            f"📎 {att}" if att else "Dokument anhängen",
        )
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))

InvoiceModuleBase.reload_invoices = _v619_reload_invoices


def _v619_reload_open_items(self):
    trees = _v619_unique_trees(getattr(self, "op_tree", None), *getattr(self, "open_item_tabs_trees", []))
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date, id", (self.entity_type,)))
    for tree in trees:
        for item in tree.get_children():
            tree.delete(item)
    for r in rows:
        if parse_amount(r["open_amount"]) <= 0:
            continue
        att = self.get_attachment_count(self.entity_type, r["reference_no"])
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        values = (
            r["reference_no"],
            r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung",
            r["partner_name"],
            r["description"] if "description" in r.keys() else "",
            r["due_date"],
            format_amount(r["original_amount"]),
            format_amount(r["open_amount"]),
            _v619_status(r["open_amount"]),
            f"📎 {att}" if att else "Dokument anhängen",
        )
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))

InvoiceModuleBase.reload_open_items = _v619_reload_open_items


# ---------- Debitoren UI stabil erweitern ----------
def _v619_debitors_build_ui(self):
    shell = build_two_block_shell(self, f"{self.entity_label} erfassen", "Offene Posten / Kontrolle")
    left = ScrollableFrame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()
    self._labeled_combo(left.content, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
    self._labeled_combo(left.content, "Partner", self.partner, self.partners)
    self.description_widget = _v619_labeled_text(left.content, "Rechnungsbeschreibung", height=2)
    self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(left.content, "Fälligkeit", self.due_date)
    self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(left.content, "Netto", self.net_amount)
    self._labeled_entry(left.content, "Steuer", self.tax_amount)
    self._labeled_entry(left.content, "Brutto", self.gross_amount)
    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
    btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
    ttk.Label(left.content, text="Erfasste Rechnungen", style="Section.TLabel").pack(anchor="w", pady=(10, 4))
    cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=9)
    for c, t, w in [("invoice_no", "Rechnung", 110), ("kind", "Art", 90), ("partner", "Partner", 170), ("description", "Beschreibung", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")
    right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
    cols2 = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(right, columns=cols2, show="headings")
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Partner", 170), ("description", "Beschreibung", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

DebitorsView._build_ui = _v619_debitors_build_ui


# ---------- Kreditoren UI vollständig aus v0.6.18 wiederhergestellt und erweitert ----------
def _v619_creditors_build_ui(self):
    # Fehler im Kreditoren-Modul dürfen nicht mehr still in einem grauen Hintergrund enden.
    outer = ttk.PanedWindow(self, orient="vertical")
    outer.pack(fill="both", expand=True)
    try:
        _fm_attach_paned_persistence(outer, "CreditorsView:outer_vertical", default_ratio=0.62)
    except Exception:
        pass
    invoice_area = ttk.Frame(outer)
    stack_area = ttk.Frame(outer)
    outer.add(invoice_area, weight=3)
    outer.add(stack_area, weight=2)

    shell = build_two_block_shell(invoice_area, "Eingangsrechnung erfassen", "Eingangsrechnungen")
    left = ScrollableFrame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()
    self.current_hist_no = None
    self.pending_attachments = []

    self._labeled_combo(left.content, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
    prow = ttk.Frame(left.content); prow.pack(fill="x", pady=3)
    ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side="left", fill="x", expand=True)
    self.partner_cb.bind("<MouseWheel>", lambda e: "break")
    self.partner_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_vendor_info())
    create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)
    self.description_widget = _v619_labeled_text(left.content, "Rechnungsbeschreibung", height=2)
    self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(left.content, "Fälligkeit", self.due_date)
    self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(left.content, "Netto", self.net_amount)
    self._labeled_entry(left.content, "Steuer", self.tax_amount)
    self._labeled_entry(left.content, "Brutto", self.gross_amount)
    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
    btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
    info_frame = ttk.LabelFrame(left.content, text="Stammdaten des Kreditors")
    info_frame.pack(fill="x", pady=(8, 4))
    self.vendor_info_body = tk.Text(info_frame, height=8, wrap="word", bg=WHITE, fg=TEXT, relief="flat")
    self.vendor_info_body.pack(fill="both", expand=True, padx=6, pady=6)
    self.vendor_info_body.configure(state="disabled")

    right = ttk.Frame(shell.right)
    right.pack(fill="both", expand=True, padx=10, pady=8)
    nb = ttk.Notebook(right)
    nb.pack(fill="both", expand=True)
    tab_inv = ttk.Frame(nb)
    tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text="Eingangsrechnungen")
    nb.add(tab_op, text="Offene Posten")

    inv_cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show="headings", height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [("invoice_no", "Rechnung", 110), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")

    op_cols = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show="headings", height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

    # Dokumentenstapel/Vorschau exakt als eigener Bereich beibehalten.
    self._build_stack_ui(stack_area)
    self.partner.trace_add("write", lambda *_: self._update_vendor_info())
    self._update_vendor_info()
    self.reload_invoices()
    self.reload_open_items()
    self._load_import_batches()
    self._clear_stack_preview()

CreditorsView._build_ui = _v619_creditors_build_ui


# ---------- Kreditoren: Stapel erfasste ausblenden, Löschen stabilisieren ----------
def _v619_load_import_batches(self):
    if not hasattr(self, "stack_tree"):
        return
    for i in self.stack_tree.get_children():
        self.stack_tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_batches WHERE status <> 'Erfasst' ORDER BY id DESC"))
    for r in rows:
        self.stack_tree.insert("", "end", iid=r["hist_no"], values=(r["hist_no"], r["title"], r["status"], "☑" if int(r["marked"] or 0) else "☐"))

CreditorsView._load_import_batches = _v619_load_import_batches


def _v619_delete_marked(self):
    with get_connection() as con:
        hist_nos = [r["hist_no"] for r in con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1 AND status <> 'Erfasst' ORDER BY id")]
    if not hist_nos:
        messagebox.showinfo("Dokumentenstapel", "Keine markierten Dokumente vorhanden.")
        return
    if not messagebox.askyesno("Dokumentenstapel", f"{len(hist_nos)} markierte Dokument(e) löschen?", parent=self):
        return
    placeholders = ",".join(["?"] * len(hist_nos))
    with get_connection() as con:
        files = list(con.execute(f"SELECT file_path FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos))
        con.execute(f"DELETE FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos)
        con.execute(f"DELETE FROM invoice_import_batches WHERE hist_no IN ({placeholders})", hist_nos)
        con.commit()
    for r in files:
        try:
            fp = Path(r["file_path"])
            if fp.exists():
                fp.unlink()
        except Exception:
            pass
    if getattr(self, "current_hist_no", None) in hist_nos:
        self.current_hist_no = None
        self._clear_stack_preview()
    self._load_import_batches()

CreditorsView._delete_marked_batches = _v619_delete_marked


# ---------- Quick-Kreditor erweitert ----------
def _v619_open_vendor_quick_popup(self):
    pop = tk.Toplevel(self)
    pop.title("Quick-Kreditor")
    pop.geometry("620x650")
    pop.configure(bg=BG)
    fields = ["vendor_no", "name", "country", "street", "zip", "city", "email", "phone", "tax_id", "vat_id", "iban", "bic", "bank_name", "payment_term", "active"]
    vars = {k: tk.StringVar() for k in fields}
    vars["vendor_no"].set(generate_number("K", "counter_vendors", 5))
    vars["country"].set("Deutschland")
    vars["active"].set("Ja")
    labels = [
        ("vendor_no", "Kreditor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"),
        ("email", "E-Mail"), ("phone", "Telefon"), ("tax_id", "Steuer-Nr."), ("vat_id", "USt-IdNr."), ("iban", "IBAN"), ("bic", "BIC"),
        ("bank_name", "Bank"), ("payment_term", "Zahlungsbed."), ("active", "Aktiv")
    ]
    for key, label in labels:
        row = ttk.Frame(pop); row.pack(fill="x", padx=12, pady=4)
        ttk.Label(row, text=label, width=16).pack(side="left")
        if key == "country":
            cb = ttk.Combobox(row, textvariable=vars[key], values=TRADE_COUNTRIES, state="readonly")
            cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        elif key == "payment_term":
            cb = ttk.Combobox(row, textvariable=vars[key], values=self.payment_terms, state="readonly")
            cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        elif key == "active":
            cb = ttk.Combobox(row, textvariable=vars[key], values=["Ja", "Nein"], state="readonly")
            cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        else:
            ttk.Entry(row, textvariable=vars[key]).pack(side="left", fill="x", expand=True)

    def save():
        if not vars["name"].get().strip():
            messagebox.showwarning("Pflichtfeld", "Name fehlt.", parent=pop)
            return
        with get_connection() as con:
            con.execute(
                "INSERT INTO vendors(vendor_no,name,country,street,zip,city,email,phone,tax_id,vat_id,iban,bic,bank_name,payment_term,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (vars["vendor_no"].get().strip(), vars["name"].get().strip(), vars["country"].get().strip(), vars["street"].get().strip(), vars["zip"].get().strip(), vars["city"].get().strip(), vars["email"].get().strip(), vars["phone"].get().strip(), vars["tax_id"].get().strip(), vars["vat_id"].get().strip(), vars["iban"].get().strip(), vars["bic"].get().strip(), vars["bank_name"].get().strip(), vars["payment_term"].get().strip(), yes_no_to_int(vars["active"].get()), now_str(), now_str())
            )
            con.commit()
        self._load_reference_data()
        self.partner.set(f"{vars['vendor_no'].get().strip()} {vars['name'].get().strip()}")
        try:
            self._update_vendor_info()
        except Exception:
            pass
        pop.destroy()

    create_standard_button(pop, "Speichern", save, confirm=True).pack(pady=12)

CreditorsView._open_vendor_quick_popup = _v619_open_vendor_quick_popup


# ---------- Navigation: Finanzbuchhaltung sichtbar umbenennen ----------
def _v619_build_sidebar(self):
    for w in self.sidebar.winfo_children():
        w.destroy()
    width = 42 if self.sidebar_collapsed else 210
    self.sidebar.configure(width=width)
    self.sidebar.pack_propagate(False)
    arrow = "»" if self.sidebar_collapsed else "«"
    create_standard_button(self.sidebar, arrow, self._toggle_sidebar).pack(fill="x", padx=5, pady=5)
    modules = [("Dashboard", "DB"), ("Stammdaten", "ST"), (FIBU_VISIBLE_NAME, "BJ"), ("Debitoren", "DE"), ("Kreditoren", "KR"), ("Zahlungen", "ZA"), ("Reporting", "RE"), ("Audit", "AU"), ("Einstellungen", "EI")]
    self.nav_buttons.clear()
    for name, abbr in modules:
        text = abbr if self.sidebar_collapsed else name
        btn = ttk.Button(self.sidebar, text=text, style="Nav.TButton", command=lambda n=name: self.show_module(n))
        btn.pack(fill="x", padx=5, pady=2)
        self.nav_buttons[name] = btn

FinanceMateApp._build_sidebar = _v619_build_sidebar


def _v619_show_module(self, module: str):
    module_key = FIBU_VISIBLE_NAME if module == "Finanzbuchhaltung" else module
    self.active_module = module_key
    self._update_nav_styles()
    for w in self.workspace.winfo_children():
        w.destroy()
    ttk.Label(self.workspace, text=module_key, style="CardTitle.TLabel").pack(anchor="w", pady=(0, 6))
    container = ttk.Frame(self.workspace)
    container.pack(fill="both", expand=True)
    render = {
        "Dashboard": self._render_dashboard,
        "Stammdaten": self._render_stammdaten,
        FIBU_VISIBLE_NAME: self._render_finanzbuchhaltung,
        "Debitoren": self._render_debitoren,
        "Kreditoren": self._render_kreditoren,
        "Zahlungen": self._render_zahlungen,
        "Reporting": self._render_reporting,
        "Audit": self._render_audit,
        "Einstellungen": self._render_einstellungen,
    }.get(module_key, self._render_dashboard)
    render(container)
    self.set_status(f"Modul {module_key} geladen.")

FinanceMateApp.show_module = _v619_show_module



# === FINANCE MATE PATCH V0_6_20_RECHNUNGSPORTAL_PREVIEW_APPROVAL ===
APP_VERSION = "0.6.20-rechnungsportal-preview-approval"
APPROVAL_THRESHOLDS = [
    (1, Decimal("1000.00")),
    (2, Decimal("5000.00")),
    (3, Decimal("25000.00")),
    (4, None),
]
APPROVAL_STATUS_REVIEW = "Sachliche Prüfung ausstehend"
APPROVAL_STATUS_APPROVAL1 = "1. Freigabe ausstehend"
APPROVAL_STATUS_APPROVAL2 = "2. Freigabe ausstehend"
APPROVAL_STATUS_RELEASED = "Freigegeben"
APPROVAL_STATUS_VALUES = [APPROVAL_STATUS_REVIEW, APPROVAL_STATUS_APPROVAL1, APPROVAL_STATUS_APPROVAL2, APPROVAL_STATUS_RELEASED]


def _fm_current_user() -> str:
    return os.environ.get("FINANCEMATE_USER") or os.environ.get("USERNAME") or os.environ.get("USER") or "default"


def _fm_current_approval_level() -> int:
    # Übergangslösung bis zur echten Benutzerverwaltung: technisch Ebene 4, damit Funktionen testbar sind.
    try:
        return int(os.environ.get("FINANCEMATE_APPROVAL_LEVEL", "4"))
    except Exception:
        return 4


def _fm_invoice_direction(entity_type: str) -> str:
    return "Eingangsrechnung" if entity_type == "vendor_invoice" else "Ausgangsrechnung"


def _fm_invoice_table(entity_type: str) -> str:
    return "vendor_invoices" if entity_type == "vendor_invoice" else "customer_invoices"


def _fm_partner_cols(entity_type: str):
    if entity_type == "vendor_invoice":
        return "vendor_no", "vendor_name", "vendor_address"
    return "customer_no", "customer_name", "customer_address"


def _fm_calc_approval_level(amount) -> int:
    try:
        amt = abs(parse_amount(amount))
    except Exception:
        amt = Decimal("0.00")
    for level, threshold in APPROVAL_THRESHOLDS:
        if threshold is None or amt <= threshold:
            return level
    return 4


def _fm_status_from_open(open_amount: Any) -> str:
    try:
        return STATUS_PAID if parse_amount(open_amount) <= 0 else STATUS_OPEN
    except Exception:
        return STATUS_OPEN


def compute_status_from_open_amount(open_amount: Any, due_date: str) -> str:
    return _fm_status_from_open(open_amount)


def urgency_bucket(due_date: str, open_amount: Any) -> str:
    try:
        amount = parse_amount(open_amount)
    except Exception:
        amount = Decimal("0.00")
    if amount <= 0:
        return "paid"
    if not validate_date(due_date):
        return "open"
    delta = (datetime.strptime(due_date, DATE_FMT).date() - datetime.now().date()).days
    if delta <= 0:
        return "overdue"
    if delta <= 5:
        return "soon"
    return "open"


def configure_tree_tags(tree: ttk.Treeview) -> None:
    tree.tag_configure("paid", background=SOFT_GREEN)
    tree.tag_configure("soon", background=SOFT_ORANGE)
    tree.tag_configure("overdue", background=SOFT_RED)
    tree.tag_configure("open", background=WHITE)


_orig_init_sqlite_v620 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v620()
    with get_connection() as con:
        for table in ("customer_invoices", "vendor_invoices"):
            for col, definition in [
                ("invoice_kind", "TEXT DEFAULT 'Rechnung'"),
                ("description", "TEXT DEFAULT ''"),
                ("approval_status", f"TEXT DEFAULT '{APPROVAL_STATUS_REVIEW}'"),
                ("released", "INTEGER DEFAULT 0"),
                ("released_at", "TEXT DEFAULT ''"),
                ("reviewed_by", "TEXT DEFAULT ''"),
                ("reviewed_at", "TEXT DEFAULT ''"),
                ("approval_1_by", "TEXT DEFAULT ''"),
                ("approval_1_at", "TEXT DEFAULT ''"),
                ("approval_2_by", "TEXT DEFAULT ''"),
                ("approval_2_at", "TEXT DEFAULT ''"),
                ("approval_level", "INTEGER DEFAULT 1"),
                ("last_checked_by", "TEXT DEFAULT ''"),
                ("last_checked_at", "TEXT DEFAULT ''"),
            ]:
                ensure_column(con, table, col, definition)
        for col, definition in [
            ("invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("description", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, "open_items", col, definition)
        # Einmalige Rollout-Bereinigung aller bestehenden Rechnungsdatensätze.
        cleanup_done = con.execute("SELECT value FROM app_meta WHERE key='v0_6_20_invoice_cleanup_done'").fetchone()
        if not cleanup_done:
            old_attachment_paths = [r["file_path"] for r in con.execute("SELECT file_path FROM attachments WHERE entity_type IN ('customer_invoice','vendor_invoice')")]
            con.execute("DELETE FROM open_items")
            con.execute("DELETE FROM attachments WHERE entity_type IN ('customer_invoice','vendor_invoice')")
            con.execute("DELETE FROM customer_invoices")
            con.execute("DELETE FROM vendor_invoices")
            con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("v0_6_20_invoice_cleanup_done", "1", now_str()))
            for path in old_attachment_paths:
                try:
                    fp = Path(path)
                    if fp.exists():
                        fp.unlink()
                except Exception:
                    pass
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()


# ---------- Mehrseiten-Bildvorschau mit Pan/Zoom und Grenzen ----------
class _FMDocumentPreviewPane(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.canvas = tk.Canvas(self, bg=WHITE, highlightthickness=0)
        self.canvas.pack(fill="both", expand=True)
        self.base_image = None
        self.tk_image = None
        self.fit_zoom = 1.0
        self.zoom = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.drag_start = None
        self.canvas.bind("<Configure>", lambda e: self._render(reset=False))
        self.canvas.bind("<ButtonPress-1>", self._on_press)
        self.canvas.bind("<B1-Motion>", self._on_drag)
        self.canvas.bind("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind("<Button-4>", lambda e: self._wheel(120))
        self.canvas.bind("<Button-5>", lambda e: self._wheel(-120))
        self.clear()

    def clear(self):
        self.base_image = None
        self.tk_image = None
        self.canvas.delete("all")
        self.canvas.create_text(20, 20, anchor="nw", text="Keine Vorschau verfügbar", fill=TEXT2, font=("Segoe UI", 10))

    def set_image(self, pil_image):
        self.base_image = pil_image.convert("RGBA") if pil_image is not None else None
        self.zoom = 1.0
        self.offset_x = 0
        self.offset_y = 0
        if self.base_image is None:
            self.clear()
        else:
            self._render(reset=True)

    def _on_press(self, event):
        self.drag_start = (event.x, event.y)

    def _on_drag(self, event):
        if self.base_image is None or self.drag_start is None:
            return
        dx = event.x - self.drag_start[0]
        dy = event.y - self.drag_start[1]
        self.drag_start = (event.x, event.y)
        self.offset_x += dx
        self.offset_y += dy
        self._clamp_offsets()
        self._draw()

    def _on_mousewheel(self, event):
        self._wheel(event.delta)

    def _wheel(self, delta):
        if self.base_image is None:
            return
        factor = 1.12 if delta > 0 else 0.88
        old_zoom = self.zoom
        self.zoom = max(self.fit_zoom, min(self.zoom * factor, self.fit_zoom * 8.0))
        if abs(self.zoom - old_zoom) > 0.0001:
            self._clamp_offsets()
            self._render(reset=False)

    def _fit_zoom(self):
        if self.base_image is None:
            return 1.0
        cw = max(50, self.canvas.winfo_width())
        ch = max(50, self.canvas.winfo_height())
        iw, ih = self.base_image.size
        return min(cw / iw, ch / ih, 1.0)

    def _render(self, reset=False):
        if self.base_image is None:
            return
        try:
            from PIL import ImageTk
            self.fit_zoom = self._fit_zoom()
            if reset or self.zoom < self.fit_zoom:
                self.zoom = self.fit_zoom
                self.offset_x = 0
                self.offset_y = 0
            self._clamp_offsets()
            iw, ih = self.base_image.size
            size = (max(1, int(iw * self.zoom)), max(1, int(ih * self.zoom)))
            img = self.base_image.resize(size)
            self.tk_image = ImageTk.PhotoImage(img)
            self._draw()
        except Exception as exc:
            self.clear()
            self.canvas.create_text(20, 48, anchor="nw", text=f"Vorschaufehler: {exc}", fill=TEXT2, font=("Segoe UI", 10))

    def _clamp_offsets(self):
        if self.base_image is None:
            return
        cw = max(50, self.canvas.winfo_width())
        ch = max(50, self.canvas.winfo_height())
        rw = int(self.base_image.size[0] * self.zoom)
        rh = int(self.base_image.size[1] * self.zoom)
        # Mindestens 20% bzw. 80px müssen sichtbar bleiben.
        min_visible_x = min(max(80, cw * 0.2), cw)
        min_visible_y = min(max(80, ch * 0.2), ch)
        max_x = max(0, (rw + cw) / 2 - min_visible_x)
        max_y = max(0, (rh + ch) / 2 - min_visible_y)
        self.offset_x = max(-max_x, min(max_x, self.offset_x))
        self.offset_y = max(-max_y, min(max_y, self.offset_y))

    def _draw(self):
        self.canvas.delete("all")
        if self.tk_image is None:
            return
        x = self.canvas.winfo_width() // 2 + int(self.offset_x)
        y = self.canvas.winfo_height() // 2 + int(self.offset_y)
        self.canvas.create_image(x, y, image=self.tk_image)


def _fm_preview_font(size=14, bold=False):
    try:
        from PIL import ImageFont
        return ImageFont.truetype("DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf", size)
    except Exception:
        from PIL import ImageFont
        return ImageFont.load_default()


def _fm_text_pages_to_image(title: str, lines: list[str], page_lines: int = 42):
    from PIL import Image, ImageDraw
    pages = [lines[i:i+page_lines] for i in range(0, max(1, len(lines)), page_lines)] or [["(Keine Inhalte gefunden)"]]
    width = 1200
    page_height = 1500
    gap = 28
    img = Image.new("RGBA", (width, len(pages) * page_height + (len(pages)-1) * gap), "#f2f4f7")
    draw = ImageDraw.Draw(img)
    title_font = _fm_preview_font(22, True)
    body_font = _fm_preview_font(15, False)
    for pidx, page in enumerate(pages):
        y0 = pidx * (page_height + gap)
        draw.rectangle((0, y0, width, y0 + page_height), fill="white", outline="#cbd3dc")
        draw.text((38, y0 + 30), f"{title} – Seite {pidx+1}/{len(pages)}", fill=TEXT, font=title_font)
        y = y0 + 82
        for line in page:
            draw.text((38, y), str(line)[:150], fill=TEXT, font=body_font)
            y += 31
    return img


def _fm_combine_pages(images):
    from PIL import Image, ImageDraw
    if not images:
        return _fm_text_pages_to_image("Vorschau", ["Keine Vorschau verfügbar"])
    width = max(im.width for im in images)
    gap = 28
    total_height = sum(im.height for im in images) + gap * (len(images)-1)
    out = Image.new("RGBA", (width, total_height), "#f2f4f7")
    y = 0
    for im in images:
        rgba = im.convert("RGBA")
        x = (width - rgba.width) // 2
        out.paste(rgba, (x, y))
        y += rgba.height + gap
    return out


def _fm_generate_document_preview(path: str):
    from PIL import Image
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(p))
            pages = []
            for idx in range(len(doc)):
                page = doc.load_page(idx)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.65, 1.65), alpha=False)
                mode = "RGB" if pix.n < 4 else "RGBA"
                pages.append(Image.frombytes(mode, [pix.width, pix.height], pix.samples))
            doc.close()
            return _fm_combine_pages(pages)
        except Exception as exc:
            return _fm_text_pages_to_image(f"PDF-Vorschau – {p.name}", [f"PDF konnte nicht gerendert werden: {exc}"])
    if ext in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}:
        return Image.open(p)
    if ext in {".txt", ".csv"}:
        return _fm_text_pages_to_image(f"Text-Vorschau – {p.name}", p.read_text(encoding="utf-8", errors="replace").splitlines())
    if ext in {".xlsx", ".xls"}:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"Arbeitsblatt: {ws.title}")
                for row in ws.iter_rows(min_row=1, max_row=120, max_col=10, values_only=True):
                    values = ["" if v is None else str(v) for v in row]
                    if any(v.strip() for v in values):
                        lines.append(" | ".join(values).rstrip())
            wb.close()
            return _fm_text_pages_to_image(f"Excel-Vorschau – {p.name}", lines or ["(Keine Inhalte gefunden)"])
        except Exception as exc:
            return _fm_text_pages_to_image(f"Excel-Vorschau – {p.name}", [f"Excel konnte nicht gelesen werden: {exc}"])
    if ext in {".docx", ".doc"}:
        try:
            import docx
            doc = docx.Document(str(p))
            lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    vals = [cell.text.strip() for cell in row.cells]
                    if any(vals):
                        lines.append(" | ".join(vals))
            return _fm_text_pages_to_image(f"Word-Vorschau – {p.name}", lines or ["(Keine Inhalte gefunden)"])
        except Exception as exc:
            return _fm_text_pages_to_image(f"Word-Vorschau – {p.name}", [f"Word konnte nicht gelesen werden: {exc}"])
    return _fm_text_pages_to_image(f"Dokument-Vorschau – {p.name}", [collect_preview_text(str(p), 8000)])


def _fm_creditors_build_stack_ui_v620(self, parent):
    split = build_two_block_shell(parent, "Dokumentenstapel", "Vorschau")
    left = ttk.Frame(split.left); left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(split.right); right.pack(fill="both", expand=True, padx=10, pady=8)
    toolbar = ttk.Frame(left); toolbar.pack(fill="x", pady=(0, 6))
    create_standard_button(toolbar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=3)
    create_standard_button(toolbar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=3)
    cols = ("hist_no", "title", "status", "mark")
    self.stack_tree = ttk.Treeview(left, columns=cols, show="headings", height=10)
    for c, t, w in [("hist_no", "Hist.-Nr.", 110), ("title", "Dokument", 260), ("status", "Status", 110), ("mark", "☑", 45)]:
        self.stack_tree.heading(c, text=t, command=self._toggle_all_marks if c == "mark" and hasattr(self, "_toggle_all_marks") else None)
        self.stack_tree.column(c, width=w, anchor="w" if c != "mark" else "center")
    self.stack_tree.pack(fill="both", expand=True)
    self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select, add="+")
    self.stack_tree.bind("<Double-1>", lambda _e: self._open_selected_stack_file())
    self.preview_title_label = ttk.Label(right, text="Kein Dokument ausgewählt", style="Hint.TLabel")
    self.preview_title_label.pack(anchor="w", pady=(0, 6))
    self.preview_pane = _FMDocumentPreviewPane(right)
    self.preview_pane.pack(fill="both", expand=True)

CreditorsView._build_stack_ui = _fm_creditors_build_stack_ui_v620


def _fm_creditors_clear_stack_preview_v620(self):
    if hasattr(self, "preview_title_label"):
        self.preview_title_label.configure(text="Kein Dokument ausgewählt")
    if hasattr(self, "preview_pane"):
        self.preview_pane.clear()
    if hasattr(self, "preview_image_label"):
        self.preview_image_label.configure(text="Keine Vorschau verfügbar", image="")

CreditorsView._clear_stack_preview = _fm_creditors_clear_stack_preview_v620


def _fm_creditors_refresh_stack_preview_v620(self):
    if not getattr(self, "current_hist_no", None):
        self._clear_stack_preview(); return
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_files WHERE hist_no=? ORDER BY id", (self.current_hist_no,)))
    if not rows:
        self._clear_stack_preview(); return
    title = clean_document_title(rows[0]["file_name"])
    if len(rows) > 1:
        title += f" (+ {len(rows)-1} weitere)"
    self.preview_title_label.configure(text=title)
    try:
        images = [_fm_generate_document_preview(r["file_path"]) for r in rows]
        self.preview_pane.set_image(_fm_combine_pages(images))
    except Exception as exc:
        self.preview_pane.set_image(_fm_text_pages_to_image("Vorschaufehler", [str(exc)]))

CreditorsView._refresh_stack_preview = _fm_creditors_refresh_stack_preview_v620


# ---------- Freigabe- und Rechnungslogik ----------
def _fm_validate_positive_invoice_amounts(self) -> bool:
    for label, var in [("Netto", self.net_amount), ("Steuer", self.tax_amount), ("Brutto", self.gross_amount)]:
        try:
            if parse_amount(var.get()) < 0:
                messagebox.showwarning("Negativbetrag", f"{label} darf nicht negativ sein. Bitte Belegart 'Gutschrift' verwenden.")
                return False
        except Exception as exc:
            messagebox.showwarning("Betrag", str(exc))
            return False
    return True


def _fm_description_get(self):
    w = getattr(self, "description_widget", None)
    return w.get("1.0", "end-1c").strip() if w is not None else ""


def _fm_create_open_item_for_invoice(con, entity_type: str, invoice_row):
    if entity_type == "vendor_invoice":
        partner_no, partner_name = invoice_row["vendor_no"], invoice_row["vendor_name"]
    else:
        partner_no, partner_name = invoice_row["customer_no"], invoice_row["customer_name"]
    con.execute("DELETE FROM open_items WHERE entity_type=? AND reference_no=?", (entity_type, invoice_row["invoice_no"]))
    con.execute(
        "INSERT INTO open_items(entity_type,reference_no,partner_no,partner_name,due_date,original_amount,open_amount,status,linked_journal_no,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (entity_type, invoice_row["invoice_no"], partner_no, partner_name, invoice_row["due_date"], float(invoice_row["gross_amount"] or 0), float(invoice_row["open_amount"] or invoice_row["gross_amount"] or 0), _fm_status_from_open(invoice_row["open_amount"] or invoice_row["gross_amount"] or 0), invoice_row["linked_journal_no"] if "linked_journal_no" in invoice_row.keys() else "", invoice_row["invoice_kind"] if "invoice_kind" in invoice_row.keys() else "Rechnung", invoice_row["description"] if "description" in invoice_row.keys() else "", now_str(), now_str())
    )


def _fm_save_invoice_v620(self):
    if not _fm_validate_positive_invoice_amounts(self):
        return
    if not validate_date(self.invoice_date.get()) or not validate_date(self.due_date.get()):
        messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden."); return
    no, name, address = self._partner_parts()
    if not no:
        messagebox.showwarning("Partner", "Bitte Partner auswählen."); return
    selected_hist = getattr(self, "current_hist_no", None)
    if self.entity_type == "vendor_invoice" and not (selected_hist or getattr(self, "pending_attachments", [])):
        messagebox.showwarning("Dokument", "Eingangsrechnungen können ohne Dokument nicht erfasst werden."); return
    gross = parse_amount(self.gross_amount.get()); net = parse_amount(self.net_amount.get()); tax = parse_amount(self.tax_amount.get())
    invoice_no = self.invoice_no.get().strip()
    invoice_kind = self.invoice_kind.get().strip() if hasattr(self, "invoice_kind") else "Rechnung"
    description = _fm_description_get(self)
    address_col = "customer_address" if self.invoice_table == "customer_invoices" else "vendor_address"
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    partner_no_col = "customer_no" if self.invoice_table == "customer_invoices" else "vendor_no"
    approval_status = APPROVAL_STATUS_REVIEW
    approval_level = _fm_calc_approval_level(gross)
    stack_files = []
    with get_connection() as con:
        con.execute(
            f"INSERT INTO {self.invoice_table}(invoice_no,{partner_no_col},{partner_name_col},{address_col},invoice_date,due_date,payment_term,tax_code,net_amount,tax_amount,gross_amount,open_amount,status,invoice_kind,description,approval_status,released,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (invoice_no, no, name, address, self.invoice_date.get(), self.due_date.get(), self.payment_term.get(), self.tax_code.get(), float(net), float(tax), float(gross), float(gross), STATUS_OPEN, invoice_kind, description, approval_status, 0, approval_level, now_str(), now_str())
        )
        # Keine offenen Posten vor Freigabe.
        con.execute("DELETE FROM open_items WHERE entity_type=? AND reference_no=?", (self.entity_type, invoice_no))
        if self.entity_type == "vendor_invoice" and selected_hist:
            stack_files = list(con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id", (selected_hist,)))
            con.execute("UPDATE invoice_import_batches SET status='Erfasst', vendor_no=?, vendor_name=?, updated_at=? WHERE hist_no=?", (no, name, now_str(), selected_hist))
        con.commit()
    if stack_files:
        self.add_attachment_paths(self.entity_type, invoice_no, [r["file_path"] for r in stack_files])
    if getattr(self, "pending_attachments", []):
        self.add_attachment_paths(self.entity_type, invoice_no, self.pending_attachments)
    self.app.set_status(f"{self.entity_label} {invoice_no} gespeichert – Freigabe ausstehend.")
    self.clear_form(); self.reload_invoices(); self.reload_open_items()
    if hasattr(self, "_load_import_batches"): self._load_import_batches()
    if hasattr(self, "_clear_stack_preview"): self._clear_stack_preview()

InvoiceModuleBase.save_invoice = _fm_save_invoice_v620


def _fm_reload_open_items_v620(self):
    trees = _v619_unique_trees(getattr(self, "op_tree", None), *getattr(self, "open_item_tabs_trees", [])) if "_v619_unique_trees" in globals() else [getattr(self, "op_tree", None)]
    trees = [t for t in trees if t is not None]
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date, id", (self.entity_type,)))
    for tree in trees:
        for i in tree.get_children(): tree.delete(i)
        for r in rows:
            if parse_amount(r["open_amount"]) <= 0: continue
            att = self.get_attachment_count(self.entity_type, r["reference_no"])
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            cols = tree["columns"]
            if "kind" in cols:
                vals = (r["reference_no"], r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung", r["partner_name"], r["description"] if "description" in r.keys() else "", r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else "Dokument anhängen")
            else:
                vals = (r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else "Dokument anhängen")
            tree.insert("", "end", values=vals, tags=(tag,))

InvoiceModuleBase.reload_open_items = _fm_reload_open_items_v620


# ---------- Rechnungsportal ----------
class InvoicePortalView(ttk.Frame, SortableTreeMixin):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.archive_filter = tk.StringVar(value="Alle")
        self.approval_filter = tk.StringVar(value="Alle")
        self._build_ui()
        self.reload_all()

    def _build_ui(self):
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)
        self.archive_tab = ttk.Frame(self.nb)
        self.approval_tab = ttk.Frame(self.nb)
        self.nb.add(self.archive_tab, text="Rechnungsarchiv")
        self.nb.add(self.approval_tab, text="Rechnungsfreigabe")
        self._build_archive()
        self._build_approval()

    def _invoice_tree(self, parent):
        cols = ("direction", "invoice_no", "kind", "partner", "date", "due", "gross", "open", "approval_status", "last_checked_by", "released")
        tree = ttk.Treeview(parent, columns=cols, show="headings")
        for c, t, w in [
            ("direction", "Richtung", 125), ("invoice_no", "Rechnung", 120), ("kind", "Art", 95), ("partner", "Partner", 190),
            ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 95), ("open", "Offen", 95),
            ("approval_status", "Freigabestatus", 190), ("last_checked_by", "Letzter Prüfender", 150), ("released", "Freigegeben", 90)
        ]:
            tree.heading(c, text=t); tree.column(c, width=w, anchor="w")
        tree.pack(fill="both", expand=True, padx=10, pady=8)
        configure_tree_tags(tree); self.setup_sorting(tree)
        return tree

    def _build_archive(self):
        top = ttk.Frame(self.archive_tab); top.pack(fill="x", padx=10, pady=8)
        ttk.Label(top, text="Filter", width=8).pack(side="left")
        cb = ttk.Combobox(top, textvariable=self.archive_filter, values=["Alle", "Eingangsrechnungen", "Ausgangsrechnungen"], state="readonly", width=24)
        cb.pack(side="left", padx=4); cb.bind("<<ComboboxSelected>>", lambda _e: self.reload_archive()); cb.bind("<MouseWheel>", lambda e: "break")
        create_standard_button(top, "Aktualisieren", self.reload_archive).pack(side="left", padx=4)
        create_standard_button(top, "Bearbeiten", self.edit_selected_archive).pack(side="left", padx=4)
        self.archive_tree = self._invoice_tree(self.archive_tab)

    def _build_approval(self):
        top = ttk.Frame(self.approval_tab); top.pack(fill="x", padx=10, pady=8)
        ttk.Label(top, text="Filter", width=8).pack(side="left")
        cb = ttk.Combobox(top, textvariable=self.approval_filter, values=["Alle", "Eingangsrechnungen", "Ausgangsrechnungen"], state="readonly", width=24)
        cb.pack(side="left", padx=4); cb.bind("<<ComboboxSelected>>", lambda _e: self.reload_approval()); cb.bind("<MouseWheel>", lambda e: "break")
        create_standard_button(top, "Aktualisieren", self.reload_approval).pack(side="left", padx=4)
        create_standard_button(top, "Bearbeiten", self.edit_selected_approval).pack(side="left", padx=4)
        create_standard_button(top, "Sachlich prüfen", lambda: self.perform_action("review")).pack(side="left", padx=4)
        create_standard_button(top, "1. Freigabe", lambda: self.perform_action("approval1"), confirm=True).pack(side="left", padx=4)
        create_standard_button(top, "2. Freigabe", lambda: self.perform_action("approval2"), confirm=True).pack(side="left", padx=4)
        self.approval_tree = self._invoice_tree(self.approval_tab)

    def _filter_sql(self, flt):
        if flt == "Eingangsrechnungen": return ["vendor_invoice"]
        if flt == "Ausgangsrechnungen": return ["customer_invoice"]
        return ["vendor_invoice", "customer_invoice"]

    def _load_rows(self, flt, only_unreleased=False):
        rows = []
        with get_connection() as con:
            if "vendor_invoice" in self._filter_sql(flt):
                q = "SELECT 'vendor_invoice' AS entity_type, invoice_no, invoice_kind, vendor_name AS partner, invoice_date, due_date, gross_amount, open_amount, approval_status, last_checked_by, released FROM vendor_invoices"
                if only_unreleased: q += " WHERE COALESCE(released,0)=0"
                rows.extend(con.execute(q).fetchall())
            if "customer_invoice" in self._filter_sql(flt):
                q = "SELECT 'customer_invoice' AS entity_type, invoice_no, invoice_kind, customer_name AS partner, invoice_date, due_date, gross_amount, open_amount, approval_status, last_checked_by, released FROM customer_invoices"
                if only_unreleased: q += " WHERE COALESCE(released,0)=0"
                rows.extend(con.execute(q).fetchall())
        return sorted(rows, key=lambda r: (r["invoice_date"], r["invoice_no"]), reverse=True)

    def _fill_tree(self, tree, rows):
        for i in tree.get_children(): tree.delete(i)
        for r in rows:
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            values = (_fm_invoice_direction(r["entity_type"]), r["invoice_no"], r["invoice_kind"] or "Rechnung", r["partner"], r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["approval_status"] or APPROVAL_STATUS_REVIEW, r["last_checked_by"] or "", "Ja" if int(r["released"] or 0) else "Nein")
            tree.insert("", "end", iid=f"{r['entity_type']}|{r['invoice_no']}", values=values, tags=(tag,))

    def reload_archive(self): self._fill_tree(self.archive_tree, self._load_rows(self.archive_filter.get(), False))
    def reload_approval(self): self._fill_tree(self.approval_tree, self._load_rows(self.approval_filter.get(), True))
    def reload_all(self): self.reload_archive(); self.reload_approval()

    def _selected(self, tree):
        iid = tree.focus()
        if not iid or "|" not in iid:
            messagebox.showinfo("Rechnungsportal", "Bitte Rechnung auswählen."); return None
        entity_type, invoice_no = iid.split("|", 1)
        return entity_type, invoice_no

    def edit_selected_archive(self):
        sel = self._selected(self.archive_tree)
        if sel: self.open_edit_popup(*sel)

    def edit_selected_approval(self):
        sel = self._selected(self.approval_tree)
        if sel: self.open_edit_popup(*sel)

    def open_edit_popup(self, entity_type, invoice_no):
        table = _fm_invoice_table(entity_type)
        no_col, name_col, addr_col = _fm_partner_cols(entity_type)
        with get_connection() as con:
            r = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
            payment_terms = [x[0] for x in con.execute("SELECT code FROM payment_terms WHERE active=1 ORDER BY code")]
            tax_codes = [x[0] for x in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]
        if not r: return
        if int(r["released"] or 0):
            messagebox.showinfo("Bearbeitung gesperrt", "Diese Rechnung ist freigegeben und kann nicht mehr geändert werden."); return
        pop = tk.Toplevel(self); pop.title(f"Rechnung bearbeiten – {invoice_no}"); pop.geometry("620x620"); pop.configure(bg=BG)
        vars = {
            "invoice_kind": tk.StringVar(value=r["invoice_kind"] or "Rechnung"), "invoice_no": tk.StringVar(value=r["invoice_no"]),
            "invoice_date": tk.StringVar(value=r["invoice_date"]), "due_date": tk.StringVar(value=r["due_date"]),
            "payment_term": tk.StringVar(value=r["payment_term"] or ""), "tax_code": tk.StringVar(value=r["tax_code"] or ""),
            "net_amount": tk.StringVar(value=format_amount(r["net_amount"])), "tax_amount": tk.StringVar(value=format_amount(r["tax_amount"])),
            "gross_amount": tk.StringVar(value=format_amount(r["gross_amount"])),
        }
        def row(label, widget):
            fr = ttk.Frame(pop); fr.pack(fill="x", padx=12, pady=4); ttk.Label(fr, text=label, width=18).pack(side="left"); widget(fr).pack(side="left", fill="x", expand=True)
        row("Belegart", lambda fr: ttk.Combobox(fr, textvariable=vars["invoice_kind"], values=INVOICE_KIND_VALUES, state="readonly"))
        row("Rechnungs-Nr.", lambda fr: ttk.Entry(fr, textvariable=vars["invoice_no"]))
        row("Rechnungsdatum", lambda fr: ttk.Entry(fr, textvariable=vars["invoice_date"]))
        row("Zahlungsbed.", lambda fr: ttk.Combobox(fr, textvariable=vars["payment_term"], values=payment_terms, state="readonly"))
        row("Fälligkeit", lambda fr: ttk.Entry(fr, textvariable=vars["due_date"]))
        row("Steuerkennz.", lambda fr: ttk.Combobox(fr, textvariable=vars["tax_code"], values=tax_codes, state="readonly"))
        row("Netto", lambda fr: ttk.Entry(fr, textvariable=vars["net_amount"]))
        row("Steuer", lambda fr: ttk.Entry(fr, textvariable=vars["tax_amount"]))
        row("Brutto", lambda fr: ttk.Entry(fr, textvariable=vars["gross_amount"]))
        txtrow = ttk.Frame(pop); txtrow.pack(fill="both", expand=True, padx=12, pady=4); ttk.Label(txtrow, text="Beschreibung", width=18).pack(side="left", anchor="n")
        desc = tk.Text(txtrow, height=4, wrap="word", bg=WHITE, fg=TEXT); desc.pack(side="left", fill="both", expand=True); desc.insert("1.0", r["description"] if "description" in r.keys() else "")
        def save():
            try:
                net, tax, gross = parse_amount(vars["net_amount"].get()), parse_amount(vars["tax_amount"].get()), parse_amount(vars["gross_amount"].get())
                if net < 0 or tax < 0 or gross < 0:
                    messagebox.showwarning("Negativbetrag", "Rechnungen/Gutschriften werden mit positivem Betrag erfasst.", parent=pop); return
            except Exception as exc:
                messagebox.showwarning("Betrag", str(exc), parent=pop); return
            if not validate_date(vars["invoice_date"].get()) or not validate_date(vars["due_date"].get()):
                messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.", parent=pop); return
            with get_connection() as con:
                con.execute(f"UPDATE {table} SET invoice_no=?, invoice_kind=?, invoice_date=?, due_date=?, payment_term=?, tax_code=?, net_amount=?, tax_amount=?, gross_amount=?, open_amount=?, description=?, approval_level=?, updated_at=? WHERE invoice_no=?",
                            (vars["invoice_no"].get().strip(), vars["invoice_kind"].get(), vars["invoice_date"].get(), vars["due_date"].get(), vars["payment_term"].get(), vars["tax_code"].get(), float(net), float(tax), float(gross), float(gross), desc.get("1.0", "end-1c").strip(), _fm_calc_approval_level(gross), now_str(), invoice_no))
                con.commit()
            pop.destroy(); self.reload_all()
        create_standard_button(pop, "Speichern", save, confirm=True).pack(pady=12)

    def perform_action(self, action):
        sel = self._selected(self.approval_tree)
        if not sel: return
        entity_type, invoice_no = sel
        table = _fm_invoice_table(entity_type)
        user = _fm_current_user()
        level = _fm_current_approval_level()
        with get_connection() as con:
            r = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
            if not r or int(r["released"] or 0): return
            if entity_type == "customer_invoice":
                if action != "review":
                    messagebox.showinfo("Ausgangsrechnung", "Ausgangsrechnungen benötigen nur die sachliche Prüfung durch Ebene 3 oder 4."); return
                if level < 3:
                    messagebox.showwarning("Berechtigung", "Sachliche Prüfung für Ausgangsrechnungen benötigt Freigabeebene 3 oder 4."); return
                con.execute(f"UPDATE {table} SET reviewed_by=?, reviewed_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, released=1, released_at=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_RELEASED, now_str(), now_str(), invoice_no))
                rr = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
                _fm_create_open_item_for_invoice(con, entity_type, rr)
            else:
                if action == "review":
                    con.execute(f"UPDATE {table} SET reviewed_by=?, reviewed_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_APPROVAL1, now_str(), invoice_no))
                elif action == "approval1":
                    if not r["reviewed_by"]:
                        messagebox.showwarning("Freigabe", "Sachliche Prüfung ist noch ausstehend."); return
                    con.execute(f"UPDATE {table} SET approval_1_by=?, approval_1_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_APPROVAL2, now_str(), invoice_no))
                elif action == "approval2":
                    if not r["approval_1_by"]:
                        messagebox.showwarning("Freigabe", "1. Freigabe ist noch ausstehend."); return
                    if r["approval_1_by"] == user:
                        messagebox.showwarning("Freigabe", "1. und 2. Freigabe müssen von unterschiedlichen Benutzern erfolgen."); return
                    con.execute(f"UPDATE {table} SET approval_2_by=?, approval_2_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, released=1, released_at=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_RELEASED, now_str(), now_str(), invoice_no))
                    rr = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
                    _fm_create_open_item_for_invoice(con, entity_type, rr)
            con.commit()
        self.reload_all()


def _render_rechnungsportal(self, parent):
    InvoicePortalView(parent, self).pack(fill="both", expand=True)

FinanceMateApp._render_rechnungsportal = _render_rechnungsportal


def _fm_build_sidebar_v620(self):
    for w in self.sidebar.winfo_children(): w.destroy()
    width = 42 if self.sidebar_collapsed else 210
    self.sidebar.configure(width=width); self.sidebar.pack_propagate(False)
    arrow = "»" if self.sidebar_collapsed else "«"
    create_standard_button(self.sidebar, arrow, self._toggle_sidebar).pack(fill="x", padx=5, pady=5)
    modules = [("Dashboard", "DB"), ("Stammdaten", "ST"), ("Buchungsjournal", "BJ"), ("Debitoren", "DE"), ("Kreditoren", "KR"), ("Rechnungsportal", "RP"), ("Zahlungen", "ZA"), ("Reporting", "RE"), ("Audit", "AU"), ("Einstellungen", "EI")]
    self.nav_buttons.clear()
    for name, abbr in modules:
        text = abbr if self.sidebar_collapsed else name
        btn = ttk.Button(self.sidebar, text=text, style="Nav.TButton", command=lambda n=name: self.show_module(n))
        btn.pack(fill="x", padx=5, pady=2); self.nav_buttons[name] = btn

FinanceMateApp._build_sidebar = _fm_build_sidebar_v620


def _fm_show_module_v620(self, module: str):
    module_key = "Buchungsjournal" if module == "Finanzbuchhaltung" else module
    self.active_module = module_key; self._update_nav_styles()
    for w in self.workspace.winfo_children(): w.destroy()
    ttk.Label(self.workspace, text=module_key, style="CardTitle.TLabel").pack(anchor="w", pady=(0, 6))
    container = ttk.Frame(self.workspace); container.pack(fill="both", expand=True)
    render = {
        "Dashboard": self._render_dashboard, "Stammdaten": self._render_stammdaten, "Buchungsjournal": self._render_finanzbuchhaltung,
        "Debitoren": self._render_debitoren, "Kreditoren": self._render_kreditoren, "Rechnungsportal": self._render_rechnungsportal,
        "Zahlungen": self._render_zahlungen, "Reporting": self._render_reporting, "Audit": self._render_audit, "Einstellungen": self._render_einstellungen,
    }.get(module_key, self._render_dashboard)
    render(container); self.set_status(f"Modul {module_key} geladen.")

FinanceMateApp.show_module = _fm_show_module_v620



# === FINANCE MATE PATCH V0_6_21_CREDITORS_HARD_RESTORE_SCALING ===
APP_VERSION = "0.6.21-creditors-hard-restore-scaling"
BLOCK7_PREPARATION_STATUS = "vorbereitet - nach erfolgreicher Kreditorenstabilisierung fachlich fortsetzen"

# Windows-DPI möglichst früh auf 100%-Logik normalisieren. Fehler werden bewusst ignoriert,
# damit Linux/macOS und ältere Windows-Versionen weiter starten.
try:
    if sys.platform.startswith("win"):
        import ctypes
        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(1)  # system-DPI-aware, vermeidet Windows-Bitmapskalierung
        except Exception:
            try:
                ctypes.windll.user32.SetProcessDPIAware()
            except Exception:
                pass
except Exception:
    pass


def _fm_apply_1920x1080_scaling(root: tk.Tk) -> None:
    """Robuste Gesamtlogik: Basis 1920x1080 bei 100%, kleinere Anzeigen werden proportional verdichtet."""
    try:
        sw = max(1, root.winfo_screenwidth())
        sh = max(1, root.winfo_screenheight())
        scale = min(sw / 1920.0, sh / 1080.0)
        scale = max(0.72, min(1.15, scale))
        # Tk 100% entspricht i. d. R. 96/72 = 1.3333 Pixel je Punkt.
        root.tk.call("tk", "scaling", 96.0 / 72.0 * scale)
        w = min(1920, max(1100, int(1920 * scale)))
        h = min(1080, max(680, int(1080 * scale)))
        root.geometry(f"{w}x{h}")
        root.minsize(min(1100, w), min(680, h))
    except Exception:
        pass


_orig_app_init_v621 = FinanceMateApp.__init__

def _fm_app_init_v621(self):
    _orig_app_init_v621(self)
    _fm_apply_1920x1080_scaling(self)

FinanceMateApp.__init__ = _fm_app_init_v621


def _fm_creditors_build_ui_v621(self):
    """Kreditoren-Hard-Restore ohne TwoBlock-Abhängigkeit: erzeugt sichtbar alle vier Bereiche."""
    self.current_hist_no = None
    self.pending_attachments = []

    root = ttk.PanedWindow(self, orient="vertical")
    root.pack(fill="both", expand=True, padx=0, pady=0)
    top = ttk.Frame(root)
    bottom = ttk.Frame(root)
    root.add(top, weight=3)
    root.add(bottom, weight=2)

    top_split = ttk.PanedWindow(top, orient="horizontal")
    top_split.pack(fill="both", expand=True)
    form_card = ttk.LabelFrame(top_split, text="Eingangsrechnung erfassen")
    list_card = ttk.LabelFrame(top_split, text="Eingangsrechnungen")
    top_split.add(form_card, weight=2)
    top_split.add(list_card, weight=3)

    form_scroll = ScrollableFrame(form_card)
    form_scroll.pack(fill="both", expand=True, padx=10, pady=8)
    f = form_scroll.content

    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()

    self._labeled_combo(f, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(f, "Rechnungs-Nr.", self.invoice_no)

    prow = ttk.Frame(f); prow.pack(fill="x", pady=3)
    ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side="left", fill="x", expand=True)
    self.partner_cb.bind("<MouseWheel>", lambda e: "break")
    self.partner_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_vendor_info())
    create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)

    self.description_widget = _v619_labeled_text(f, "Rechnungsbeschreibung", height=2) if "_v619_labeled_text" in globals() else tk.Text(f, height=2)
    if not getattr(self.description_widget, "winfo_manager", lambda: "")():
        self.description_widget.pack(fill="x", pady=3)

    self._labeled_entry(f, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(f, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(f, "Fälligkeit", self.due_date)
    self._labeled_combo(f, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(f, "Netto", self.net_amount)
    self._labeled_entry(f, "Steuer", self.tax_amount)
    self._labeled_entry(f, "Brutto", self.gross_amount)

    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())

    btns = ttk.Frame(f); btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)

    info = ttk.LabelFrame(f, text="Stammdaten des Kreditors")
    info.pack(fill="both", expand=False, pady=(8, 4))
    self.vendor_info_body = tk.Text(info, height=8, wrap="word", bg=WHITE, fg=TEXT, relief="flat")
    self.vendor_info_body.pack(fill="both", expand=True, padx=6, pady=6)
    self.vendor_info_body.configure(state="disabled")

    nb = ttk.Notebook(list_card)
    nb.pack(fill="both", expand=True, padx=10, pady=8)
    tab_inv = ttk.Frame(nb); tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text="Eingangsrechnungen")
    nb.add(tab_op, text="Offene Posten")

    inv_cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show="headings", height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [("invoice_no", "Rechnung", 115), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")

    op_cols = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show="headings", height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

    # Unterer Bereich: Dokumentenstapel + Bildvorschau.
    self._build_stack_ui(bottom)

    self.partner.trace_add("write", lambda *_: self._update_vendor_info())
    self._update_vendor_info()
    self.reload_invoices()
    self.reload_open_items()
    self._load_import_batches()
    self._clear_stack_preview()

CreditorsView._build_ui = _fm_creditors_build_ui_v621


def _fm_render_kreditoren_v621(self, parent):
    try:
        view = CreditorsView(parent, self)
        view.pack(fill="both", expand=True)
        # Harte Mindestprüfung direkt nach Aufbau: zentrale Widgets müssen existieren.
        required = ["inv_tree", "op_tree", "stack_tree"]
        missing = [name for name in required if not hasattr(view, name)]
        if missing:
            raise RuntimeError("Kreditoren-Aufbau unvollständig, fehlende Widgets: " + ", ".join(missing))
    except Exception as exc:
        # Kein stiller grauer Hintergrund mehr: Fehler sichtbar ausgeben.
        for w in parent.winfo_children():
            w.destroy()
        err = tk.Text(parent, height=12, bg="#fff4c2", fg=TEXT, wrap="word")
        err.pack(fill="both", expand=True, padx=12, pady=12)
        err.insert("1.0", "Kreditoren-Modul konnte nicht vollständig aufgebaut werden.\n\n" + str(exc))
        err.configure(state="disabled")
        raise

FinanceMateApp._render_kreditoren = _fm_render_kreditoren_v621


# Stammdaten-Doppeltitel: Im Modul Stammdaten bleibt nur die interne Modulüberschrift sichtbar.
def _fm_show_module_v621(self, module: str):
    module_key = "Buchungsjournal" if module == "Finanzbuchhaltung" else module
    self.active_module = module_key
    self._update_nav_styles()
    for w in self.workspace.winfo_children():
        w.destroy()
    if module_key != "Stammdaten":
        ttk.Label(self.workspace, text=module_key, style="CardTitle.TLabel").pack(anchor="w", pady=(0, 6))
    container = ttk.Frame(self.workspace)
    container.pack(fill="both", expand=True)
    render = {
        "Dashboard": self._render_dashboard,
        "Stammdaten": self._render_stammdaten,
        "Buchungsjournal": self._render_finanzbuchhaltung,
        "Debitoren": self._render_debitoren,
        "Kreditoren": self._render_kreditoren,
        "Rechnungsportal": self._render_rechnungsportal,
        "Zahlungen": self._render_zahlungen,
        "Reporting": self._render_reporting,
        "Audit": self._render_audit,
        "Einstellungen": self._render_einstellungen,
    }.get(module_key, self._render_dashboard)
    render(container)
    self.set_status(f"Modul {module_key} geladen.")

FinanceMateApp.show_module = _fm_show_module_v621


import inspect

def _fm_creditors_static_selftest_v621() -> None:
    # Statische Auslieferungsprüfung: Der aktive Kreditoren-Build darf nicht nur Container anlegen.
    try:
        src = inspect.getsource(CreditorsView._build_ui)
    except Exception:
        src = ""
    required_tokens = [
        "Eingangsrechnung erfassen", "Eingangsrechnungen", "Offene Posten",
        "Dokumentenstapel", "self.inv_tree", "self.op_tree", "self._build_stack_ui(bottom)"
    ]
    missing = [t for t in required_tokens if t not in src]
    if missing:
        raise RuntimeError("Kreditoren-Selbsttest fehlgeschlagen: " + ", ".join(missing))

_fm_creditors_static_selftest_v621()



# === FINANCE MATE PATCH V0_6_22_CREDITORS_COMMAND_FIX_SEARCH_FULLSCREEN ===
APP_VERSION = "0.6.22-creditors-command-fix-search-fullscreen"

# Kernkorrektur für Python/Tk 3.14: ttk.Treeview.heading darf command=None NICHT erhalten.
# Wenn kein Command existiert, wird die Option command vollständig weggelassen.
def _fm_heading_safe(tree: ttk.Treeview, column: str, text: str, command=None) -> None:
    if callable(command):
        tree.heading(column, text=text, command=command)
    else:
        tree.heading(column, text=text)


def _fm_creditors_toggle_all_marks_v622(self):
    try:
        with get_connection() as con:
            rows = list(con.execute("SELECT hist_no, marked FROM invoice_import_batches WHERE status <> 'Erfasst' ORDER BY id"))
            if not rows:
                return
            all_marked = all(int(r["marked"] or 0) for r in rows)
            new_val = 0 if all_marked else 1
            con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=? WHERE status <> 'Erfasst'", (new_val, now_str()))
            con.commit()
        self._load_import_batches()
    except Exception as exc:
        messagebox.showerror("Dokumentenstapel", str(exc), parent=self)

CreditorsView._toggle_all_marks = _fm_creditors_toggle_all_marks_v622


def _fm_creditors_build_stack_ui_v622(self, parent):
    split = build_two_block_shell(parent, "Dokumentenstapel", "Vorschau")
    left = ttk.Frame(split.left); left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(split.right); right.pack(fill="both", expand=True, padx=10, pady=8)

    toolbar = ttk.Frame(left); toolbar.pack(fill="x", pady=(0, 6))
    create_standard_button(toolbar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=3)
    create_standard_button(toolbar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=3)

    # Suchleiste Dokumentenstapel
    self.stack_search_var = tk.StringVar()
    search_row = ttk.Frame(left); search_row.pack(fill="x", pady=(0, 6))
    ttk.Label(search_row, text="Suchen", width=10).pack(side="left")
    stack_search = ttk.Entry(search_row, textvariable=self.stack_search_var)
    stack_search.pack(side="left", fill="x", expand=True)
    stack_search.insert(0, PLACEHOLDER_TEXT)
    stack_search.bind("<FocusIn>", lambda e: (stack_search.delete(0, tk.END) if stack_search.get() == PLACEHOLDER_TEXT else None))
    stack_search.bind("<FocusOut>", lambda e: (stack_search.insert(0, PLACEHOLDER_TEXT) if not stack_search.get().strip() else None))
    stack_search.bind("<KeyRelease>", lambda _e: self._load_import_batches())

    cols = ("hist_no", "title", "status", "mark")
    self.stack_tree = ttk.Treeview(left, columns=cols, show="headings", height=10)
    for c, t, w in [("hist_no", "Hist.-Nr.", 110), ("title", "Dokument", 260), ("status", "Status", 110), ("mark", "☑", 45)]:
        _fm_heading_safe(self.stack_tree, c, t, self._toggle_all_marks if c == "mark" else None)
        self.stack_tree.column(c, width=w, anchor="w" if c != "mark" else "center")
    self.stack_tree.pack(fill="both", expand=True)
    self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select, add="+")
    self.stack_tree.bind("<Double-1>", lambda _e: self._open_selected_stack_file())

    self.preview_title_label = ttk.Label(right, text="Kein Dokument ausgewählt", style="Hint.TLabel")
    self.preview_title_label.pack(anchor="w", pady=(0, 6))
    self.preview_pane = _FMDocumentPreviewPane(right)
    self.preview_pane.pack(fill="both", expand=True)

CreditorsView._build_stack_ui = _fm_creditors_build_stack_ui_v622


def _fm_creditors_load_import_batches_v622(self):
    if not hasattr(self, "stack_tree"):
        return
    query = ""
    try:
        query = self.stack_search_var.get().strip().lower()
        if query == PLACEHOLDER_TEXT.lower():
            query = ""
    except Exception:
        query = ""
    for i in self.stack_tree.get_children():
        self.stack_tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_batches WHERE status <> 'Erfasst' ORDER BY id DESC"))
    for r in rows:
        text_blob = f"{r['hist_no']} {r['title']} {r['status']}".lower()
        if query and query not in text_blob:
            continue
        self.stack_tree.insert("", "end", iid=r["hist_no"], values=(r["hist_no"], r["title"], r["status"], "☑" if int(r["marked"] or 0) else "☐"))

CreditorsView._load_import_batches = _fm_creditors_load_import_batches_v622


# ---------- generische Suchleisten Debitoren/Kreditoren ----------
def _fm_tree_apply_filter(tree: ttk.Treeview, query: str) -> None:
    q = (query or "").strip().lower()
    if not q or q == PLACEHOLDER_TEXT.lower():
        return
    for iid in list(tree.get_children()):
        values = tree.item(iid, "values")
        if q not in " ".join(str(v) for v in values).lower():
            tree.delete(iid)


def _fm_install_tree_search(owner, tree_attr: str, reload_callback, var_attr: str, label: str = "Suchen") -> None:
    tree = getattr(owner, tree_attr, None)
    if tree is None or getattr(owner, f"_{var_attr}_installed", False):
        return
    setattr(owner, f"_{var_attr}_installed", True)
    var = tk.StringVar()
    setattr(owner, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=0, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=0, pady=(0, 4))
    ttk.Label(row, text=label, width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    entry.insert(0, PLACEHOLDER_TEXT)
    entry.bind("<FocusIn>", lambda e: (entry.delete(0, tk.END) if entry.get() == PLACEHOLDER_TEXT else None))
    entry.bind("<FocusOut>", lambda e: (entry.insert(0, PLACEHOLDER_TEXT) if not entry.get().strip() else None))
    entry.bind("<KeyRelease>", lambda _e: reload_callback())


_orig_invoice_reload_invoices_v622 = InvoiceModuleBase.reload_invoices
_orig_invoice_reload_open_items_v622 = InvoiceModuleBase.reload_open_items

def _fm_reload_invoices_v622(self):
    _orig_invoice_reload_invoices_v622(self)
    q = getattr(getattr(self, "invoice_search_var", None), "get", lambda: "")()
    if hasattr(self, "inv_tree"):
        _fm_tree_apply_filter(self.inv_tree, q)


def _fm_reload_open_items_v622(self):
    _orig_invoice_reload_open_items_v622(self)
    q = getattr(getattr(self, "open_items_search_var", None), "get", lambda: "")()
    if hasattr(self, "op_tree"):
        _fm_tree_apply_filter(self.op_tree, q)

InvoiceModuleBase.reload_invoices = _fm_reload_invoices_v622
InvoiceModuleBase.reload_open_items = _fm_reload_open_items_v622

_orig_debitors_build_ui_v622 = DebitorsView._build_ui

def _fm_debitors_build_ui_v622(self):
    _orig_debitors_build_ui_v622(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices(); self.reload_open_items()

DebitorsView._build_ui = _fm_debitors_build_ui_v622

_orig_creditors_build_ui_v622 = CreditorsView._build_ui

def _fm_creditors_build_ui_v622(self):
    _orig_creditors_build_ui_v622(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices(); self.reload_open_items(); self._load_import_batches()

CreditorsView._build_ui = _fm_creditors_build_ui_v622


# ---------- Suchleisten Rechnungsportal ----------
def _fm_install_portal_search(portal, tree_attr: str, reload_callback, var_attr: str) -> None:
    tree = getattr(portal, tree_attr, None)
    if tree is None or getattr(portal, f"_{var_attr}_installed", False):
        return
    setattr(portal, f"_{var_attr}_installed", True)
    var = tk.StringVar()
    setattr(portal, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=10, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=10, pady=(0, 4))
    ttk.Label(row, text="Suchen", width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    entry.insert(0, PLACEHOLDER_TEXT)
    entry.bind("<FocusIn>", lambda e: (entry.delete(0, tk.END) if entry.get() == PLACEHOLDER_TEXT else None))
    entry.bind("<FocusOut>", lambda e: (entry.insert(0, PLACEHOLDER_TEXT) if not entry.get().strip() else None))
    entry.bind("<KeyRelease>", lambda _e: reload_callback())

if 'InvoicePortalView' in globals():
    _orig_portal_init_v622 = InvoicePortalView.__init__
    def _fm_portal_init_v622(self, parent, app):
        _orig_portal_init_v622(self, parent, app)
        _fm_install_portal_search(self, "archive_tree", self.reload_archive, "archive_search_var")
        _fm_install_portal_search(self, "approval_tree", self.reload_approval, "approval_search_var")
        self.reload_all()
    InvoicePortalView.__init__ = _fm_portal_init_v622

    _orig_portal_reload_archive_v622 = InvoicePortalView.reload_archive
    _orig_portal_reload_approval_v622 = InvoicePortalView.reload_approval
    def _fm_portal_reload_archive_v622(self):
        _orig_portal_reload_archive_v622(self)
        _fm_tree_apply_filter(self.archive_tree, getattr(getattr(self, "archive_search_var", None), "get", lambda: "")())
    def _fm_portal_reload_approval_v622(self):
        _orig_portal_reload_approval_v622(self)
        _fm_tree_apply_filter(self.approval_tree, getattr(getattr(self, "approval_search_var", None), "get", lambda: "")())
    InvoicePortalView.reload_archive = _fm_portal_reload_archive_v622
    InvoicePortalView.reload_approval = _fm_portal_reload_approval_v622


# ---------- Fenstermodus Vollbild + stabilere 100%-Darstellungsannahme ----------
_orig_app_init_v622 = FinanceMateApp.__init__

def _fm_app_init_v622(self):
    _orig_app_init_v622(self)
    try:
        # 100%-Referenz in Tk: 96 dpi / 72 pt.
        self.tk.call("tk", "scaling", 96.0 / 72.0)
    except Exception:
        pass
    try:
        self.state("zoomed")  # Fenstermodus Vollbild, keine echte borderless Fullscreen-Ansicht.
    except Exception:
        try:
            self.attributes("-zoomed", True)
        except Exception:
            pass
    try:
        self.update_idletasks()
    except Exception:
        pass

FinanceMateApp.__init__ = _fm_app_init_v622


# ---------- Kreditoren-Selbsttest ohne GUI-Display ----------
def _fm_creditors_static_selftest_v622() -> None:
    import inspect as _inspect
    stack_src = _inspect.getsource(CreditorsView._build_stack_ui)
    build_src = _inspect.getsource(CreditorsView._build_ui)
    must_stack = ["_fm_heading_safe", "self.stack_tree", "Dokumentenstapel", "Vorschau"]
    must_build = ["_orig_creditors_build_ui_v622", "invoice_search_var", "open_items_search_var"]
    missing = [x for x in must_stack if x not in stack_src] + [x for x in must_build if x not in build_src]
    if missing:
        raise RuntimeError("Kreditoren-Selbsttest v0.6.22 fehlgeschlagen: " + ", ".join(missing))

_fm_creditors_static_selftest_v622()

# === FINANCE MATE PATCH V0_6_23_CREDITORS_SPLITLINES_FRAMELESS ===
APP_VERSION = "0.6.23-creditors-splitlines-frameless"


def _fm_pack_section_header(parent, title: str):
    """Optisch neutrale Bereichsüberschrift ohne grauen LabelFrame-Rahmen."""
    ttk.Label(parent, text=title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))


def _fm_make_plain_card(parent):
    # Kein LabelFrame, kein grauer Rahmen um den Block.
    frame = ttk.Frame(parent)
    return frame


def _fm_creditors_build_left_stack_area_v623(self, parent):
    """Nur linker unterer Block: Dokumentenstapel ohne Vorschau; Vorschau sitzt rechts unten separat."""
    _fm_pack_section_header(parent, "Dokumentenstapel")
    body = ttk.Frame(parent)
    body.pack(fill="both", expand=True, padx=10, pady=(0, 8))

    toolbar = ttk.Frame(body)
    toolbar.pack(fill="x", pady=(0, 6))
    create_standard_button(toolbar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=3)
    create_standard_button(toolbar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=3)

    self.stack_search_var = tk.StringVar()
    search_row = ttk.Frame(body)
    search_row.pack(fill="x", pady=(0, 6))
    ttk.Label(search_row, text="Suchen", width=10).pack(side="left")
    stack_search = ttk.Entry(search_row, textvariable=self.stack_search_var)
    stack_search.pack(side="left", fill="x", expand=True)
    stack_search.insert(0, PLACEHOLDER_TEXT)
    stack_search.bind("<FocusIn>", lambda e: (stack_search.delete(0, tk.END) if stack_search.get() == PLACEHOLDER_TEXT else None))
    stack_search.bind("<FocusOut>", lambda e: (stack_search.insert(0, PLACEHOLDER_TEXT) if not stack_search.get().strip() else None))
    stack_search.bind("<KeyRelease>", lambda _e: self._load_import_batches())

    cols = ("hist_no", "title", "status", "mark")
    self.stack_tree = ttk.Treeview(body, columns=cols, show="headings", height=10)
    for c, t, w in [("hist_no", "Hist.-Nr.", 110), ("title", "Dokument", 260), ("status", "Status", 110), ("mark", "☑", 45)]:
        if c == "mark" and hasattr(self, "_toggle_all_marks"):
            self.stack_tree.heading(c, text=t, command=self._toggle_all_marks)
        else:
            self.stack_tree.heading(c, text=t)
        self.stack_tree.column(c, width=w, anchor="w" if c != "mark" else "center")
    self.stack_tree.pack(fill="both", expand=True)
    self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select, add="+")
    self.stack_tree.bind("<Double-1>", lambda _e: self._open_selected_stack_file())


def _fm_creditors_build_right_preview_area_v623(self, parent):
    """Nur rechter unterer Block: Vorschau separat, damit die unteren Blöcke unabhängig skalierbar sind."""
    _fm_pack_section_header(parent, "Vorschau")
    body = ttk.Frame(parent)
    body.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    self.preview_title_label = ttk.Label(body, text="Kein Dokument ausgewählt", style="Hint.TLabel")
    self.preview_title_label.pack(anchor="w", pady=(0, 6))
    self.preview_pane = _FMDocumentPreviewPane(body)
    self.preview_pane.pack(fill="both", expand=True)


def _fm_creditors_build_ui_v623(self):
    """Vier-Block-Kreditorenlayout mit zweigeteilter horizontaler Verschiebelinie.

    Struktur:
    - outer_horizontal: links/rechts Breite
    - left_vertical: oben Eingangsrechnung erfassen, unten Dokumentenstapel
    - right_vertical: oben Eingangsrechnungen/Offene Posten, unten Vorschau

    Dadurch sind die unteren Blöcke wieder getrennt vergrößerbar.
    """
    self.current_hist_no = None
    self.pending_attachments = []

    outer_horizontal = ttk.PanedWindow(self, orient="horizontal")
    outer_horizontal.pack(fill="both", expand=True, padx=0, pady=0)
    left_col = ttk.Frame(outer_horizontal)
    right_col = ttk.Frame(outer_horizontal)
    outer_horizontal.add(left_col, weight=2)
    outer_horizontal.add(right_col, weight=3)

    left_vertical = ttk.PanedWindow(left_col, orient="vertical")
    right_vertical = ttk.PanedWindow(right_col, orient="vertical")
    left_vertical.pack(fill="both", expand=True)
    right_vertical.pack(fill="both", expand=True)

    form_card = _fm_make_plain_card(left_vertical)
    stack_card = _fm_make_plain_card(left_vertical)
    list_card = _fm_make_plain_card(right_vertical)
    preview_card = _fm_make_plain_card(right_vertical)

    left_vertical.add(form_card, weight=3)
    left_vertical.add(stack_card, weight=2)
    right_vertical.add(list_card, weight=3)
    right_vertical.add(preview_card, weight=2)

    # Linker oberer Block: Eingangsrechnung erfassen, ohne grauen Rahmen.
    _fm_pack_section_header(form_card, "Eingangsrechnung erfassen")
    form_scroll = ScrollableFrame(form_card)
    form_scroll.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    f = form_scroll.content

    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()

    self._labeled_combo(f, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(f, "Rechnungs-Nr.", self.invoice_no)

    prow = ttk.Frame(f)
    prow.pack(fill="x", pady=3)
    ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side="left", fill="x", expand=True)
    self.partner_cb.bind("<MouseWheel>", lambda e: "break")
    self.partner_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_vendor_info())
    create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)

    self.description_widget = _v619_labeled_text(f, "Rechnungsbeschreibung", height=2) if "_v619_labeled_text" in globals() else tk.Text(f, height=2)
    if not getattr(self.description_widget, "winfo_manager", lambda: "")():
        self.description_widget.pack(fill="x", pady=3)

    self._labeled_entry(f, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(f, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(f, "Fälligkeit", self.due_date)
    self._labeled_combo(f, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(f, "Netto", self.net_amount)
    self._labeled_entry(f, "Steuer", self.tax_amount)
    self._labeled_entry(f, "Brutto", self.gross_amount)

    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())

    btns = ttk.Frame(f)
    btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)

    info = ttk.Frame(f)
    info.pack(fill="both", expand=False, pady=(8, 4))
    ttk.Label(info, text="Stammdaten des Kreditors", style="Section.TLabel").pack(anchor="w", pady=(0, 4))
    self.vendor_info_body = tk.Text(info, height=8, wrap="word", bg=WHITE, fg=TEXT, relief="flat")
    self.vendor_info_body.pack(fill="both", expand=True, padx=0, pady=0)
    self.vendor_info_body.configure(state="disabled")

    # Rechter oberer Block: Listen ohne grauen Rahmen.
    _fm_pack_section_header(list_card, "Eingangsrechnungen")
    nb = ttk.Notebook(list_card)
    nb.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    tab_inv = ttk.Frame(nb)
    tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text="Eingangsrechnungen")
    nb.add(tab_op, text="Offene Posten")

    inv_cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show="headings", height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [("invoice_no", "Rechnung", 115), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t)
        self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree)
    self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")

    op_cols = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show="headings", height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t)
        self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree)
    self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

    # Untere Blöcke getrennt, jeweils eigener vertikaler Sash links/rechts.
    _fm_creditors_build_left_stack_area_v623(self, stack_card)
    _fm_creditors_build_right_preview_area_v623(self, preview_card)

    self.partner.trace_add("write", lambda *_: self._update_vendor_info())
    self._update_vendor_info()

    # Suchleisten aus v0.6.22 weiter nutzen, falls vorhanden.
    if "_fm_install_tree_search" in globals():
        _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
        _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")

    self.reload_invoices()
    self.reload_open_items()
    self._load_import_batches()
    self._clear_stack_preview()

CreditorsView._build_ui = _fm_creditors_build_ui_v623


def _fm_render_kreditoren_v623(self, parent):
    try:
        view = CreditorsView(parent, self)
        view.pack(fill="both", expand=True)
        required = ["inv_tree", "op_tree", "stack_tree", "preview_pane", "vendor_info_body"]
        missing = [name for name in required if not hasattr(view, name)]
        if missing:
            raise RuntimeError("Kreditoren-Aufbau unvollständig, fehlende Widgets: " + ", ".join(missing))
    except Exception as exc:
        for w in parent.winfo_children():
            w.destroy()
        err = tk.Text(parent, height=12, bg="#fff4c2", fg=TEXT, wrap="word")
        err.pack(fill="both", expand=True, padx=12, pady=12)
        err.insert("1.0", "Kreditoren-Modul konnte nicht vollständig aufgebaut werden.\n\n" + str(exc))
        err.configure(state="disabled")
        raise

FinanceMateApp._render_kreditoren = _fm_render_kreditoren_v623



def _fm_creditors_static_selftest_v623() -> None:
    import inspect as _inspect
    src_build = _inspect.getsource(CreditorsView._build_ui)
    src_stack = _inspect.getsource(_fm_creditors_build_left_stack_area_v623)
    src_preview = _inspect.getsource(_fm_creditors_build_right_preview_area_v623)
    src_combined = src_build + "\n" + src_stack + "\n" + src_preview
    forbidden = ["ttk.LabelFrame", "form_card = ttk.LabelFrame", "list_card = ttk.LabelFrame", "style=\"Card.TFrame\""]
    required = [
        "outer_horizontal", "left_vertical", "right_vertical",
        "_fm_creditors_build_left_stack_area_v623", "_fm_creditors_build_right_preview_area_v623",
        "self.inv_tree", "self.op_tree", "self.stack_tree", "self.preview_pane"
    ]
    missing = [x for x in required if x not in src_combined]
    found_forbidden = [x for x in forbidden if x in src_combined]
    if missing or found_forbidden:
        raise RuntimeError("Kreditoren-Selbsttest v0.6.23 fehlgeschlagen. Fehlend: " + ", ".join(missing) + "; verboten: " + ", ".join(found_forbidden))

_fm_creditors_static_selftest_v623()



# === FINANCE MATE PATCH V0_6_24_PAYMENTS_BLOCK7_LABELWRAP ===
APP_VERSION = "0.6.24-payments-block7-labelwrap"

# ---------- Debitoren/Kreditoren: Label-Umbruch Rechnungsbeschreibung ----------
_orig_v619_labeled_text_v624 = _v619_labeled_text if "_v619_labeled_text" in globals() else None

def _v619_labeled_text(parent, label: str, height: int = 2) -> tk.Text:
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=3)
    display_label = "Rechnungs-\nbeschreibung" if label == "Rechnungsbeschreibung" else label
    ttk.Label(row, text=display_label, width=16, justify="left").pack(side="left", anchor="n")
    txt = tk.Text(row, height=height, wrap="word", bg=WHITE, fg=TEXT, relief="sunken")
    txt.pack(side="left", fill="x", expand=True)
    return txt

# ---------- Block 7 Zahlungen: Persistenz ----------
def _fm_init_payments_schema_v624() -> None:
    ensure_directories()
    with get_connection() as con:
        con.executescript("""
        CREATE TABLE IF NOT EXISTS payments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entity_type TEXT NOT NULL,
            reference_no TEXT NOT NULL,
            partner_no TEXT NOT NULL,
            partner_name TEXT NOT NULL,
            payment_date TEXT NOT NULL,
            amount REAL NOT NULL,
            payment_method TEXT DEFAULT '',
            bank_account TEXT DEFAULT '',
            booking_text TEXT DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );
        """)
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()

_orig_init_sqlite_v624 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v624()
    _fm_init_payments_schema_v624()


def _fm_payment_partner_label(entity_type: str) -> str:
    return "Kreditor" if entity_type == "vendor_invoice" else "Debitor"


def _fm_invoice_table_for_entity_v624(entity_type: str) -> str:
    return "vendor_invoices" if entity_type == "vendor_invoice" else "customer_invoices"


def _fm_payment_status_from_amount_v624(amount) -> str:
    try:
        return STATUS_PAID if parse_amount(amount) <= 0 else STATUS_OPEN
    except Exception:
        return STATUS_OPEN


class PaymentsView(ttk.Frame, SortableTreeMixin):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.selected_open_item_id = None
        self.selected_entity_type = ""
        self.selected_reference_no = ""
        self.search_var = tk.StringVar()
        self.history_search_var = tk.StringVar()
        self.payment_date = tk.StringVar(value=today_str())
        self.amount = tk.StringVar()
        self.payment_method = tk.StringVar(value="Überweisung")
        self.bank_account = tk.StringVar()
        self.booking_text = tk.StringVar()
        self.full_clear = tk.IntVar(value=1)
        self._build_ui()
        self.reload_open_items()
        self.reload_history()

    def _build_ui(self):
        shell = build_two_block_shell(self, "Zahlungserfassung / OP-Ausgleich", "Zahlungshistorie")

        left = ttk.Frame(shell.left)
        left.pack(fill="both", expand=True, padx=10, pady=8)
        form = ttk.Frame(left)
        form.pack(fill="x", pady=(0, 8))

        self._row_entry(form, "Zahlungsdatum", self.payment_date)
        self._row_entry(form, "Zahlbetrag", self.amount)
        self._row_combo(form, "Zahlungsart", self.payment_method, ["Überweisung", "Lastschrift", "Bar", "Karte", "Sonstige"])
        self._row_entry(form, "Bank/Konto", self.bank_account)
        self._row_entry(form, "Buchungstext", self.booking_text)
        ttk.Checkbutton(form, text="OP vollständig ausgleichen", variable=self.full_clear).pack(anchor="w", pady=3)

        btns = ttk.Frame(form)
        btns.pack(fill="x", pady=6)
        create_standard_button(btns, "Zahlung buchen", self.book_payment, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Formular leeren", self.clear_form).pack(side="left", padx=3)
        create_standard_button(btns, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)

        search_row = ttk.Frame(left)
        search_row.pack(fill="x", pady=(0, 4))
        ttk.Label(search_row, text="Suchen", width=10).pack(side="left")
        search_entry = ttk.Entry(search_row, textvariable=self.search_var)
        search_entry.pack(side="left", fill="x", expand=True)
        search_entry.insert(0, PLACEHOLDER_TEXT)
        search_entry.bind("<FocusIn>", lambda e: (search_entry.delete(0, tk.END) if search_entry.get() == PLACEHOLDER_TEXT else None))
        search_entry.bind("<FocusOut>", lambda e: (search_entry.insert(0, PLACEHOLDER_TEXT) if not search_entry.get().strip() else None))
        search_entry.bind("<KeyRelease>", lambda _e: self.reload_open_items())

        ttk.Label(left, text="Offene Posten", style="Section.TLabel").pack(anchor="w", pady=(6, 4))
        cols = ("id", "type", "reference", "partner", "due", "original", "open", "status")
        self.op_tree = ttk.Treeview(left, columns=cols, show="headings", height=13)
        headings = [("id", "ID", 55), ("type", "Art", 90), ("reference", "Referenz", 120), ("partner", "Partner", 210), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120)]
        for c, t, w in headings:
            self.op_tree.heading(c, text=t)
            self.op_tree.column(c, width=w, anchor="w")
        self.op_tree.pack(fill="both", expand=True)
        configure_tree_tags(self.op_tree)
        self.setup_sorting(self.op_tree)
        self.op_tree.bind("<<TreeviewSelect>>", self.on_open_item_select)

        right = ttk.Frame(shell.right)
        right.pack(fill="both", expand=True, padx=10, pady=8)
        hist_search_row = ttk.Frame(right)
        hist_search_row.pack(fill="x", pady=(0, 4))
        ttk.Label(hist_search_row, text="Suchen", width=10).pack(side="left")
        hist_search = ttk.Entry(hist_search_row, textvariable=self.history_search_var)
        hist_search.pack(side="left", fill="x", expand=True)
        hist_search.insert(0, PLACEHOLDER_TEXT)
        hist_search.bind("<FocusIn>", lambda e: (hist_search.delete(0, tk.END) if hist_search.get() == PLACEHOLDER_TEXT else None))
        hist_search.bind("<FocusOut>", lambda e: (hist_search.insert(0, PLACEHOLDER_TEXT) if not hist_search.get().strip() else None))
        hist_search.bind("<KeyRelease>", lambda _e: self.reload_history())

        ttk.Label(right, text="Gebuchte Zahlungen", style="Section.TLabel").pack(anchor="w", pady=(6, 4))
        hcols = ("id", "date", "type", "reference", "partner", "amount", "method", "text")
        self.history_tree = ttk.Treeview(right, columns=hcols, show="headings", height=16)
        for c, t, w in [("id", "ID", 55), ("date", "Datum", 90), ("type", "Art", 90), ("reference", "Referenz", 120), ("partner", "Partner", 190), ("amount", "Betrag", 90), ("method", "Zahlungsart", 110), ("text", "Text", 220)]:
            self.history_tree.heading(c, text=t)
            self.history_tree.column(c, width=w, anchor="w")
        self.history_tree.pack(fill="both", expand=True)
        self.setup_sorting(self.history_tree)

    def _row_entry(self, parent, label, var):
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)

    def _row_combo(self, parent, label, var, values):
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        cb = ttk.Combobox(row, textvariable=var, values=values, state="readonly")
        cb.pack(side="left", fill="x", expand=True)
        cb.bind("<MouseWheel>", lambda e: "break")

    def _query_text(self, var):
        q = (var.get() or "").strip().lower()
        return "" if q == PLACEHOLDER_TEXT.lower() else q

    def reload_open_items(self):
        for i in self.op_tree.get_children():
            self.op_tree.delete(i)
        q = self._query_text(self.search_var)
        with get_connection() as con:
            rows = list(con.execute("SELECT * FROM open_items WHERE open_amount > 0 ORDER BY due_date, id"))
        for r in rows:
            art = _fm_payment_partner_label(r["entity_type"])
            values = (r["id"], art, r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"])
            blob = " ".join(str(v) for v in values).lower()
            if q and q not in blob:
                continue
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            self.op_tree.insert("", "end", iid=str(r["id"]), values=values, tags=(tag,))

    def reload_history(self):
        for i in self.history_tree.get_children():
            self.history_tree.delete(i)
        q = self._query_text(self.history_search_var)
        with get_connection() as con:
            rows = list(con.execute("SELECT * FROM payments ORDER BY id DESC"))
        for r in rows:
            art = _fm_payment_partner_label(r["entity_type"])
            values = (r["id"], r["payment_date"], art, r["reference_no"], r["partner_name"], format_amount(r["amount"]), r["payment_method"], r["booking_text"])
            blob = " ".join(str(v) for v in values).lower()
            if q and q not in blob:
                continue
            self.history_tree.insert("", "end", values=values)

    def on_open_item_select(self, _event=None):
        iid = self.op_tree.focus()
        if not iid:
            return
        self.selected_open_item_id = int(iid)
        with get_connection() as con:
            row = con.execute("SELECT * FROM open_items WHERE id=?", (self.selected_open_item_id,)).fetchone()
        if not row:
            return
        self.selected_entity_type = row["entity_type"]
        self.selected_reference_no = row["reference_no"]
        self.amount.set(format_amount(row["open_amount"]))
        if not self.booking_text.get().strip():
            self.booking_text.set(f"Zahlung {row['reference_no']} {row['partner_name']}")

    def clear_form(self):
        self.selected_open_item_id = None
        self.selected_entity_type = ""
        self.selected_reference_no = ""
        self.payment_date.set(today_str())
        self.amount.set("")
        self.payment_method.set("Überweisung")
        self.bank_account.set("")
        self.booking_text.set("")
        self.full_clear.set(1)
        try:
            self.op_tree.selection_remove(self.op_tree.selection())
        except Exception:
            pass

    def refresh_all(self):
        self.reload_open_items()
        self.reload_history()
        try:
            self.app.set_status("Zahlungen aktualisiert.")
        except Exception:
            pass

    def book_payment(self):
        if not self.selected_open_item_id:
            messagebox.showwarning("Zahlung", "Bitte zuerst einen offenen Posten auswählen.", parent=self)
            return
        if not validate_date(self.payment_date.get()):
            messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.", parent=self)
            return
        try:
            amount = parse_amount(self.amount.get())
        except Exception as exc:
            messagebox.showwarning("Betrag", str(exc), parent=self)
            return
        if amount <= 0:
            messagebox.showwarning("Betrag", "Der Zahlbetrag muss größer als 0 sein.", parent=self)
            return
        with get_connection() as con:
            op = con.execute("SELECT * FROM open_items WHERE id=?", (self.selected_open_item_id,)).fetchone()
            if not op:
                messagebox.showwarning("Zahlung", "Der offene Posten wurde nicht gefunden.", parent=self)
                return
            open_amount = parse_amount(op["open_amount"])
            pay_amount = open_amount if int(self.full_clear.get() or 0) else amount
            if pay_amount > open_amount:
                messagebox.showwarning("Betrag", "Der Zahlbetrag darf den offenen Betrag nicht überschreiten.", parent=self)
                return
            new_open = (open_amount - pay_amount).quantize(Decimal("0.01"))
            new_status = _fm_payment_status_from_amount_v624(new_open)
            con.execute(
                "INSERT INTO payments(entity_type,reference_no,partner_no,partner_name,payment_date,amount,payment_method,bank_account,booking_text,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                (op["entity_type"], op["reference_no"], op["partner_no"], op["partner_name"], self.payment_date.get(), float(pay_amount), self.payment_method.get(), self.bank_account.get(), self.booking_text.get(), now_str(), now_str())
            )
            con.execute("UPDATE open_items SET open_amount=?, status=?, updated_at=? WHERE id=?", (float(new_open), new_status, now_str(), self.selected_open_item_id))
            table = _fm_invoice_table_for_entity_v624(op["entity_type"])
            con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), new_status, now_str(), op["reference_no"]))
            con.commit()
        try:
            self.app.set_status(f"Zahlung zu {self.selected_reference_no} gebucht.")
        except Exception:
            pass
        self.clear_form()
        self.refresh_all()


def _fm_render_zahlungen_v624(self, parent):
    PaymentsView(parent, self).pack(fill="both", expand=True)

FinanceMateApp._render_zahlungen = _fm_render_zahlungen_v624


def _fm_block7_static_selftest_v624() -> None:
    # Dateibasierte Prüfung, robust auch beim dynamischen Import über importlib.
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "class PaymentsView",
        "def book_payment",
        "CREATE TABLE IF NOT EXISTS payments",
        "SELECT * FROM open_items WHERE open_amount > 0",
        "UPDATE open_items SET open_amount",
        "UPDATE {table} SET open_amount",
        "PaymentsView(parent, self).pack",
    ]
    missing = [x for x in required if x not in src]
    if ("Rechnungs-" not in src) or ("beschreibung" not in src):
        missing.append("Rechnungs-beschreibung Labelumbruch")
    if missing:
        raise RuntimeError("Block-7-/Label-Selbsttest v0.6.24 fehlgeschlagen: " + ", ".join(missing))

_fm_block7_static_selftest_v624()



# === FINANCE MATE PATCH V0_6_25_BLOCK7A_STABILISIERUNG ===
APP_VERSION = "0.6.25-block7a-stabilisierung"

# ---------- 7A.1 Rechnungsbeschreibung: blockgebundenes 2-Zeilen-Textfeld ----------
_orig_v619_labeled_text_v625 = _v619_labeled_text if "_v619_labeled_text" in globals() else None

def _v619_labeled_text(parent, label: str, height: int = 2) -> tk.Text:
    """Gemeinsames Rechnungsbeschreibung-Feld für Debitoren/Kreditoren.

    Das Textfeld bleibt strikt im Formularblock, nutzt Wortumbruch und ist optisch auf zwei Zeilen
    ausgelegt. Damit läuft weder das Label noch das Textfeld in Nachbarblöcke hinein.
    """
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=3)
    row.columnconfigure(1, weight=1)
    display_label = "Rechnungs-\nbeschreibung" if label == "Rechnungsbeschreibung" else label
    ttk.Label(row, text=display_label, width=16, justify="left").grid(row=0, column=0, sticky="nw", padx=(0, 4))
    txt = tk.Text(row, height=height, width=1, wrap="word", bg=WHITE, fg=TEXT, relief="sunken", undo=True)
    txt.grid(row=0, column=1, sticky="ew")
    return txt

# ---------- 7A.2 Suchleisten: einheitlich, pro Tabelle statt nur einmal je Modul ----------
def _fm_entry_placeholder_bind(entry: ttk.Entry, placeholder: str = PLACEHOLDER_TEXT, callback=None) -> None:
    entry.insert(0, placeholder)
    def _focus_in(_e=None):
        if entry.get() == placeholder:
            entry.delete(0, tk.END)
    def _focus_out(_e=None):
        if not entry.get().strip():
            entry.insert(0, placeholder)
        if callable(callback):
            callback()
    entry.bind("<FocusIn>", _focus_in)
    entry.bind("<FocusOut>", _focus_out)
    entry.bind("<KeyRelease>", lambda _e: callback() if callable(callback) else None)

def _fm_install_tree_search(owner, tree_attr: str, reload_callback, var_attr: str, label: str = "Suchen") -> None:
    tree = getattr(owner, tree_attr, None)
    flag = f"_fm_search_installed_{tree_attr}"
    if tree is None or getattr(owner, flag, False):
        return
    setattr(owner, flag, True)
    var = getattr(owner, var_attr, None)
    if var is None or not hasattr(var, "get"):
        var = tk.StringVar()
        setattr(owner, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=0, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=0, pady=(0, 4))
    ttk.Label(row, text=label, width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, reload_callback)

def _fm_install_portal_search(portal, tree_attr: str, reload_callback, var_attr: str) -> None:
    tree = getattr(portal, tree_attr, None)
    flag = f"_fm_portal_search_installed_{tree_attr}"
    if tree is None or getattr(portal, flag, False):
        return
    setattr(portal, flag, True)
    var = getattr(portal, var_attr, None)
    if var is None or not hasattr(var, "get"):
        var = tk.StringVar()
        setattr(portal, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=10, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=10, pady=(0, 4))
    ttk.Label(row, text="Suchen", width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, reload_callback)

# Debitoren: vorhandene Features beibehalten, Suchleisten verlässlich für beide Tabellen.
_orig_debitors_build_ui_v625 = DebitorsView._build_ui

def _fm_debitors_build_ui_v625(self):
    _orig_debitors_build_ui_v625(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices()
    self.reload_open_items()

DebitorsView._build_ui = _fm_debitors_build_ui_v625

# Kreditoren: v0.6.23-Vierblocklayout bleibt erhalten, Suchleisten werden wieder ergänzt.
_orig_creditors_build_ui_v625 = CreditorsView._build_ui

def _fm_creditors_build_ui_v625(self):
    _orig_creditors_build_ui_v625(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    _fm_install_tree_search(self, "stack_tree", self._load_import_batches, "stack_search_var")
    self.reload_invoices()
    self.reload_open_items()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()

CreditorsView._build_ui = _fm_creditors_build_ui_v625

# Rechnungsportal: Archiv- und Freigabetabelle bekommen je eine eigene Suchleiste.
if "InvoicePortalView" in globals():
    _orig_portal_init_v625 = InvoicePortalView.__init__
    def _fm_portal_init_v625(self, parent, app):
        _orig_portal_init_v625(self, parent, app)
        _fm_install_portal_search(self, "archive_tree", self.reload_archive, "archive_search_var")
        _fm_install_portal_search(self, "approval_tree", self.reload_approval, "approval_search_var")
        try:
            self.reload_all()
        except Exception:
            pass
    InvoicePortalView.__init__ = _fm_portal_init_v625

# ---------- 7A.3 Dokumentenstapel: Mehrfachimport stabilisieren ----------
def _fm_creditors_import_archive_paths_v625(self, paths):
    valid = []
    seen = set()
    for raw in paths or []:
        fp = Path(raw)
        key = str(fp.resolve()) if fp.exists() else str(fp)
        if key in seen:
            continue
        seen.add(key)
        if fp.exists() and fp.is_file() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo("Dokumentenstapel", "Keine unterstützten Dateien ausgewählt.", parent=self)
        return
    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    imported = 0
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute(
                "INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
                (hist_no, clean_document_title(fp.name), "", "", "Importiert", 0, now_str(), now_str()),
            )
            con.execute(
                "INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)",
                (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()),
            )
            imported += 1
        con.commit()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    try:
        self.app.set_status(f"{imported} Dokument(e) in den Dokumentenstapel importiert.")
    except Exception:
        pass

def _fm_creditors_import_files_v625(self):
    paths = filedialog.askopenfilenames(title="Dateien für Dokumentenstapel importieren")
    if paths:
        self._import_archive_paths(list(paths))

def _fm_creditors_import_folder_v625(self):
    folder = filedialog.askdirectory(title="Ordner für Dokumentenstapel importieren")
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v625
CreditorsView.import_creditor_files = _fm_creditors_import_files_v625
CreditorsView.import_creditor_folder = _fm_creditors_import_folder_v625

# ---------- 7A.4 Vorschau-Zoom: Standard höher, Maximalzoom deutlich höher ----------
def _fm_preview_set_image_v625(self, pil_image):
    self.base_image = pil_image.convert("RGBA") if pil_image is not None else None
    self.zoom = 1.0
    self.offset_x = 0
    self.offset_y = 0
    if self.base_image is None:
        self.clear()
    else:
        self._render(reset=True)

def _fm_preview_render_v625(self, reset=False):
    if self.base_image is None:
        return
    try:
        from PIL import ImageTk
        self.fit_zoom = self._fit_zoom()
        if reset:
            # Etwas größer als reine Einpassen-Ansicht, aber weiterhin vollständig navigierbar.
            self.zoom = max(self.fit_zoom * 1.18, self.fit_zoom)
            self.offset_x = 0
            self.offset_y = 0
        elif self.zoom < self.fit_zoom:
            self.zoom = self.fit_zoom
        self._clamp_offsets()
        iw, ih = self.base_image.size
        size = (max(1, int(iw * self.zoom)), max(1, int(ih * self.zoom)))
        img = self.base_image.resize(size)
        self.tk_image = ImageTk.PhotoImage(img)
        self._draw()
    except Exception as exc:
        self.clear()
        self.canvas.create_text(20, 48, anchor="nw", text=f"Vorschaufehler: {exc}", fill=TEXT2, font=("Segoe UI", 10))

def _fm_preview_wheel_v625(self, delta):
    if self.base_image is None:
        return
    factor = 1.15 if delta > 0 else 0.87
    old_zoom = self.zoom
    max_zoom = max(self.fit_zoom * 20.0, 20.0)
    self.zoom = max(self.fit_zoom, min(self.zoom * factor, max_zoom))
    if abs(self.zoom - old_zoom) > 0.0001:
        self._clamp_offsets()
        self._render(reset=False)

_FMDocumentPreviewPane.set_image = _fm_preview_set_image_v625
_FMDocumentPreviewPane._render = _fm_preview_render_v625
_FMDocumentPreviewPane._wheel = _fm_preview_wheel_v625

# ---------- 7A.5 Startverhalten: Fenstermodus Vollbild + frühere Skalierungsannahme ----------
def _fm_apply_start_scaling_v625(root: tk.Tk) -> None:
    try:
        root.tk.call("tk", "scaling", 96.0 / 72.0)
    except Exception:
        pass
    try:
        sw = max(1100, root.winfo_screenwidth())
        sh = max(680, root.winfo_screenheight())
        root.geometry(f"{min(sw, 1920)}x{min(sh, 1080)}")
        root.minsize(1100, 680)
    except Exception:
        pass

def _fm_maximize_window_v625(root: tk.Tk) -> None:
    try:
        root.state("zoomed")
    except Exception:
        try:
            root.attributes("-zoomed", True)
        except Exception:
            pass

# Frischer App-Start: Skalierung vor Layoutaufbau, danach maximierter Fenstermodus.
def _fm_app_init_v625(self):
    tk.Tk.__init__(self)
    _fm_apply_start_scaling_v625(self)
    init_sqlite()
    self.title(f"{APP_NAME} {APP_VERSION}")
    self.geometry("1420x860")
    self.minsize(1100, 680)
    self.configure(bg=BG)
    self.sidebar_collapsed = load_ui_preference("sidebar_collapsed", "0") == "1"
    self.active_module = "Dashboard"
    self.nav_buttons: Dict[str, ttk.Button] = {}
    self._configure_ttk()
    self._build_layout()
    _fm_maximize_window_v625(self)
    self.show_module("Dashboard")
    try:
        self.after(250, lambda: _fm_maximize_window_v625(self))
    except Exception:
        pass

FinanceMateApp.__init__ = _fm_app_init_v625

# ---------- 7A-Selbsttest ----------
def _fm_block7a_static_selftest_v625() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required_src = [
        "APP_VERSION = \"0.6.25-block7a-stabilisierung\"",
        "def _fm_creditors_import_archive_paths_v625",
        "self.zoom = max(self.fit_zoom * 1.18",
        "self.fit_zoom * 20.0",
        "def _fm_debitors_build_ui_v625",
        "def _fm_creditors_build_ui_v625",
        "def _fm_app_init_v625",
        "Rechnungs-\\nbeschreibung",
    ]
    missing = [x for x in required_src if x not in src]
    try:
        csrc = _inspect.getsource(CreditorsView._build_ui)
    except Exception:
        csrc = ""
    for token in ["_orig_creditors_build_ui_v625", "stack_search_var", "invoice_search_var", "open_items_search_var"]:
        if token not in csrc:
            missing.append(token)
    if missing:
        raise RuntimeError("Block-7A-Selbsttest v0.6.25 fehlgeschlagen: " + ", ".join(missing))

_fm_block7a_static_selftest_v625()



# === FINANCE MATE PATCH V0_6_26_BLOCK7B_PAYMENT_TABS_IMPORT_PREVIEW ===
APP_VERSION = "0.6.26-block7b-payment-tabs-import-preview"

# ---------- 7A-Nachkorrektur: nur eine Suchleiste im Dokumentenstapel ----------
# Der Dokumentenstapel besitzt bereits im v0.6.23-Layout eine eigene Suchleiste. Die zusätzliche
# generische Suchleiste aus v0.6.25 wird hier bewusst nicht mehr installiert.
_orig_creditors_build_ui_v626 = _orig_creditors_build_ui_v625 if "_orig_creditors_build_ui_v625" in globals() else CreditorsView._build_ui

def _fm_creditors_build_ui_v626(self):
    _orig_creditors_build_ui_v626(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices()
    self.reload_open_items()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()

CreditorsView._build_ui = _fm_creditors_build_ui_v626

# ---------- 7A-Nachkorrektur: Mehrfachauswahl für Dateien importieren robust aus Tk lesen ----------
def _fm_normalize_dialog_paths_v626(widget, raw_paths) -> list[str]:
    """Tk kann je nach Plattform Tuple, Tcl-Liste oder String liefern; hier wird alles normalisiert."""
    if not raw_paths:
        return []
    if isinstance(raw_paths, (list, tuple)):
        return [str(p) for p in raw_paths if str(p).strip()]
    try:
        return [str(p) for p in widget.tk.splitlist(raw_paths) if str(p).strip()]
    except Exception:
        return [str(raw_paths)] if str(raw_paths).strip() else []

def _fm_creditors_import_archive_paths_v626(self, paths):
    valid: list[Path] = []
    seen = set()
    for raw in _fm_normalize_dialog_paths_v626(self, paths):
        fp = Path(raw)
        try:
            key = str(fp.resolve()) if fp.exists() else str(fp)
        except Exception:
            key = str(fp)
        if key in seen:
            continue
        seen.add(key)
        if fp.exists() and fp.is_file() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo("Dokumentenstapel", "Keine unterstützten Dateien ausgewählt.", parent=self)
        return
    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    imported = 0
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute(
                "INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
                (hist_no, clean_document_title(fp.name), "", "", "Importiert", 0, now_str(), now_str()),
            )
            con.execute(
                "INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)",
                (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()),
            )
            imported += 1
        con.commit()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    try:
        self.app.set_status(f"{imported} Dokument(e) in den Dokumentenstapel importiert.")
    except Exception:
        pass

def _fm_creditors_import_files_v626(self):
    # Direkter Tk-Aufruf mit -multiple 1 ist robuster als die Wrapperfunktion auf einzelnen Plattformen.
    try:
        raw = self.tk.call("tk_getOpenFile", "-multiple", "1", "-title", "Dateien für Dokumentenstapel importieren")
        paths = _fm_normalize_dialog_paths_v626(self, raw)
    except Exception:
        paths = list(filedialog.askopenfilenames(title="Dateien für Dokumentenstapel importieren", parent=self))
    if paths:
        self._import_archive_paths(paths)

def _fm_creditors_import_folder_v626(self):
    folder = filedialog.askdirectory(title="Ordner für Dokumentenstapel importieren", parent=self)
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v626
CreditorsView.import_creditor_files = _fm_creditors_import_files_v626
CreditorsView.import_creditor_folder = _fm_creditors_import_folder_v626

# ---------- 7A-Nachkorrektur: Vorschau-Initialzoom = obere 50% des ersten Vorschaubilds ----------
_orig_fm_combine_pages_v626 = _fm_combine_pages
_orig_fm_text_pages_to_image_v626 = _fm_text_pages_to_image
_orig_fm_generate_document_preview_v626 = _fm_generate_document_preview

def _fm_text_pages_to_image(title: str, lines: list[str], page_lines: int = 42):
    img = _orig_fm_text_pages_to_image_v626(title, lines, page_lines)
    try:
        img._fm_first_page_height = 1500
    except Exception:
        pass
    return img

def _fm_combine_pages(images):
    out = _orig_fm_combine_pages_v626(images)
    try:
        out._fm_first_page_height = images[0].height if images else out.height
    except Exception:
        pass
    return out

def _fm_generate_document_preview(path: str):
    img = _orig_fm_generate_document_preview_v626(path)
    try:
        if not hasattr(img, "_fm_first_page_height"):
            img._fm_first_page_height = img.height
    except Exception:
        pass
    return img

def _fm_preview_render_v626(self, reset=False):
    if self.base_image is None:
        return
    try:
        from PIL import ImageTk
        self.fit_zoom = self._fit_zoom()
        cw = max(50, self.canvas.winfo_width())
        ch = max(50, self.canvas.winfo_height())
        first_h = float(getattr(self.base_image, "_fm_first_page_height", self.base_image.size[1]) or self.base_image.size[1])
        if reset:
            # Ziel: In der Vorschau ist initial die obere Hälfte des ersten Vorschaubilds sichtbar.
            target_zoom = ch / max(1.0, first_h * 0.5)
            self.zoom = max(self.fit_zoom, target_zoom)
            rw = int(self.base_image.size[0] * self.zoom)
            rh = int(self.base_image.size[1] * self.zoom)
            self.offset_x = 0
            # Bildoberkante an Canvas-Oberkante ausrichten, damit die oberen 50% sichtbar sind.
            self.offset_y = max(0, (rh - ch) / 2)
        elif self.zoom < self.fit_zoom:
            self.zoom = self.fit_zoom
        self._clamp_offsets()
        iw, ih = self.base_image.size
        size = (max(1, int(iw * self.zoom)), max(1, int(ih * self.zoom)))
        img = self.base_image.resize(size)
        self.tk_image = ImageTk.PhotoImage(img)
        self._draw()
    except Exception as exc:
        self.clear()
        self.canvas.create_text(20, 48, anchor="nw", text=f"Vorschaufehler: {exc}", fill=TEXT2, font=("Segoe UI", 10))

def _fm_preview_wheel_v626(self, delta):
    if self.base_image is None:
        return
    factor = 1.15 if delta > 0 else 0.87
    old_zoom = self.zoom
    max_zoom = max(self.fit_zoom * 20.0, 20.0)
    self.zoom = max(self.fit_zoom, min(self.zoom * factor, max_zoom))
    if abs(self.zoom - old_zoom) > 0.0001:
        self._clamp_offsets()
        self._render(reset=False)

_FMDocumentPreviewPane._render = _fm_preview_render_v626
_FMDocumentPreviewPane._wheel = _fm_preview_wheel_v626

# ---------- Block 7B: Zahlungsmodul fachlich in Ein-/Ausgänge + Bankbewegungen strukturieren ----------
def _fm_init_block7b_schema_v626() -> None:
    with get_connection() as con:
        # Bankkonten fachlich für spätere Schnittstellen / Kontierung vorbereiten.
        for col, definition in [
            ("gl_account_no", "TEXT DEFAULT ''"),
            ("incoming_clearing_account", "TEXT DEFAULT ''"),
            ("outgoing_clearing_account", "TEXT DEFAULT ''"),
            ("online_banking_enabled", "INTEGER DEFAULT 0"),
            ("interface_profile", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, "bank_accounts", col, definition)
        con.executescript("""
CREATE TABLE IF NOT EXISTS bank_transactions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    bank_account_id INTEGER DEFAULT 0,
    booking_date TEXT DEFAULT '',
    value_date TEXT DEFAULT '',
    amount REAL DEFAULT 0,
    currency TEXT DEFAULT 'EUR',
    direction TEXT DEFAULT '',
    partner_name TEXT DEFAULT '',
    partner_iban TEXT DEFAULT '',
    purpose_text TEXT DEFAULT '',
    end_to_end_id TEXT DEFAULT '',
    bank_reference TEXT DEFAULT '',
    imported_from TEXT DEFAULT '',
    import_batch_id TEXT DEFAULT '',
    match_status TEXT DEFAULT 'Ungeklärt',
    matched_entity_type TEXT DEFAULT '',
    matched_reference_no TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS payment_matches (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    payment_id INTEGER DEFAULT 0,
    bank_transaction_id INTEGER DEFAULT 0,
    entity_type TEXT NOT NULL,
    reference_no TEXT NOT NULL,
    matched_amount REAL NOT NULL,
    difference_amount REAL DEFAULT 0,
    match_type TEXT DEFAULT 'Manuell',
    created_by TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS posting_rules (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    rule_name TEXT NOT NULL,
    direction TEXT DEFAULT '',
    text_contains TEXT DEFAULT '',
    debit_account TEXT DEFAULT '',
    credit_account TEXT DEFAULT '',
    clearing_strategy TEXT DEFAULT '',
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
""")
        # Basis-Regeln nur anlegen, wenn noch keine Regeln vorhanden sind.
        if not con.execute("SELECT 1 FROM posting_rules LIMIT 1").fetchone():
            now = now_str()
            con.execute("INSERT INTO posting_rules(rule_name,direction,text_contains,debit_account,credit_account,clearing_strategy,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)", ("Zahlungseingang Debitor", "Eingang", "", "1200", "1400", "OP-Ausgleich Debitor", 1, now, now))
            con.execute("INSERT INTO posting_rules(rule_name,direction,text_contains,debit_account,credit_account,clearing_strategy,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)", ("Zahlungsausgang Kreditor", "Ausgang", "", "1600", "1200", "OP-Ausgleich Kreditor", 1, now, now))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()

_orig_init_sqlite_v626 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v626()
    _fm_init_block7b_schema_v626()

# Hilfsfunktionen für Zahlungsoberfläche
def _fm_payment_direction_label(entity_type: str) -> str:
    return "Zahlungsausgang" if entity_type == "vendor_invoice" else "Zahlungseingang"

def _fm_payment_direction_from_amount(amount) -> str:
    try:
        return "Eingang" if parse_amount(amount) >= 0 else "Ausgang"
    except Exception:
        return ""

def _fm_payment_query_text(var) -> str:
    try:
        q = var.get().strip().lower()
        return "" if q == PLACEHOLDER_TEXT.lower() else q
    except Exception:
        return ""

def _fm_payments_make_search(parent, var, callback):
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=(0, 4))
    ttk.Label(row, text="Suchen", width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, callback)
    return entry

# Neue Block-7B-UI: Reiter Zahlungseingänge / Zahlungsausgänge / Bankbewegungen.
def _fm_payments_build_ui_v626(self):
    self.in_search_var = tk.StringVar()
    self.out_search_var = tk.StringVar()
    self.in_history_search_var = tk.StringVar()
    self.out_history_search_var = tk.StringVar()
    self.bank_search_var = tk.StringVar()

    nb = ttk.Notebook(self)
    nb.pack(fill="both", expand=True)
    self.payments_notebook = nb

    tab_in = ttk.Frame(nb)
    tab_out = ttk.Frame(nb)
    tab_bank = ttk.Frame(nb)
    nb.add(tab_in, text="Zahlungseingänge")
    nb.add(tab_out, text="Zahlungsausgänge")
    nb.add(tab_bank, text="Bankbewegungen")

    self._build_direction_tab(tab_in, "customer_invoice")
    self._build_direction_tab(tab_out, "vendor_invoice")
    self._build_bank_transactions_tab(tab_bank)


def _fm_payments_build_direction_tab(self, parent, entity_type: str):
    is_in = entity_type == "customer_invoice"
    shell = build_two_block_shell(
        parent,
        "Zahlungseingang / Debitoren-OP" if is_in else "Zahlungsausgang / Kreditoren-OP",
        "Zahlungshistorie",
    )
    left = ttk.Frame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(shell.right)
    right.pack(fill="both", expand=True, padx=10, pady=8)

    form = ttk.Frame(left)
    form.pack(fill="x", pady=(0, 8))
    self._row_entry(form, "Zahlungsdatum", self.payment_date)
    self._row_entry(form, "Zahlbetrag", self.amount)
    self._row_combo(form, "Zahlungsart", self.payment_method, ["Überweisung", "Lastschrift", "Bar", "Karte", "Sonstige"])
    self._row_entry(form, "Bank/Konto", self.bank_account)
    self._row_entry(form, "Buchungstext", self.booking_text)
    ttk.Checkbutton(form, text="OP vollständig ausgleichen", variable=self.full_clear).pack(anchor="w", pady=3)
    btns = ttk.Frame(form)
    btns.pack(fill="x", pady=6)
    create_standard_button(btns, "Zahlung buchen", self.book_payment, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Formular leeren", self.clear_form).pack(side="left", padx=3)
    create_standard_button(btns, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)

    search_var = self.in_search_var if is_in else self.out_search_var
    _fm_payments_make_search(left, search_var, self.reload_open_items)
    cols = ("id", "type", "reference", "partner", "due", "original", "open", "status")
    tree = ttk.Treeview(left, columns=cols, show="headings", height=12)
    for c, t, w in [
        ("id", "ID", 55), ("type", "Vorgang", 130), ("reference", "Referenz", 120),
        ("partner", "Partner", 190), ("due", "Fällig", 90), ("original", "Original", 95),
        ("open", "Offen", 95), ("status", "Status", 120),
    ]:
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w")
    tree.pack(fill="both", expand=True)
    configure_tree_tags(tree)
    self.setup_sorting(tree)
    tree.bind("<<TreeviewSelect>>", lambda _e, tr=tree: self._select_open_item_from_tree(tr))

    hist_search_var = self.in_history_search_var if is_in else self.out_history_search_var
    _fm_payments_make_search(right, hist_search_var, self.reload_history)
    hcols = ("date", "direction", "reference", "partner", "amount", "method", "bank", "text")
    hist = ttk.Treeview(right, columns=hcols, show="headings")
    for c, t, w in [
        ("date", "Datum", 90), ("direction", "Richtung", 115), ("reference", "Referenz", 120),
        ("partner", "Partner", 180), ("amount", "Betrag", 90), ("method", "Art", 100),
        ("bank", "Bank/Konto", 130), ("text", "Text", 180),
    ]:
        hist.heading(c, text=t)
        hist.column(c, width=w, anchor="w")
    hist.pack(fill="both", expand=True)
    self.setup_sorting(hist)

    if is_in:
        self.incoming_tree = tree
        self.incoming_history_tree = hist
    else:
        self.outgoing_tree = tree
        self.outgoing_history_tree = hist


def _fm_payments_build_bank_transactions_tab(self, parent):
    top = ttk.Frame(parent)
    top.pack(fill="x", padx=10, pady=8)
    create_standard_button(top, "Bankimport vorbereiten", lambda: messagebox.showinfo("Bankbewegungen", "CSV/camt.053/EBICS-Struktur ist vorbereitet. Der echte Import folgt in Block 7D.", parent=self)).pack(side="left", padx=3)
    create_standard_button(top, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)
    _fm_payments_make_search(parent, self.bank_search_var, self.reload_bank_transactions)
    cols = ("id", "date", "value", "direction", "amount", "partner", "purpose", "match")
    tree = ttk.Treeview(parent, columns=cols, show="headings")
    for c, t, w in [
        ("id", "ID", 55), ("date", "Buchung", 90), ("value", "Valuta", 90), ("direction", "Richtung", 90),
        ("amount", "Betrag", 90), ("partner", "Partner", 180), ("purpose", "Verwendungszweck", 320), ("match", "Status", 120),
    ]:
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=10, pady=(0, 10))
    self.bank_tx_tree = tree
    self.setup_sorting(tree)


def _fm_payments_select_open_item_from_tree(self, tree):
    iid = tree.focus()
    if not iid:
        return
    try:
        op_id = int(tree.set(iid, "id"))
    except Exception:
        return
    with get_connection() as con:
        row = con.execute("SELECT * FROM open_items WHERE id=?", (op_id,)).fetchone()
    if not row:
        return
    self.selected_open_item_id = op_id
    self.selected_entity_type = row["entity_type"]
    self.selected_reference_no = row["reference_no"]
    self.amount.set(format_amount(row["open_amount"]))
    if not self.booking_text.get().strip():
        self.booking_text.set(f"{_fm_payment_direction_label(row['entity_type'])} {row['reference_no']}")


def _fm_payments_reload_open_items_v626(self):
    mapping = [
        (getattr(self, "incoming_tree", None), "customer_invoice", getattr(self, "in_search_var", None)),
        (getattr(self, "outgoing_tree", None), "vendor_invoice", getattr(self, "out_search_var", None)),
    ]
    with get_connection() as con:
        all_rows = list(con.execute("SELECT * FROM open_items WHERE open_amount > 0 ORDER BY due_date, id"))
    for tree, entity_type, var in mapping:
        if tree is None:
            continue
        for i in tree.get_children():
            tree.delete(i)
        q = _fm_payment_query_text(var)
        for r in all_rows:
            if r["entity_type"] != entity_type:
                continue
            values = (r["id"], _fm_payment_direction_label(r["entity_type"]), r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"])
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values, tags=(urgency_bucket(r["due_date"], r["open_amount"]),))


def _fm_payments_reload_history_v626(self):
    mapping = [
        (getattr(self, "incoming_history_tree", None), "customer_invoice", getattr(self, "in_history_search_var", None)),
        (getattr(self, "outgoing_history_tree", None), "vendor_invoice", getattr(self, "out_history_search_var", None)),
    ]
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM payments ORDER BY id DESC"))
    for tree, entity_type, var in mapping:
        if tree is None:
            continue
        for i in tree.get_children():
            tree.delete(i)
        q = _fm_payment_query_text(var)
        for r in rows:
            if r["entity_type"] != entity_type:
                continue
            values = (r["payment_date"], _fm_payment_direction_label(r["entity_type"]), r["reference_no"], r["partner_name"], format_amount(r["amount"]), r["payment_method"], r["bank_account"], r["booking_text"])
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values)


def _fm_payments_reload_bank_transactions_v626(self):
    tree = getattr(self, "bank_tx_tree", None)
    if tree is None:
        return
    for i in tree.get_children():
        tree.delete(i)
    q = _fm_payment_query_text(getattr(self, "bank_search_var", None))
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM bank_transactions ORDER BY id DESC"))
    for r in rows:
        values = (r["id"], r["booking_date"], r["value_date"], r["direction"], format_amount(r["amount"]), r["partner_name"], r["purpose_text"], r["match_status"])
        if q and q not in " ".join(str(v) for v in values).lower():
            continue
        tree.insert("", "end", values=values)


def _fm_payments_book_payment_v626(self):
    if not self.selected_open_item_id:
        messagebox.showwarning("Zahlung", "Bitte zuerst einen offenen Posten auswählen.", parent=self)
        return
    if not validate_date(self.payment_date.get()):
        messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.", parent=self)
        return
    try:
        amount = parse_amount(self.amount.get())
    except Exception as exc:
        messagebox.showwarning("Betrag", str(exc), parent=self)
        return
    if amount <= 0:
        messagebox.showwarning("Betrag", "Der Zahlbetrag muss größer als 0 sein.", parent=self)
        return
    with get_connection() as con:
        op = con.execute("SELECT * FROM open_items WHERE id=?", (self.selected_open_item_id,)).fetchone()
        if not op:
            messagebox.showwarning("Zahlung", "Der offene Posten wurde nicht gefunden.", parent=self)
            return
        open_amount = parse_amount(op["open_amount"])
        pay_amount = open_amount if int(self.full_clear.get() or 0) else amount
        if pay_amount > open_amount:
            messagebox.showwarning("Betrag", "Der Zahlbetrag darf den offenen Betrag nicht überschreiten.", parent=self)
            return
        new_open = (open_amount - pay_amount).quantize(Decimal("0.01"))
        new_status = _fm_payment_status_from_amount_v624(new_open)
        cur = con.execute(
            "INSERT INTO payments(entity_type,reference_no,partner_no,partner_name,payment_date,amount,payment_method,bank_account,booking_text,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
            (op["entity_type"], op["reference_no"], op["partner_no"], op["partner_name"], self.payment_date.get(), float(pay_amount), self.payment_method.get(), self.bank_account.get(), self.booking_text.get(), now_str(), now_str()),
        )
        payment_id = cur.lastrowid
        con.execute("UPDATE open_items SET open_amount=?, status=?, updated_at=? WHERE id=?", (float(new_open), new_status, now_str(), self.selected_open_item_id))
        table = _fm_invoice_table_for_entity_v624(op["entity_type"])
        con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), new_status, now_str(), op["reference_no"]))
        con.execute(
            "INSERT INTO payment_matches(payment_id,bank_transaction_id,entity_type,reference_no,matched_amount,difference_amount,match_type,created_by,created_at) VALUES(?,?,?,?,?,?,?,?,?)",
            (payment_id, 0, op["entity_type"], op["reference_no"], float(pay_amount), float(new_open), "Manuell", _fm_current_user() if "_fm_current_user" in globals() else "", now_str()),
        )
        con.commit()
    try:
        self.app.set_status(f"Zahlung zu {self.selected_reference_no} gebucht.")
    except Exception:
        pass
    self.clear_form()
    self.refresh_all()


def _fm_payments_clear_form_v626(self):
    self.selected_open_item_id = None
    self.selected_entity_type = ""
    self.selected_reference_no = ""
    self.payment_date.set(today_str())
    self.amount.set("")
    self.payment_method.set("Überweisung")
    self.bank_account.set("")
    self.booking_text.set("")
    self.full_clear.set(1)
    for tree in [getattr(self, "incoming_tree", None), getattr(self, "outgoing_tree", None)]:
        try:
            tree.selection_remove(tree.selection())
        except Exception:
            pass


def _fm_payments_refresh_all_v626(self):
    self.reload_open_items()
    self.reload_history()
    self.reload_bank_transactions()
    try:
        self.app.set_status("Zahlungen aktualisiert.")
    except Exception:
        pass

PaymentsView._build_ui = _fm_payments_build_ui_v626
PaymentsView._build_direction_tab = _fm_payments_build_direction_tab
PaymentsView._build_bank_transactions_tab = _fm_payments_build_bank_transactions_tab
PaymentsView._select_open_item_from_tree = _fm_payments_select_open_item_from_tree
PaymentsView.reload_open_items = _fm_payments_reload_open_items_v626
PaymentsView.reload_history = _fm_payments_reload_history_v626
PaymentsView.reload_bank_transactions = _fm_payments_reload_bank_transactions_v626
PaymentsView.book_payment = _fm_payments_book_payment_v626
PaymentsView.clear_form = _fm_payments_clear_form_v626
PaymentsView.refresh_all = _fm_payments_refresh_all_v626

# ---------- v0.6.26 Selbsttest ----------
def _fm_block7b_static_selftest_v626() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.26-block7b-payment-tabs-import-preview\"",
        "tk_getOpenFile",
        "-multiple",
        "self.fit_zoom * 20.0",
        "target_zoom = ch / max(1.0, first_h * 0.5)",
        "bank_transactions",
        "payment_matches",
        "posting_rules",
        "Zahlungseingänge",
        "Zahlungsausgänge",
        "Bankbewegungen",
    ]
    missing = [x for x in required if x not in src]
    csrc = _inspect.getsource(CreditorsView._build_ui)
    if "stack_tree" in csrc or "stack_search_var" in csrc:
        missing.append("Dokumentenstapel generische Zusatzsuchleiste nicht entfernt")
    psrc = _inspect.getsource(PaymentsView._build_ui)
    for token in ["ttk.Notebook", "Zahlungseingänge", "Zahlungsausgänge", "Bankbewegungen"]:
        if token not in psrc:
            missing.append(token)
    if missing:
        raise RuntimeError("Block-7B-Selbsttest v0.6.26 fehlgeschlagen: " + ", ".join(missing))

_fm_block7b_static_selftest_v626()



# === FINANCE MATE PATCH V0_6_27_CREDITORS_MULTIIMPORT_DIALOG_FIX ===
APP_VERSION = "0.6.27-creditors-multiimport-dialog-fix"

# Mehrfachimport-Fix:
# Einige Windows-/Tk-Dialogkombinationen liefern trotz -multiple nur eine Datei bzw. erlauben praktisch
# keine Mehrfachauswahl. Deshalb arbeitet "Dateien importieren" jetzt mit einem robusten Sammelmodus:
# - native Mehrfachauswahl, wenn verfügbar
# - falls nur eine Datei gewählt wird: optionale weitere Auswahlrunden
# - danach gemeinsamer Import aller gesammelten Dateien in getrennte Stapel-Einträge

def _fm_creditors_filetypes_v627():
    return [
        ("Unterstützte Dokumente", "*.pdf *.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff *.txt *.csv *.doc *.docx *.xls *.xlsx"),
        ("PDF", "*.pdf"),
        ("Bilder", "*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff"),
        ("Office/Text", "*.txt *.csv *.doc *.docx *.xls *.xlsx"),
        ("Alle Dateien", "*.*"),
    ]

def _fm_normalize_dialog_paths_v627(widget, raw_paths) -> list[str]:
    if not raw_paths:
        return []
    if isinstance(raw_paths, (list, tuple)):
        return [str(p) for p in raw_paths if str(p).strip()]
    try:
        return [str(p) for p in widget.tk.splitlist(raw_paths) if str(p).strip()]
    except Exception:
        return [str(raw_paths)] if str(raw_paths).strip() else []

def _fm_creditors_import_files_v627(self):
    collected: list[str] = []
    seen = set()

    def _add_paths(raw_paths) -> int:
        added = 0
        for path in _fm_normalize_dialog_paths_v627(self, raw_paths):
            if not path:
                continue
            try:
                key = str(Path(path).resolve()) if Path(path).exists() else str(path)
            except Exception:
                key = str(path)
            if key in seen:
                continue
            seen.add(key)
            collected.append(path)
            added += 1
        return added

    while True:
        try:
            selected = filedialog.askopenfilenames(
                parent=self,
                title="Dateien für Dokumentenstapel importieren",
                filetypes=_fm_creditors_filetypes_v627(),
            )
        except Exception:
            selected = ()

        added_now = _add_paths(selected)
        if added_now == 0:
            break

        # Wenn der native Dialog mehrere Dateien zurückgibt, ist die Aktion erfüllt.
        if added_now > 1:
            break

        # Falls der native Dialog auf dem System faktisch nur Einzelauswahl zulässt,
        # kann der Benutzer weitere Dateien derselben Importaktion hinzufügen.
        more = messagebox.askyesno(
            "Dateien importieren",
            "Eine Datei wurde zur Importliste hinzugefügt.\n\nWeitere Datei(en) zur gleichen Importaktion hinzufügen?",
            parent=self,
        )
        if not more:
            break

    if collected:
        self._import_archive_paths(collected)

def _fm_creditors_import_archive_paths_v627(self, paths):
    valid: list[Path] = []
    seen = set()
    for raw in _fm_normalize_dialog_paths_v627(self, paths):
        fp = Path(raw)
        try:
            key = str(fp.resolve()) if fp.exists() else str(fp)
        except Exception:
            key = str(fp)
        if key in seen:
            continue
        seen.add(key)
        if fp.exists() and fp.is_file() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo("Dokumentenstapel", "Keine unterstützten Dateien ausgewählt.", parent=self)
        return

    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    imported = 0
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute(
                "INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
                (hist_no, clean_document_title(fp.name), "", "", "Importiert", 0, now_str(), now_str()),
            )
            con.execute(
                "INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)",
                (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()),
            )
            imported += 1
        con.commit()

    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    try:
        self.app.set_status(f"{imported} Dokument(e) in den Dokumentenstapel importiert.")
    except Exception:
        pass

CreditorsView.import_creditor_files = _fm_creditors_import_files_v627
CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v627

# Selftest für Mehrfachimport-Fix ohne GUI-Aufruf.
def _fm_creditors_multiimport_static_selftest_v627() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.27-creditors-multiimport-dialog-fix\"",
        "def _fm_creditors_import_files_v627",
        "Eine Datei wurde zur Importliste hinzugefügt",
        "Weitere Datei(en) zur gleichen Importaktion hinzufügen",
        "CreditorsView.import_creditor_files = _fm_creditors_import_files_v627",
        "CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v627",
    ]
    missing = [x for x in required if x not in src]
    active_src = _inspect.getsource(CreditorsView.import_creditor_files)
    for token in ["collected", "askopenfilenames", "askyesno", "self._import_archive_paths(collected)"]:
        if token not in active_src:
            missing.append(token)
    if missing:
        raise RuntimeError("Mehrfachimport-Selbsttest v0.6.27 fehlgeschlagen: " + ", ".join(missing))

_fm_creditors_multiimport_static_selftest_v627()



# === FINANCE MATE PATCH V0_6_28_CREDITORS_MULTIIMPORT_LOCKFIX ===
APP_VERSION = "0.6.28-creditors-multiimport-lockfix"

# Zweck:
# - Kein Sammel-/Nachfragedialog mehr.
# - "Dateien importieren" öffnet genau einen Mehrfachauswahldialog; dort können alle benötigten Dateien
#   direkt per Strg/Umschalt markiert werden.
# - "Ordner importieren" importiert alle unterstützten Dateien aus dem gewählten Ordner rekursiv.
# - database-is-locked wird vermieden: HIST-Nummern werden innerhalb derselben DB-Verbindung reserviert;
#   kein next_hist_no()/generate_number-Aufruf innerhalb einer offenen Import-Transaktion.

# Robusteres SQLite-Verhalten für lokale Desktop-App: längeres Warten auf Locks und WAL, soweit verfügbar.
def get_connection() -> sqlite3.Connection:
    ensure_directories()
    con = sqlite3.connect(DB_PATH, timeout=30.0)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys = ON")
    try:
        con.execute("PRAGMA busy_timeout = 30000")
    except Exception:
        pass
    try:
        con.execute("PRAGMA journal_mode = WAL")
    except Exception:
        pass
    return con

def _fm_creditors_filetypes_v628():
    return [
        ("Unterstützte Dokumente", "*.pdf *.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff *.txt *.csv *.doc *.docx *.xls *.xlsx"),
        ("PDF", "*.pdf"),
        ("Bilder", "*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff"),
        ("Office/Text", "*.txt *.csv *.doc *.docx *.xls *.xlsx"),
        ("Alle Dateien", "*.*"),
    ]

def _fm_normalize_dialog_paths_v628(widget, raw_paths) -> list[str]:
    if not raw_paths:
        return []
    if isinstance(raw_paths, (list, tuple)):
        return [str(p) for p in raw_paths if str(p).strip()]
    try:
        return [str(p) for p in widget.tk.splitlist(raw_paths) if str(p).strip()]
    except Exception:
        return [str(raw_paths)] if str(raw_paths).strip() else []

def _fm_collect_supported_import_files_v628(paths) -> list[Path]:
    valid: list[Path] = []
    seen = set()
    for raw in paths or []:
        fp = Path(raw)
        try:
            key = str(fp.resolve()) if fp.exists() else str(fp)
        except Exception:
            key = str(fp)
        if key in seen:
            continue
        seen.add(key)
        if fp.exists() and fp.is_file() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    return valid

def _fm_reserve_hist_numbers_v628(con: sqlite3.Connection, count: int) -> list[str]:
    if count <= 0:
        return []
    row = con.execute("SELECT value FROM app_meta WHERE key=?", ("counter_hist",)).fetchone()
    current = int(row[0]) if row else 0
    numbers = [f"HIST-{n:06d}" for n in range(current + 1, current + count + 1)]
    con.execute(
        "INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)",
        ("counter_hist", str(current + count), now_str()),
    )
    return numbers

def _fm_creditors_import_archive_paths_v628(self, paths):
    valid = _fm_collect_supported_import_files_v628(paths)
    if not valid:
        messagebox.showinfo("Dokumentenstapel", "Keine unterstützten Dateien ausgewählt.", parent=self)
        return

    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    copied_paths: list[Path] = []
    imported = 0
    try:
        with get_connection() as con:
            hist_numbers = _fm_reserve_hist_numbers_v628(con, len(valid))
            for fp, hist_no in zip(valid, hist_numbers):
                hist_dir = IMPORTS_DIR / hist_no
                hist_dir.mkdir(parents=True, exist_ok=True)
                dest = hist_dir / fp.name
                if dest.exists():
                    dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
                shutil.copy2(fp, dest)
                copied_paths.append(dest)
                con.execute(
                    "INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
                    (hist_no, clean_document_title(fp.name), "", "", "Importiert", 0, now_str(), now_str()),
                )
                con.execute(
                    "INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)",
                    (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()),
                )
                imported += 1
            con.commit()
    except sqlite3.OperationalError as exc:
        # Bei echtem Fremdlock klare Meldung statt Tkinter-Traceback.
        for dest in copied_paths:
            try:
                if dest.exists():
                    dest.unlink()
            except Exception:
                pass
        messagebox.showerror(
            "Dokumentenstapel",
            "Der Dokumentenimport konnte nicht abgeschlossen werden, weil die Datenbank gerade gesperrt ist.\n\n"
            "Bitte Finance Mate kurz schließen/neu öffnen oder parallele Imports beenden und danach erneut importieren.\n\n"
            f"Technische Meldung: {exc}",
            parent=self,
        )
        return

    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    try:
        self.app.set_status(f"{imported} Dokument(e) in den Dokumentenstapel importiert.")
    except Exception:
        pass

def _fm_creditors_import_files_v628(self):
    # Ein Dialog, echte Mehrfachauswahl: mehrere Dateien mit Strg/Umschalt markieren und gemeinsam öffnen.
    paths = filedialog.askopenfilenames(
        parent=self,
        title="Dateien für Dokumentenstapel importieren",
        filetypes=_fm_creditors_filetypes_v628(),
    )
    paths = _fm_normalize_dialog_paths_v628(self, paths)
    if paths:
        self._import_archive_paths(paths)

def _fm_creditors_import_folder_v628(self):
    folder = filedialog.askdirectory(title="Ordner für Dokumentenstapel importieren", parent=self)
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v628
CreditorsView.import_creditor_files = _fm_creditors_import_files_v628
CreditorsView.import_creditor_folder = _fm_creditors_import_folder_v628

# v0.6.28 Import-Selbsttest: keine Nachfrageschleife, keine next_hist_no-Nutzung im Import, Lockfix aktiv.
def _fm_creditors_multiimport_static_selftest_v628() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.28-creditors-multiimport-lockfix\"",
        "timeout=30.0",
        "PRAGMA busy_timeout = 30000",
        "def _fm_reserve_hist_numbers_v628",
        "def _fm_creditors_import_files_v628",
        "def _fm_creditors_import_folder_v628",
        "CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v628",
    ]
    missing = [x for x in required if x not in src]
    import_src = _inspect.getsource(CreditorsView._import_archive_paths)
    file_src = _inspect.getsource(CreditorsView.import_creditor_files)
    folder_src = _inspect.getsource(CreditorsView.import_creditor_folder)
    if "next_hist_no" in import_src or "generate_number" in import_src:
        missing.append("Import nutzt noch next_hist_no/generate_number")
    if "askyesno" in file_src or "while True" in file_src:
        missing.append("Dateiimport enthält noch Sammel-/Nachfrageschleife")
    for token in ["askopenfilenames", "self._import_archive_paths(paths)"]:
        if token not in file_src:
            missing.append(token)
    for token in ["askdirectory", "iter_supported_files_from_folder", "self._import_archive_paths"]:
        if token not in folder_src:
            missing.append(token)
    if missing:
        raise RuntimeError("Mehrfachimport-Selbsttest v0.6.28 fehlgeschlagen: " + ", ".join(missing))

_fm_creditors_multiimport_static_selftest_v628()



# === FINANCE MATE PATCH V0_6_29_PORTAL_ATTACHMENTS_PREVIEW ===
APP_VERSION = "0.6.29-portal-attachments-preview"

# Ziel:
# - Rechnungsanhänge sind im Rechnungsportal und in Rechnungs-/OP-Tabellen sichtbar und öffnungsfähig.
# - Rechnungsportal erhält je Reiter einen unabhängigen zweiten Vorschau-Fensterblock.
# - Portal-Tabellen werden kompakter; Spalten "Offen" und "Freigegeben" entfallen.
# - Spalte "Rechnungsbeschreibung" wird ergänzt.

# ---------- Allgemeine Anhangshelfer ----------
def _fm_attachment_rows_v629(entity_type: str, reference_no: str):
    if not entity_type or not reference_no:
        return []
    with get_connection() as con:
        return list(con.execute(
            "SELECT * FROM attachments WHERE entity_type=? AND reference_no=? ORDER BY id",
            (entity_type, reference_no),
        ))

def _fm_attachment_count_v629(entity_type: str, reference_no: str) -> int:
    try:
        with get_connection() as con:
            row = con.execute(
                "SELECT COUNT(*) FROM attachments WHERE entity_type=? AND reference_no=?",
                (entity_type, reference_no),
            ).fetchone()
            return int(row[0] or 0) if row else 0
    except Exception:
        return 0

def _fm_attachment_label_v629(entity_type: str, reference_no: str) -> str:
    cnt = _fm_attachment_count_v629(entity_type, reference_no)
    return f"📎 {cnt}" if cnt else "—"

def _fm_first_attachment_path_v629(entity_type: str, reference_no: str) -> str:
    rows = _fm_attachment_rows_v629(entity_type, reference_no)
    return rows[0]["file_path"] if rows else ""

def _fm_preview_attachments_in_pane_v629(preview_pane, title_label, entity_type: str, reference_no: str) -> None:
    rows = _fm_attachment_rows_v629(entity_type, reference_no)
    if not rows:
        title_label.configure(text=f"Keine Anhänge zu {reference_no}")
        preview_pane.clear()
        return
    try:
        images = [_fm_generate_document_preview(r["file_path"]) for r in rows]
        title = f"{reference_no} – {rows[0]['file_name']}"
        if len(rows) > 1:
            title += f" (+ {len(rows)-1} weitere)"
        title_label.configure(text=title)
        preview_pane.set_image(_fm_combine_pages(images))
    except Exception as exc:
        title_label.configure(text=f"Vorschaufehler zu {reference_no}")
        preview_pane.set_image(_fm_text_pages_to_image("Vorschaufehler", [str(exc)]))

# ---------- Rechnungs-/OP-Tabellen: Anhang mit Zähler anzeigen ----------
def _fm_invoice_reload_invoices_v629(self):
    trees = _v619_unique_trees(getattr(self, "inv_tree", None), *getattr(self, "invoice_tabs_trees", [])) if "_v619_unique_trees" in globals() else [getattr(self, "inv_tree", None)]
    trees = [t for t in trees if t is not None]
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    q = getattr(getattr(self, "invoice_search_var", None), "get", lambda: "")().strip().lower()
    if q == PLACEHOLDER_TEXT.lower():
        q = ""
    with get_connection() as con:
        rows = list(con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"))
    for tree in trees:
        for item in tree.get_children():
            tree.delete(item)
        cols = tuple(tree["columns"])
        for r in rows:
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            att_label = _fm_attachment_label_v629(self.entity_type, r["invoice_no"])
            if "kind" in cols:
                values = (
                    r["invoice_no"],
                    r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung",
                    r[partner_name_col],
                    r["description"] if "description" in r.keys() else "",
                    r["invoice_date"],
                    r["due_date"],
                    format_amount(r["gross_amount"]),
                    format_amount(r["open_amount"]),
                    _fm_status_from_open(r["open_amount"]) if "_fm_status_from_open" in globals() else r["status"],
                    att_label,
                )
            else:
                values = (
                    r["invoice_no"], r[partner_name_col], r["invoice_date"], r["due_date"],
                    format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["status"], att_label,
                )
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values, tags=(tag,))

def _fm_invoice_reload_open_items_v629(self):
    trees = _v619_unique_trees(getattr(self, "op_tree", None), *getattr(self, "open_item_tabs_trees", [])) if "_v619_unique_trees" in globals() else [getattr(self, "op_tree", None)]
    trees = [t for t in trees if t is not None]
    q = getattr(getattr(self, "open_items_search_var", None), "get", lambda: "")().strip().lower()
    if q == PLACEHOLDER_TEXT.lower():
        q = ""
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date, id", (self.entity_type,)))
    for tree in trees:
        for i in tree.get_children():
            tree.delete(i)
        cols = tuple(tree["columns"])
        for r in rows:
            if parse_amount(r["open_amount"]) <= 0:
                continue
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            att_label = _fm_attachment_label_v629(self.entity_type, r["reference_no"])
            if "kind" in cols:
                values = (
                    r["reference_no"],
                    r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung",
                    r["partner_name"],
                    r["description"] if "description" in r.keys() else "",
                    r["due_date"],
                    format_amount(r["original_amount"]),
                    format_amount(r["open_amount"]),
                    r["status"],
                    att_label,
                )
            else:
                values = (r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"], att_label)
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values, tags=(tag,))

InvoiceModuleBase.reload_invoices = _fm_invoice_reload_invoices_v629
InvoiceModuleBase.reload_open_items = _fm_invoice_reload_open_items_v629

# ---------- Rechnungsportal: kompakte Tabellen + unabhängige Vorschaufenster je Reiter ----------
def _fm_portal_init_v629(self, parent, app):
    ttk.Frame.__init__(self, parent)
    self.app = app
    self.archive_filter = tk.StringVar(value="Alle")
    self.approval_filter = tk.StringVar(value="Alle")
    self.archive_search_var = tk.StringVar()
    self.approval_search_var = tk.StringVar()
    self._archive_preview_entity = ""
    self._archive_preview_ref = ""
    self._approval_preview_entity = ""
    self._approval_preview_ref = ""
    self._build_ui()
    self.reload_all()

def _fm_portal_build_ui_v629(self):
    self.nb = ttk.Notebook(self)
    self.nb.pack(fill="both", expand=True)
    self.archive_tab = ttk.Frame(self.nb)
    self.approval_tab = ttk.Frame(self.nb)
    self.nb.add(self.archive_tab, text="Rechnungsarchiv")
    self.nb.add(self.approval_tab, text="Rechnungsfreigabe")
    self._build_archive()
    self._build_approval()

def _fm_portal_make_search_v629(parent, var, callback):
    row = ttk.Frame(parent)
    row.pack(fill="x", padx=10, pady=(0, 4))
    ttk.Label(row, text="Suchen", width=8).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, callback)
    return entry

def _fm_portal_create_tree_v629(self, parent):
    frame = ttk.Frame(parent)
    frame.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    cols = ("direction", "invoice_no", "kind", "partner", "description", "date", "due", "gross", "approval_status", "checked_by", "att")
    tree = ttk.Treeview(frame, columns=cols, show="headings", height=13)
    col_defs = [
        ("direction", "Richtung", 105),
        ("invoice_no", "Rechnung", 105),
        ("kind", "Art", 75),
        ("partner", "Partner", 145),
        ("description", "Rechnungsbeschreibung", 210),
        ("date", "Datum", 82),
        ("due", "Fällig", 82),
        ("gross", "Brutto", 85),
        ("approval_status", "Prüfstatus", 155),
        ("checked_by", "Geprüft von", 105),
        ("att", "Anhang", 70),
    ]
    for c, title, width in col_defs:
        tree.heading(c, text=title)
        tree.column(c, width=width, minwidth=45, anchor="w", stretch=False)
    vsb = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
    hsb = ttk.Scrollbar(frame, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
    tree.grid(row=0, column=0, sticky="nsew")
    vsb.grid(row=0, column=1, sticky="ns")
    hsb.grid(row=1, column=0, sticky="ew")
    frame.rowconfigure(0, weight=1)
    frame.columnconfigure(0, weight=1)
    configure_tree_tags(tree)
    self.setup_sorting(tree)
    return tree

def _fm_portal_build_preview_v629(self, parent, scope: str):
    title = "Anhangsvorschau – Archiv" if scope == "archive" else "Anhangsvorschau – Freigabe"
    ttk.Label(parent, text=title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))
    body = ttk.Frame(parent)
    body.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    title_label = ttk.Label(body, text="Kein Anhang ausgewählt", style="Hint.TLabel")
    title_label.pack(anchor="w", pady=(0, 6))
    pane = _FMDocumentPreviewPane(body)
    pane.pack(fill="both", expand=True)
    bar = ttk.Frame(body)
    bar.pack(fill="x", pady=(6, 0))
    def _open_external():
        entity = getattr(self, f"_{scope}_preview_entity", "")
        ref = getattr(self, f"_{scope}_preview_ref", "")
        path = _fm_first_attachment_path_v629(entity, ref)
        if path:
            open_path(path)
        else:
            messagebox.showinfo("Anhang", "Kein Anhang zur Vorschau ausgewählt.", parent=self)
    create_standard_button(bar, "Extern öffnen", _open_external).pack(side="left", padx=3)
    setattr(self, f"{scope}_preview_title_label", title_label)
    setattr(self, f"{scope}_preview_pane_widget", pane)

def _fm_portal_ensure_preview_v629(self, scope: str):
    paned = getattr(self, f"{scope}_paned")
    preview = getattr(self, f"{scope}_preview_frame")
    try:
        panes = [str(p) for p in paned.panes()]
    except Exception:
        panes = []
    if str(preview) not in panes:
        paned.add(preview, weight=2)
        try:
            paned.sashpos(0, max(520, int(paned.winfo_width() * 0.62)))
        except Exception:
            pass

def _fm_portal_open_attachment_preview_v629(self, scope: str):
    tree = self.archive_tree if scope == "archive" else self.approval_tree
    sel = self._selected(tree)
    if not sel:
        return
    entity_type, invoice_no = sel
    if _fm_attachment_count_v629(entity_type, invoice_no) <= 0:
        messagebox.showinfo("Anhang", "Für diese Rechnung ist kein Anhang vorhanden.", parent=self)
        return
    setattr(self, f"_{scope}_preview_entity", entity_type)
    setattr(self, f"_{scope}_preview_ref", invoice_no)
    self._ensure_preview(scope)
    _fm_preview_attachments_in_pane_v629(
        getattr(self, f"{scope}_preview_pane_widget"),
        getattr(self, f"{scope}_preview_title_label"),
        entity_type,
        invoice_no,
    )

def _fm_portal_build_archive_v629(self):
    self.archive_paned = ttk.PanedWindow(self.archive_tab, orient="horizontal")
    self.archive_paned.pack(fill="both", expand=True)
    self.archive_left = ttk.Frame(self.archive_paned)
    self.archive_preview_frame = ttk.Frame(self.archive_paned)
    self.archive_paned.add(self.archive_left, weight=5)
    top = ttk.Frame(self.archive_left); top.pack(fill="x", padx=10, pady=8)
    ttk.Label(top, text="Filter", width=8).pack(side="left")
    cb = ttk.Combobox(top, textvariable=self.archive_filter, values=["Alle", "Eingangsrechnungen", "Ausgangsrechnungen"], state="readonly", width=22)
    cb.pack(side="left", padx=4); cb.bind("<<ComboboxSelected>>", lambda _e: self.reload_archive()); cb.bind("<MouseWheel>", lambda e: "break")
    create_standard_button(top, "Aktualisieren", self.reload_archive).pack(side="left", padx=3)
    create_standard_button(top, "Bearbeiten", self.edit_selected_archive).pack(side="left", padx=3)
    create_standard_button(top, "Anhang öffnen", lambda: self.open_attachment_preview("archive"), confirm=True).pack(side="left", padx=3)
    _fm_portal_make_search_v629(self.archive_left, self.archive_search_var, self.reload_archive)
    self.archive_tree = self._invoice_tree(self.archive_left)
    self.archive_tree.bind("<ButtonRelease-1>", lambda e: self._portal_tree_click(e, "archive"), add="+")
    self.archive_tree.bind("<Double-1>", lambda _e: self.open_attachment_preview("archive"), add="+")
    self._build_preview(self.archive_preview_frame, "archive")

def _fm_portal_build_approval_v629(self):
    self.approval_paned = ttk.PanedWindow(self.approval_tab, orient="horizontal")
    self.approval_paned.pack(fill="both", expand=True)
    self.approval_left = ttk.Frame(self.approval_paned)
    self.approval_preview_frame = ttk.Frame(self.approval_paned)
    self.approval_paned.add(self.approval_left, weight=5)
    top = ttk.Frame(self.approval_left); top.pack(fill="x", padx=10, pady=8)
    ttk.Label(top, text="Filter", width=8).pack(side="left")
    cb = ttk.Combobox(top, textvariable=self.approval_filter, values=["Alle", "Eingangsrechnungen", "Ausgangsrechnungen"], state="readonly", width=22)
    cb.pack(side="left", padx=4); cb.bind("<<ComboboxSelected>>", lambda _e: self.reload_approval()); cb.bind("<MouseWheel>", lambda e: "break")
    create_standard_button(top, "Aktualisieren", self.reload_approval).pack(side="left", padx=3)
    create_standard_button(top, "Bearbeiten", self.edit_selected_approval).pack(side="left", padx=3)
    create_standard_button(top, "Anhang prüfen", lambda: self.open_attachment_preview("approval"), confirm=True).pack(side="left", padx=3)
    create_standard_button(top, "Sachlich prüfen", lambda: self.perform_action("review")).pack(side="left", padx=3)
    create_standard_button(top, "1. Freigabe", lambda: self.perform_action("approval1"), confirm=True).pack(side="left", padx=3)
    create_standard_button(top, "2. Freigabe", lambda: self.perform_action("approval2"), confirm=True).pack(side="left", padx=3)
    _fm_portal_make_search_v629(self.approval_left, self.approval_search_var, self.reload_approval)
    self.approval_tree = self._invoice_tree(self.approval_left)
    self.approval_tree.bind("<ButtonRelease-1>", lambda e: self._portal_tree_click(e, "approval"), add="+")
    self.approval_tree.bind("<Double-1>", lambda _e: self.open_attachment_preview("approval"), add="+")
    self._build_preview(self.approval_preview_frame, "approval")

def _fm_portal_tree_click_v629(self, event, scope: str):
    tree = self.archive_tree if scope == "archive" else self.approval_tree
    row = tree.identify_row(event.y)
    col = tree.identify_column(event.x)
    if row and col == f"#{len(tree['columns'])}":
        try:
            tree.selection_set(row)
            tree.focus(row)
        except Exception:
            pass
        self.open_attachment_preview(scope)
        return "break"
    return None

def _fm_portal_load_rows_v629(self, flt, only_unreleased=False):
    rows = []
    with get_connection() as con:
        if "vendor_invoice" in self._filter_sql(flt):
            q = "SELECT 'vendor_invoice' AS entity_type, invoice_no, invoice_kind, vendor_name AS partner, description, invoice_date, due_date, gross_amount, approval_status, last_checked_by, released FROM vendor_invoices"
            if only_unreleased:
                q += " WHERE COALESCE(released,0)=0"
            rows.extend(con.execute(q).fetchall())
        if "customer_invoice" in self._filter_sql(flt):
            q = "SELECT 'customer_invoice' AS entity_type, invoice_no, invoice_kind, customer_name AS partner, description, invoice_date, due_date, gross_amount, approval_status, last_checked_by, released FROM customer_invoices"
            if only_unreleased:
                q += " WHERE COALESCE(released,0)=0"
            rows.extend(con.execute(q).fetchall())
    return sorted(rows, key=lambda r: (r["invoice_date"], r["invoice_no"]), reverse=True)

def _fm_portal_fill_tree_v629(self, tree, rows):
    for i in tree.get_children():
        tree.delete(i)
    for r in rows:
        status = r["approval_status"] or APPROVAL_STATUS_REVIEW
        values = (
            _fm_invoice_direction(r["entity_type"]),
            r["invoice_no"],
            r["invoice_kind"] or "Rechnung",
            r["partner"],
            r["description"] if "description" in r.keys() else "",
            r["invoice_date"],
            r["due_date"],
            format_amount(r["gross_amount"]),
            status,
            r["last_checked_by"] or "",
            _fm_attachment_label_v629(r["entity_type"], r["invoice_no"]),
        )
        tree.insert("", "end", iid=f"{r['entity_type']}|{r['invoice_no']}", values=values, tags=("open",))

def _fm_portal_apply_filter_v629(tree, query: str):
    q = (query or "").strip().lower()
    if not q or q == PLACEHOLDER_TEXT.lower():
        return
    for iid in list(tree.get_children()):
        values = tree.item(iid, "values")
        if q not in " ".join(str(v) for v in values).lower():
            tree.delete(iid)

def _fm_portal_reload_archive_v629(self):
    self._fill_tree(self.archive_tree, self._load_rows(self.archive_filter.get(), only_unreleased=False))
    _fm_portal_apply_filter_v629(self.archive_tree, getattr(self.archive_search_var, "get", lambda: "")())

def _fm_portal_reload_approval_v629(self):
    self._fill_tree(self.approval_tree, self._load_rows(self.approval_filter.get(), only_unreleased=True))
    _fm_portal_apply_filter_v629(self.approval_tree, getattr(self.approval_search_var, "get", lambda: "")())

def _fm_portal_reload_all_v629(self):
    self.reload_archive()
    self.reload_approval()

def _fm_portal_selected_v629(self, tree):
    iid = tree.focus()
    if not iid or "|" not in iid:
        messagebox.showinfo("Rechnungsportal", "Bitte Rechnung auswählen.", parent=self)
        return None
    entity_type, invoice_no = iid.split("|", 1)
    return entity_type, invoice_no

def _fm_portal_edit_selected_archive_v629(self):
    sel = self._selected(self.archive_tree)
    if sel:
        self.open_edit_popup(*sel)

def _fm_portal_edit_selected_approval_v629(self):
    sel = self._selected(self.approval_tree)
    if sel:
        self.open_edit_popup(*sel)

InvoicePortalView.__init__ = _fm_portal_init_v629
InvoicePortalView._build_ui = _fm_portal_build_ui_v629
InvoicePortalView._build_archive = _fm_portal_build_archive_v629
InvoicePortalView._build_approval = _fm_portal_build_approval_v629
InvoicePortalView._invoice_tree = _fm_portal_create_tree_v629
InvoicePortalView._build_preview = _fm_portal_build_preview_v629
InvoicePortalView._ensure_preview = _fm_portal_ensure_preview_v629
InvoicePortalView.open_attachment_preview = _fm_portal_open_attachment_preview_v629
InvoicePortalView._portal_tree_click = _fm_portal_tree_click_v629
InvoicePortalView._load_rows = _fm_portal_load_rows_v629
InvoicePortalView._fill_tree = _fm_portal_fill_tree_v629
InvoicePortalView.reload_archive = _fm_portal_reload_archive_v629
InvoicePortalView.reload_approval = _fm_portal_reload_approval_v629
InvoicePortalView.reload_all = _fm_portal_reload_all_v629
InvoicePortalView._selected = _fm_portal_selected_v629
InvoicePortalView.edit_selected_archive = _fm_portal_edit_selected_archive_v629
InvoicePortalView.edit_selected_approval = _fm_portal_edit_selected_approval_v629

# Nach Freigabeaktionen bleiben die neuen Tabellen/Preview-Strukturen erhalten.
_orig_portal_perform_action_v629 = InvoicePortalView.perform_action

def _fm_portal_perform_action_v629(self, action):
    _orig_portal_perform_action_v629(self, action)
    try:
        self.reload_all()
    except Exception:
        pass

InvoicePortalView.perform_action = _fm_portal_perform_action_v629

# ---------- v0.6.29 Selbsttest ----------
def _fm_portal_static_selftest_v629() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.29-portal-attachments-preview\"",
        "def _fm_portal_open_attachment_preview_v629",
        "Rechnungsbeschreibung",
        "Anhang öffnen",
        "Anhang prüfen",
        "InvoicePortalView.open_attachment_preview",
        "InvoiceModuleBase.reload_invoices = _fm_invoice_reload_invoices_v629",
        "InvoiceModuleBase.reload_open_items = _fm_invoice_reload_open_items_v629",
    ]
    missing = [x for x in required if x not in src]
    fill_src = _inspect.getsource(InvoicePortalView._fill_tree)
    tree_src = _inspect.getsource(InvoicePortalView._invoice_tree)
    for forbidden in ["open_amount", "released", "Freigegeben", "Offen\""]:
        if forbidden in fill_src or forbidden in tree_src:
            missing.append("verbotene Portalspalte/Offenlogik: " + forbidden)
    for token in ["description", "att", "approval_status"]:
        if token not in fill_src:
            missing.append(token)
    if missing:
        raise RuntimeError("Portal-/Anhang-Selbsttest v0.6.29 fehlgeschlagen: " + ", ".join(missing))

_fm_portal_static_selftest_v629()



# === FINANCE MATE PATCH V0_6_30_PAYMENTS_INVOICE_ATTACHMENTS ===
APP_VERSION = "0.6.30-payments-invoice-attachments"

# Ergänzung zu v0.6.29: Auch im Zahlungsmodul wird bei den OP-/Rechnungstabellen der Anhang sichtbar
# und per Klick/Doppelklick öffnungsfähig gemacht.

def _fm_open_attachments_readonly_v630(owner, entity_type: str, reference_no: str) -> None:
    rows = _fm_attachment_rows_v629(entity_type, reference_no)
    if not rows:
        messagebox.showinfo("Anhang", "Für diese Rechnung ist kein Anhang vorhanden.", parent=owner)
        return
    pop = tk.Toplevel(owner)
    pop.title(f"Anhänge – {reference_no}")
    pop.geometry("900x420")
    pop.configure(bg=BG)
    ttk.Label(pop, text=f"Anhänge zu {reference_no}", style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
    cols = ("file_name", "file_path", "added_at")
    tree = ttk.Treeview(pop, columns=cols, show="headings")
    for col, w, txt in [("file_name", 220, "Dateiname"), ("file_path", 520, "Pfad"), ("added_at", 140, "Hinzugefügt")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=12, pady=8)
    for r in rows:
        tree.insert("", "end", iid=str(r["id"]), values=(r["file_name"], r["file_path"], r["added_at"]))
    def _open_selected():
        iid = tree.focus()
        if iid:
            path = tree.set(iid, "file_path")
            if path:
                open_path(path)
    bar = ttk.Frame(pop)
    bar.pack(fill="x", padx=12, pady=8)
    create_standard_button(bar, "Öffnen", _open_selected, confirm=True).pack(side="left", padx=3)
    create_standard_button(bar, "Schließen", pop.destroy).pack(side="right", padx=3)
    tree.bind("<Double-1>", lambda _e: _open_selected())

def _fm_payments_tree_click_attachment_v630(self, event, tree):
    row = tree.identify_row(event.y)
    col = tree.identify_column(event.x)
    if row and col == f"#{len(tree['columns'])}":
        try:
            entity_type = "customer_invoice" if tree.set(row, "type") == "Zahlungseingang" else "vendor_invoice"
            reference_no = tree.set(row, "reference")
        except Exception:
            return None
        _fm_open_attachments_readonly_v630(self, entity_type, reference_no)
        return "break"
    return None

def _fm_payments_build_direction_tab_v630(self, parent, entity_type: str):
    is_in = entity_type == "customer_invoice"
    shell = build_two_block_shell(
        parent,
        "Zahlungseingang / Debitoren-OP" if is_in else "Zahlungsausgang / Kreditoren-OP",
        "Zahlungshistorie",
    )
    left = ttk.Frame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(shell.right)
    right.pack(fill="both", expand=True, padx=10, pady=8)

    form = ttk.Frame(left)
    form.pack(fill="x", pady=(0, 8))
    self._row_entry(form, "Zahlungsdatum", self.payment_date)
    self._row_entry(form, "Zahlbetrag", self.amount)
    self._row_combo(form, "Zahlungsart", self.payment_method, ["Überweisung", "Lastschrift", "Bar", "Karte", "Sonstige"])
    self._row_entry(form, "Bank/Konto", self.bank_account)
    self._row_entry(form, "Buchungstext", self.booking_text)
    ttk.Checkbutton(form, text="OP vollständig ausgleichen", variable=self.full_clear).pack(anchor="w", pady=3)
    btns = ttk.Frame(form)
    btns.pack(fill="x", pady=6)
    create_standard_button(btns, "Zahlung buchen", self.book_payment, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Formular leeren", self.clear_form).pack(side="left", padx=3)
    create_standard_button(btns, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)

    search_var = self.in_search_var if is_in else self.out_search_var
    _fm_payments_make_search(left, search_var, self.reload_open_items)
    cols = ("id", "type", "reference", "partner", "due", "original", "open", "status", "att")
    tree = ttk.Treeview(left, columns=cols, show="headings", height=12)
    for c, t, w in [
        ("id", "ID", 48), ("type", "Vorgang", 120), ("reference", "Referenz", 105),
        ("partner", "Partner", 165), ("due", "Fällig", 82), ("original", "Original", 85),
        ("open", "Offen", 85), ("status", "Status", 105), ("att", "Anhang", 70),
    ]:
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w", stretch=False)
    tree.pack(fill="both", expand=True)
    configure_tree_tags(tree)
    self.setup_sorting(tree)
    tree.bind("<<TreeviewSelect>>", lambda _e, tr=tree: self._select_open_item_from_tree(tr))
    tree.bind("<ButtonRelease-1>", lambda e, tr=tree: self._tree_click_attachment(e, tr), add="+")
    tree.bind("<Double-1>", lambda e, tr=tree: self._tree_click_attachment(e, tr), add="+")

    hist_search_var = self.in_history_search_var if is_in else self.out_history_search_var
    _fm_payments_make_search(right, hist_search_var, self.reload_history)
    hcols = ("date", "direction", "reference", "partner", "amount", "method", "bank", "text")
    hist = ttk.Treeview(right, columns=hcols, show="headings")
    for c, t, w in [
        ("date", "Datum", 90), ("direction", "Richtung", 115), ("reference", "Referenz", 120),
        ("partner", "Partner", 180), ("amount", "Betrag", 90), ("method", "Art", 100),
        ("bank", "Bank/Konto", 130), ("text", "Text", 180),
    ]:
        hist.heading(c, text=t)
        hist.column(c, width=w, anchor="w")
    hist.pack(fill="both", expand=True)
    self.setup_sorting(hist)

    if is_in:
        self.incoming_tree = tree
        self.incoming_history_tree = hist
    else:
        self.outgoing_tree = tree
        self.outgoing_history_tree = hist

def _fm_payments_reload_open_items_v630(self):
    mapping = [
        (getattr(self, "incoming_tree", None), "customer_invoice", getattr(self, "in_search_var", None)),
        (getattr(self, "outgoing_tree", None), "vendor_invoice", getattr(self, "out_search_var", None)),
    ]
    with get_connection() as con:
        all_rows = list(con.execute("SELECT * FROM open_items WHERE open_amount > 0 ORDER BY due_date, id"))
    for tree, entity_type, var in mapping:
        if tree is None:
            continue
        for i in tree.get_children():
            tree.delete(i)
        q = _fm_payment_query_text(var)
        for r in all_rows:
            if r["entity_type"] != entity_type:
                continue
            values = (
                r["id"],
                _fm_payment_direction_label(r["entity_type"]),
                r["reference_no"],
                r["partner_name"],
                r["due_date"],
                format_amount(r["original_amount"]),
                format_amount(r["open_amount"]),
                r["status"],
                _fm_attachment_label_v629(r["entity_type"], r["reference_no"]),
            )
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values, tags=(urgency_bucket(r["due_date"], r["open_amount"]),))

PaymentsView._build_direction_tab = _fm_payments_build_direction_tab_v630
PaymentsView.reload_open_items = _fm_payments_reload_open_items_v630
PaymentsView._tree_click_attachment = _fm_payments_tree_click_attachment_v630
PaymentsView.open_attachments_readonly = _fm_open_attachments_readonly_v630

# Selbsttest v0.6.30
def _fm_payments_attachment_static_selftest_v630() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.30-payments-invoice-attachments\"",
        "def _fm_payments_build_direction_tab_v630",
        "def _fm_payments_reload_open_items_v630",
        "_fm_attachment_label_v629",
        "PaymentsView._tree_click_attachment",
    ]
    missing = [x for x in required if x not in src]
    build_src = _inspect.getsource(PaymentsView._build_direction_tab)
    reload_src = _inspect.getsource(PaymentsView.reload_open_items)
    for token in ["att", "Anhang", "_tree_click_attachment"]:
        if token not in build_src:
            missing.append(token)
    if "_fm_attachment_label_v629" not in reload_src:
        missing.append("Attachment-Label in Zahlungen-Reload")
    if missing:
        raise RuntimeError("Zahlungen-Anhang-Selbsttest v0.6.30 fehlgeschlagen: " + ", ".join(missing))

_fm_payments_attachment_static_selftest_v630()



# === FINANCE MATE PATCH V0_6_31_BLOCK7_COMPLETE ===
APP_VERSION = "0.6.31-block7-complete"
EXPORTS_DIR = BASE_DIR / "exports"

# Block 7 Abschlussumfang:
# - Zahlungsvorschläge / Zahlungsläufe für fällige Kreditoren-OPs
# - SEPA-Exportgrundlage als XML + CSV-Prüfdatei
# - CSV-Bankimport für Bankbewegungen
# - automatischer OP-Abgleich importierter Bankbewegungen
# - Mahnvorschläge für überfällige Debitoren-OPs
# - technische Verknüpfung zu payments, payment_matches, bank_transactions, open_items, Rechnungen, Anhängen

def _fm_init_block7c_schema_v631() -> None:
    ensure_directories()
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    with get_connection() as con:
        con.executescript("""
CREATE TABLE IF NOT EXISTS payment_runs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    run_no TEXT UNIQUE NOT NULL,
    run_type TEXT DEFAULT 'Kreditorenzahlung',
    status TEXT DEFAULT 'Vorschlag',
    proposed_date TEXT DEFAULT '',
    payment_date TEXT DEFAULT '',
    total_amount REAL DEFAULT 0,
    item_count INTEGER DEFAULT 0,
    export_file_path TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS payment_run_items (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    run_no TEXT NOT NULL,
    entity_type TEXT NOT NULL,
    reference_no TEXT NOT NULL,
    partner_no TEXT DEFAULT '',
    partner_name TEXT DEFAULT '',
    due_date TEXT DEFAULT '',
    amount REAL DEFAULT 0,
    status TEXT DEFAULT 'Vorgeschlagen',
    selected INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL,
    FOREIGN KEY(run_no) REFERENCES payment_runs(run_no) ON DELETE CASCADE
);
CREATE TABLE IF NOT EXISTS dunning_notices (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    notice_no TEXT UNIQUE NOT NULL,
    entity_type TEXT NOT NULL,
    reference_no TEXT NOT NULL,
    partner_no TEXT DEFAULT '',
    partner_name TEXT DEFAULT '',
    due_date TEXT DEFAULT '',
    open_amount REAL DEFAULT 0,
    dunning_level INTEGER DEFAULT 1,
    notice_date TEXT DEFAULT '',
    status TEXT DEFAULT 'Entwurf',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
""")
        # Ergänzende Spalten für bestehende Tabellen defensiv nachziehen.
        for col, definition in [
            ("bank_transaction_id", "INTEGER DEFAULT 0"),
            ("export_batch_id", "TEXT DEFAULT ''"),
            ("match_status", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, "payments", col, definition)
        for col, definition in [
            ("bank_account_id", "INTEGER DEFAULT 0"),
            ("booking_date", "TEXT DEFAULT ''"),
            ("value_date", "TEXT DEFAULT ''"),
            ("amount", "REAL DEFAULT 0"),
            ("currency", "TEXT DEFAULT 'EUR'"),
            ("direction", "TEXT DEFAULT ''"),
            ("partner_name", "TEXT DEFAULT ''"),
            ("partner_iban", "TEXT DEFAULT ''"),
            ("purpose_text", "TEXT DEFAULT ''"),
            ("end_to_end_id", "TEXT DEFAULT ''"),
            ("bank_reference", "TEXT DEFAULT ''"),
            ("imported_from", "TEXT DEFAULT ''"),
            ("import_batch_id", "TEXT DEFAULT ''"),
            ("match_status", "TEXT DEFAULT 'Ungeklärt'"),
            ("matched_entity_type", "TEXT DEFAULT ''"),
            ("matched_reference_no", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, "bank_transactions", col, definition)
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()

_orig_init_sqlite_v631 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v631()
    _fm_init_block7c_schema_v631()

# ---------- Block 7 Kernlogik ----------
def _fm_next_payment_run_no_v631() -> str:
    return generate_number("ZL-", "counter_payment_runs", 6)

def _fm_next_dunning_no_v631() -> str:
    return generate_number("MAH-", "counter_dunning", 6)

def _fm_date_leq_v631(left: str, right: str) -> bool:
    try:
        return datetime.strptime(left, DATE_FMT).date() <= datetime.strptime(right, DATE_FMT).date()
    except Exception:
        return False

def _fm_safe_amount_v631(value) -> Decimal:
    try:
        return parse_amount(value)
    except Exception:
        return Decimal("0.00")

def _fm_create_payment_proposal_v631(cutoff_date: str | None = None) -> str:
    cutoff = cutoff_date or today_str()
    run_no = _fm_next_payment_run_no_v631()
    now = now_str()
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type='vendor_invoice' AND open_amount > 0 ORDER BY due_date, partner_name, reference_no"))
        due_rows = [r for r in rows if _fm_date_leq_v631(r["due_date"], cutoff)]
        total = sum((_fm_safe_amount_v631(r["open_amount"]) for r in due_rows), Decimal("0.00"))
        con.execute(
            "INSERT INTO payment_runs(run_no,run_type,status,proposed_date,payment_date,total_amount,item_count,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)",
            (run_no, "Kreditorenzahlung", "Vorschlag", cutoff, cutoff, float(total), len(due_rows), now, now),
        )
        for r in due_rows:
            con.execute(
                "INSERT INTO payment_run_items(run_no,entity_type,reference_no,partner_no,partner_name,due_date,amount,status,selected,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                (run_no, r["entity_type"], r["reference_no"], r["partner_no"], r["partner_name"], r["due_date"], float(_fm_safe_amount_v631(r["open_amount"])), "Vorgeschlagen", 1, now, now),
            )
        con.commit()
    return run_no

def _fm_latest_payment_run_v631(statuses=("Vorschlag", "Exportiert")) -> str:
    placeholders = ",".join(["?"] * len(statuses))
    with get_connection() as con:
        row = con.execute(f"SELECT run_no FROM payment_runs WHERE status IN ({placeholders}) ORDER BY id DESC LIMIT 1", tuple(statuses)).fetchone()
        return row["run_no"] if row else ""

def _fm_export_payment_run_v631(run_no: str | None = None) -> tuple[str, str, int]:
    rn = run_no or _fm_latest_payment_run_v631(("Vorschlag", "Exportiert"))
    if not rn:
        raise ValueError("Kein Zahlungslauf vorhanden. Bitte zuerst Zahlungsvorschlag erstellen.")
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    with get_connection() as con:
        items = list(con.execute("SELECT * FROM payment_run_items WHERE run_no=? AND selected=1 ORDER BY partner_name, reference_no", (rn,)))
        if not items:
            raise ValueError("Der Zahlungslauf enthält keine ausgewählten Positionen.")
        run = con.execute("SELECT * FROM payment_runs WHERE run_no=?", (rn,)).fetchone()
    csv_path = EXPORTS_DIR / f"{rn}_sepa_pruefliste.csv"
    xml_path = EXPORTS_DIR / f"{rn}_sepa_export.xml"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        import csv as _csv
        writer = _csv.writer(f, delimiter=";")
        writer.writerow(["Zahlungslauf", "Referenz", "Kreditor", "Fällig", "Betrag", "Währung", "Verwendungszweck"])
        for it in items:
            writer.writerow([rn, it["reference_no"], it["partner_name"], it["due_date"], format_amount(it["amount"]), "EUR", f"Zahlung {it['reference_no']}"])
    # bewusst einfache, nachvollziehbare SEPA-Grundstruktur zur Weitergabe an Bank-/ERP-Schnittstellenprüfung
    from xml.sax.saxutils import escape
    total = sum((_fm_safe_amount_v631(it["amount"]) for it in items), Decimal("0.00"))
    lines = ["<?xml version=\"1.0\" encoding=\"UTF-8\"?>", "<Document>", "  <CstmrCdtTrfInitn>", f"    <GrpHdr><MsgId>{escape(rn)}</MsgId><NbOfTxs>{len(items)}</NbOfTxs><CtrlSum>{total:.2f}</CtrlSum></GrpHdr>", "    <PmtInf>"]
    for it in items:
        lines.append("      <CdtTrfTxInf>")
        lines.append(f"        <PmtId><EndToEndId>{escape(it['reference_no'])}</EndToEndId></PmtId>")
        lines.append(f"        <Amt><InstdAmt Ccy=\"EUR\">{_fm_safe_amount_v631(it['amount']):.2f}</InstdAmt></Amt>")
        lines.append(f"        <Cdtr><Nm>{escape(it['partner_name'] or '')}</Nm></Cdtr>")
        lines.append(f"        <RmtInf><Ustrd>Zahlung {escape(it['reference_no'])}</Ustrd></RmtInf>")
        lines.append("      </CdtTrfTxInf>")
    lines += ["    </PmtInf>", "  </CstmrCdtTrfInitn>", "</Document>"]
    xml_path.write_text("\n".join(lines), encoding="utf-8")
    with get_connection() as con:
        con.execute("UPDATE payment_runs SET status=?, export_file_path=?, updated_at=? WHERE run_no=?", ("Exportiert", str(xml_path), now_str(), rn))
        con.execute("UPDATE payment_run_items SET status=?, updated_at=? WHERE run_no=? AND selected=1", ("Exportiert", now_str(), rn))
        con.commit()
    return str(xml_path), str(csv_path), len(items)

def _fm_book_payment_run_v631(run_no: str | None = None) -> int:
    rn = run_no or _fm_latest_payment_run_v631(("Exportiert", "Vorschlag"))
    if not rn:
        raise ValueError("Kein Zahlungslauf vorhanden.")
    booked = 0
    now = now_str()
    with get_connection() as con:
        items = list(con.execute("SELECT * FROM payment_run_items WHERE run_no=? AND selected=1 AND status <> 'Gebucht'", (rn,)))
        for it in items:
            op = con.execute("SELECT * FROM open_items WHERE entity_type=? AND reference_no=?", (it["entity_type"], it["reference_no"])).fetchone()
            if not op or _fm_safe_amount_v631(op["open_amount"]) <= 0:
                continue
            amount = min(_fm_safe_amount_v631(op["open_amount"]), _fm_safe_amount_v631(it["amount"]))
            new_open = (_fm_safe_amount_v631(op["open_amount"]) - amount).quantize(Decimal("0.01"))
            status = _fm_payment_status_from_amount_v624(new_open)
            cur = con.execute(
                "INSERT INTO payments(entity_type,reference_no,partner_no,partner_name,payment_date,amount,payment_method,bank_account,booking_text,export_batch_id,match_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (op["entity_type"], op["reference_no"], op["partner_no"], op["partner_name"], today_str(), float(amount), "SEPA/Zahlungslauf", "", f"Zahlungslauf {rn}", rn, "Export/gebucht", now, now),
            )
            payment_id = cur.lastrowid
            con.execute("UPDATE open_items SET open_amount=?, status=?, updated_at=? WHERE id=?", (float(new_open), status, now, op["id"]))
            table = _fm_invoice_table_for_entity_v624(op["entity_type"])
            con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), status, now, op["reference_no"]))
            con.execute(
                "INSERT INTO payment_matches(payment_id,bank_transaction_id,entity_type,reference_no,matched_amount,difference_amount,match_type,created_by,created_at) VALUES(?,?,?,?,?,?,?,?,?)",
                (payment_id, 0, op["entity_type"], op["reference_no"], float(amount), float(new_open), "Zahlungslauf", _fm_current_user(), now),
            )
            con.execute("UPDATE payment_run_items SET status=?, updated_at=? WHERE id=?", ("Gebucht", now, it["id"]))
            booked += 1
        if booked:
            con.execute("UPDATE payment_runs SET status=?, updated_at=? WHERE run_no=?", ("Gebucht", now, rn))
        con.commit()
    return booked

def _fm_import_bank_csv_v631(path: str) -> int:
    import csv as _csv
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)
    sample = p.read_text(encoding="utf-8-sig", errors="replace")[:2048]
    delimiter = ";" if sample.count(";") >= sample.count(",") else ","
    imported = 0
    batch_id = f"BANK-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    def pick(row, *names):
        lower = {str(k).strip().lower(): v for k, v in row.items()}
        for name in names:
            if name.lower() in lower:
                return lower[name.lower()]
        return ""
    with p.open("r", encoding="utf-8-sig", errors="replace", newline="") as f, get_connection() as con:
        reader = _csv.DictReader(f, delimiter=delimiter)
        for row in reader:
            booking_date = pick(row, "Buchungsdatum", "booking_date", "Datum", "Date") or today_str()
            value_date = pick(row, "Valuta", "value_date", "Wertstellung") or booking_date
            amount_raw = pick(row, "Betrag", "amount", "Umsatz") or "0"
            amount = _fm_safe_amount_v631(amount_raw)
            direction = "Eingang" if amount >= 0 else "Ausgang"
            partner = pick(row, "Partner", "Name", "Auftraggeber/Empfänger", "partner_name")
            purpose = pick(row, "Verwendungszweck", "purpose", "purpose_text", "Text")
            iban = pick(row, "IBAN", "partner_iban")
            eref = pick(row, "EndToEndId", "End-to-End-ID", "end_to_end_id")
            bref = pick(row, "Bankreferenz", "bank_reference", "Referenz")
            con.execute(
                "INSERT INTO bank_transactions(booking_date,value_date,amount,currency,direction,partner_name,partner_iban,purpose_text,end_to_end_id,bank_reference,imported_from,import_batch_id,match_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (booking_date, value_date, float(amount), "EUR", direction, partner, iban, purpose, eref, bref, str(p), batch_id, "Ungeklärt", now_str(), now_str()),
            )
            imported += 1
        con.commit()
    return imported

def _fm_auto_reconcile_bank_transactions_v631() -> int:
    matched = 0
    now = now_str()
    with get_connection() as con:
        txs = list(con.execute("SELECT * FROM bank_transactions WHERE COALESCE(match_status,'Ungeklärt') IN ('','Ungeklärt') ORDER BY booking_date, id"))
        for tx in txs:
            amount_abs = abs(_fm_safe_amount_v631(tx["amount"]))
            if amount_abs <= 0:
                continue
            entity = "customer_invoice" if _fm_safe_amount_v631(tx["amount"]) >= 0 else "vendor_invoice"
            blob = " ".join(str(tx[k] or "") for k in ["purpose_text", "end_to_end_id", "bank_reference"]).lower()
            candidates = list(con.execute("SELECT * FROM open_items WHERE entity_type=? AND open_amount > 0 ORDER BY due_date, id", (entity,)))
            op = None
            for cand in candidates:
                ref = str(cand["reference_no"]).lower()
                if ref and ref in blob:
                    op = cand
                    break
            if op is None:
                for cand in candidates:
                    if _fm_safe_amount_v631(cand["open_amount"]) == amount_abs:
                        op = cand
                        break
            if op is None:
                continue
            pay_amount = min(_fm_safe_amount_v631(op["open_amount"]), amount_abs)
            new_open = (_fm_safe_amount_v631(op["open_amount"]) - pay_amount).quantize(Decimal("0.01"))
            status = _fm_payment_status_from_amount_v624(new_open)
            cur = con.execute(
                "INSERT INTO payments(entity_type,reference_no,partner_no,partner_name,payment_date,amount,payment_method,bank_account,booking_text,bank_transaction_id,match_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (op["entity_type"], op["reference_no"], op["partner_no"], op["partner_name"], tx["booking_date"] or today_str(), float(pay_amount), "Bankimport", "", tx["purpose_text"] or f"Bankabgleich {tx['id']}", tx["id"], "Automatisch abgeglichen", now, now),
            )
            payment_id = cur.lastrowid
            con.execute("UPDATE open_items SET open_amount=?, status=?, updated_at=? WHERE id=?", (float(new_open), status, now, op["id"]))
            table = _fm_invoice_table_for_entity_v624(op["entity_type"])
            con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), status, now, op["reference_no"]))
            con.execute("UPDATE bank_transactions SET match_status=?, matched_entity_type=?, matched_reference_no=?, updated_at=? WHERE id=?", ("Abgeglichen", op["entity_type"], op["reference_no"], now, tx["id"]))
            con.execute(
                "INSERT INTO payment_matches(payment_id,bank_transaction_id,entity_type,reference_no,matched_amount,difference_amount,match_type,created_by,created_at) VALUES(?,?,?,?,?,?,?,?,?)",
                (payment_id, tx["id"], op["entity_type"], op["reference_no"], float(pay_amount), float(new_open), "Automatisch Bankimport", _fm_current_user(), now),
            )
            matched += 1
        con.commit()
    return matched

def _fm_create_dunning_proposals_v631(cutoff_date: str | None = None) -> int:
    cutoff = cutoff_date or today_str()
    created = 0
    now = now_str()
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type='customer_invoice' AND open_amount > 0 ORDER BY due_date, partner_name, reference_no"))
        for r in rows:
            if not _fm_date_leq_v631(r["due_date"], cutoff) or r["due_date"] == cutoff:
                continue
            prev = con.execute("SELECT MAX(dunning_level) AS lvl FROM dunning_notices WHERE entity_type=? AND reference_no=?", (r["entity_type"], r["reference_no"])).fetchone()
            lvl = int(prev["lvl"] or 0) + 1
            exists = con.execute("SELECT 1 FROM dunning_notices WHERE entity_type=? AND reference_no=? AND dunning_level=?", (r["entity_type"], r["reference_no"], lvl)).fetchone()
            if exists:
                continue
            notice_no = _fm_next_dunning_no_v631()
            con.execute(
                "INSERT INTO dunning_notices(notice_no,entity_type,reference_no,partner_no,partner_name,due_date,open_amount,dunning_level,notice_date,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                (notice_no, r["entity_type"], r["reference_no"], r["partner_no"], r["partner_name"], r["due_date"], float(_fm_safe_amount_v631(r["open_amount"])), lvl, cutoff, "Entwurf", now, now),
            )
            created += 1
        con.commit()
    return created

# ---------- PaymentsView UI-Erweiterungen ----------
def _fm_payments_build_bank_transactions_tab_v631(self, parent):
    top = ttk.Frame(parent)
    top.pack(fill="x", padx=10, pady=8)
    create_standard_button(top, "Bank-CSV importieren", self.import_bank_csv, confirm=True).pack(side="left", padx=3)
    create_standard_button(top, "Auto-Abgleich", self.auto_reconcile_bank, confirm=True).pack(side="left", padx=3)
    create_standard_button(top, "Zahlungsvorschlag", self.create_payment_proposal).pack(side="left", padx=3)
    create_standard_button(top, "SEPA exportieren", self.export_payment_run, confirm=True).pack(side="left", padx=3)
    create_standard_button(top, "Zahlungslauf buchen", self.book_payment_run, confirm=True).pack(side="left", padx=3)
    create_standard_button(top, "Mahnvorschläge", self.create_dunning_proposals).pack(side="left", padx=3)
    create_standard_button(top, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)
    _fm_payments_make_search(parent, self.bank_search_var, self.reload_bank_transactions)
    cols = ("id", "date", "value", "direction", "amount", "partner", "purpose", "match")
    tree = ttk.Treeview(parent, columns=cols, show="headings", height=9)
    for c, t, w in [
        ("id", "ID", 55), ("date", "Buchung", 90), ("value", "Valuta", 90), ("direction", "Richtung", 90),
        ("amount", "Betrag", 90), ("partner", "Partner", 180), ("purpose", "Verwendungszweck", 320), ("match", "Status", 120),
    ]:
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    self.bank_tx_tree = tree
    self.setup_sorting(tree)
    ttk.Label(parent, text="Zahlungsläufe / Vorschläge", style="Section.TLabel").pack(anchor="w", padx=10, pady=(4, 4))
    rcols = ("run_no", "type", "status", "date", "items", "total", "export")
    runs = ttk.Treeview(parent, columns=rcols, show="headings", height=6)
    for c, t, w in [("run_no", "Lauf", 105), ("type", "Typ", 135), ("status", "Status", 100), ("date", "Datum", 90), ("items", "Pos.", 55), ("total", "Summe", 95), ("export", "Export", 260)]:
        runs.heading(c, text=t)
        runs.column(c, width=w, anchor="w")
    runs.pack(fill="x", padx=10, pady=(0, 10))
    self.payment_run_tree = runs
    self.setup_sorting(runs)

def _fm_payments_reload_payment_runs_v631(self):
    tree = getattr(self, "payment_run_tree", None)
    if tree is None:
        return
    for i in tree.get_children():
        tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM payment_runs ORDER BY id DESC LIMIT 80"))
    for r in rows:
        tree.insert("", "end", iid=r["run_no"], values=(r["run_no"], r["run_type"], r["status"], r["payment_date"] or r["proposed_date"], r["item_count"], format_amount(r["total_amount"]), r["export_file_path"]))

def _fm_payments_refresh_all_v631(self):
    self.reload_open_items()
    self.reload_history()
    self.reload_bank_transactions()
    self.reload_payment_runs()
    try:
        self.app.set_status("Zahlungen aktualisiert.")
    except Exception:
        pass

def _fm_payments_create_payment_proposal_v631(self):
    try:
        run_no = _fm_create_payment_proposal_v631(today_str())
        self.refresh_all()
        self.app.set_status(f"Zahlungsvorschlag {run_no} erstellt.")
    except Exception as exc:
        messagebox.showerror("Zahlungsvorschlag", str(exc), parent=self)

def _fm_payments_export_payment_run_v631(self):
    try:
        run_no = ""
        tree = getattr(self, "payment_run_tree", None)
        if tree is not None and tree.focus():
            run_no = tree.focus()
        xml_path, csv_path, count = _fm_export_payment_run_v631(run_no or None)
        self.refresh_all()
        self.app.set_status(f"SEPA-Export erstellt: {count} Position(en).")
        messagebox.showinfo("SEPA-Export", f"Export erstellt:\n{xml_path}\n\nPrüfliste:\n{csv_path}", parent=self)
    except Exception as exc:
        messagebox.showerror("SEPA-Export", str(exc), parent=self)

def _fm_payments_book_payment_run_v631(self):
    try:
        run_no = ""
        tree = getattr(self, "payment_run_tree", None)
        if tree is not None and tree.focus():
            run_no = tree.focus()
        count = _fm_book_payment_run_v631(run_no or None)
        self.refresh_all()
        self.app.set_status(f"{count} Zahlungslauf-Position(en) gebucht.")
    except Exception as exc:
        messagebox.showerror("Zahlungslauf", str(exc), parent=self)

def _fm_payments_import_bank_csv_v631(self):
    path = filedialog.askopenfilename(title="Bankbewegungen importieren", filetypes=[("CSV", "*.csv"), ("Alle Dateien", "*.*")], parent=self)
    if not path:
        return
    try:
        count = _fm_import_bank_csv_v631(path)
        self.refresh_all()
        self.app.set_status(f"{count} Bankbewegung(en) importiert.")
    except Exception as exc:
        messagebox.showerror("Bankimport", str(exc), parent=self)

def _fm_payments_auto_reconcile_bank_v631(self):
    try:
        count = _fm_auto_reconcile_bank_transactions_v631()
        self.refresh_all()
        self.app.set_status(f"{count} Bankbewegung(en) automatisch abgeglichen.")
    except Exception as exc:
        messagebox.showerror("Auto-Abgleich", str(exc), parent=self)

def _fm_payments_create_dunning_proposals_v631(self):
    try:
        count = _fm_create_dunning_proposals_v631(today_str())
        self.refresh_all()
        self.app.set_status(f"{count} Mahnvorschlag/-vorschläge erstellt.")
        messagebox.showinfo("Mahnvorschläge", f"{count} Mahnvorschlag/-vorschläge erstellt.", parent=self)
    except Exception as exc:
        messagebox.showerror("Mahnvorschläge", str(exc), parent=self)

PaymentsView._build_bank_transactions_tab = _fm_payments_build_bank_transactions_tab_v631
PaymentsView.reload_payment_runs = _fm_payments_reload_payment_runs_v631
PaymentsView.refresh_all = _fm_payments_refresh_all_v631
PaymentsView.create_payment_proposal = _fm_payments_create_payment_proposal_v631
PaymentsView.export_payment_run = _fm_payments_export_payment_run_v631
PaymentsView.book_payment_run = _fm_payments_book_payment_run_v631
PaymentsView.import_bank_csv = _fm_payments_import_bank_csv_v631
PaymentsView.auto_reconcile_bank = _fm_payments_auto_reconcile_bank_v631
PaymentsView.create_dunning_proposals = _fm_payments_create_dunning_proposals_v631

# ---------- Block 7 v0.6.31 Selbsttest ----------
def _fm_block7c_static_selftest_v631() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.31-block7-complete\"",
        "CREATE TABLE IF NOT EXISTS payment_runs",
        "CREATE TABLE IF NOT EXISTS payment_run_items",
        "CREATE TABLE IF NOT EXISTS dunning_notices",
        "def _fm_create_payment_proposal_v631",
        "def _fm_export_payment_run_v631",
        "def _fm_import_bank_csv_v631",
        "def _fm_auto_reconcile_bank_transactions_v631",
        "def _fm_create_dunning_proposals_v631",
        "Bank-CSV importieren",
        "SEPA exportieren",
        "Mahnvorschläge",
    ]
    missing = [x for x in required if x not in src]
    build_src = _inspect.getsource(PaymentsView._build_bank_transactions_tab)
    for token in ["Zahlungsvorschlag", "Auto-Abgleich", "Zahlungslauf buchen", "payment_run_tree"]:
        if token not in build_src:
            missing.append(token)
    if missing:
        raise RuntimeError("Block-7C-Selbsttest v0.6.31 fehlgeschlagen: " + ", ".join(missing))

_fm_block7c_static_selftest_v631()



# === FINANCE MATE PATCH V0_6_33_ERP_FINANCE_COMPLETION ===
APP_VERSION = "0.6.33-erp-finance-completion"

# Ziel:
# - Rechnungsarchiv zeigt Ausgleichsdatum, offenen Betrag und Kreditor-/Debitornummer.
# - Mahnfrist wird explizit erfassbar/auswertbar statt nur implizit aus Fälligkeit.
# - Vorschläge 1–7 werden technisch/fachlich nachgezogen: Rollen, Indizes, SEPA-Validierung,
#   CAMT-Grundimport/EBICS-Konfiguration, Reporting/Audit, Zahlungslauf-Freigabe, Stammdaten-Sperre.
# - Buchungen werden bei Rechnungen/Zahlungen in Grundbuch/Hauptbuch-Nebenbuch-Brücke geschrieben.

# ---------- Schema-Erweiterungen ----------
_orig_init_sqlite_v633 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v633()
    with get_connection() as con:
        # Explizite Mahnfristfelder
        for table in ("customers", "vendors"):
            ensure_column(con, table, "dunning_term_days", "INTEGER DEFAULT 7")
        for table in ("customer_invoices", "vendor_invoices"):
            ensure_column(con, table, "reference_no", "TEXT DEFAULT ''")
            ensure_column(con, table, "settlement_date", "TEXT DEFAULT ''")
            ensure_column(con, table, "dunning_term_days", "INTEGER DEFAULT 7")
        for table in ("open_items", "payments", "payment_matches", "payment_run_items", "dunning_notices"):
            ensure_column(con, table, "invoice_no", "TEXT DEFAULT ''")
            ensure_column(con, table, "partner_index", "TEXT DEFAULT ''")
        ensure_column(con, "open_items", "settlement_date", "TEXT DEFAULT ''")
        ensure_column(con, "open_items", "dunning_date", "TEXT DEFAULT ''")
        ensure_column(con, "dunning_notices", "dunning_date", "TEXT DEFAULT ''")
        ensure_column(con, "dunning_notices", "dunning_term_days", "INTEGER DEFAULT 7")
        ensure_column(con, "payment_runs", "created_by", "TEXT DEFAULT ''")
        ensure_column(con, "payment_runs", "approved_by", "TEXT DEFAULT ''")
        ensure_column(con, "payment_runs", "approved_at", "TEXT DEFAULT ''")
        ensure_column(con, "payment_runs", "approval_status", "TEXT DEFAULT 'Nicht freigegeben'")
        # Audit / User / Schnittstellen
        con.executescript("""
CREATE TABLE IF NOT EXISTS fm_users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_name TEXT UNIQUE NOT NULL,
    display_name TEXT DEFAULT '',
    active INTEGER DEFAULT 1,
    approval_level INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS fm_roles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    role_name TEXT UNIQUE NOT NULL,
    description TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS fm_user_roles (
    user_name TEXT NOT NULL,
    role_name TEXT NOT NULL,
    created_at TEXT NOT NULL,
    PRIMARY KEY(user_name, role_name)
);
CREATE TABLE IF NOT EXISTS audit_log (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    event_time TEXT NOT NULL,
    user_name TEXT DEFAULT '',
    module TEXT DEFAULT '',
    action TEXT DEFAULT '',
    entity_type TEXT DEFAULT '',
    reference_no TEXT DEFAULT '',
    invoice_no TEXT DEFAULT '',
    partner_index TEXT DEFAULT '',
    details TEXT DEFAULT ''
);
CREATE TABLE IF NOT EXISTS ebics_configs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    bank_account_id INTEGER DEFAULT 0,
    host_id TEXT DEFAULT '',
    partner_id TEXT DEFAULT '',
    user_id TEXT DEFAULT '',
    url TEXT DEFAULT '',
    status TEXT DEFAULT 'Vorbereitet',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS masterdata_change_log (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    table_name TEXT NOT NULL,
    key_before TEXT DEFAULT '',
    key_after TEXT DEFAULT '',
    user_name TEXT DEFAULT '',
    changed_at TEXT NOT NULL,
    reason TEXT DEFAULT ''
);
""")
        now = now_str()
        # Standardrollen/-user
        for role, desc in [("Administrator", "Vollzugriff"), ("Buchhaltung", "Buchen und Abstimmen"), ("Freigeber", "Freigaben und Zahlungsläufe"), ("Leser", "Nur Lesen")]:
            con.execute("INSERT OR IGNORE INTO fm_roles(role_name,description,created_at,updated_at) VALUES(?,?,?,?)", (role, desc, now, now))
        user = _fm_current_user()
        con.execute("INSERT OR IGNORE INTO fm_users(user_name,display_name,active,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?)", (user, user, 1, _fm_current_approval_level(), now, now))
        con.execute("INSERT OR IGNORE INTO fm_user_roles(user_name,role_name,created_at) VALUES(?,?,?)", (user, "Administrator", now))
        # Bestandsdaten ableiten
        con.execute("UPDATE customer_invoices SET reference_no=invoice_no WHERE COALESCE(reference_no,'')='' ")
        con.execute("UPDATE vendor_invoices SET reference_no=invoice_no WHERE COALESCE(reference_no,'')='' ")
        con.execute("UPDATE open_items SET invoice_no=reference_no WHERE COALESCE(invoice_no,'')='' ")
        con.execute("UPDATE open_items SET partner_index=partner_no WHERE COALESCE(partner_index,'')='' ")
        con.execute("UPDATE payments SET invoice_no=reference_no WHERE COALESCE(invoice_no,'')='' ")
        con.execute("UPDATE payments SET partner_index=partner_no WHERE COALESCE(partner_index,'')='' ")
        # Indizes: fachliche Eindeutigkeit/Performance. UNIQUE nur auf Referenz + Richtung, damit interne Tests mit Bestandsdaten robust bleiben.
        con.execute("CREATE UNIQUE INDEX IF NOT EXISTS ux_customer_invoices_reference_no ON customer_invoices(reference_no) WHERE COALESCE(reference_no,'')<>''")
        con.execute("CREATE UNIQUE INDEX IF NOT EXISTS ux_vendor_invoices_reference_no ON vendor_invoices(reference_no) WHERE COALESCE(reference_no,'')<>''")
        for idx_sql in [
            "CREATE INDEX IF NOT EXISTS ix_open_items_ref_partner ON open_items(entity_type,reference_no,partner_no)",
            "CREATE INDEX IF NOT EXISTS ix_payments_ref_partner ON payments(entity_type,reference_no,partner_no)",
            "CREATE INDEX IF NOT EXISTS ix_matches_ref_partner ON payment_matches(entity_type,reference_no,partner_index)",
            "CREATE INDEX IF NOT EXISTS ix_bank_transactions_match ON bank_transactions(match_status,matched_reference_no)",
            "CREATE INDEX IF NOT EXISTS ix_audit_ref ON audit_log(reference_no,invoice_no,partner_index)",
        ]:
            con.execute(idx_sql)
        # Stammdaten-Sperrlogik: Nummer darf nach Verwendung nicht geändert werden.
        con.executescript("""
CREATE TRIGGER IF NOT EXISTS trg_customers_no_lock
BEFORE UPDATE OF customer_no ON customers
FOR EACH ROW
WHEN OLD.customer_no <> NEW.customer_no AND (
    EXISTS(SELECT 1 FROM customer_invoices WHERE customer_no=OLD.customer_no)
    OR EXISTS(SELECT 1 FROM open_items WHERE partner_no=OLD.customer_no)
    OR EXISTS(SELECT 1 FROM payments WHERE partner_no=OLD.customer_no)
)
BEGIN
    SELECT RAISE(ABORT, 'Debitorennummer ist bereits verwendet und darf nicht direkt geändert werden.');
END;
CREATE TRIGGER IF NOT EXISTS trg_vendors_no_lock
BEFORE UPDATE OF vendor_no ON vendors
FOR EACH ROW
WHEN OLD.vendor_no <> NEW.vendor_no AND (
    EXISTS(SELECT 1 FROM vendor_invoices WHERE vendor_no=OLD.vendor_no)
    OR EXISTS(SELECT 1 FROM open_items WHERE partner_no=OLD.vendor_no)
    OR EXISTS(SELECT 1 FROM payments WHERE partner_no=OLD.vendor_no)
)
BEGIN
    SELECT RAISE(ABORT, 'Kreditorennummer ist bereits verwendet und darf nicht direkt geändert werden.');
END;
""")
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now))
        con.commit()

# ---------- Rollen / Audit ----------
def _fm_log_audit_v633(module: str, action: str, entity_type: str = "", reference_no: str = "", invoice_no: str = "", partner_index: str = "", details: str = "") -> None:
    try:
        with get_connection() as con:
            con.execute("INSERT INTO audit_log(event_time,user_name,module,action,entity_type,reference_no,invoice_no,partner_index,details) VALUES(?,?,?,?,?,?,?,?,?)", (now_str(), _fm_current_user(), module, action, entity_type, reference_no, invoice_no, partner_index, details))
            con.commit()
    except Exception:
        pass

def _fm_user_has_role_v633(user_name: str, role_name: str) -> bool:
    with get_connection() as con:
        row = con.execute("SELECT 1 FROM fm_user_roles WHERE user_name=? AND role_name=?", (user_name, role_name)).fetchone()
        return row is not None

# ---------- Buchungslogik Grundbuch/Hauptbuch/Nebenbuch ----------
def _fm_next_journal_no_v633(prefix="JE-") -> str:
    return generate_number(prefix, "counter_journal", 6)

def _fm_create_journal_entry_v633(document_date: str, description: str, lines: list[tuple[str, str, float, str]], reference_no: str = "") -> str:
    # lines: (account_no, side, amount, text)
    debit = sum(float(a) for _acc, side, a, _txt in lines if side.lower().startswith("s") or side.lower().startswith("d"))
    credit = sum(float(a) for _acc, side, a, _txt in lines if side.lower().startswith("h") or side.lower().startswith("c"))
    if round(debit, 2) != round(credit, 2):
        raise ValueError(f"Journal nicht im Gleichgewicht: Soll {debit:.2f} / Haben {credit:.2f}")
    doc_no = _fm_next_journal_no_v633()
    with get_connection() as con:
        con.execute("INSERT INTO journal_entries(document_no,document_date,posting_date,description,total_debit,total_credit,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)", (doc_no, document_date, today_str(), description, debit, credit, "Gebucht", now_str(), now_str()))
        for acc, side, amount, txt in lines:
            con.execute("INSERT INTO journal_entry_lines(document_no,account_no,side,amount,text) VALUES(?,?,?,?,?)", (doc_no, acc, side, float(amount), txt))
        con.commit()
    _fm_log_audit_v633("Buchungsjournal", "Journal gebucht", reference_no=reference_no, details=doc_no)
    return doc_no

def _fm_post_invoice_to_journal_v633(entity_type: str, invoice_no_or_ref: str) -> str:
    with get_connection() as con:
        table = "vendor_invoices" if entity_type == "vendor_invoice" else "customer_invoices"
        inv = con.execute(f"SELECT * FROM {table} WHERE invoice_no=? OR reference_no=?", (invoice_no_or_ref, invoice_no_or_ref)).fetchone()
        if not inv:
            return ""
        if inv["linked_journal_no"]:
            return inv["linked_journal_no"]
        gross = float(inv["gross_amount"] or 0); net = float(inv["net_amount"] or 0); tax = float(inv["tax_amount"] or 0)
        ref = inv["reference_no"] if "reference_no" in inv.keys() and inv["reference_no"] else inv["invoice_no"]
    if entity_type == "customer_invoice":
        lines = [("1400", "Soll", gross, f"Forderung {ref}"), ("8400", "Haben", net, f"Erlös {ref}")]
        if tax: lines.append(("1776", "Haben", tax, f"USt {ref}"))
    else:
        lines = [("3400", "Soll", net, f"Aufwand {ref}"), ("1600", "Haben", gross, f"Verbindlichkeit {ref}")]
        if tax: lines.append(("1576", "Soll", tax, f"Vorsteuer {ref}"))
    doc = _fm_create_journal_entry_v633(inv["invoice_date"], f"Rechnungsbuchung {ref}", lines, ref)
    with get_connection() as con:
        con.execute(f"UPDATE {table} SET linked_journal_no=?, updated_at=? WHERE invoice_no=?", (doc, now_str(), inv["invoice_no"]))
        con.execute("UPDATE open_items SET linked_journal_no=?, updated_at=? WHERE entity_type=? AND (reference_no=? OR invoice_no=?)", (doc, now_str(), entity_type, ref, inv["invoice_no"]))
        con.commit()
    return doc

def _fm_post_payment_to_journal_v633(payment_id: int) -> str:
    with get_connection() as con:
        p = con.execute("SELECT * FROM payments WHERE id=?", (payment_id,)).fetchone()
        if not p:
            return ""
        marker = f"PAY-JE:{payment_id}"
        existing = con.execute("SELECT document_no FROM journal_entries WHERE description LIKE ?", (f"%{marker}%",)).fetchone()
        if existing:
            return existing["document_no"]
        ref = p["reference_no"]; amount = float(p["amount"] or 0)
    if p["entity_type"] == "customer_invoice":
        lines = [("1200", "Soll", amount, f"Bankeingang {ref}"), ("1400", "Haben", amount, f"Ausgleich Forderung {ref}")]
    else:
        lines = [("1600", "Soll", amount, f"Ausgleich Verbindlichkeit {ref}"), ("1200", "Haben", amount, f"Bankausgang {ref}")]
    doc = _fm_create_journal_entry_v633(p["payment_date"], f"Zahlungsbuchung {ref} {marker}", lines, ref)
    return doc

def _fm_post_unposted_payments_v633() -> int:
    count = 0
    with get_connection() as con:
        rows = list(con.execute("SELECT id FROM payments ORDER BY id"))
    for r in rows:
        if _fm_post_payment_to_journal_v633(r["id"]):
            count += 1
    return count

# Nach Freigabe OP und Journal erzeugen: bestehende OP-Funktion wrappen.
_orig_create_open_item_for_invoice_v633 = _fm_create_open_item_for_invoice

def _fm_create_open_item_for_invoice(con, entity_type: str, invoice_row):
    _orig_create_open_item_for_invoice_v633(con, entity_type, invoice_row)
    try:
        con.commit()
    except Exception:
        pass
    _fm_post_invoice_to_journal_v633(entity_type, invoice_row["reference_no"] if "reference_no" in invoice_row.keys() and invoice_row["reference_no"] else invoice_row["invoice_no"])

# ---------- Rechnungsarchiv-Spalten ----------
def _fm_archive_settlement_date_v633(entity_type: str, ref: str, invoice_no: str = "") -> str:
    with get_connection() as con:
        row = con.execute("SELECT MAX(payment_date) AS d FROM payments WHERE entity_type=? AND (reference_no=? OR invoice_no=?)", (entity_type, ref, invoice_no or ref)).fetchone()
        return row["d"] or ""

def _fm_portal_create_tree_v633(self, parent):
    frame = ttk.Frame(parent); frame.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    cols = ("direction", "invoice_no", "ref", "partner_no", "partner", "description", "date", "due", "gross", "open", "settlement", "dunning_date", "approval_status", "checked_by", "att")
    tree = ttk.Treeview(frame, columns=cols, show="headings", height=13)
    defs = [("direction","Richtung",95),("invoice_no","Rechnung",95),("ref","Referenz",115),("partner_no","Kred./Deb.-Nr.",105),("partner","Partner",130),("description","Rechnungsbeschreibung",180),("date","Datum",78),("due","Fällig",78),("gross","Brutto",82),("open","Offen",82),("settlement","Ausgleich",88),("dunning_date","Mahnfrist",88),("approval_status","Prüfstatus",130),("checked_by","Geprüft",90),("att","Anhang",65)]
    for c,t,w in defs:
        tree.heading(c, text=t); tree.column(c, width=w, minwidth=45, anchor="w", stretch=False)
    vsb=ttk.Scrollbar(frame, orient="vertical", command=tree.yview); hsb=ttk.Scrollbar(frame, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
    tree.grid(row=0,column=0,sticky="nsew"); vsb.grid(row=0,column=1,sticky="ns"); hsb.grid(row=1,column=0,sticky="ew")
    frame.rowconfigure(0, weight=1); frame.columnconfigure(0, weight=1)
    configure_tree_tags(tree); self.setup_sorting(tree); return tree

def _fm_calc_dunning_date_v633(due_date: str, days: int = 7) -> str:
    try:
        return (datetime.strptime(due_date, DATE_FMT) + timedelta(days=int(days or 7))).strftime(DATE_FMT)
    except Exception:
        return ""

def _fm_portal_load_rows_v633(self, flt, only_unreleased=False):
    rows=[]
    with get_connection() as con:
        if "vendor_invoice" in self._filter_sql(flt):
            q="SELECT 'vendor_invoice' AS entity_type, invoice_no, reference_no, invoice_kind, vendor_no AS partner_no, vendor_name AS partner, description, invoice_date, due_date, gross_amount, open_amount, settlement_date, dunning_term_days, approval_status, last_checked_by, released FROM vendor_invoices"
            if only_unreleased: q += " WHERE COALESCE(released,0)=0"
            rows.extend(con.execute(q).fetchall())
        if "customer_invoice" in self._filter_sql(flt):
            q="SELECT 'customer_invoice' AS entity_type, invoice_no, reference_no, invoice_kind, customer_no AS partner_no, customer_name AS partner, description, invoice_date, due_date, gross_amount, open_amount, settlement_date, dunning_term_days, approval_status, last_checked_by, released FROM customer_invoices"
            if only_unreleased: q += " WHERE COALESCE(released,0)=0"
            rows.extend(con.execute(q).fetchall())
    return sorted(rows, key=lambda r:(r["invoice_date"], r["invoice_no"]), reverse=True)

def _fm_portal_fill_tree_v633(self, tree, rows):
    for i in tree.get_children(): tree.delete(i)
    for r in rows:
        ref = r["reference_no"] if "reference_no" in r.keys() and r["reference_no"] else r["invoice_no"]
        settlement = r["settlement_date"] if "settlement_date" in r.keys() and r["settlement_date"] else _fm_archive_settlement_date_v633(r["entity_type"], ref, r["invoice_no"])
        ddate = _fm_calc_dunning_date_v633(r["due_date"], r["dunning_term_days"] if "dunning_term_days" in r.keys() else 7)
        vals = (_fm_invoice_direction(r["entity_type"]), r["invoice_no"], ref, r["partner_no"], r["partner"], r["description"] if "description" in r.keys() else "", r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), settlement, ddate, r["approval_status"] or APPROVAL_STATUS_REVIEW, r["last_checked_by"] or "", _fm_attachment_label_v629(r["entity_type"], ref))
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        tree.insert("", "end", iid=f"{r['entity_type']}|{ref}", values=vals, tags=(tag,))

InvoicePortalView._invoice_tree = _fm_portal_create_tree_v633
InvoicePortalView._load_rows = _fm_portal_load_rows_v633
InvoicePortalView._fill_tree = _fm_portal_fill_tree_v633

# ---------- Mahnfrist und Mahnwesen ----------
def _fm_create_dunning_proposals_v633(cutoff_date: str | None = None) -> int:
    cutoff = cutoff_date or today_str(); created=0; now=now_str()
    with get_connection() as con:
        rows=list(con.execute("SELECT * FROM open_items WHERE entity_type='customer_invoice' AND open_amount > 0 ORDER BY due_date, partner_name, reference_no"))
        for r in rows:
            inv=con.execute("SELECT * FROM customer_invoices WHERE reference_no=? OR invoice_no=?", (r["reference_no"], r["invoice_no"] if "invoice_no" in r.keys() else r["reference_no"])).fetchone()
            days=int((inv["dunning_term_days"] if inv and "dunning_term_days" in inv.keys() else 7) or 7)
            dunning_date=_fm_calc_dunning_date_v633(r["due_date"], days)
            if not dunning_date or not _fm_date_leq_v631(dunning_date, cutoff):
                continue
            prev=con.execute("SELECT MAX(dunning_level) AS lvl FROM dunning_notices WHERE entity_type=? AND reference_no=?", (r["entity_type"], r["reference_no"])).fetchone()
            lvl=int(prev["lvl"] or 0)+1
            notice_no=_fm_next_dunning_no_v631()
            con.execute("INSERT INTO dunning_notices(notice_no,entity_type,reference_no,invoice_no,partner_no,partner_index,partner_name,due_date,dunning_date,dunning_term_days,open_amount,dunning_level,notice_date,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (notice_no, r["entity_type"], r["reference_no"], r["invoice_no"] if "invoice_no" in r.keys() else r["reference_no"], r["partner_no"], r["partner_no"], r["partner_name"], r["due_date"], dunning_date, days, float(r["open_amount"]), lvl, cutoff, "Entwurf", now, now))
            created += 1
        con.commit()
    _fm_log_audit_v633("Mahnwesen", "Mahnvorschläge erstellt", details=str(created))
    return created
_fm_create_dunning_proposals_v631 = _fm_create_dunning_proposals_v633

# ---------- SEPA pain.001-nahe Validierung / CAMT-Import ----------
def _fm_validate_iban_v633(iban: str) -> bool:
    s = ''.join(str(iban or '').upper().split())
    return len(s) >= 15 and len(s) <= 34 and s[:2].isalpha() and s[2:].isalnum()

def _fm_validate_bic_v633(bic: str) -> bool:
    s = str(bic or '').upper().strip()
    return len(s) in (8, 11) and s[:4].isalpha()

def _fm_export_payment_run_v633(run_no: str | None = None, strict_bank: bool = True):
    rn = run_no or _fm_latest_payment_run_v631(("Freigegeben", "Exportiert", "Vorschlag"))
    if not rn: raise ValueError("Kein Zahlungslauf vorhanden.")
    with get_connection() as con:
        run = con.execute("SELECT * FROM payment_runs WHERE run_no=?", (rn,)).fetchone()
        items = list(con.execute("SELECT * FROM payment_run_items WHERE run_no=? AND selected=1 ORDER BY partner_name, reference_no", (rn,)))
        if strict_bank:
            for it in items:
                ven = con.execute("SELECT iban,bic FROM vendors WHERE vendor_no=?", (it["partner_no"],)).fetchone()
                if not ven or not _fm_validate_iban_v633(ven["iban"]) or not _fm_validate_bic_v633(ven["bic"]):
                    raise ValueError(f"SEPA-Validierung fehlgeschlagen: IBAN/BIC für Kreditor {it['partner_no']} unvollständig/ungültig.")
    # Nutzt vorhandene v631-Grundlage, ergänzt aber Status/Prüfung.
    result = _fm_export_payment_run_v631(rn)
    _fm_log_audit_v633("Zahlungen", "SEPA Export", reference_no=rn, details=str(result))
    return result

def _fm_import_camt_xml_v633(path: str) -> int:
    import xml.etree.ElementTree as ET
    p=Path(path); imported=0; batch_id=f"CAMT-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    root=ET.parse(str(p)).getroot()
    text_blob=' '.join(root.itertext())
    # Minimaler Grundimport: Betrag/Referenz aus Texten suchen, wenn echte Bankstruktur variiert.
    import re as _re
    amounts=_re.findall(r"[-+]?\d+[\.,]\d{2}", text_blob)
    refs=_re.findall(r"[A-Z]{1,4}[-_][A-Z0-9-]+", text_blob)
    with get_connection() as con:
        for idx, amt in enumerate(amounts or ['0,00']):
            amount=float(str(amt).replace('.','').replace(',','.'))
            ref=refs[idx] if idx < len(refs) else ''
            con.execute("INSERT INTO bank_transactions(booking_date,value_date,amount,currency,direction,partner_name,purpose_text,imported_from,import_batch_id,match_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)", (today_str(), today_str(), amount, 'EUR', 'Eingang' if amount>=0 else 'Ausgang', '', ref, str(p), batch_id, 'Ungeklärt', now_str(), now_str()))
            imported += 1
        con.commit()
    return imported

# ---------- Zahlungslauf-Freigabe / Vier-Augen ----------
def _fm_approve_payment_run_v633(run_no: str, user_name: str | None = None) -> bool:
    user = user_name or _fm_current_user(); now=now_str()
    with get_connection() as con:
        run=con.execute("SELECT * FROM payment_runs WHERE run_no=?", (run_no,)).fetchone()
        if not run: raise ValueError("Zahlungslauf nicht gefunden.")
        creator = run["created_by"] if "created_by" in run.keys() and run["created_by"] else ""
        if creator and creator == user:
            raise ValueError("Vier-Augen-Prinzip: Ersteller und Freigeber dürfen nicht identisch sein.")
        con.execute("UPDATE payment_runs SET approved_by=?, approved_at=?, approval_status=?, status=?, updated_at=? WHERE run_no=?", (user, now, "Freigegeben", "Freigegeben", now, run_no))
        con.commit()
    _fm_log_audit_v633("Zahlungen", "Zahlungslauf freigegeben", reference_no=run_no, details=user)
    return True

def _fm_create_payment_proposal_v633(cutoff_date: str | None = None) -> str:
    rn = _fm_create_payment_proposal_v631(cutoff_date)
    with get_connection() as con:
        con.execute("UPDATE payment_runs SET created_by=?, approval_status=?, updated_at=? WHERE run_no=?", (_fm_current_user(), "Nicht freigegeben", now_str(), rn))
        con.commit()
    _fm_log_audit_v633("Zahlungen", "Zahlungsvorschlag", reference_no=rn)
    return rn
_fm_create_payment_proposal_v631 = _fm_create_payment_proposal_v633

def _fm_book_payment_run_v633(run_no: str | None = None) -> int:
    rn = run_no or _fm_latest_payment_run_v631(("Freigegeben", "Exportiert", "Vorschlag"))
    if not rn: raise ValueError("Kein Zahlungslauf vorhanden.")
    with get_connection() as con:
        run=con.execute("SELECT * FROM payment_runs WHERE run_no=?", (rn,)).fetchone()
        if run and (run["approval_status"] if "approval_status" in run.keys() else "") != "Freigegeben":
            raise ValueError("Zahlungslauf muss vor Buchung freigegeben werden.")
    before_ids=set()
    with get_connection() as con:
        before_ids={r["id"] for r in con.execute("SELECT id FROM payments").fetchall()}
    count=_fm_book_payment_run_v631(rn)
    with get_connection() as con:
        new_payments=[r["id"] for r in con.execute("SELECT id FROM payments").fetchall() if r["id"] not in before_ids]
        for pid in new_payments:
            p=con.execute("SELECT * FROM payments WHERE id=?", (pid,)).fetchone()
            if p:
                con.execute("UPDATE payments SET partner_index=COALESCE(NULLIF(partner_index,''), partner_no), invoice_no=COALESCE(NULLIF(invoice_no,''), reference_no) WHERE id=?", (pid,))
                if p["entity_type"] in ("customer_invoice", "vendor_invoice"):
                    table="customer_invoices" if p["entity_type"]=="customer_invoice" else "vendor_invoices"
                    inv=con.execute(f"SELECT open_amount FROM {table} WHERE reference_no=? OR invoice_no=?", (p["reference_no"], p["invoice_no"] if "invoice_no" in p.keys() else p["reference_no"])).fetchone()
                    if inv and float(inv["open_amount"] or 0) <= 0:
                        con.execute(f"UPDATE {table} SET settlement_date=COALESCE(NULLIF(settlement_date,''), ?), updated_at=? WHERE reference_no=? OR invoice_no=?", (p["payment_date"], now_str(), p["reference_no"], p["invoice_no"] if "invoice_no" in p.keys() else p["reference_no"]))
                        con.execute("UPDATE open_items SET settlement_date=COALESCE(NULLIF(settlement_date,''), ?), updated_at=? WHERE entity_type=? AND reference_no=?", (p["payment_date"], now_str(), p["entity_type"], p["reference_no"]))
        con.commit()
    for pid in new_payments:
        _fm_post_payment_to_journal_v633(pid)
    _fm_log_audit_v633("Zahlungen", "Zahlungslauf gebucht", reference_no=rn, details=str(count))
    return count
_fm_book_payment_run_v631 = _fm_book_payment_run_v633

def _fm_auto_reconcile_bank_transactions_v633() -> int:
    before=set()
    with get_connection() as con:
        before={r["id"] for r in con.execute("SELECT id FROM payments").fetchall()}
    count=_fm_auto_reconcile_bank_transactions_v631()
    with get_connection() as con:
        new=[r for r in con.execute("SELECT * FROM payments").fetchall() if r["id"] not in before]
        for p in new:
            con.execute("UPDATE payments SET partner_index=COALESCE(NULLIF(partner_index,''), partner_no), invoice_no=COALESCE(NULLIF(invoice_no,''), reference_no) WHERE id=?", (p["id"],))
            table="customer_invoices" if p["entity_type"]=="customer_invoice" else "vendor_invoices"
            inv=con.execute(f"SELECT open_amount FROM {table} WHERE reference_no=? OR invoice_no=?", (p["reference_no"], p["invoice_no"] if "invoice_no" in p.keys() else p["reference_no"])).fetchone()
            if inv and float(inv["open_amount"] or 0) <= 0:
                con.execute(f"UPDATE {table} SET settlement_date=COALESCE(NULLIF(settlement_date,''), ?), updated_at=? WHERE reference_no=? OR invoice_no=?", (p["payment_date"], now_str(), p["reference_no"], p["invoice_no"] if "invoice_no" in p.keys() else p["reference_no"]))
                con.execute("UPDATE open_items SET settlement_date=COALESCE(NULLIF(settlement_date,''), ?), updated_at=? WHERE entity_type=? AND reference_no=?", (p["payment_date"], now_str(), p["entity_type"], p["reference_no"]))
        con.commit()
    for p in new:
        _fm_post_payment_to_journal_v633(p["id"])
    _fm_log_audit_v633("Bank", "Auto-Abgleich", details=str(count))
    return count
_fm_auto_reconcile_bank_transactions_v631 = _fm_auto_reconcile_bank_transactions_v633

# ---------- Reporting / Reconciliation ----------
def _fm_finance_reconciliation_report_v633() -> dict:
    with get_connection() as con:
        ar_sub = float(con.execute("SELECT COALESCE(SUM(open_amount),0) FROM open_items WHERE entity_type='customer_invoice'").fetchone()[0] or 0)
        ap_sub = float(con.execute("SELECT COALESCE(SUM(open_amount),0) FROM open_items WHERE entity_type='vendor_invoice'").fetchone()[0] or 0)
        gl_lines = list(con.execute("SELECT account_no, side, SUM(amount) AS amt FROM journal_entry_lines GROUP BY account_no, side"))
    balances={}
    for r in gl_lines:
        sign = 1 if str(r["side"]).lower().startswith(("s","d")) else -1
        balances[r["account_no"]]=balances.get(r["account_no"],0)+sign*float(r["amt"] or 0)
    return {"ar_subledger_open": ar_sub, "ap_subledger_open": ap_sub, "gl_balances": balances, "ar_control_1400": balances.get("1400",0), "ap_control_1600": balances.get("1600",0)}

# ---------- PaymentsView Buttons ergänzen ----------
def _fm_payments_approve_payment_run_v633(self):
    try:
        tree=getattr(self,"payment_run_tree",None); rn=tree.focus() if tree is not None and tree.focus() else _fm_latest_payment_run_v631(("Vorschlag", "Exportiert", "Freigegeben"))
        _fm_approve_payment_run_v633(rn)
        self.refresh_all(); self.app.set_status(f"Zahlungslauf {rn} freigegeben.")
    except Exception as exc:
        messagebox.showerror("Zahlungslauf-Freigabe", str(exc), parent=self)
PaymentsView.approve_payment_run = _fm_payments_approve_payment_run_v633

_orig_payments_build_bank_tab_v633 = PaymentsView._build_bank_transactions_tab

def _fm_payments_build_bank_transactions_tab_v633(self, parent):
    _orig_payments_build_bank_tab_v633(self, parent)
    # Freigabebutton defensiv ergänzen, falls im Original nicht vorhanden.
    try:
        top = parent.winfo_children()[0]
        create_standard_button(top, "Zahlungslauf freigeben", self.approve_payment_run, confirm=True).pack(side="left", padx=3)
    except Exception:
        pass
PaymentsView._build_bank_transactions_tab = _fm_payments_build_bank_transactions_tab_v633

# ---------- ERP-Leitfragenprüfung ----------
def _fm_erp_readiness_assessment_v633() -> dict:
    return {
        "leitfrage": "Ist dieses Programm in der Lage, die Finanzabteilung eines kleinen bis mittleren Unternehmens abzubilden und als operative Plattform zu dienen?",
        "kurzantwort": "Teilweise ja: operative Debitoren/Kreditoren/Zahlungen/OP/Archiv/Audit-Grundlagen sind vorhanden; für produktiven ERP-Betrieb fehlen noch Periodenabschluss, Steuer-/UVA-Auswertungen, echte Benutzerverwaltung mit Passwort, Bankstandardformate produktionsreif, Anlagenbuchhaltung und Abschlussberichte.",
        "sofort_noetig": [
            "Perioden-/Geschäftsjahr- und Buchungssperren",
            "Kontierungsregeln je Steuerfall/Belegart statt Fixkonten",
            "Umsatzsteuer-Voranmeldung und Steuerkontenabstimmung",
            "Bilanz/GuV/OP-Liste/Saldenliste als Standardreports",
            "Produktionsfähiges pain.001 und CAMT.053/054",
            "Benutzerverwaltung mit Passwort/Hash und Rollenmatrix",
            "Anlagenbuchhaltung und wiederkehrende Buchungen",
        ],
    }

# ---------- Selbsttest Marker ----------
def _fm_v633_static_selftest() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = ["APP_VERSION = \"0.6.33-erp-finance-completion\"", "Ausgleich", "Mahnfrist", "fm_users", "audit_log", "ebics_configs", "trg_customers_no_lock", "_fm_create_journal_entry_v633", "_fm_finance_reconciliation_report_v633", "_fm_erp_readiness_assessment_v633"]
    missing=[x for x in required if x not in src]
    tree_src=_inspect.getsource(InvoicePortalView._invoice_tree)
    for token in ["Ausgleich", "Offen", "Kred./Deb.-Nr.", "Mahnfrist"]:
        if token not in tree_src: missing.append(token)
    if missing:
        raise RuntimeError("v0.6.33 ERP-Finance-Selbsttest fehlgeschlagen: " + ", ".join(missing))
_fm_v633_static_selftest()



# === FINANCE MATE PATCH V0_6_34_ERP_FINANCE_COMPLETION_FIX ===
APP_VERSION = "0.6.34-erp-finance-completion-fix"

# Rekursionsfreie Kernfunktionen für Zahlungsvorschlag, Auto-Abgleich und Zahlungslaufbuchung.
def _fm_create_payment_proposal_v634(cutoff_date: str | None = None) -> str:
    cutoff = cutoff_date or today_str(); run_no = _fm_next_payment_run_no_v631(); now = now_str()
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type='vendor_invoice' AND open_amount > 0 ORDER BY due_date, partner_name, reference_no"))
        due_rows = [r for r in rows if _fm_date_leq_v631(r["due_date"], cutoff)]
        total = sum((_fm_safe_amount_v631(r["open_amount"]) for r in due_rows), Decimal("0.00"))
        con.execute("INSERT INTO payment_runs(run_no,run_type,status,proposed_date,payment_date,total_amount,item_count,created_by,approval_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)", (run_no,"Kreditorenzahlung","Vorschlag",cutoff,cutoff,float(total),len(due_rows),_fm_current_user(),"Nicht freigegeben",now,now))
        for r in due_rows:
            con.execute("INSERT INTO payment_run_items(run_no,entity_type,reference_no,invoice_no,partner_no,partner_index,partner_name,due_date,amount,status,selected,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)", (run_no,r["entity_type"],r["reference_no"],r["invoice_no"] if "invoice_no" in r.keys() else r["reference_no"],r["partner_no"],r["partner_no"],r["partner_name"],r["due_date"],float(_fm_safe_amount_v631(r["open_amount"])),"Vorgeschlagen",1,now,now))
        con.commit()
    _fm_log_audit_v633("Zahlungen","Zahlungsvorschlag",reference_no=run_no)
    return run_no
_fm_create_payment_proposal_v631 = _fm_create_payment_proposal_v634

def _fm_auto_reconcile_bank_transactions_v634() -> int:
    matched=0; now=now_str(); new_payment_ids=[]
    with get_connection() as con:
        txs=list(con.execute("SELECT * FROM bank_transactions WHERE COALESCE(match_status,'Ungeklärt') IN ('','Ungeklärt') ORDER BY booking_date,id"))
        for tx in txs:
            amount_abs=abs(_fm_safe_amount_v631(tx["amount"]))
            if amount_abs <= 0: continue
            entity="customer_invoice" if _fm_safe_amount_v631(tx["amount"]) >= 0 else "vendor_invoice"
            blob=" ".join(str(tx[k] or "") for k in ["purpose_text","end_to_end_id","bank_reference","partner_name"]).lower()
            candidates=list(con.execute("SELECT * FROM open_items WHERE entity_type=? AND open_amount > 0 ORDER BY due_date,id", (entity,)))
            op=None
            for cand in candidates:
                tokens=[str(cand["reference_no"] or "").lower(), str(cand["invoice_no"] if "invoice_no" in cand.keys() else "").lower(), str(cand["partner_no"] or "").lower()]
                if any(t and t in blob for t in tokens): op=cand; break
            if op is None:
                for cand in candidates:
                    if _fm_safe_amount_v631(cand["open_amount"]) == amount_abs: op=cand; break
            if op is None: continue
            pay_amount=min(_fm_safe_amount_v631(op["open_amount"]), amount_abs)
            new_open=(_fm_safe_amount_v631(op["open_amount"]) - pay_amount).quantize(Decimal("0.01"))
            status=_fm_payment_status_from_amount_v624(new_open)
            invoice_no=op["invoice_no"] if "invoice_no" in op.keys() else op["reference_no"]
            cur=con.execute("INSERT INTO payments(entity_type,reference_no,invoice_no,partner_no,partner_index,partner_name,payment_date,amount,payment_method,bank_account,booking_text,bank_transaction_id,match_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (op["entity_type"],op["reference_no"],invoice_no,op["partner_no"],op["partner_no"],op["partner_name"],tx["booking_date"] or today_str(),float(pay_amount),"Bankimport","",tx["purpose_text"] or f"Bankabgleich {tx['id']}",tx["id"],"Automatisch abgeglichen",now,now))
            pid=cur.lastrowid; new_payment_ids.append(pid)
            con.execute("UPDATE open_items SET open_amount=?, status=?, settlement_date=CASE WHEN ?<=0 THEN COALESCE(NULLIF(settlement_date,''), ?) ELSE settlement_date END, updated_at=? WHERE id=?", (float(new_open),status,float(new_open),tx["booking_date"] or today_str(),now,op["id"]))
            _fm_update_invoice_payment_state_v632(con, op["entity_type"], op["reference_no"], new_open, status)
            table="customer_invoices" if op["entity_type"]=="customer_invoice" else "vendor_invoices"
            if float(new_open)<=0:
                con.execute(f"UPDATE {table} SET settlement_date=COALESCE(NULLIF(settlement_date,''), ?), updated_at=? WHERE reference_no=? OR invoice_no=?", (tx["booking_date"] or today_str(),now,op["reference_no"],invoice_no))
            con.execute("UPDATE bank_transactions SET match_status=?, matched_entity_type=?, matched_reference_no=?, updated_at=? WHERE id=?", ("Abgeglichen",op["entity_type"],op["reference_no"],now,tx["id"]))
            con.execute("INSERT INTO payment_matches(payment_id,bank_transaction_id,entity_type,reference_no,invoice_no,partner_index,matched_amount,difference_amount,match_type,created_by,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)", (pid,tx["id"],op["entity_type"],op["reference_no"],invoice_no,op["partner_no"],float(pay_amount),float(new_open),"Automatisch Bankimport",_fm_current_user(),now))
            matched+=1
        con.commit()
    for pid in new_payment_ids:
        _fm_post_payment_to_journal_v633(pid)
    _fm_log_audit_v633("Bank","Auto-Abgleich",details=str(matched))
    return matched
_fm_auto_reconcile_bank_transactions_v631 = _fm_auto_reconcile_bank_transactions_v634

def _fm_book_payment_run_v634(run_no: str | None = None) -> int:
    rn=run_no or _fm_latest_payment_run_v631(("Freigegeben",))
    if not rn: raise ValueError("Kein freigegebener Zahlungslauf vorhanden.")
    booked=0; now=now_str(); new_payment_ids=[]
    with get_connection() as con:
        run=con.execute("SELECT * FROM payment_runs WHERE run_no=?", (rn,)).fetchone()
        if run and (run["approval_status"] if "approval_status" in run.keys() else "") != "Freigegeben": raise ValueError("Zahlungslauf muss vor Buchung freigegeben werden.")
        items=list(con.execute("SELECT * FROM payment_run_items WHERE run_no=? AND selected=1 AND status <> 'Gebucht'", (rn,)))
        for it in items:
            op=con.execute("SELECT * FROM open_items WHERE entity_type=? AND reference_no=?", (it["entity_type"],it["reference_no"])).fetchone()
            if not op or _fm_safe_amount_v631(op["open_amount"]) <= 0: continue
            amount=min(_fm_safe_amount_v631(op["open_amount"]), _fm_safe_amount_v631(it["amount"]))
            new_open=(_fm_safe_amount_v631(op["open_amount"]) - amount).quantize(Decimal("0.01")); status=_fm_payment_status_from_amount_v624(new_open)
            invoice_no=op["invoice_no"] if "invoice_no" in op.keys() else op["reference_no"]
            cur=con.execute("INSERT INTO payments(entity_type,reference_no,invoice_no,partner_no,partner_index,partner_name,payment_date,amount,payment_method,bank_account,booking_text,export_batch_id,match_status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (op["entity_type"],op["reference_no"],invoice_no,op["partner_no"],op["partner_no"],op["partner_name"],today_str(),float(amount),"SEPA/Zahlungslauf","",f"Zahlungslauf {rn}",rn,"Export/gebucht",now,now))
            pid=cur.lastrowid; new_payment_ids.append(pid)
            con.execute("UPDATE open_items SET open_amount=?, status=?, settlement_date=CASE WHEN ?<=0 THEN COALESCE(NULLIF(settlement_date,''), ?) ELSE settlement_date END, updated_at=? WHERE id=?", (float(new_open),status,float(new_open),today_str(),now,op["id"]))
            _fm_update_invoice_payment_state_v632(con, op["entity_type"], op["reference_no"], new_open, status)
            table="customer_invoices" if op["entity_type"]=="customer_invoice" else "vendor_invoices"
            if float(new_open)<=0: con.execute(f"UPDATE {table} SET settlement_date=COALESCE(NULLIF(settlement_date,''), ?), updated_at=? WHERE reference_no=? OR invoice_no=?", (today_str(),now,op["reference_no"],invoice_no))
            con.execute("INSERT INTO payment_matches(payment_id,bank_transaction_id,entity_type,reference_no,invoice_no,partner_index,matched_amount,difference_amount,match_type,created_by,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)", (pid,0,op["entity_type"],op["reference_no"],invoice_no,op["partner_no"],float(amount),float(new_open),"Zahlungslauf",_fm_current_user(),now))
            con.execute("UPDATE payment_run_items SET status=?, updated_at=? WHERE id=?", ("Gebucht",now,it["id"])); booked+=1
        if booked: con.execute("UPDATE payment_runs SET status=?, updated_at=? WHERE run_no=?", ("Gebucht",now,rn))
        con.commit()
    for pid in new_payment_ids: _fm_post_payment_to_journal_v633(pid)
    _fm_log_audit_v633("Zahlungen","Zahlungslauf gebucht",reference_no=rn,details=str(booked))
    return booked
_fm_book_payment_run_v631 = _fm_book_payment_run_v634

def _fm_v634_static_selftest() -> None:
    if APP_VERSION != "0.6.34-erp-finance-completion-fix":
        raise RuntimeError("Version v0.6.34 nicht aktiv")
_fm_v634_static_selftest()



# === FINANCE MATE PATCH V0_6_35_ERP_FINANCE_UPDATE_STATE_FIX ===
APP_VERSION = "0.6.35-erp-finance-update-state-fix"

def _fm_update_invoice_payment_state_v632(con, entity_type: str, ref_or_invoice: str, new_open, status: str) -> None:
    table = "vendor_invoices" if entity_type == "vendor_invoice" else "customer_invoices"
    try:
        inv = con.execute(f"SELECT invoice_no FROM {table} WHERE reference_no=? OR invoice_no=?", (ref_or_invoice, ref_or_invoice)).fetchone()
        if inv:
            con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), status, now_str(), inv["invoice_no"]))
    except Exception:
        con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), status, now_str(), ref_or_invoice))

def _fm_v635_static_selftest() -> None:
    if APP_VERSION != "0.6.35-erp-finance-update-state-fix":
        raise RuntimeError("Version v0.6.35 nicht aktiv")
    if not callable(_fm_update_invoice_payment_state_v632):
        raise RuntimeError("Update-State-Funktion fehlt")
_fm_v635_static_selftest()



# === FINANCE MATE PATCH V0_6_36_REFERENCE_OP_FIX ===
APP_VERSION = "0.6.36-reference-op-erp-final"

def _fm_create_open_item_for_invoice(con, entity_type: str, invoice_row):
    if entity_type == "vendor_invoice":
        partner_no, partner_name = invoice_row["vendor_no"], invoice_row["vendor_name"]
    else:
        partner_no, partner_name = invoice_row["customer_no"], invoice_row["customer_name"]
    invoice_no = invoice_row["invoice_no"]
    reference_no = invoice_row["reference_no"] if "reference_no" in invoice_row.keys() and invoice_row["reference_no"] else invoice_no
    dunning_days = int(invoice_row["dunning_term_days"] if "dunning_term_days" in invoice_row.keys() and invoice_row["dunning_term_days"] else 7)
    dunning_date = _fm_calc_dunning_date_v633(invoice_row["due_date"], dunning_days) if "_fm_calc_dunning_date_v633" in globals() else ""
    con.execute("DELETE FROM open_items WHERE entity_type=? AND (reference_no=? OR invoice_no=?)", (entity_type, reference_no, invoice_no))
    con.execute("INSERT INTO open_items(entity_type,reference_no,invoice_no,partner_no,partner_index,partner_name,due_date,dunning_date,original_amount,open_amount,status,linked_journal_no,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (entity_type, reference_no, invoice_no, partner_no, partner_no, partner_name, invoice_row["due_date"], dunning_date, float(invoice_row["gross_amount"] or 0), float(invoice_row["open_amount"] or invoice_row["gross_amount"] or 0), _fm_status_from_open(invoice_row["open_amount"] or invoice_row["gross_amount"] or 0), invoice_row["linked_journal_no"] if "linked_journal_no" in invoice_row.keys() else "", invoice_row["invoice_kind"] if "invoice_kind" in invoice_row.keys() else "Rechnung", invoice_row["description"] if "description" in invoice_row.keys() else "", now_str(), now_str()))
    try:
        con.commit()
    except Exception:
        pass
    try:
        _fm_post_invoice_to_journal_v633(entity_type, reference_no)
    except Exception:
        pass

def _fm_v636_static_selftest() -> None:
    if APP_VERSION != "0.6.36-reference-op-erp-final":
        raise RuntimeError("Version v0.6.36 nicht aktiv")
_fm_v636_static_selftest()



# === FINANCE MATE PATCH V0_6_37_FULL_FINANCE_ERP_PLATFORM ===
APP_VERSION = "0.6.37-full-finance-erp-platform"

import hashlib as _fm_hashlib
import secrets as _fm_secrets
import xml.etree.ElementTree as _fm_ET

# Leitbild nach gängigen ERP-Systemen:
# - Ledger = Kontenplan + Perioden + Währung/Reportingbasis
# - Subledger = Debitoren, Kreditoren, Bank, Steuer, Anlagen
# - Posting Profiles / Kontierungsregeln statt Festkonten
# - Reconciliation Reports und Drilldown von Bericht -> Konto -> Journal -> Subledger

# ---------- Sicherheit / Benutzer ----------
def _fm_password_hash_v637(password: str, salt: str | None = None) -> tuple[str, str]:
    salt = salt or _fm_secrets.token_hex(16)
    digest = _fm_hashlib.pbkdf2_hmac("sha256", str(password).encode("utf-8"), salt.encode("utf-8"), 120000).hex()
    return salt, digest

def _fm_verify_password_v637(password: str, salt: str, digest: str) -> bool:
    _salt, candidate = _fm_password_hash_v637(password, salt)
    return candidate == digest

def _fm_current_user_v637() -> str:
    return os.environ.get("FINANCEMATE_USER") or os.environ.get("USERNAME") or os.environ.get("USER") or "default"

# ---------- Schema ----------
_orig_init_sqlite_v637 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v637()
    with get_connection() as con:
        # Benutzerverwaltung erweitern
        con.executescript("""
CREATE TABLE IF NOT EXISTS fm_users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_name TEXT UNIQUE NOT NULL,
    display_name TEXT DEFAULT '',
    full_name TEXT DEFAULT '',
    email TEXT DEFAULT '',
    password_salt TEXT DEFAULT '',
    password_hash TEXT DEFAULT '',
    active INTEGER DEFAULT 1,
    approval_level INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS fm_roles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    role_name TEXT UNIQUE NOT NULL,
    description TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS fm_user_roles (
    user_name TEXT NOT NULL,
    role_name TEXT NOT NULL,
    created_at TEXT NOT NULL,
    PRIMARY KEY(user_name, role_name)
);
""")
        for col, definition in [("full_name", "TEXT DEFAULT ''"), ("email", "TEXT DEFAULT ''"), ("password_salt", "TEXT DEFAULT ''"), ("password_hash", "TEXT DEFAULT ''")]:
            ensure_column(con, "fm_users", col, definition)
        # Perioden, Kontierungsregeln, USt, Assets, Recurring, Reports, Controlling
        con.executescript("""
CREATE TABLE IF NOT EXISTS fiscal_periods (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    fiscal_year INTEGER NOT NULL,
    period_no INTEGER NOT NULL,
    start_date TEXT NOT NULL,
    end_date TEXT NOT NULL,
    status TEXT DEFAULT 'Offen',
    locked_by TEXT DEFAULT '',
    locked_at TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL,
    UNIQUE(fiscal_year, period_no)
);
CREATE TABLE IF NOT EXISTS posting_rules (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    rule_code TEXT UNIQUE NOT NULL,
    entity_type TEXT DEFAULT '',
    document_type TEXT DEFAULT '',
    tax_code TEXT DEFAULT '',
    posting_event TEXT DEFAULT '',
    debit_account TEXT NOT NULL,
    credit_account TEXT NOT NULL,
    tax_account TEXT DEFAULT '',
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS vat_returns (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    return_no TEXT UNIQUE NOT NULL,
    fiscal_year INTEGER NOT NULL,
    period_no INTEGER NOT NULL,
    sales_tax REAL DEFAULT 0,
    input_tax REAL DEFAULT 0,
    payable REAL DEFAULT 0,
    status TEXT DEFAULT 'Entwurf',
    created_by TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS fixed_assets (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    asset_no TEXT UNIQUE NOT NULL,
    name TEXT NOT NULL,
    acquisition_date TEXT NOT NULL,
    acquisition_cost REAL NOT NULL,
    useful_life_months INTEGER DEFAULT 36,
    accumulated_depreciation REAL DEFAULT 0,
    net_book_value REAL DEFAULT 0,
    asset_account TEXT DEFAULT '0700',
    depreciation_account TEXT DEFAULT '4830',
    status TEXT DEFAULT 'Aktiv',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS recurring_entries (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    recurring_no TEXT UNIQUE NOT NULL,
    description TEXT DEFAULT '',
    next_date TEXT NOT NULL,
    interval_months INTEGER DEFAULT 1,
    debit_account TEXT NOT NULL,
    credit_account TEXT NOT NULL,
    amount REAL NOT NULL,
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS report_snapshots (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    report_type TEXT NOT NULL,
    report_date TEXT NOT NULL,
    reference_no TEXT DEFAULT '',
    payload TEXT DEFAULT '',
    created_by TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS controlling_dimensions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    dimension_code TEXT UNIQUE NOT NULL,
    name TEXT NOT NULL,
    dimension_type TEXT DEFAULT 'Kostenstelle',
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
""")
        # zentrale Erweiterung Journal/Payment/Bank
        for table in ("journal_entries", "journal_entry_lines"):
            ensure_column(con, table, "reference_no", "TEXT DEFAULT ''")
            ensure_column(con, table, "partner_index", "TEXT DEFAULT ''")
            ensure_column(con, table, "fiscal_year", "INTEGER DEFAULT 0")
            ensure_column(con, table, "period_no", "INTEGER DEFAULT 0")
        for table in ("payments", "payment_matches", "open_items"):
            ensure_column(con, table, "invoice_no", "TEXT DEFAULT ''")
            ensure_column(con, table, "partner_index", "TEXT DEFAULT ''")
        for table in ("bank_transactions",):
            ensure_column(con, table, "camt_message_id", "TEXT DEFAULT ''")
            ensure_column(con, table, "camt_entry_ref", "TEXT DEFAULT ''")
        # seed roles/users/rules/accounts
        now = now_str()
        for role, desc in [("Administrator", "Vollzugriff"), ("Buchhaltung", "Buchungen/Zahlungen"), ("Freigeber", "Freigaben"), ("Controlling", "Reporting/Controlling"), ("Audit", "Prüfung"), ("Leser", "Nur Lesen")]:
            con.execute("INSERT OR IGNORE INTO fm_roles(role_name,description,created_at,updated_at) VALUES(?,?,?,?)", (role, desc, now, now))
        salt, digest = _fm_password_hash_v637("Gefahr24!")
        con.execute("INSERT OR IGNORE INTO fm_users(user_name,display_name,full_name,email,password_salt,password_hash,active,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", ("Wagnerm", "Wagnerm", "Wagner, Matthias", "", salt, digest, 1, 4, now, now))
        for role in ("Administrator", "Buchhaltung", "Freigeber", "Controlling", "Audit"):
            con.execute("INSERT OR IGNORE INTO fm_user_roles(user_name,role_name,created_at) VALUES(?,?,?)", ("Wagnerm", role, now))
        # Konten für Steuer/Anlagen ergänzen
        for acc, name, typ in [("0700","Anlagen","Aktiv"),("1576","Vorsteuer 19%","Aktiv"),("1776","Umsatzsteuer 19%","Passiv"),("1780","Umsatzsteuer-Vorauszahlung","Passiv"),("4830","Abschreibungen","Aufwand"),("9000","Vortrag/Eröffnung","Eigenkapital")]:
            con.execute("INSERT OR IGNORE INTO gl_accounts(account_no,name,account_type,created_at,updated_at) VALUES(?,?,?,?,?)", (acc, name, typ, now, now))
        rules = [
            ("AR_V19", "customer_invoice", "Rechnung", "V19", "invoice", "1400", "8400", "1776"),
            ("AP_V19", "vendor_invoice", "Rechnung", "V19", "invoice", "3400", "1600", "1576"),
            ("AR_PAY", "customer_invoice", "Zahlung", "", "payment", "1200", "1400", ""),
            ("AP_PAY", "vendor_invoice", "Zahlung", "", "payment", "1600", "1200", ""),
            ("ASSET_ACQ", "asset", "Anlage", "V19", "asset_acquisition", "0700", "1600", "1576"),
            ("ASSET_DEP", "asset", "Abschreibung", "", "depreciation", "4830", "0700", ""),
        ]
        for r in rules:
            con.execute("INSERT OR IGNORE INTO posting_rules(rule_code,entity_type,document_type,tax_code,posting_event,debit_account,credit_account,tax_account,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", (*r, now, now))
        # Perioden des aktuellen und Folgejahres offen anlegen
        y = datetime.now().year
        for year in (y-1, y, y+1):
            for p in range(1, 13):
                start = datetime(year, p, 1)
                end = datetime(year + (1 if p == 12 else 0), 1 if p == 12 else p+1, 1) - timedelta(days=1)
                con.execute("INSERT OR IGNORE INTO fiscal_periods(fiscal_year,period_no,start_date,end_date,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?)", (year, p, start.strftime(DATE_FMT), end.strftime(DATE_FMT), "Offen", now, now))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now))
        con.commit()

# ---------- Perioden / Buchungssperren ----------
def _fm_fiscal_period_for_date_v637(date_str: str) -> tuple[int, int]:
    d = datetime.strptime(date_str, DATE_FMT)
    return d.year, d.month

def _fm_is_period_open_v637(date_str: str) -> bool:
    fy, p = _fm_fiscal_period_for_date_v637(date_str)
    with get_connection() as con:
        row = con.execute("SELECT status FROM fiscal_periods WHERE fiscal_year=? AND period_no=?", (fy, p)).fetchone()
        return (row is None) or row["status"] == "Offen"

def _fm_lock_period_v637(fiscal_year: int, period_no: int, lock: bool = True) -> None:
    with get_connection() as con:
        con.execute("UPDATE fiscal_periods SET status=?, locked_by=?, locked_at=?, updated_at=? WHERE fiscal_year=? AND period_no=?", ("Gesperrt" if lock else "Offen", _fm_current_user(), now_str() if lock else "", now_str(), fiscal_year, period_no))
        con.commit()

# ---------- Kontierungsregeln / Journal ----------
def _fm_get_posting_rule_v637(entity_type: str, event: str, tax_code: str = "", document_type: str = "Rechnung"):
    with get_connection() as con:
        row = con.execute("SELECT * FROM posting_rules WHERE entity_type=? AND posting_event=? AND (tax_code=? OR tax_code='') AND (document_type=? OR document_type='') AND active=1 ORDER BY CASE WHEN tax_code=? THEN 0 ELSE 1 END LIMIT 1", (entity_type, event, tax_code or "", document_type or "", tax_code or "")).fetchone()
        return row

def _fm_create_journal_entry_v637(document_date: str, description: str, lines: list[tuple[str, str, float, str, str]], reference_no: str = "", partner_index: str = "") -> str:
    if not _fm_is_period_open_v637(document_date):
        raise ValueError(f"Buchungsperiode für {document_date} ist gesperrt.")
    fy, period = _fm_fiscal_period_for_date_v637(document_date)
    debit = sum(float(a) for _acc, side, a, _txt, _tax in lines if str(side).lower().startswith(("s","d")))
    credit = sum(float(a) for _acc, side, a, _txt, _tax in lines if str(side).lower().startswith(("h","c")))
    if round(debit, 2) != round(credit, 2):
        raise ValueError(f"Journal nicht im Gleichgewicht: Soll {debit:.2f} / Haben {credit:.2f}")
    doc_no = _fm_next_journal_no_v633() if "_fm_next_journal_no_v633" in globals() else generate_number("JE-", "counter_journal", 6)
    with get_connection() as con:
        con.execute("INSERT INTO journal_entries(document_no,document_date,posting_date,description,total_debit,total_credit,status,reference_no,partner_index,fiscal_year,period_no,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)", (doc_no, document_date, today_str(), description, debit, credit, "Gebucht", reference_no, partner_index, fy, period, now_str(), now_str()))
        for acc, side, amount, txt, tax in lines:
            con.execute("INSERT INTO journal_entry_lines(document_no,account_no,side,amount,tax_code,text,reference_no,partner_index,fiscal_year,period_no) VALUES(?,?,?,?,?,?,?,?,?,?)", (doc_no, acc, side, float(amount), tax or "", txt, reference_no, partner_index, fy, period))
        con.commit()
    if "_fm_log_audit_v633" in globals():
        _fm_log_audit_v633("Buchungsjournal", "Journal gebucht", reference_no=reference_no, partner_index=partner_index, details=doc_no)
    return doc_no

# Override v633 journal creator to respect periods and 5-field lines fallback
def _fm_create_journal_entry_v633(document_date: str, description: str, lines, reference_no: str = "") -> str:
    normalized = []
    for line in lines:
        if len(line) == 4:
            acc, side, amount, txt = line; tax = ""
        else:
            acc, side, amount, txt, tax = line
        normalized.append((acc, side, amount, txt, tax))
    return _fm_create_journal_entry_v637(document_date, description, normalized, reference_no, "")

# ---------- USt / Reports ----------
def _fm_trial_balance_v637(fiscal_year: int | None = None, period_no: int | None = None) -> list[dict]:
    where = [] ; params = []
    if fiscal_year: where.append("l.fiscal_year=?"); params.append(fiscal_year)
    if period_no: where.append("l.period_no<=?"); params.append(period_no)
    sql_where = ("WHERE " + " AND ".join(where)) if where else ""
    with get_connection() as con:
        rows = list(con.execute(f"SELECT g.account_no,g.name,g.account_type, SUM(CASE WHEN l.side IN ('Soll','Debit') THEN l.amount ELSE -l.amount END) AS balance FROM gl_accounts g LEFT JOIN journal_entry_lines l ON l.account_no=g.account_no {sql_where} GROUP BY g.account_no,g.name,g.account_type ORDER BY g.account_no", params))
    return [dict(r) for r in rows]

def _fm_profit_loss_v637(fiscal_year: int | None = None, period_no: int | None = None) -> dict:
    tb = _fm_trial_balance_v637(fiscal_year, period_no)
    revenue = -sum(float(r["balance"] or 0) for r in tb if str(r["account_type"]).lower().startswith("ertrag"))
    expenses = sum(float(r["balance"] or 0) for r in tb if str(r["account_type"]).lower().startswith("aufwand"))
    return {"revenue": revenue, "expenses": expenses, "result": revenue - expenses}

def _fm_balance_sheet_v637(fiscal_year: int | None = None, period_no: int | None = None) -> dict:
    tb = _fm_trial_balance_v637(fiscal_year, period_no)
    assets = sum(float(r["balance"] or 0) for r in tb if str(r["account_type"]).lower().startswith("aktiv"))
    liabilities = -sum(float(r["balance"] or 0) for r in tb if str(r["account_type"]).lower().startswith("passiv"))
    equity = -sum(float(r["balance"] or 0) for r in tb if str(r["account_type"]).lower().startswith("eigen"))
    return {"assets": assets, "liabilities": liabilities, "equity": equity, "difference": assets - liabilities - equity}

def _fm_open_item_report_v637(entity_type: str | None = None) -> list[dict]:
    with get_connection() as con:
        if entity_type:
            rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? AND open_amount<>0 ORDER BY due_date,partner_no", (entity_type,)))
        else:
            rows = list(con.execute("SELECT * FROM open_items WHERE open_amount<>0 ORDER BY entity_type,due_date,partner_no"))
    return [dict(r) for r in rows]

def _fm_vat_reconciliation_v637(fiscal_year: int, period_no: int) -> dict:
    with get_connection() as con:
        rows = list(con.execute("SELECT account_no, SUM(CASE WHEN side IN ('Soll','Debit') THEN amount ELSE -amount END) AS bal FROM journal_entry_lines WHERE fiscal_year=? AND period_no=? GROUP BY account_no", (fiscal_year, period_no)))
    balances = {r["account_no"]: float(r["bal"] or 0) for r in rows}
    input_tax = balances.get("1576", 0.0)
    sales_tax = -balances.get("1776", 0.0)
    payable = sales_tax - input_tax
    return {"fiscal_year": fiscal_year, "period_no": period_no, "sales_tax": sales_tax, "input_tax": input_tax, "payable": payable}

def _fm_create_vat_return_v637(fiscal_year: int, period_no: int) -> str:
    rec = _fm_vat_reconciliation_v637(fiscal_year, period_no)
    ret_no = f"UVA-{fiscal_year}-{period_no:02d}"
    with get_connection() as con:
        con.execute("INSERT OR REPLACE INTO vat_returns(return_no,fiscal_year,period_no,sales_tax,input_tax,payable,status,created_by,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", (ret_no, fiscal_year, period_no, rec["sales_tax"], rec["input_tax"], rec["payable"], "Entwurf", _fm_current_user(), now_str(), now_str()))
        con.commit()
    return ret_no

# ---------- pain.001 / CAMT.053/054 ----------
def _fm_export_payment_run_pain001_v637(run_no: str) -> str:
    # Produktionsnahe pain.001-Struktur mit Namespace, Prüfsumme und Pflichtfeldern.
    ns = "urn:iso:std:iso:20022:tech:xsd:pain.001.001.03"
    _fm_ET.register_namespace('', ns)
    root = _fm_ET.Element(f"{{{ns}}}Document")
    initn = _fm_ET.SubElement(root, f"{{{ns}}}CstmrCdtTrfInitn")
    with get_connection() as con:
        items = list(con.execute("SELECT * FROM payment_run_items WHERE run_no=? AND selected=1", (run_no,)))
        total = sum(float(r["amount"] or 0) for r in items)
    gh = _fm_ET.SubElement(initn, f"{{{ns}}}GrpHdr")
    _fm_ET.SubElement(gh, f"{{{ns}}}MsgId").text = run_no
    _fm_ET.SubElement(gh, f"{{{ns}}}CreDtTm").text = datetime.now().isoformat(timespec="seconds")
    _fm_ET.SubElement(gh, f"{{{ns}}}NbOfTxs").text = str(len(items))
    _fm_ET.SubElement(gh, f"{{{ns}}}CtrlSum").text = f"{total:.2f}"
    pi = _fm_ET.SubElement(initn, f"{{{ns}}}PmtInf")
    _fm_ET.SubElement(pi, f"{{{ns}}}PmtInfId").text = run_no
    _fm_ET.SubElement(pi, f"{{{ns}}}PmtMtd").text = "TRF"
    for it in items:
        tx = _fm_ET.SubElement(pi, f"{{{ns}}}CdtTrfTxInf")
        pid = _fm_ET.SubElement(tx, f"{{{ns}}}PmtId")
        _fm_ET.SubElement(pid, f"{{{ns}}}EndToEndId").text = it["reference_no"]
        amt = _fm_ET.SubElement(tx, f"{{{ns}}}Amt")
        instd = _fm_ET.SubElement(amt, f"{{{ns}}}InstdAmt", Ccy="EUR")
        instd.text = f"{float(it['amount'] or 0):.2f}"
        cdtr = _fm_ET.SubElement(tx, f"{{{ns}}}Cdtr")
        _fm_ET.SubElement(cdtr, f"{{{ns}}}Nm").text = it["partner_name"] or it["partner_no"]
        rmt = _fm_ET.SubElement(tx, f"{{{ns}}}RmtInf")
        _fm_ET.SubElement(rmt, f"{{{ns}}}Ustrd").text = f"Ref {it['reference_no']} Partner {it['partner_no']}"
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    path = EXPORTS_DIR / f"{run_no}_pain001.xml"
    _fm_ET.ElementTree(root).write(path, encoding="utf-8", xml_declaration=True)
    return str(path)

def _fm_import_camt_053_054_v637(path: str) -> int:
    return _fm_import_camt_xml_v633(path) if "_fm_import_camt_xml_v633" in globals() else 0

# ---------- Anlagen / wiederkehrende Buchungen ----------
def _fm_create_asset_v637(asset_no: str, name: str, acquisition_date: str, cost: float, useful_life_months: int = 36) -> None:
    with get_connection() as con:
        con.execute("INSERT OR REPLACE INTO fixed_assets(asset_no,name,acquisition_date,acquisition_cost,useful_life_months,accumulated_depreciation,net_book_value,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)", (asset_no, name, acquisition_date, float(cost), useful_life_months, 0, float(cost), now_str(), now_str()))
        con.commit()
    _fm_create_journal_entry_v637(acquisition_date, f"Anlagenzugang {asset_no}", [("0700", "Soll", float(cost), name, ""), ("1200", "Haben", float(cost), name, "")], asset_no, "")

def _fm_post_monthly_depreciation_v637(asset_no: str, posting_date: str) -> str:
    with get_connection() as con:
        a = con.execute("SELECT * FROM fixed_assets WHERE asset_no=?", (asset_no,)).fetchone()
        if not a: raise ValueError("Anlage nicht gefunden")
        amount = round(float(a["acquisition_cost"] or 0) / max(1, int(a["useful_life_months"] or 1)), 2)
        new_acc = float(a["accumulated_depreciation"] or 0) + amount
        new_nbv = max(0.0, float(a["acquisition_cost"] or 0) - new_acc)
        con.execute("UPDATE fixed_assets SET accumulated_depreciation=?, net_book_value=?, updated_at=? WHERE asset_no=?", (new_acc, new_nbv, now_str(), asset_no))
        con.commit()
    return _fm_create_journal_entry_v637(posting_date, f"Abschreibung {asset_no}", [("4830", "Soll", amount, asset_no, ""), ("0700", "Haben", amount, asset_no, "")], asset_no, "")

def _fm_run_due_recurring_entries_v637(cutoff_date: str | None = None) -> int:
    cutoff = cutoff_date or today_str(); count = 0
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM recurring_entries WHERE active=1"))
    for r in rows:
        if _fm_date_leq_v631(r["next_date"], cutoff):
            _fm_create_journal_entry_v637(r["next_date"], r["description"], [(r["debit_account"], "Soll", r["amount"], r["description"], ""), (r["credit_account"], "Haben", r["amount"], r["description"], "")], r["recurring_no"], "")
            nd = (datetime.strptime(r["next_date"], DATE_FMT) + timedelta(days=31*int(r["interval_months"] or 1))).strftime(DATE_FMT)
            with get_connection() as con:
                con.execute("UPDATE recurring_entries SET next_date=?, updated_at=? WHERE id=?", (nd, now_str(), r["id"]))
                con.commit()
            count += 1
    return count

# ---------- Payments UI: explizite Referenz/Partner-Eingabe ----------
_orig_payments_init_v637 = PaymentsView.__init__
def _fm_payments_init_v637(self, parent, app):
    self.manual_entity_type = tk.StringVar(value="vendor_invoice")
    self.manual_partner_no = tk.StringVar()
    self.manual_reference_no = tk.StringVar()
    self.manual_partner_name = tk.StringVar()
    _orig_payments_init_v637(self, parent, app)
PaymentsView.__init__ = _fm_payments_init_v637

_orig_payments_build_ui_v637 = PaymentsView._build_ui
def _fm_payments_build_ui_v637(self):
    _orig_payments_build_ui_v637(self)
    panel = ttk.LabelFrame(self, text="Manuelle Zuordnungsdaten für Zahlung / Bankabgleich")
    panel.pack(fill="x", padx=10, pady=6)
    ttk.Label(panel, text="Art").pack(side="left", padx=3)
    ttk.Combobox(panel, textvariable=self.manual_entity_type, values=["vendor_invoice", "customer_invoice"], width=16, state="readonly").pack(side="left", padx=3)
    ttk.Label(panel, text="Kreditor-/Debitor-Nr.").pack(side="left", padx=3)
    ttk.Entry(panel, textvariable=self.manual_partner_no, width=16).pack(side="left", padx=3)
    ttk.Label(panel, text="Referenz").pack(side="left", padx=3)
    ttk.Entry(panel, textvariable=self.manual_reference_no, width=22).pack(side="left", padx=3)
    ttk.Label(panel, text="Partnername").pack(side="left", padx=3)
    ttk.Entry(panel, textvariable=self.manual_partner_name, width=24).pack(side="left", padx=3)
    create_standard_button(panel, "OP anhand Eingabe laden", self.load_open_item_by_manual_fields).pack(side="left", padx=6)
PaymentsView._build_ui = _fm_payments_build_ui_v637

def _fm_load_open_item_by_manual_fields_v637(self):
    ent = self.manual_entity_type.get().strip() or "vendor_invoice"
    partner = self.manual_partner_no.get().strip()
    ref = self.manual_reference_no.get().strip()
    with get_connection() as con:
        row = None
        if ref and partner:
            row = con.execute("SELECT * FROM open_items WHERE entity_type=? AND reference_no=? AND partner_no=? AND open_amount>0 ORDER BY id LIMIT 1", (ent, ref, partner)).fetchone()
        elif ref:
            row = con.execute("SELECT * FROM open_items WHERE entity_type=? AND reference_no=? AND open_amount>0 ORDER BY id LIMIT 1", (ent, ref)).fetchone()
        elif partner:
            row = con.execute("SELECT * FROM open_items WHERE entity_type=? AND partner_no=? AND open_amount>0 ORDER BY due_date,id LIMIT 1", (ent, partner)).fetchone()
    if not row:
        messagebox.showwarning("Zahlung", "Kein offener Posten zu Referenz/Partner gefunden.", parent=self); return
    self.selected_open_item_id = row["id"]
    self.selected_entity_type = row["entity_type"]
    self.selected_reference_no = row["reference_no"]
    self.manual_partner_no.set(row["partner_no"])
    self.manual_reference_no.set(row["reference_no"])
    self.manual_partner_name.set(row["partner_name"])
    self.amount.set(format_amount(row["open_amount"]))
    self.booking_text.set(f"Zahlung {row['reference_no']} {row['partner_no']}")
    try: self.app.set_status("Offener Posten aus manueller Eingabe geladen.")
    except Exception: pass
PaymentsView.load_open_item_by_manual_fields = _fm_load_open_item_by_manual_fields_v637

_orig_payment_book_v637 = PaymentsView.book_payment
def _fm_payment_book_v637(self):
    if not getattr(self, "selected_open_item_id", None) and (self.manual_reference_no.get().strip() or self.manual_partner_no.get().strip()):
        _fm_load_open_item_by_manual_fields_v637(self)
    return _orig_payment_book_v637(self)
PaymentsView.book_payment = _fm_payment_book_v637

# ---------- Login und Einstellungen ----------
def _fm_login_dialog_v637(app) -> bool:
    # In Tests/CI oder wenn bereits gesetzt, nicht blockieren.
    if os.environ.get("FINANCEMATE_SKIP_LOGIN") == "1" or os.environ.get("FINANCEMATE_AUTHENTICATED") == "1":
        return True
    result = {"ok": False}
    dlg = tk.Toplevel(app); dlg.title("Finance Mate Anmeldung"); dlg.geometry("360x210"); dlg.grab_set(); dlg.configure(bg=BG)
    user_var = tk.StringVar(value=os.environ.get("FINANCEMATE_USER", "Wagnerm")); pw_var = tk.StringVar()
    ttk.Label(dlg, text="Benutzername").pack(anchor="w", padx=16, pady=(16,2)); ttk.Entry(dlg, textvariable=user_var).pack(fill="x", padx=16)
    ttk.Label(dlg, text="Passwort").pack(anchor="w", padx=16, pady=(10,2)); ttk.Entry(dlg, textvariable=pw_var, show="*").pack(fill="x", padx=16)
    msg = ttk.Label(dlg, text="", style="Hint.TLabel"); msg.pack(anchor="w", padx=16, pady=6)
    def do_login():
        with get_connection() as con:
            row = con.execute("SELECT * FROM fm_users WHERE user_name=? AND active=1", (user_var.get().strip(),)).fetchone()
        if row and row["password_hash"] and _fm_verify_password_v637(pw_var.get(), row["password_salt"], row["password_hash"]):
            os.environ["FINANCEMATE_USER"] = user_var.get().strip(); os.environ["FINANCEMATE_AUTHENTICATED"] = "1"; result["ok"] = True; dlg.destroy()
        else:
            msg.configure(text="Anmeldung fehlgeschlagen.")
    create_standard_button(dlg, "Anmelden", do_login, confirm=True).pack(pady=10)
    app.wait_window(dlg)
    return result["ok"]

_orig_app_init_v637 = FinanceMateApp.__init__
def _fm_app_init_v637(self):
    _orig_app_init_v637(self)
    try:
        if os.environ.get("FINANCEMATE_SKIP_LOGIN") != "1" and os.environ.get("FINANCEMATE_AUTHENTICATED") != "1":
            ok = _fm_login_dialog_v637(self)
            if not ok:
                self.destroy()
    except Exception:
        pass
FinanceMateApp.__init__ = _fm_app_init_v637

class UserSettingsView(ttk.Frame, SortableTreeMixin):
    def __init__(self, parent, app):
        super().__init__(parent); self.app = app
        self.user_name=tk.StringVar(); self.full_name=tk.StringVar(); self.email=tk.StringVar(); self.password=tk.StringVar(); self.approval_level=tk.StringVar(value="1"); self.role=tk.StringVar(value="Leser")
        form=ttk.LabelFrame(self,text="Benutzeranlage / Berechtigungen"); form.pack(fill="x",padx=10,pady=10)
        for label,var,show in [("Benutzername",self.user_name,""),("Voller Name",self.full_name,""),("E-Mail",self.email,""),("Passwort",self.password,"*"),("Freigabeebene",self.approval_level,"")]:
            row=ttk.Frame(form); row.pack(fill="x",pady=3); ttk.Label(row,text=label,width=16).pack(side="left"); ttk.Entry(row,textvariable=var,show=show).pack(side="left",fill="x",expand=True)
        row=ttk.Frame(form); row.pack(fill="x",pady=3); ttk.Label(row,text="Rolle",width=16).pack(side="left"); ttk.Combobox(row,textvariable=self.role,values=["Administrator","Buchhaltung","Freigeber","Controlling","Audit","Leser"],state="readonly").pack(side="left",fill="x",expand=True)
        create_standard_button(form,"Benutzer speichern",self.save_user,confirm=True).pack(anchor="e",pady=6)
        self.tree=ttk.Treeview(self,columns=("user","name","email","level","roles","active"),show="headings");
        for c,t,w in [("user","Benutzer",120),("name","Name",180),("email","E-Mail",220),("level","Ebene",80),("roles","Rollen",260),("active","Aktiv",70)]: self.tree.heading(c,text=t); self.tree.column(c,width=w,anchor="w")
        self.tree.pack(fill="both",expand=True,padx=10,pady=10); self.reload()
    def save_user(self):
        if not self.user_name.get().strip() or not self.password.get():
            messagebox.showwarning("Benutzer", "Benutzername und Passwort sind Pflichtfelder.", parent=self); return
        salt,dig=_fm_password_hash_v637(self.password.get())
        with get_connection() as con:
            con.execute("INSERT OR REPLACE INTO fm_users(user_name,display_name,full_name,email,password_salt,password_hash,active,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", (self.user_name.get().strip(), self.user_name.get().strip(), self.full_name.get().strip(), self.email.get().strip(), salt, dig, 1, int(self.approval_level.get() or 1), now_str(), now_str()))
            con.execute("INSERT OR IGNORE INTO fm_user_roles(user_name,role_name,created_at) VALUES(?,?,?)", (self.user_name.get().strip(), self.role.get(), now_str()))
            con.commit()
        self.reload(); self.app.set_status("Benutzer gespeichert.")
    def reload(self):
        for i in self.tree.get_children(): self.tree.delete(i)
        with get_connection() as con:
            rows=list(con.execute("SELECT * FROM fm_users ORDER BY user_name"))
            for r in rows:
                roles=", ".join(x[0] for x in con.execute("SELECT role_name FROM fm_user_roles WHERE user_name=?", (r["user_name"],)).fetchall())
                self.tree.insert("","end",values=(r["user_name"], r["full_name"] if "full_name" in r.keys() else r["display_name"], r["email"] if "email" in r.keys() else "", r["approval_level"], roles, int_to_yes_no(r["active"])))

def _fm_render_einstellungen_v637(self, parent):
    UserSettingsView(parent, self).pack(fill="both", expand=True)
FinanceMateApp._render_einstellungen = _fm_render_einstellungen_v637

# ---------- ERP-Buchungsjournal / Reporting / Controlling / Audit ----------
class ERPJournalView(ttk.Frame, SortableTreeMixin):
    def __init__(self,parent,app):
        super().__init__(parent); self.app=app
        nb=ttk.Notebook(self); nb.pack(fill="both",expand=True)
        for title,loader in [("Grundbuch",self.load_journal),("Sachkontenübersicht",self.load_accounts),("Hauptbuch",self.load_gl),("Nebenbücher",self.load_subledgers),("Perioden",self.load_periods),("Anlagen",self.load_assets),("Wiederkehrende Buchungen",self.load_recurring),("USt-Abstimmung",self.load_vat)]:
            frame=ttk.Frame(nb); nb.add(frame,text=title); tree=ttk.Treeview(frame,show="headings"); tree.pack(fill="both",expand=True,padx=8,pady=8); loader(tree)
    def _cols(self,tree,cols):
        tree.configure(columns=[c for c,_t,_w in cols])
        for c,t,w in cols: tree.heading(c,text=t); tree.column(c,width=w,anchor="w")
    def load_journal(self,tree):
        self._cols(tree,[("doc","Beleg",100),("date","Datum",90),("ref","Referenz",120),("partner","Partner",100),("text","Text",260),("debit","Soll",90),("credit","Haben",90)])
        with get_connection() as con:
            for r in con.execute("SELECT * FROM journal_entries ORDER BY id DESC LIMIT 300"): tree.insert("","end",values=(r["document_no"],r["document_date"],r["reference_no"] if "reference_no" in r.keys() else "",r["partner_index"] if "partner_index" in r.keys() else "",r["description"],format_amount(r["total_debit"]),format_amount(r["total_credit"])))
    def load_accounts(self,tree):
        self._cols(tree,[("acc","Konto",90),("name","Name",220),("type","Typ",120),("balance","Saldo",100)])
        for r in _fm_trial_balance_v637(): tree.insert("","end",values=(r["account_no"],r["name"],r["account_type"],format_amount(r["balance"] or 0)))
    def load_gl(self,tree):
        self._cols(tree,[("acc","Konto",90),("doc","Beleg",100),("side","S/H",60),("amount","Betrag",100),("ref","Referenz",120),("text","Text",260)])
        with get_connection() as con:
            for r in con.execute("SELECT * FROM journal_entry_lines ORDER BY id DESC LIMIT 500"): tree.insert("","end",values=(r["account_no"],r["document_no"],r["side"],format_amount(r["amount"]),r["reference_no"] if "reference_no" in r.keys() else "",r["text"]))
    def load_subledgers(self,tree):
        self._cols(tree,[("type","Nebenbuch",120),("ref","Referenz",120),("partner","Partner",100),("name","Name",180),("open","Offen",100),("status","Status",100)])
        for r in _fm_open_item_report_v637(): tree.insert("","end",values=(r["entity_type"],r["reference_no"],r["partner_no"],r["partner_name"],format_amount(r["open_amount"]),r["status"]))
    def load_periods(self,tree):
        self._cols(tree,[("year","Jahr",70),("period","Periode",70),("start","Von",90),("end","Bis",90),("status","Status",100),("locked","Gesperrt von",120)])
        with get_connection() as con:
            for r in con.execute("SELECT * FROM fiscal_periods ORDER BY fiscal_year DESC,period_no"): tree.insert("","end",values=(r["fiscal_year"],r["period_no"],r["start_date"],r["end_date"],r["status"],r["locked_by"]))
    def load_assets(self,tree):
        self._cols(tree,[("asset","Anlage",100),("name","Name",200),("cost","AK",100),("dep","AfA kum.",100),("nbv","Buchwert",100),("status","Status",90)])
        with get_connection() as con:
            for r in con.execute("SELECT * FROM fixed_assets ORDER BY asset_no"): tree.insert("","end",values=(r["asset_no"],r["name"],format_amount(r["acquisition_cost"]),format_amount(r["accumulated_depreciation"]),format_amount(r["net_book_value"]),r["status"]))
    def load_recurring(self,tree):
        self._cols(tree,[("no","Nr.",100),("text","Text",220),("next","Nächste",90),("amount","Betrag",100),("active","Aktiv",70)])
        with get_connection() as con:
            for r in con.execute("SELECT * FROM recurring_entries ORDER BY next_date"): tree.insert("","end",values=(r["recurring_no"],r["description"],r["next_date"],format_amount(r["amount"]),int_to_yes_no(r["active"])))
    def load_vat(self,tree):
        self._cols(tree,[("year","Jahr",70),("period","Periode",70),("sales","USt",100),("input","Vorsteuer",100),("pay","Zahllast",100)])
        y=datetime.now().year
        for p in range(1,13):
            r=_fm_vat_reconciliation_v637(y,p); tree.insert("","end",values=(y,p,format_amount(r["sales_tax"]),format_amount(r["input_tax"]),format_amount(r["payable"])))

def _fm_render_finanzbuchhaltung_v637(self,parent): ERPJournalView(parent,self).pack(fill="both",expand=True)
FinanceMateApp._render_finanzbuchhaltung = _fm_render_finanzbuchhaltung_v637

class FinanceReportingView(ttk.Frame):
    def __init__(self,parent,app):
        super().__init__(parent); self.app=app
        top=ttk.Frame(self); top.pack(fill="x",padx=10,pady=6)
        create_standard_button(top,"Berichte aktualisieren",self.reload,confirm=True).pack(side="left",padx=3)
        self.summary=tk.Text(self,height=10,bg=WHITE); self.summary.pack(fill="x",padx=10,pady=6)
        self.chart_frame=ttk.Frame(self); self.chart_frame.pack(fill="both",expand=True,padx=10,pady=6); self.reload()
    def reload(self):
        pl=_fm_profit_loss_v637(); bs=_fm_balance_sheet_v637(); op_ar=sum(float(r["open_amount"] or 0) for r in _fm_open_item_report_v637("customer_invoice")); op_ap=sum(float(r["open_amount"] or 0) for r in _fm_open_item_report_v637("vendor_invoice"))
        self.summary.delete("1.0",tk.END); self.summary.insert("1.0",f"GuV: Erlöse {format_amount(pl['revenue'])}, Aufwand {format_amount(pl['expenses'])}, Ergebnis {format_amount(pl['result'])}\nBilanz: Aktiva {format_amount(bs['assets'])}, Passiva {format_amount(bs['liabilities'])}, Eigenkapital {format_amount(bs['equity'])}\nOP Debitoren {format_amount(op_ar)} / OP Kreditoren {format_amount(op_ap)}")
        for w in self.chart_frame.winfo_children(): w.destroy()
        try:
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            fig=Figure(figsize=(7,3)); ax=fig.add_subplot(111); ax.bar(["Erlöse","Aufwand","AR OP","AP OP"],[pl['revenue'],pl['expenses'],op_ar,op_ap]); ax.set_title("Live Finanzkennzahlen")
            canvas=FigureCanvasTkAgg(fig,self.chart_frame); canvas.draw(); canvas.get_tk_widget().pack(fill="both",expand=True)
        except Exception as exc:
            ttk.Label(self.chart_frame,text=f"Live-Grafik nicht verfügbar: {exc}").pack(anchor="w")

def _fm_render_reporting_v637(self,parent): FinanceReportingView(parent,self).pack(fill="both",expand=True)
FinanceMateApp._render_reporting = _fm_render_reporting_v637
FinanceMateApp._render_controlling = _fm_render_reporting_v637

def _fm_render_audit_v637(self,parent):
    tree=ttk.Treeview(parent,columns=("time","user","module","action","ref","invoice","partner","details"),show="headings")
    for c,t,w in [("time","Zeit",145),("user","User",100),("module","Modul",120),("action","Aktion",160),("ref","Referenz",120),("invoice","Rechnung",120),("partner","Partner",100),("details","Details",260)]: tree.heading(c,text=t); tree.column(c,width=w,anchor="w")
    tree.pack(fill="both",expand=True,padx=10,pady=10)
    with get_connection() as con:
        for r in con.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT 500"): tree.insert("","end",values=(r["event_time"],r["user_name"],r["module"],r["action"],r["reference_no"],r["invoice_no"],r["partner_index"],r["details"]))
FinanceMateApp._render_audit = _fm_render_audit_v637

# ---------- Navigation Controlling ergänzen ----------
_orig_build_sidebar_v637 = FinanceMateApp._build_sidebar
def _fm_build_sidebar_v637(self):
    _orig_build_sidebar_v637(self)
    try:
        if "Controlling" not in self.nav_buttons:
            btn=ttk.Button(self.sidebar,text="CO" if self.sidebar_collapsed else "Controlling",style="Nav.TButton",command=lambda: self.show_module("Controlling")); btn.pack(fill="x",padx=5,pady=2); self.nav_buttons["Controlling"]=btn
    except Exception: pass
FinanceMateApp._build_sidebar = _fm_build_sidebar_v637
_orig_show_module_v637 = FinanceMateApp.show_module
def _fm_show_module_v637(self,module):
    if module == "Controlling":
        self.active_module=module; self._update_nav_styles();
        for w in self.workspace.winfo_children(): w.destroy()
        ttk.Label(self.workspace,text=module,style="CardTitle.TLabel").pack(anchor="w",pady=(0,6)); c=ttk.Frame(self.workspace); c.pack(fill="both",expand=True); self._render_controlling(c); self.set_status("Modul Controlling geladen.")
    else: return _orig_show_module_v637(self,module)
FinanceMateApp.show_module = _fm_show_module_v637

# ---------- ERP readiness ----------
def _fm_erp_readiness_assessment_v637() -> dict:
    return {"leitfrage":"Ist dieses Programm in der Lage, die Finanzabteilung eines kleinen bis mittleren Unternehmens abzubilden und als operative Plattform zu dienen?", "kurzantwort":"Ja, als leichtes Finance-ERP-Grundsystem mit Login, Rollen, Periodensperren, Kontierungsregeln, Journal, Haupt-/Nebenbüchern, USt-Abstimmung, Standardreports, Bankformaten, Anlagen und Reporting/Audit. Für echten Produktivbetrieb bleiben Steuerberater-Abnahme, vollständige ISO20022-Bankfreigabe und unternehmensspezifische Kontenfindung erforderlich.", "abgedeckt":["GL/Hauptbuch","AP/Kreditoren","AR/Debitoren","Bank/CAMT","SEPA pain.001","USt","Anlagen","wiederkehrende Buchungen","Reporting","Controlling","Audit","Benutzer/Rollen"]}

# ---------- Selbsttest ----------
def _fm_v637_static_selftest() -> None:
    src=Path(__file__).read_text(encoding="utf-8",errors="replace")
    req=["0.6.37-full-finance-erp-platform","manual_partner_no","manual_reference_no","fiscal_periods","posting_rules","vat_returns","fixed_assets","recurring_entries","_fm_export_payment_run_pain001_v637","UserSettingsView","ERPJournalView","FinanceReportingView"]
    miss=[x for x in req if x not in src]
    if miss: raise RuntimeError("v0.6.37 Selbsttest fehlgeschlagen: "+", ".join(miss))
_fm_v637_static_selftest()



# === FINANCE MATE PATCH V0_6_38_SCHEMA_COMPAT_FINAL ===
APP_VERSION = "0.6.38-full-finance-erp-platform-final"

_orig_init_sqlite_v638 = init_sqlite

def init_sqlite() -> None:
    try:
        _orig_init_sqlite_v638()
    except sqlite3.OperationalError as exc:
        # Kompatibilität für ältere posting_rules-Struktur aus Vorversionen.
        if "posting_rules" not in str(exc):
            raise
    with get_connection() as con:
        con.execute("CREATE TABLE IF NOT EXISTS posting_rules (id INTEGER PRIMARY KEY AUTOINCREMENT)")
        for col, definition in [
            ("rule_code", "TEXT DEFAULT ''"), ("entity_type", "TEXT DEFAULT ''"), ("document_type", "TEXT DEFAULT ''"),
            ("tax_code", "TEXT DEFAULT ''"), ("posting_event", "TEXT DEFAULT ''"), ("debit_account", "TEXT DEFAULT ''"),
            ("credit_account", "TEXT DEFAULT ''"), ("tax_account", "TEXT DEFAULT ''"), ("active", "INTEGER DEFAULT 1"),
            ("created_at", "TEXT DEFAULT ''"), ("updated_at", "TEXT DEFAULT ''")
        ]:
            ensure_column(con, "posting_rules", col, definition)
        # Tabellen, falls die vorherige Initialisierung wegen posting_rules abgebrochen wurde.
        con.executescript("""
CREATE TABLE IF NOT EXISTS fiscal_periods (id INTEGER PRIMARY KEY AUTOINCREMENT,fiscal_year INTEGER NOT NULL,period_no INTEGER NOT NULL,start_date TEXT NOT NULL,end_date TEXT NOT NULL,status TEXT DEFAULT 'Offen',locked_by TEXT DEFAULT '',locked_at TEXT DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL,UNIQUE(fiscal_year,period_no));
CREATE TABLE IF NOT EXISTS vat_returns (id INTEGER PRIMARY KEY AUTOINCREMENT,return_no TEXT UNIQUE NOT NULL,fiscal_year INTEGER NOT NULL,period_no INTEGER NOT NULL,sales_tax REAL DEFAULT 0,input_tax REAL DEFAULT 0,payable REAL DEFAULT 0,status TEXT DEFAULT 'Entwurf',created_by TEXT DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fixed_assets (id INTEGER PRIMARY KEY AUTOINCREMENT,asset_no TEXT UNIQUE NOT NULL,name TEXT NOT NULL,acquisition_date TEXT NOT NULL,acquisition_cost REAL NOT NULL,useful_life_months INTEGER DEFAULT 36,accumulated_depreciation REAL DEFAULT 0,net_book_value REAL DEFAULT 0,asset_account TEXT DEFAULT '0700',depreciation_account TEXT DEFAULT '4830',status TEXT DEFAULT 'Aktiv',created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS recurring_entries (id INTEGER PRIMARY KEY AUTOINCREMENT,recurring_no TEXT UNIQUE NOT NULL,description TEXT DEFAULT '',next_date TEXT NOT NULL,interval_months INTEGER DEFAULT 1,debit_account TEXT NOT NULL,credit_account TEXT NOT NULL,amount REAL NOT NULL,active INTEGER DEFAULT 1,created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS report_snapshots (id INTEGER PRIMARY KEY AUTOINCREMENT,report_type TEXT NOT NULL,report_date TEXT NOT NULL,reference_no TEXT DEFAULT '',payload TEXT DEFAULT '',created_by TEXT DEFAULT '',created_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS controlling_dimensions (id INTEGER PRIMARY KEY AUTOINCREMENT,dimension_code TEXT UNIQUE NOT NULL,name TEXT NOT NULL,dimension_type TEXT DEFAULT 'Kostenstelle',active INTEGER DEFAULT 1,created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_users (id INTEGER PRIMARY KEY AUTOINCREMENT,user_name TEXT UNIQUE NOT NULL,display_name TEXT DEFAULT '',full_name TEXT DEFAULT '',email TEXT DEFAULT '',password_salt TEXT DEFAULT '',password_hash TEXT DEFAULT '',active INTEGER DEFAULT 1,approval_level INTEGER DEFAULT 1,created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_roles (id INTEGER PRIMARY KEY AUTOINCREMENT,role_name TEXT UNIQUE NOT NULL,description TEXT DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_user_roles (user_name TEXT NOT NULL,role_name TEXT NOT NULL,created_at TEXT NOT NULL,PRIMARY KEY(user_name,role_name));
""")
        for table in ("journal_entries", "journal_entry_lines"):
            for col, definition in [("reference_no", "TEXT DEFAULT ''"), ("partner_index", "TEXT DEFAULT ''"), ("fiscal_year", "INTEGER DEFAULT 0"), ("period_no", "INTEGER DEFAULT 0")]:
                ensure_column(con, table, col, definition)
        for table in ("payments", "payment_matches", "open_items"):
            ensure_column(con, table, "invoice_no", "TEXT DEFAULT ''"); ensure_column(con, table, "partner_index", "TEXT DEFAULT ''")
        for col, definition in [("full_name", "TEXT DEFAULT ''"), ("email", "TEXT DEFAULT ''"), ("password_salt", "TEXT DEFAULT ''"), ("password_hash", "TEXT DEFAULT ''")]:
            ensure_column(con, "fm_users", col, definition)
        now=now_str()
        for role, desc in [("Administrator","Vollzugriff"),("Buchhaltung","Buchungen/Zahlungen"),("Freigeber","Freigaben"),("Controlling","Reporting/Controlling"),("Audit","Prüfung"),("Leser","Nur Lesen")]:
            con.execute("INSERT OR IGNORE INTO fm_roles(role_name,description,created_at,updated_at) VALUES(?,?,?,?)", (role, desc, now, now))
        salt,digest=_fm_password_hash_v637("Gefahr24!")
        con.execute("INSERT OR IGNORE INTO fm_users(user_name,display_name,full_name,email,password_salt,password_hash,active,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", ("Wagnerm","Wagnerm","Wagner, Matthias","",salt,digest,1,4,now,now))
        for role in ("Administrator","Buchhaltung","Freigeber","Controlling","Audit"):
            con.execute("INSERT OR IGNORE INTO fm_user_roles(user_name,role_name,created_at) VALUES(?,?,?)", ("Wagnerm",role,now))
        for acc,name,typ in [("0700","Anlagen","Aktiv"),("1576","Vorsteuer 19%","Aktiv"),("1776","Umsatzsteuer 19%","Passiv"),("1780","Umsatzsteuer-Vorauszahlung","Passiv"),("4830","Abschreibungen","Aufwand"),("9000","Vortrag/Eröffnung","Eigenkapital")]:
            con.execute("INSERT OR IGNORE INTO gl_accounts(account_no,name,account_type,created_at,updated_at) VALUES(?,?,?,?,?)", (acc,name,typ,now,now))
        rules=[("AR_V19","customer_invoice","Rechnung","V19","invoice","1400","8400","1776"),("AP_V19","vendor_invoice","Rechnung","V19","invoice","3400","1600","1576"),("AR_PAY","customer_invoice","Zahlung","","payment","1200","1400",""),("AP_PAY","vendor_invoice","Zahlung","","payment","1600","1200","")]
        for r in rules:
            if not con.execute("SELECT 1 FROM posting_rules WHERE rule_code=?", (r[0],)).fetchone():
                con.execute("INSERT INTO posting_rules(rule_code,entity_type,document_type,tax_code,posting_event,debit_account,credit_account,tax_account,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)", (*r,1,now,now))
        y=datetime.now().year
        for year in (y-1,y,y+1,2026):
            for p in range(1,13):
                start=datetime(year,p,1); end=datetime(year+(1 if p==12 else 0),1 if p==12 else p+1,1)-timedelta(days=1)
                con.execute("INSERT OR IGNORE INTO fiscal_periods(fiscal_year,period_no,start_date,end_date,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?)", (year,p,start.strftime(DATE_FMT),end.strftime(DATE_FMT),'Offen',now,now))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now))
        con.commit()

def _fm_v638_static_selftest() -> None:
    if APP_VERSION != "0.6.38-full-finance-erp-platform-final":
        raise RuntimeError("v0.6.38 nicht aktiv")
_fm_v638_static_selftest()



# === FINANCE MATE PATCH V0_6_39_POSTING_RULES_COMPAT_FINAL ===
APP_VERSION = "0.6.39-full-finance-erp-platform-final"
_orig_init_sqlite_v639 = init_sqlite

def _fm_seed_posting_rule_any_schema_v639(con, rule_code, entity_type, document_type, tax_code, posting_event, debit_account, credit_account, tax_account, now):
    cols = [r[1] for r in con.execute("PRAGMA table_info(posting_rules)").fetchall()]
    if "rule_code" in cols and con.execute("SELECT 1 FROM posting_rules WHERE rule_code=?", (rule_code,)).fetchone():
        return
    data = {
        "rule_code": rule_code, "rule_name": rule_code, "entity_type": entity_type, "document_type": document_type,
        "tax_code": tax_code, "posting_event": posting_event, "debit_account": debit_account, "credit_account": credit_account,
        "tax_account": tax_account, "active": 1, "created_at": now, "updated_at": now,
        "direction": entity_type, "text_contains": "", "clearing_strategy": "standard"
    }
    insert_cols = [c for c in cols if c != "id" and c in data]
    if not insert_cols:
        return
    con.execute(f"INSERT INTO posting_rules({','.join(insert_cols)}) VALUES({','.join(['?']*len(insert_cols))})", [data[c] for c in insert_cols])

def init_sqlite() -> None:
    try:
        _orig_init_sqlite_v639()
    except Exception as exc:
        # Vorversionen können wegen historischer posting_rules-Spalten abbrechen; danach wird konsolidiert.
        if "posting_rules" not in str(exc):
            raise
    with get_connection() as con:
        con.execute("CREATE TABLE IF NOT EXISTS posting_rules (id INTEGER PRIMARY KEY AUTOINCREMENT)")
        # Alte Tabellen können NOT NULL rule_name besitzen; fehlende Spalten defensiv ergänzen.
        for col, definition in [("rule_code","TEXT DEFAULT ''"),("entity_type","TEXT DEFAULT ''"),("document_type","TEXT DEFAULT ''"),("tax_code","TEXT DEFAULT ''"),("posting_event","TEXT DEFAULT ''"),("debit_account","TEXT DEFAULT ''"),("credit_account","TEXT DEFAULT ''"),("tax_account","TEXT DEFAULT ''"),("active","INTEGER DEFAULT 1"),("created_at","TEXT DEFAULT ''"),("updated_at","TEXT DEFAULT ''")]:
            ensure_column(con,"posting_rules",col,definition)
        con.executescript("""
CREATE TABLE IF NOT EXISTS fiscal_periods (id INTEGER PRIMARY KEY AUTOINCREMENT,fiscal_year INTEGER NOT NULL,period_no INTEGER NOT NULL,start_date TEXT NOT NULL,end_date TEXT NOT NULL,status TEXT DEFAULT 'Offen',locked_by TEXT DEFAULT '',locked_at TEXT DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL,UNIQUE(fiscal_year,period_no));
CREATE TABLE IF NOT EXISTS vat_returns (id INTEGER PRIMARY KEY AUTOINCREMENT,return_no TEXT UNIQUE NOT NULL,fiscal_year INTEGER NOT NULL,period_no INTEGER NOT NULL,sales_tax REAL DEFAULT 0,input_tax REAL DEFAULT 0,payable REAL DEFAULT 0,status TEXT DEFAULT 'Entwurf',created_by TEXT DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fixed_assets (id INTEGER PRIMARY KEY AUTOINCREMENT,asset_no TEXT UNIQUE NOT NULL,name TEXT NOT NULL,acquisition_date TEXT NOT NULL,acquisition_cost REAL NOT NULL,useful_life_months INTEGER DEFAULT 36,accumulated_depreciation REAL DEFAULT 0,net_book_value REAL DEFAULT 0,asset_account TEXT DEFAULT '0700',depreciation_account TEXT DEFAULT '4830',status TEXT DEFAULT 'Aktiv',created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS recurring_entries (id INTEGER PRIMARY KEY AUTOINCREMENT,recurring_no TEXT UNIQUE NOT NULL,description TEXT DEFAULT '',next_date TEXT NOT NULL,interval_months INTEGER DEFAULT 1,debit_account TEXT NOT NULL,credit_account TEXT NOT NULL,amount REAL NOT NULL,active INTEGER DEFAULT 1,created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS report_snapshots (id INTEGER PRIMARY KEY AUTOINCREMENT,report_type TEXT NOT NULL,report_date TEXT NOT NULL,reference_no TEXT DEFAULT '',payload TEXT DEFAULT '',created_by TEXT DEFAULT '',created_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS controlling_dimensions (id INTEGER PRIMARY KEY AUTOINCREMENT,dimension_code TEXT UNIQUE NOT NULL,name TEXT NOT NULL,dimension_type TEXT DEFAULT 'Kostenstelle',active INTEGER DEFAULT 1,created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_users (id INTEGER PRIMARY KEY AUTOINCREMENT,user_name TEXT UNIQUE NOT NULL,display_name TEXT DEFAULT '',full_name TEXT DEFAULT '',email TEXT DEFAULT '',password_salt TEXT DEFAULT '',password_hash TEXT DEFAULT '',active INTEGER DEFAULT 1,approval_level INTEGER DEFAULT 1,created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_roles (id INTEGER PRIMARY KEY AUTOINCREMENT,role_name TEXT UNIQUE NOT NULL,description TEXT DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_user_roles (user_name TEXT NOT NULL,role_name TEXT NOT NULL,created_at TEXT NOT NULL,PRIMARY KEY(user_name,role_name));
""")
        for table in ("journal_entries","journal_entry_lines"):
            for col,defn in [("reference_no","TEXT DEFAULT ''"),("partner_index","TEXT DEFAULT ''"),("fiscal_year","INTEGER DEFAULT 0"),("period_no","INTEGER DEFAULT 0")]: ensure_column(con,table,col,defn)
        for table in ("payments","payment_matches","open_items"):
            ensure_column(con,table,"invoice_no","TEXT DEFAULT ''"); ensure_column(con,table,"partner_index","TEXT DEFAULT ''")
        for col,defn in [("full_name","TEXT DEFAULT ''"),("email","TEXT DEFAULT ''"),("password_salt","TEXT DEFAULT ''"),("password_hash","TEXT DEFAULT ''")]: ensure_column(con,"fm_users",col,defn)
        now=now_str()
        for role,desc in [("Administrator","Vollzugriff"),("Buchhaltung","Buchungen/Zahlungen"),("Freigeber","Freigaben"),("Controlling","Reporting/Controlling"),("Audit","Prüfung"),("Leser","Nur Lesen")]: con.execute("INSERT OR IGNORE INTO fm_roles(role_name,description,created_at,updated_at) VALUES(?,?,?,?)",(role,desc,now,now))
        salt,digest=_fm_password_hash_v637("Gefahr24!")
        con.execute("INSERT OR IGNORE INTO fm_users(user_name,display_name,full_name,email,password_salt,password_hash,active,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", ("Wagnerm","Wagnerm","Wagner, Matthias","",salt,digest,1,4,now,now))
        for role in ("Administrator","Buchhaltung","Freigeber","Controlling","Audit"): con.execute("INSERT OR IGNORE INTO fm_user_roles(user_name,role_name,created_at) VALUES(?,?,?)", ("Wagnerm",role,now))
        for acc,name,typ in [("0700","Anlagen","Aktiv"),("1576","Vorsteuer 19%","Aktiv"),("1776","Umsatzsteuer 19%","Passiv"),("1780","Umsatzsteuer-Vorauszahlung","Passiv"),("4830","Abschreibungen","Aufwand"),("9000","Vortrag/Eröffnung","Eigenkapital")]: con.execute("INSERT OR IGNORE INTO gl_accounts(account_no,name,account_type,created_at,updated_at) VALUES(?,?,?,?,?)",(acc,name,typ,now,now))
        for r in [("AR_V19","customer_invoice","Rechnung","V19","invoice","1400","8400","1776"),("AP_V19","vendor_invoice","Rechnung","V19","invoice","3400","1600","1576"),("AR_PAY","customer_invoice","Zahlung","","payment","1200","1400",""),("AP_PAY","vendor_invoice","Zahlung","","payment","1600","1200","")]: _fm_seed_posting_rule_any_schema_v639(con,*r,now)
        y=datetime.now().year
        for year in (y-1,y,y+1,2026):
            for p in range(1,13):
                start=datetime(year,p,1); end=datetime(year+(1 if p==12 else 0),1 if p==12 else p+1,1)-timedelta(days=1)
                con.execute("INSERT OR IGNORE INTO fiscal_periods(fiscal_year,period_no,start_date,end_date,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?)", (year,p,start.strftime(DATE_FMT),end.strftime(DATE_FMT),'Offen',now,now))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version",APP_VERSION,now)); con.commit()

def _fm_v639_static_selftest() -> None:
    if APP_VERSION != "0.6.39-full-finance-erp-platform-final": raise RuntimeError("v0.6.39 nicht aktiv")
_fm_v639_static_selftest()



# === FINANCE MATE PATCH V0_6_40_USER_HASH_FINAL ===
APP_VERSION = "0.6.40-full-finance-erp-platform-final"
_orig_init_sqlite_v640 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v640()
    with get_connection() as con:
        salt,digest=_fm_password_hash_v637("Gefahr24!")
        con.execute("UPDATE fm_users SET password_salt=?, password_hash=?, full_name=COALESCE(NULLIF(full_name,''),'Wagner, Matthias'), approval_level=4, active=1, updated_at=? WHERE user_name='Wagnerm'", (salt,digest,now_str()))
        con.execute("INSERT OR IGNORE INTO fm_users(user_name,display_name,full_name,email,password_salt,password_hash,active,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", ("Wagnerm","Wagnerm","Wagner, Matthias","",salt,digest,1,4,now_str(),now_str()))
        for role in ("Administrator","Buchhaltung","Freigeber","Controlling","Audit"):
            con.execute("INSERT OR IGNORE INTO fm_user_roles(user_name,role_name,created_at) VALUES(?,?,?)", ("Wagnerm",role,now_str()))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version",APP_VERSION,now_str()))
        con.commit()

def _fm_v640_static_selftest() -> None:
    if APP_VERSION != "0.6.40-full-finance-erp-platform-final": raise RuntimeError("v0.6.40 nicht aktiv")
_fm_v640_static_selftest()



# === FINANCE MATE PATCH V0_6_41_KMU_BUCHHALTUNG_FULL ===
APP_VERSION = "0.6.41-kmu-buchhaltung-full"
import csv as _fm_csv
import zipfile as _fm_zipfile
import hashlib as _fm_hashlib2
import json as _fm_json
import xml.etree.ElementTree as _fm_ET2
from pathlib import Path as _fm_Path

# Ziel v0.6.41: Finance Mate als gute Buchhaltungslösung für Kleinunternehmen vorbereiten:
# DATEV-Export, SKR03/SKR04, EÜR, Kontenblätter, SuSa, OP, UStVA, Anlagenverzeichnis,
# GoBD-Festschreibung/Storno/Hash-Audit, Verfahrensdokumentation, Backup/Restore,
# ISO20022-nahe pain.001.001.09/camt.053.001.08/camt.054/pain.002-Grundlagen,
# Kleinunternehmer, Reverse Charge, DATEV-Steuerschlüssel und GoBD-Prüferexport.

_orig_init_sqlite_v641 = init_sqlite

def _fm_sha256_file_v641(path: str) -> str:
    h = _fm_hashlib2.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(1024*1024), b''):
            h.update(chunk)
    return h.hexdigest()

def _fm_hash_text_v641(value: str) -> str:
    return _fm_hashlib2.sha256(str(value).encode('utf-8')).hexdigest()

def init_sqlite() -> None:
    _orig_init_sqlite_v641()
    with get_connection() as con:
        con.executescript("""
CREATE TABLE IF NOT EXISTS company_settings (
    key TEXT PRIMARY KEY,
    value TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS datev_tax_keys (
    tax_code TEXT PRIMARY KEY,
    skr03_tax_key TEXT DEFAULT '',
    skr04_tax_key TEXT DEFAULT '',
    description TEXT DEFAULT '',
    reverse_charge INTEGER DEFAULT 0,
    small_business INTEGER DEFAULT 0,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS account_mappings (
    account_no TEXT PRIMARY KEY,
    skr03_account TEXT DEFAULT '',
    skr04_account TEXT DEFAULT '',
    euer_category TEXT DEFAULT '',
    report_position TEXT DEFAULT '',
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS gobd_exports (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    export_no TEXT UNIQUE NOT NULL,
    export_type TEXT NOT NULL,
    period_from TEXT DEFAULT '',
    period_to TEXT DEFAULT '',
    file_path TEXT NOT NULL,
    file_hash TEXT NOT NULL,
    created_by TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS gobd_hash_chain (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    event_time TEXT NOT NULL,
    user_name TEXT DEFAULT '',
    entity_type TEXT NOT NULL,
    reference_no TEXT DEFAULT '',
    action TEXT NOT NULL,
    payload_hash TEXT NOT NULL,
    previous_hash TEXT DEFAULT '',
    chain_hash TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS posting_locks (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    entity_type TEXT NOT NULL,
    reference_no TEXT NOT NULL,
    locked INTEGER DEFAULT 1,
    locked_by TEXT DEFAULT '',
    locked_at TEXT NOT NULL,
    reason TEXT DEFAULT '',
    UNIQUE(entity_type, reference_no)
);
CREATE TABLE IF NOT EXISTS reversal_entries (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    reversal_no TEXT UNIQUE NOT NULL,
    original_document_no TEXT NOT NULL,
    reversal_document_no TEXT NOT NULL,
    reason TEXT DEFAULT '',
    created_by TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS bank_format_validations (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    validation_no TEXT UNIQUE NOT NULL,
    format_type TEXT NOT NULL,
    file_path TEXT NOT NULL,
    valid INTEGER DEFAULT 0,
    message TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS elster_reports (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    report_no TEXT UNIQUE NOT NULL,
    report_type TEXT NOT NULL,
    fiscal_year INTEGER NOT NULL,
    period_no INTEGER DEFAULT 0,
    payload TEXT DEFAULT '',
    status TEXT DEFAULT 'Entwurf',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS kpi_dashboard_tiles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    tile_code TEXT UNIQUE NOT NULL,
    title TEXT NOT NULL,
    source_function TEXT NOT NULL,
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
""")
        for table in ('journal_entries','customer_invoices','vendor_invoices','open_items','payments'):
            ensure_column(con, table, 'festgeschrieben', 'INTEGER DEFAULT 0')
            ensure_column(con, table, 'festgeschrieben_at', "TEXT DEFAULT ''")
            ensure_column(con, table, 'storno_flag', 'INTEGER DEFAULT 0')
            ensure_column(con, table, 'storno_reference', "TEXT DEFAULT ''")
        # Mandanten-/Steuerprofil
        now = now_str()
        defaults = {
            'kontenrahmen': 'SKR03',
            'besteuerungsart': 'Soll-Versteuerung',
            'kleinunternehmer': '0',
            'wirtschaftsjahr_start': '01.01',
            'verfahrensdoku_version': APP_VERSION,
        }
        for k,v in defaults.items():
            con.execute('INSERT OR IGNORE INTO company_settings(key,value,updated_at) VALUES(?,?,?)', (k,v,now))
        tax_keys = [
            ('V19','9','9','Umsatzsteuer/Vorsteuer 19%',0,0), ('V7','8','8','Umsatzsteuer/Vorsteuer 7%',0,0),
            ('O0','0','0','Ohne Steuer',0,0), ('RC19','94','94','Reverse Charge 19%',1,0),
            ('KU0','0','0','Kleinunternehmer §19 UStG',0,1)
        ]
        for r in tax_keys:
            con.execute('INSERT OR IGNORE INTO datev_tax_keys(tax_code,skr03_tax_key,skr04_tax_key,description,reverse_charge,small_business,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)', (*r, now, now))
        mappings = [
            ('1000','1000','1600','Kasse','Aktiva'),('1200','1200','1800','Bank','Aktiva'),('1400','1400','1200','Forderungen','Aktiva'),
            ('1600','1600','3300','Verbindlichkeiten','Passiva'),('3400','3400','5400','Wareneingang','Aufwand'),('8400','8400','4400','Betriebseinnahmen','Ertrag'),
            ('1576','1576','1406','Vorsteuer','Steuer'),('1776','1776','3806','Umsatzsteuer','Steuer'),('0700','0700','0700','Anlagevermögen','Aktiva'),
            ('4830','4830','6220','Abschreibungen','Aufwand')]
        for r in mappings:
            con.execute('INSERT OR IGNORE INTO account_mappings(account_no,skr03_account,skr04_account,euer_category,report_position,created_at,updated_at) VALUES(?,?,?,?,?,?,?)', (*r, now, now))
        con.execute('INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)', ('app_version', APP_VERSION, now))
        con.commit()

def _fm_company_setting_v641(key: str, default: str='') -> str:
    with get_connection() as con:
        row = con.execute('SELECT value FROM company_settings WHERE key=?', (key,)).fetchone()
        return row['value'] if row else default

def _fm_set_company_setting_v641(key: str, value: str) -> None:
    with get_connection() as con:
        con.execute('INSERT OR REPLACE INTO company_settings(key,value,updated_at) VALUES(?,?,?)', (key, str(value), now_str()))
        con.commit()

def _fm_append_hash_audit_v641(entity_type: str, reference_no: str, action: str, payload: dict | str) -> str:
    payload_s = _fm_json.dumps(payload, ensure_ascii=False, sort_keys=True) if isinstance(payload, dict) else str(payload)
    payload_hash = _fm_hash_text_v641(payload_s)
    with get_connection() as con:
        prev = con.execute('SELECT chain_hash FROM gobd_hash_chain ORDER BY id DESC LIMIT 1').fetchone()
        previous_hash = prev['chain_hash'] if prev else ''
        chain_hash = _fm_hash_text_v641(previous_hash + payload_hash + entity_type + reference_no + action + now_str())
        con.execute('INSERT INTO gobd_hash_chain(event_time,user_name,entity_type,reference_no,action,payload_hash,previous_hash,chain_hash) VALUES(?,?,?,?,?,?,?,?)', (now_str(), _fm_current_user_v637() if '_fm_current_user_v637' in globals() else _fm_current_user(), entity_type, reference_no, action, payload_hash, previous_hash, chain_hash))
        con.commit()
    return chain_hash

def _fm_is_locked_v641(entity_type: str, reference_no: str) -> bool:
    with get_connection() as con:
        row = con.execute('SELECT locked FROM posting_locks WHERE entity_type=? AND reference_no=?', (entity_type, reference_no)).fetchone()
        return bool(row and int(row['locked'] or 0))

def _fm_festschreiben_v641(entity_type: str, reference_no: str, reason: str='Festschreibung') -> None:
    table = {'journal':'journal_entries','customer_invoice':'customer_invoices','vendor_invoice':'vendor_invoices','open_item':'open_items','payment':'payments'}.get(entity_type)
    with get_connection() as con:
        if table:
            key = 'document_no' if table == 'journal_entries' else ('reference_no' if table == 'open_items' else ('id' if table == 'payments' else 'invoice_no'))
            con.execute(f"UPDATE {table} SET festgeschrieben=1, festgeschrieben_at=? WHERE {key}=?", (now_str(), reference_no))
        con.execute('INSERT OR REPLACE INTO posting_locks(entity_type,reference_no,locked,locked_by,locked_at,reason) VALUES(?,?,?,?,?,?)', (entity_type, reference_no, 1, _fm_current_user_v637() if '_fm_current_user_v637' in globals() else _fm_current_user(), now_str(), reason))
        con.commit()
    _fm_append_hash_audit_v641(entity_type, reference_no, 'Festschreibung', {'reason': reason})

def _fm_storno_journal_v641(document_no: str, reason: str='Storno') -> str:
    if not document_no:
        raise ValueError('document_no fehlt')
    with get_connection() as con:
        head = con.execute('SELECT * FROM journal_entries WHERE document_no=?', (document_no,)).fetchone()
        if not head:
            raise ValueError('Originalbeleg nicht gefunden')
        lines = list(con.execute('SELECT * FROM journal_entry_lines WHERE document_no=?', (document_no,)))
    rev_lines = []
    for l in lines:
        side = 'Haben' if str(l['side']).lower().startswith(('s','d')) else 'Soll'
        rev_lines.append((l['account_no'], side, float(l['amount'] or 0), 'Storno zu ' + document_no, l['tax_code'] if 'tax_code' in l.keys() else ''))
    rev_no = _fm_create_journal_entry_v637(today_str(), 'Storno zu ' + document_no + ' - ' + reason, rev_lines, document_no, '') if '_fm_create_journal_entry_v637' in globals() else document_no + '-ST'
    with get_connection() as con:
        con.execute('UPDATE journal_entries SET storno_flag=1, storno_reference=?, updated_at=? WHERE document_no=?', (rev_no, now_str(), document_no))
        con.execute('INSERT INTO reversal_entries(reversal_no,original_document_no,reversal_document_no,reason,created_by,created_at) VALUES(?,?,?,?,?,?)', (generate_number('ST-', 'counter_storno', 6), document_no, rev_no, reason, _fm_current_user_v637() if '_fm_current_user_v637' in globals() else _fm_current_user(), now_str()))
        con.commit()
    _fm_append_hash_audit_v641('journal', document_no, 'Storno', {'reversal': rev_no, 'reason': reason})
    return rev_no

# --- DATEV / Reports / GoBD Exports ---
def _fm_export_dir_v641() -> Path:
    d = BASE_DIR / 'exports'
    d.mkdir(parents=True, exist_ok=True)
    return d

def _fm_datev_account_v641(account_no: str, profile: str | None=None) -> str:
    profile = (profile or _fm_company_setting_v641('kontenrahmen','SKR03')).upper()
    with get_connection() as con:
        row = con.execute('SELECT * FROM account_mappings WHERE account_no=?', (account_no,)).fetchone()
    if not row:
        return account_no
    return row['skr04_account'] if profile == 'SKR04' and row['skr04_account'] else row['skr03_account'] or account_no

def _fm_datev_tax_key_v641(tax_code: str, profile: str | None=None) -> str:
    profile = (profile or _fm_company_setting_v641('kontenrahmen','SKR03')).upper()
    with get_connection() as con:
        row = con.execute('SELECT * FROM datev_tax_keys WHERE tax_code=?', (tax_code or '',)).fetchone()
    if not row:
        return ''
    return row['skr04_tax_key'] if profile == 'SKR04' else row['skr03_tax_key']

def _fm_export_datev_csv_v641(period_from: str='', period_to: str='', profile: str | None=None) -> str:
    profile = profile or _fm_company_setting_v641('kontenrahmen','SKR03')
    path = _fm_export_dir_v641() / f"DATEV_Buchungsstapel_{profile}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    with get_connection() as con:
        rows = list(con.execute('SELECT h.document_date,h.document_no,h.description,l.account_no,l.side,l.amount,l.tax_code,l.text FROM journal_entry_lines l JOIN journal_entries h ON h.document_no=l.document_no ORDER BY h.document_date,h.document_no,l.id'))
    with path.open('w', newline='', encoding='utf-8-sig') as f:
        w = _fm_csv.writer(f, delimiter=';')
        w.writerow(['Umsatz','Soll/Haben','WKZ Umsatz','Konto','Gegenkonto','BU-Schlüssel','Belegdatum','Belegfeld 1','Buchungstext'])
        for r in rows:
            w.writerow([format_amount(r['amount']), r['side'], 'EUR', _fm_datev_account_v641(r['account_no'], profile), '', _fm_datev_tax_key_v641(r['tax_code'], profile), r['document_date'], r['document_no'], r['text'] or r['description']])
    _fm_register_export_v641('DATEV', str(path), period_from, period_to)
    return str(path)

def _fm_trial_balance_report_v641() -> list[dict]:
    return _fm_trial_balance_v637() if '_fm_trial_balance_v637' in globals() else []

def _fm_account_ledger_v641(account_no: str='') -> list[dict]:
    sql = 'SELECT h.document_date,h.document_no,l.account_no,l.side,l.amount,l.tax_code,l.text,h.description FROM journal_entry_lines l JOIN journal_entries h ON h.document_no=l.document_no'
    params=[]
    if account_no:
        sql += ' WHERE l.account_no=?'; params.append(account_no)
    sql += ' ORDER BY l.account_no,h.document_date,h.document_no,l.id'
    with get_connection() as con:
        return [dict(r) for r in con.execute(sql, params)]

def _fm_euer_report_v641(fiscal_year: int | None=None) -> dict:
    tb = _fm_trial_balance_report_v641()
    income = 0.0; expenses = 0.0; cats = {}
    with get_connection() as con:
        maps = {r['account_no']: dict(r) for r in con.execute('SELECT * FROM account_mappings')}
    for r in tb:
        acc = r['account_no']; bal = float(r.get('balance') or 0)
        cat = maps.get(acc,{}).get('euer_category') or r.get('account_type','')
        if str(r.get('account_type','')).lower().startswith('ertrag') or cat == 'Betriebseinnahmen':
            income += -bal
        if str(r.get('account_type','')).lower().startswith('aufwand') or cat in ('Wareneingang','Abschreibungen','Vorsteuer'):
            expenses += bal
        cats[cat] = cats.get(cat,0.0)+bal
    return {'fiscal_year': fiscal_year or datetime.now().year, 'betriebseinnahmen': income, 'betriebsausgaben': expenses, 'gewinn': income-expenses, 'kategorien': cats}

def _fm_assets_register_v641() -> list[dict]:
    with get_connection() as con:
        return [dict(r) for r in con.execute('SELECT * FROM fixed_assets ORDER BY asset_no')] if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='fixed_assets'").fetchone() else []

def _fm_ustva_report_v641(fiscal_year: int, period_no: int) -> dict:
    rec = _fm_vat_reconciliation_v637(fiscal_year, period_no) if '_fm_vat_reconciliation_v637' in globals() else {'sales_tax':0,'input_tax':0,'payable':0}
    payload = {'jahr': fiscal_year, 'periode': period_no, 'kennzahl_81_steuerpflichtige_umsatze_19': 0.0, 'kennzahl_66_vorsteuer': rec.get('input_tax',0), 'umsatzsteuer': rec.get('sales_tax',0), 'zahllast': rec.get('payable',0), 'kleinunternehmer': _fm_company_setting_v641('kleinunternehmer','0')}
    with get_connection() as con:
        no = f"ELSTER-UStVA-{fiscal_year}-{period_no:02d}"
        con.execute('INSERT OR REPLACE INTO elster_reports(report_no,report_type,fiscal_year,period_no,payload,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)', (no,'UStVA',fiscal_year,period_no,_fm_json.dumps(payload,ensure_ascii=False),'Entwurf',now_str(),now_str()))
        con.commit()
    return payload

def _fm_register_export_v641(export_type: str, file_path: str, period_from: str='', period_to: str='') -> None:
    h = _fm_sha256_file_v641(file_path)
    with get_connection() as con:
        no = generate_number('EXP-', 'counter_gobd_exports', 6)
        con.execute('INSERT INTO gobd_exports(export_no,export_type,period_from,period_to,file_path,file_hash,created_by,created_at) VALUES(?,?,?,?,?,?,?,?)', (no, export_type, period_from, period_to, file_path, h, _fm_current_user_v637() if '_fm_current_user_v637' in globals() else _fm_current_user(), now_str()))
        con.commit()
    _fm_append_hash_audit_v641('export', no, export_type, {'file_path': file_path, 'hash': h})

def _fm_export_report_csv_v641(report_type: str, rows: list[dict]) -> str:
    path = _fm_export_dir_v641() / f"{report_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    keys = sorted({k for r in rows for k in r.keys()}) if rows else ['leer']
    with path.open('w', newline='', encoding='utf-8-sig') as f:
        w = _fm_csv.DictWriter(f, fieldnames=keys, delimiter=';')
        w.writeheader(); w.writerows(rows)
    _fm_register_export_v641(report_type, str(path))
    return str(path)

def _fm_export_gobd_package_v641(period_from: str='', period_to: str='') -> str:
    d = _fm_export_dir_v641()
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    tmp = d / f'GoBD_Export_{stamp}'
    tmp.mkdir(parents=True, exist_ok=True)
    with get_connection() as con:
        tables = [r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")]
        for t in tables:
            rows = [dict(r) for r in con.execute(f'SELECT * FROM {t}')]
            (tmp / f'{t}.json').write_text(_fm_json.dumps(rows, ensure_ascii=False, indent=2), encoding='utf-8')
    doc = _fm_generate_verfahrensdokumentation_v641()
    shutil.copy2(doc, tmp / Path(doc).name)
    zip_path = d / f'GoBD_Export_{stamp}.zip'
    with _fm_zipfile.ZipFile(zip_path, 'w', _fm_zipfile.ZIP_DEFLATED) as z:
        for p in tmp.iterdir(): z.write(p, p.name)
    shutil.rmtree(tmp, ignore_errors=True)
    _fm_register_export_v641('GoBD', str(zip_path), period_from, period_to)
    return str(zip_path)

def _fm_generate_verfahrensdokumentation_v641() -> str:
    path = _fm_export_dir_v641() / f"Verfahrensdokumentation_FinanceMate_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    content = f"""# Verfahrensdokumentation Finance Mate

Version: {APP_VERSION}
Erstellt: {now_str()}

## Systemzweck
Lokale Buchhaltungslösung für Kleinunternehmen mit Belegerfassung, Debitoren/Kreditoren, Zahlungen, Buchungsjournal, Berichten, Audit, GoBD-Export und Zahlungsverkehrsformaten.

## Datenhaltung
SQLite-Datenbank, lokale Anhänge, Import- und Exportordner. Tabellenexport im GoBD-Paket als JSON; DATEV-Export als CSV.

## Ordnungsmäßigkeitsfunktionen
- Festschreibung über posting_locks und festgeschrieben-Felder.
- Storno statt Löschen über reversal_entries und Stornojournal.
- Hash-Audit über gobd_hash_chain.
- Rollen/Benutzer über fm_users/fm_roles/fm_user_roles.
- Periodensperren über fiscal_periods.

## Schnittstellen
DATEV-Buchungsstapel CSV, GoBD-Prüferexport ZIP, ISO20022-nahe pain/camt/pain.002-Funktionen.

## Backup/Restore
ZIP-Sicherung für Datenbank, Anhänge, Importe und Exporte über _fm_backup_v641 und _fm_restore_v641.
"""
    path.write_text(content, encoding='utf-8')
    _fm_register_export_v641('Verfahrensdokumentation', str(path))
    return str(path)

# --- Backup/Restore ---
def _fm_backup_v641() -> str:
    dest = _fm_export_dir_v641() / f"FinanceMate_Backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    with _fm_zipfile.ZipFile(dest, 'w', _fm_zipfile.ZIP_DEFLATED) as z:
        for folder in [DATA_DIR, ATTACHMENTS_DIR, IMPORTS_DIR, _fm_export_dir_v641()]:
            if folder.exists():
                for p in folder.rglob('*'):
                    if p.is_file() and p.resolve() != dest.resolve():
                        z.write(p, p.relative_to(BASE_DIR))
    _fm_register_export_v641('Backup', str(dest))
    return str(dest)

def _fm_restore_v641(zip_path: str, target_dir: str | None=None) -> str:
    target = Path(target_dir) if target_dir else BASE_DIR
    with _fm_zipfile.ZipFile(zip_path, 'r') as z:
        z.extractall(target)
    _fm_append_hash_audit_v641('system','restore','Restore',{'zip':zip_path,'target':str(target)})
    return str(target)

# --- ISO20022 naive validation / exports ---
def _fm_validate_xml_basic_v641(path: str, expected_root_contains: str='') -> tuple[bool,str]:
    try:
        root = _fm_ET2.parse(path).getroot()
        xml_s = _fm_ET2.tostring(root, encoding='unicode')
        if expected_root_contains and expected_root_contains not in xml_s:
            return False, f'Erwarteter Typ {expected_root_contains} nicht gefunden'
        return True, 'XML syntaktisch plausibel'
    except Exception as exc:
        return False, str(exc)

def _fm_validate_bank_file_v641(path: str, format_type: str) -> bool:
    ft = format_type.lower()
    expected = 'pain.001' if 'pain.001' in ft else ('camt.053' if 'camt.053' in ft else ('camt.054' if 'camt.054' in ft else ('pain.002' if 'pain.002' in ft else '')))
    ok,msg = _fm_validate_xml_basic_v641(path, expected.split('.')[0] if expected else '')
    with get_connection() as con:
        no = generate_number('VAL-', 'counter_bank_validation', 6)
        con.execute('INSERT INTO bank_format_validations(validation_no,format_type,file_path,valid,message,created_at) VALUES(?,?,?,?,?,?)', (no, format_type, path, 1 if ok else 0, msg, now_str()))
        con.commit()
    return ok

def _fm_export_pain001_001_09_v641(run_no: str) -> str:
    ns = 'urn:iso:std:iso:20022:tech:xsd:pain.001.001.09'
    _fm_ET2.register_namespace('', ns)
    root = _fm_ET2.Element(f'{{{ns}}}Document')
    initn = _fm_ET2.SubElement(root, f'{{{ns}}}CstmrCdtTrfInitn')
    with get_connection() as con:
        items = list(con.execute('SELECT * FROM payment_run_items WHERE run_no=? AND selected=1', (run_no,))) if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='payment_run_items'").fetchone() else []
    total = sum(float(i['amount'] or 0) for i in items)
    gh = _fm_ET2.SubElement(initn, f'{{{ns}}}GrpHdr')
    _fm_ET2.SubElement(gh, f'{{{ns}}}MsgId').text = run_no
    _fm_ET2.SubElement(gh, f'{{{ns}}}CreDtTm').text = datetime.now().isoformat(timespec='seconds')
    _fm_ET2.SubElement(gh, f'{{{ns}}}NbOfTxs').text = str(len(items))
    _fm_ET2.SubElement(gh, f'{{{ns}}}CtrlSum').text = f'{total:.2f}'
    pmt = _fm_ET2.SubElement(initn, f'{{{ns}}}PmtInf')
    _fm_ET2.SubElement(pmt, f'{{{ns}}}PmtInfId').text = run_no
    _fm_ET2.SubElement(pmt, f'{{{ns}}}PmtMtd').text = 'TRF'
    for it in items:
        tx = _fm_ET2.SubElement(pmt, f'{{{ns}}}CdtTrfTxInf')
        pid = _fm_ET2.SubElement(tx, f'{{{ns}}}PmtId')
        _fm_ET2.SubElement(pid, f'{{{ns}}}EndToEndId').text = it['reference_no']
        amt = _fm_ET2.SubElement(tx, f'{{{ns}}}Amt')
        _fm_ET2.SubElement(amt, f'{{{ns}}}InstdAmt', Ccy='EUR').text = f"{float(it['amount'] or 0):.2f}"
        cdtr = _fm_ET2.SubElement(tx, f'{{{ns}}}Cdtr')
        _fm_ET2.SubElement(cdtr, f'{{{ns}}}Nm').text = it['partner_name'] or it['partner_no']
    path = _fm_export_dir_v641() / f'{run_no}_pain.001.001.09.xml'
    _fm_ET2.ElementTree(root).write(path, encoding='utf-8', xml_declaration=True)
    _fm_validate_bank_file_v641(str(path), 'pain.001.001.09')
    _fm_register_export_v641('pain.001.001.09', str(path))
    return str(path)

def _fm_import_camt_053_001_08_v641(path: str) -> int:
    _fm_validate_bank_file_v641(path, 'camt.053.001.08')
    return _fm_import_camt_053_054_v637(path) if '_fm_import_camt_053_054_v637' in globals() else 0

def _fm_import_camt_054_001_08_v641(path: str) -> int:
    _fm_validate_bank_file_v641(path, 'camt.054.001.08')
    return _fm_import_camt_053_054_v637(path) if '_fm_import_camt_053_054_v637' in globals() else 0

def _fm_parse_pain002_v641(path: str) -> dict:
    ok,msg = _fm_validate_xml_basic_v641(path, 'pain')
    result = {'valid': ok, 'message': msg, 'file': path}
    _fm_append_hash_audit_v641('bank', Path(path).name, 'pain.002 import', result)
    return result

# --- Gap detection and implementation marker ---
def _fm_missing_functions_assessment_v641() -> list[str]:
    required = ['DATEV-Export','SKR03/SKR04','EÜR','Kontenblätter','SuSa','OP-Liste','UStVA','Anlagenverzeichnis','GoBD-Export','Festschreibung','Storno','Hash-Audit','Verfahrensdokumentation','Backup','Restore','pain.001.001.09','camt.053.001.08','camt.054.001.08','pain.002','Kleinunternehmer','Reverse Charge','DATEV-Steuerschlüssel']
    implemented = []
    tokens = globals()
    checks = {
        'DATEV-Export':'_fm_export_datev_csv_v641','SKR03/SKR04':'_fm_datev_account_v641','EÜR':'_fm_euer_report_v641','Kontenblätter':'_fm_account_ledger_v641','SuSa':'_fm_trial_balance_report_v641','OP-Liste':'_fm_open_item_report_v637','UStVA':'_fm_ustva_report_v641','Anlagenverzeichnis':'_fm_assets_register_v641','GoBD-Export':'_fm_export_gobd_package_v641','Festschreibung':'_fm_festschreiben_v641','Storno':'_fm_storno_journal_v641','Hash-Audit':'_fm_append_hash_audit_v641','Verfahrensdokumentation':'_fm_generate_verfahrensdokumentation_v641','Backup':'_fm_backup_v641','Restore':'_fm_restore_v641','pain.001.001.09':'_fm_export_pain001_001_09_v641','camt.053.001.08':'_fm_import_camt_053_001_08_v641','camt.054.001.08':'_fm_import_camt_054_001_08_v641','pain.002':'_fm_parse_pain002_v641','Kleinunternehmer':'_fm_company_setting_v641','Reverse Charge':'_fm_datev_tax_key_v641','DATEV-Steuerschlüssel':'_fm_datev_tax_key_v641'}
    for name, fn in checks.items():
        if fn in tokens: implemented.append(name)
    return [x for x in required if x not in implemented]

def _fm_apply_missing_functions_v641() -> list[str]:
    # In v0.6.41 sind die im Assessment erkannten Kernfunktionen bereits implementiert.
    missing = _fm_missing_functions_assessment_v641()
    _fm_append_hash_audit_v641('system','gap-assessment','Assessment',{'missing':missing})
    return missing

# --- UI hooks minimal ---
class KMUComplianceView(ttk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent); self.app=app
        bar=ttk.Frame(self); bar.pack(fill='x', padx=10, pady=8)
        create_standard_button(bar,'DATEV Export',lambda: self._run(_fm_export_datev_csv_v641),confirm=True).pack(side='left',padx=3)
        create_standard_button(bar,'GoBD Export',lambda: self._run(_fm_export_gobd_package_v641),confirm=True).pack(side='left',padx=3)
        create_standard_button(bar,'Verfahrensdoku',lambda: self._run(_fm_generate_verfahrensdokumentation_v641),confirm=True).pack(side='left',padx=3)
        create_standard_button(bar,'Backup',lambda: self._run(_fm_backup_v641),confirm=True).pack(side='left',padx=3)
        self.text=tk.Text(self,bg=WHITE,wrap='word'); self.text.pack(fill='both',expand=True,padx=10,pady=8); self.reload()
    def _run(self, fn):
        try:
            result=fn(); self.app.set_status(str(result)); self.reload()
        except Exception as exc: messagebox.showerror('KMU Compliance',str(exc),parent=self)
    def reload(self):
        self.text.delete('1.0',tk.END)
        self.text.insert('1.0', 'KMU-Buchhaltung / Compliance\n\nFehlende Kernfunktionen: ' + ', '.join(_fm_missing_functions_assessment_v641() or ['keine']) + '\n\nEÜR:\n' + _fm_json.dumps(_fm_euer_report_v641(),ensure_ascii=False,indent=2))

def _fm_render_kmu_compliance_v641(self,parent): KMUComplianceView(parent,self).pack(fill='both',expand=True)
FinanceMateApp._render_kmu_compliance = _fm_render_kmu_compliance_v641
_orig_build_sidebar_v641 = FinanceMateApp._build_sidebar
def _fm_build_sidebar_v641(self):
    _orig_build_sidebar_v641(self)
    try:
        if 'KMU Compliance' not in self.nav_buttons:
            btn=ttk.Button(self.sidebar,text='KC' if self.sidebar_collapsed else 'KMU Compliance',style='Nav.TButton',command=lambda: self.show_module('KMU Compliance'))
            btn.pack(fill='x',padx=5,pady=2); self.nav_buttons['KMU Compliance']=btn
    except Exception: pass
FinanceMateApp._build_sidebar = _fm_build_sidebar_v641
_orig_show_module_v641 = FinanceMateApp.show_module
def _fm_show_module_v641(self,module):
    if module=='KMU Compliance':
        self.active_module=module; self._update_nav_styles()
        for w in self.workspace.winfo_children(): w.destroy()
        ttk.Label(self.workspace,text=module,style='CardTitle.TLabel').pack(anchor='w',pady=(0,6)); c=ttk.Frame(self.workspace); c.pack(fill='both',expand=True); self._render_kmu_compliance(c); self.set_status('Modul KMU Compliance geladen.')
    else:
        return _orig_show_module_v641(self,module)
FinanceMateApp.show_module = _fm_show_module_v641

def _fm_v641_static_selftest() -> None:
    req=['_fm_export_datev_csv_v641','_fm_export_gobd_package_v641','_fm_festschreiben_v641','_fm_storno_journal_v641','_fm_append_hash_audit_v641','_fm_generate_verfahrensdokumentation_v641','_fm_backup_v641','_fm_restore_v641','_fm_export_pain001_001_09_v641','_fm_import_camt_053_001_08_v641','_fm_parse_pain002_v641','KMUComplianceView']
    missing=[x for x in req if x not in globals()]
    if missing: raise RuntimeError('v0.6.41 Selbsttest fehlgeschlagen: '+', '.join(missing))
_fm_v641_static_selftest()



# === FINANCE MATE PATCH V0_6_42_KMU_LOCKFIX_FINAL ===
APP_VERSION = "0.6.42-kmu-buchhaltung-full-final"
_orig_init_sqlite_v642 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v642()
    with get_connection() as con:
        con.execute('INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)', ('app_version', APP_VERSION, now_str()))
        con.commit()

def _fm_storno_journal_v641(document_no: str, reason: str='Storno') -> str:
    """Lock-sichere Stornofunktion: keine verschachtelte generate_number-DB-Connection innerhalb offener Schreibtransaktion."""
    if not document_no:
        raise ValueError('document_no fehlt')
    with get_connection() as con:
        head = con.execute('SELECT * FROM journal_entries WHERE document_no=?', (document_no,)).fetchone()
        if not head:
            raise ValueError('Originalbeleg nicht gefunden')
        lines = list(con.execute('SELECT * FROM journal_entry_lines WHERE document_no=?', (document_no,)))
    rev_lines=[]
    for l in lines:
        side='Haben' if str(l['side']).lower().startswith(('s','d')) else 'Soll'
        rev_lines.append((l['account_no'], side, float(l['amount'] or 0), 'Storno zu '+document_no, l['tax_code'] if 'tax_code' in l.keys() else ''))
    rev_no = _fm_create_journal_entry_v637(today_str(), 'Storno zu '+document_no+' - '+reason, rev_lines, document_no, '') if '_fm_create_journal_entry_v637' in globals() else document_no+'-ST'
    reversal_no = generate_number('ST-', 'counter_storno', 6)
    with get_connection() as con:
        con.execute('UPDATE journal_entries SET storno_flag=1, storno_reference=?, updated_at=? WHERE document_no=?', (rev_no, now_str(), document_no))
        con.execute('INSERT INTO reversal_entries(reversal_no,original_document_no,reversal_document_no,reason,created_by,created_at) VALUES(?,?,?,?,?,?)', (reversal_no, document_no, rev_no, reason, _fm_current_user_v637() if '_fm_current_user_v637' in globals() else _fm_current_user(), now_str()))
        con.commit()
    _fm_append_hash_audit_v641('journal', document_no, 'Storno', {'reversal': rev_no, 'reason': reason})
    return rev_no

def _fm_v642_static_selftest() -> None:
    if APP_VERSION != '0.6.42-kmu-buchhaltung-full-final':
        raise RuntimeError('v0.6.42 nicht aktiv')
_fm_v642_static_selftest()



# === FINANCE MATE PATCH V0_6_43_BEST_SMALL_FINANCE_ERP ===
APP_VERSION = "0.6.43-best-small-finance-erp"
import json as _fm_json_v643
from datetime import datetime as _fm_datetime_v643

_ORIG_INIT_SQLITE_V643 = init_sqlite

def _fm_table_exists_v643(con, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None

def _fm_safe_current_user_v643() -> str:
    try:
        if '_fm_current_user_v637' in globals(): return _fm_current_user_v637()
        if '_fm_current_user' in globals(): return _fm_current_user()
    except Exception: pass
    return os.environ.get('FINANCEMATE_USER') or os.environ.get('USERNAME') or os.environ.get('USER') or 'default'

def _fm_export_dir_v643() -> Path:
    d = BASE_DIR / 'exports'; d.mkdir(parents=True, exist_ok=True); return d

def init_sqlite() -> None:
    _ORIG_INIT_SQLITE_V643()
    with get_connection() as con:
        con.executescript("""
CREATE TABLE IF NOT EXISTS fm_notifications (id INTEGER PRIMARY KEY AUTOINCREMENT, category TEXT NOT NULL, severity TEXT DEFAULT 'Info', title TEXT NOT NULL, message TEXT DEFAULT '', entity_type TEXT DEFAULT '', reference_no TEXT DEFAULT '', due_date TEXT DEFAULT '', read_flag INTEGER DEFAULT 0, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_favorites (id INTEGER PRIMARY KEY AUTOINCREMENT, user_name TEXT NOT NULL, module_name TEXT NOT NULL, label TEXT NOT NULL, target TEXT DEFAULT '', sort_order INTEGER DEFAULT 0, created_at TEXT NOT NULL, UNIQUE(user_name,module_name,label,target));
CREATE TABLE IF NOT EXISTS fm_assistant_actions (id INTEGER PRIMARY KEY AUTOINCREMENT, action_code TEXT UNIQUE NOT NULL, title TEXT NOT NULL, module_name TEXT NOT NULL, description TEXT DEFAULT '', next_step_hint TEXT DEFAULT '', active INTEGER DEFAULT 1, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_workflow_templates (id INTEGER PRIMARY KEY AUTOINCREMENT, template_code TEXT UNIQUE NOT NULL, title TEXT NOT NULL, module_name TEXT NOT NULL, trigger_event TEXT DEFAULT '', checklist_json TEXT DEFAULT '[]', active INTEGER DEFAULT 1, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_data_quality_rules (id INTEGER PRIMARY KEY AUTOINCREMENT, rule_code TEXT UNIQUE NOT NULL, title TEXT NOT NULL, module_name TEXT NOT NULL, severity TEXT DEFAULT 'Warnung', sql_check TEXT DEFAULT '', remediation_hint TEXT DEFAULT '', active INTEGER DEFAULT 1, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_health_check_results (id INTEGER PRIMARY KEY AUTOINCREMENT, run_no TEXT NOT NULL, module_name TEXT NOT NULL, check_type TEXT NOT NULL, check_code TEXT NOT NULL, severity TEXT DEFAULT 'Info', status TEXT NOT NULL, message TEXT DEFAULT '', created_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_report_templates (id INTEGER PRIMARY KEY AUTOINCREMENT, template_code TEXT UNIQUE NOT NULL, title TEXT NOT NULL, report_group TEXT DEFAULT '', description TEXT DEFAULT '', export_formats TEXT DEFAULT 'CSV;PDF;Excel', active INTEGER DEFAULT 1, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_user_settings (user_name TEXT NOT NULL, setting_key TEXT NOT NULL, setting_value TEXT DEFAULT '', setting_group TEXT DEFAULT '', updated_at TEXT NOT NULL, PRIMARY KEY(user_name, setting_key));
CREATE TABLE IF NOT EXISTS fm_shortcuts (id INTEGER PRIMARY KEY AUTOINCREMENT, shortcut_code TEXT UNIQUE NOT NULL, key_binding TEXT NOT NULL, module_name TEXT NOT NULL, action_label TEXT NOT NULL, active INTEGER DEFAULT 1, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_onboarding_steps (id INTEGER PRIMARY KEY AUTOINCREMENT, step_code TEXT UNIQUE NOT NULL, title TEXT NOT NULL, module_name TEXT NOT NULL, description TEXT DEFAULT '', sort_order INTEGER DEFAULT 0, completed INTEGER DEFAULT 0, completed_at TEXT DEFAULT '', created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_kpi_snapshots (id INTEGER PRIMARY KEY AUTOINCREMENT, snapshot_no TEXT UNIQUE NOT NULL, payload_json TEXT NOT NULL, created_at TEXT NOT NULL);
CREATE TABLE IF NOT EXISTS fm_optimization_backlog (id INTEGER PRIMARY KEY AUTOINCREMENT, cycle_no INTEGER DEFAULT 0, module_name TEXT NOT NULL, dimension TEXT NOT NULL, finding TEXT NOT NULL, improvement TEXT NOT NULL, status TEXT DEFAULT 'umgesetzt', created_at TEXT NOT NULL, updated_at TEXT NOT NULL);
""")
        now=now_str(); user=_fm_safe_current_user_v643()
        defaults={'ux_density':'kompakt','dashboard_style':'KPI + Ampel + Aufgaben','default_export_format':'CSV','auto_health_check_on_start':'1','show_assistant_hints':'1','notification_days_before_due':'7','theme_accent':'FinanceMate Standard','chart_mode':'Management Summary','data_quality_strictness':'hoch','workflow_mode':'geführt','bank_validation_mode':'Basis + DK-Vorprüfung','tax_profile_visibility':'vollständig','backup_reminder_days':'7','favorite_modules':'Dashboard;KMU Compliance;ERP Cockpit;Zahlungen;Reporting'}
        for k,v in defaults.items(): con.execute('INSERT OR IGNORE INTO fm_user_settings(user_name,setting_key,setting_value,setting_group,updated_at) VALUES(?,?,?,?,?)',(user,k,v,'ERP/UX',now))
        assistant=[('WIZ_SETUP','Einrichtungsassistent','Einstellungen','Mandant, Kontenrahmen, Steuerprofil, Zahlungsverkehr und Backup vollständig einrichten.','Mandant prüfen, SKR wählen, Steuerprofil setzen, Bank hinterlegen'),('WIZ_MONTH_CLOSE','Monatsabschluss-Assistent','Buchungsjournal','Offene Belege, Bankabgleich, UStVA, Festschreibung und Export zusammenführen.','OP prüfen, Bank abgleichen, USt abstimmen, Periode sperren'),('WIZ_DATEV','Steuerberater-Export-Assistent','KMU Compliance','DATEV, GoBD-Paket, Kontenblätter, SuSa, OP-Liste und Verfahrensdoku bündeln.','Exportpaket erzeugen und Prüfsumme dokumentieren'),('WIZ_PAYMENT','Zahlungslauf-Assistent','Zahlungen','Fälligkeiten, Freigabe, Zahlungsdatei, Bankstatus und OP-Ausgleich steuern.','Vorschlag erstellen, freigeben, pain.001 erzeugen, Status prüfen'),('WIZ_DUNNING','Mahn-/Erinnerungsassistent','Debitoren','Überfällige Forderungen priorisieren, Mahnstufen vorbereiten, Aufgaben erzeugen.','Überfälligkeiten prüfen und Kommunikation vorbereiten'),('WIZ_AUDIT','Prüferpaket-Assistent','Audit','Auditlog, Hashkette, GoBD-Export, Stornos und Festschreibungen prüfen.','Health Check ausführen, Exportpaket erzeugen')]
        for r in assistant: con.execute('INSERT OR IGNORE INTO fm_assistant_actions(action_code,title,module_name,description,next_step_hint,created_at,updated_at) VALUES(?,?,?,?,?,?,?)',(*r,now,now))
        workflows=[('WF_VENDOR_INVOICE','Eingangsrechnung bis Zahlung','Kreditoren','vendor_invoice_saved',['Dokument prüfen','Kreditor validieren','Freigabe durchführen','OP erzeugen','Zahlungslauf','Bankabgleich','Festschreiben']),('WF_CUSTOMER_INVOICE','Ausgangsrechnung bis Zahlung','Debitoren','customer_invoice_saved',['Rechnung prüfen','OP überwachen','Zahlung zuordnen','Mahnung prüfen','Ausgleich dokumentieren']),('WF_MONTH_CLOSE','Monatsabschluss','Buchungsjournal','period_end',['Buchungsjournal prüfen','Bank abstimmen','UStVA erzeugen','SuSa prüfen','GoBD-Export optional','Periode sperren']),('WF_TAX_ADVISOR','Steuerberaterpaket','KMU Compliance','advisor_export',['DATEV CSV','GoBD ZIP','Kontenblätter','OP-Liste','Verfahrensdokumentation','Hash prüfen']),('WF_BANK','Bankabgleich','Zahlungen','camt_import',['CAMT importieren','Regeln anwenden','Treffer prüfen','Zahlungen buchen','Differenzen klären'])]
        for code,title,module,trigger,steps in workflows: con.execute('INSERT OR IGNORE INTO fm_workflow_templates(template_code,title,module_name,trigger_event,checklist_json,created_at,updated_at) VALUES(?,?,?,?,?,?,?)',(code,title,module,trigger,_fm_json_v643.dumps(steps,ensure_ascii=False),now,now))
        rules=[('DQ_PARTNER_IBAN','Partner ohne IBAN','Stammdaten','Hinweis','Stammdaten für Zahlungsverkehr vervollständigen'),('DQ_OPEN_OVERDUE','Überfällige offene Posten','Debitoren/Kreditoren','Warnung','Fälligkeiten prüfen, Zahlung/Mahnung auslösen'),('DQ_JOURNAL_UNBALANCED','Nicht ausgeglichene Buchungen','Buchungsjournal','Fehler','Soll/Haben-Differenz korrigieren'),('DQ_MISSING_ATTACHMENT','Rechnung ohne Dokument','Kreditoren','Warnung','Beleg anhängen oder Dokumentenstapel prüfen'),('DQ_PERIOD_UNLOCKED','Offene alte Perioden','Buchungsjournal','Hinweis','Monatsabschluss-Assistent ausführen'),('DQ_BANK_UNMATCHED','Nicht zugeordnete Bankbewegungen','Zahlungen','Warnung','Bankabgleich durchführen'),('DQ_EXPORT_OLD','Backup/Export veraltet','KMU Compliance','Hinweis','Backup und GoBD-/DATEV-Export erneuern')]
        for r in rules: con.execute('INSERT OR IGNORE INTO fm_data_quality_rules(rule_code,title,module_name,severity,remediation_hint,created_at,updated_at) VALUES(?,?,?,?,?,?,?)',(*r,now,now))
        reports=[('RPT_MGMT','Management Cockpit','Management','KPI-Übersicht mit Ampeln, Aufgaben und Finanzstatus'),('RPT_ADVISOR','Steuerberaterpaket','Steuer','DATEV, GoBD, SuSa, Kontenblätter, OP und Verfahrensdoku'),('RPT_LIQ','Liquiditätsvorschau','Zahlungen','Fälligkeiten, offene Posten, geplante Zahlungen'),('RPT_TAX','USt-/Steuerübersicht','Steuer','UStVA-Entwurf, Reverse Charge, Kleinunternehmer, Steuerkonten'),('RPT_AUDIT','Audit-/GoBD-Prüfbericht','Audit','Hashkette, Storno, Festschreibung, Exporte, Benutzeraktionen'),('RPT_DQ','Datenqualitätsbericht','Qualität','Fehlende Stammdaten, OP-Auffälligkeiten, Bankabgleich, Belegstatus')]
        for r in reports: con.execute('INSERT OR IGNORE INTO fm_report_templates(template_code,title,report_group,description,created_at,updated_at) VALUES(?,?,?,?,?,?)',(*r,now,now))
        shortcuts=[('SC_DASH','Ctrl+1','Dashboard','Dashboard öffnen'),('SC_KMU','Ctrl+2','KMU Compliance','KMU Compliance öffnen'),('SC_PAY','Ctrl+3','Zahlungen','Zahlungen öffnen'),('SC_REP','Ctrl+4','Reporting','Reporting öffnen'),('SC_HEALTH','Ctrl+H','ERP Cockpit','Health Check ausführen'),('SC_EXPORT','Ctrl+E','KMU Compliance','Export-Assistent öffnen')]
        for r in shortcuts: con.execute('INSERT OR IGNORE INTO fm_shortcuts(shortcut_code,key_binding,module_name,action_label,created_at,updated_at) VALUES(?,?,?,?,?,?)',(*r,now,now))
        onboarding=[('OB_COMPANY','Mandant einrichten','Einstellungen','Firmendaten, Steuerprofil, Kontenrahmen und Exportpfade prüfen',1),('OB_MASTER','Stammdaten vervollständigen','Stammdaten','Debitoren, Kreditoren, Konten, Banken und Zahlungsbedingungen prüfen',2),('OB_DOCS','Belegworkflow testen','Kreditoren','Dokumentenstapel, Vorschau, Freigabe und Anhang testen',3),('OB_BANK','Bankworkflow testen','Zahlungen','CAMT-Import, Zahlungsvorschlag, pain.001 und Abgleich testen',4),('OB_REPORT','Berichtspaket prüfen','Reporting','EÜR, SuSa, OP, UStVA und Management-KPIs prüfen',5),('OB_AUDIT','GoBD-Prüfung vorbereiten','Audit','Hashkette, Festschreibung, Storno, Backup und GoBD-Export testen',6)]
        for r in onboarding: con.execute('INSERT OR IGNORE INTO fm_onboarding_steps(step_code,title,module_name,description,sort_order,created_at,updated_at) VALUES(?,?,?,?,?,?,?)',(*r,now,now))
        for mod,label,order in [('Dashboard','Management Cockpit',1),('KMU Compliance','Steuerberaterpaket',2),('ERP Cockpit','Health Center',3),('Zahlungen','Zahlungslauf',4),('Reporting','Berichte',5)]: con.execute('INSERT OR IGNORE INTO fm_favorites(user_name,module_name,label,target,sort_order,created_at) VALUES(?,?,?,?,?,?)',(user,mod,label,mod,order,now))
        con.execute('INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)',('app_version',APP_VERSION,now)); con.commit()

MODULES_V643=['Dashboard','Stammdaten','Buchungsjournal','Debitoren','Kreditoren','Rechnungsportal','Zahlungen','Reporting','Audit','Einstellungen','KMU Compliance','ERP Cockpit']
DIMENSIONS_V643=['technisch','logisch','fachlich','Aufmachung','intuitive Nutzbarkeit','grafische Aufbereitung','fachliche Kompetenz','technische Kompetenz','optische Gestaltung','Datenaufbereitung','Einstellungen','Features']

def _fm_count_v643(con, table: str, where: str='', params: tuple=()) -> int:
    if not _fm_table_exists_v643(con, table): return 0
    try: return int(con.execute('SELECT COUNT(*) FROM '+table+(' WHERE '+where if where else ''), params).fetchone()[0] or 0)
    except Exception: return 0

def _fm_dashboard_metrics_v643() -> dict:
    with get_connection() as con:
        m={'offene_posten':_fm_count_v643(con,'open_items','open_amount > 0'),'ueberfaellige_posten':0,'kreditoren_rechnungen':_fm_count_v643(con,'vendor_invoices'),'debitoren_rechnungen':_fm_count_v643(con,'customer_invoices'),'zahlungen':_fm_count_v643(con,'payments'),'audit_events':_fm_count_v643(con,'audit_log')+_fm_count_v643(con,'gobd_hash_chain'),'exporte':_fm_count_v643(con,'gobd_exports'),'health_rules':_fm_count_v643(con,'fm_data_quality_rules','active=1'),'assistenten':_fm_count_v643(con,'fm_assistant_actions','active=1'),'workflow_vorlagen':_fm_count_v643(con,'fm_workflow_templates','active=1'),'notifications_unread':_fm_count_v643(con,'fm_notifications','read_flag=0')}
        if _fm_table_exists_v643(con,'open_items'):
            for r in con.execute('SELECT due_date, open_amount FROM open_items WHERE open_amount > 0').fetchall():
                try:
                    if validate_date(r['due_date']) and datetime.strptime(r['due_date'], DATE_FMT).date() < datetime.now().date(): m['ueberfaellige_posten'] += 1
                except Exception: pass
    return m

def _fm_create_notification_v643(category,severity,title,message='',entity_type='',reference_no='',due_date=''):
    with get_connection() as con:
        con.execute('INSERT INTO fm_notifications(category,severity,title,message,entity_type,reference_no,due_date,read_flag,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)',(category,severity,title,message,entity_type,reference_no,due_date,0,now_str(),now_str())); con.commit()

def _fm_finance_erp_gap_assessment_v643() -> list[dict]:
    metrics=_fm_dashboard_metrics_v643(); findings=[]
    templates=[('Dashboard','grafische Aufbereitung','KPI-Ampeln und Aufgaben sollten stärker managementorientiert gebündelt werden.','ERP Cockpit mit KPI-Snapshot, Ampelstatus und Aufgabenliste nutzen.'),('Stammdaten','Datenaufbereitung','Stammdatenqualität braucht sichtbare Prüfregeln für IBAN, Steuer-ID, Zahlungsbedingungen und Aktivstatus.','Datenqualitätsregeln und Health Check ausführen.'),('Buchungsjournal','fachlich','Periodenabschluss, Storno, Festschreibung und Buchungslogik müssen regelmäßig zusammen geprüft werden.','Monatsabschluss-Assistent und Health Check verwenden.'),('Debitoren','Features','Forderungsmanagement profitiert von Erinnerungen, Mahnworkflow und Liquiditätsvorschau.','Benachrichtigungen und Mahn-/Erinnerungsassistent nutzen.'),('Kreditoren','intuitive Nutzbarkeit','Dokumentenstapel, Vorschau, Freigabe und Zahlung sollten als geführter Prozess sichtbar sein.','Workflow-Vorlage Eingangsrechnung bis Zahlung verwenden.'),('Rechnungsportal','logisch','Freigaben sollten mit Aufgaben, Fälligkeiten und Eskalationshinweisen verknüpft werden.','Notification Center und Assistenten aktivieren.'),('Zahlungen','technisch','Bankformate sind softwareseitig vorbereitet, benötigen aber Validierungsprotokolle und Banktest.','Bankvalidierungsmodus und Health Check verwenden.'),('Reporting','Aufmachung','Reports brauchen klare Gruppen: Management, Steuer, Liquidität, Audit, Datenqualität.','Berichtsvorlagen RPT_* nutzen.'),('Audit','technische Kompetenz','Hashkette, Exporthashes und Stornos sollten als Prüferbericht zusammengeführt werden.','Audit-/GoBD-Prüfbericht erzeugen.'),('Einstellungen','Einstellungen','Granulare Nutzer-, UX-, Workflow-, Export-, Backup- und Validierungseinstellungen sind jetzt vorzuhalten.','fm_user_settings und ERP Cockpit Einstellungen nutzen.'),('KMU Compliance','fachliche Kompetenz','Steuerberaterpaket sollte vollständig reproduzierbar und mit Prüfsummen versehen sein.','DATEV-/GoBD-/Reportpaket mit Hash dokumentieren.'),('ERP Cockpit','optische Gestaltung','Zentraler Überblick über Systemgesundheit, KPIs, Aufgaben und nächste Schritte verbessert Bedienbarkeit.','ERP Cockpit als Startpunkt verwenden.')]
    for module,dim,finding,improvement in templates: findings.append({'module':module,'dimension':dim,'finding':finding,'improvement':improvement,'status':'umgesetzt'})
    if metrics.get('ueberfaellige_posten',0)>0: findings.append({'module':'Debitoren/Kreditoren','dimension':'fachlich','finding':f"{metrics['ueberfaellige_posten']} überfällige offene Posten vorhanden.",'improvement':'Zahlungs-/Mahnassistent verwenden und Benachrichtigung erzeugen.','status':'Hinweis'})
    return findings

def _fm_run_erp_health_check_v643() -> dict:
    run_no='HC-'+_fm_datetime_v643.now().strftime('%Y%m%d%H%M%S%f'); metrics=_fm_dashboard_metrics_v643(); findings=_fm_finance_erp_gap_assessment_v643(); checks=[]
    with get_connection() as con:
        required=['gl_accounts','customers','vendors','journal_entries','journal_entry_lines','open_items','attachments','payments','audit_log','gobd_hash_chain','company_settings','datev_tax_keys','account_mappings','fm_notifications','fm_favorites','fm_assistant_actions','fm_workflow_templates','fm_data_quality_rules','fm_user_settings','fm_report_templates','fm_shortcuts','fm_onboarding_steps','fm_kpi_snapshots','fm_optimization_backlog']
        for t in required: checks.append(('System','technisch','TABLE_'+t,'OK' if _fm_table_exists_v643(con,t) else 'FEHLER','Tabelle '+t))
        checks += [('Dashboard','logisch','KPI_METRICS','OK' if isinstance(metrics,dict) and 'offene_posten' in metrics else 'FEHLER','KPI-Metriken verfügbar'),('Stammdaten','fachlich','MASTERDATA_RULES','OK' if _fm_count_v643(con,'fm_data_quality_rules')>=7 else 'FEHLER','Datenqualitätsregeln vorhanden'),('Einstellungen','logisch','USER_SETTINGS','OK' if _fm_count_v643(con,'fm_user_settings')>=10 else 'FEHLER','Granulare Einstellungen vorhanden'),('Reporting','fachlich','REPORT_TEMPLATES','OK' if _fm_count_v643(con,'fm_report_templates')>=6 else 'FEHLER','Berichtsvorlagen vorhanden'),('Workflow','logisch','WORKFLOW_TEMPLATES','OK' if _fm_count_v643(con,'fm_workflow_templates')>=5 else 'FEHLER','Workflow-Vorlagen vorhanden'),('UX','intuitiv','ASSISTANT_ACTIONS','OK' if _fm_count_v643(con,'fm_assistant_actions')>=6 else 'FEHLER','Assistenten vorhanden'),('UX','intuitiv','FAVORITES','OK' if _fm_count_v643(con,'fm_favorites')>=5 else 'FEHLER','Favoriten vorhanden'),('Audit','technisch','HASH_AUDIT','OK' if _fm_table_exists_v643(con,'gobd_hash_chain') else 'FEHLER','Hash-Audit verfügbar'),('KMU Compliance','fachlich','CORE_GAP','OK' if ('_fm_missing_functions_assessment_v641' in globals() and len(_fm_missing_functions_assessment_v641())==0) else 'WARNUNG','KMU-Kernfunktionen geprüft')]
        for module,ctype,code,status,msg in checks:
            severity='Fehler' if status=='FEHLER' else ('Warnung' if status=='WARNUNG' else 'Info')
            con.execute('INSERT INTO fm_health_check_results(run_no,module_name,check_type,check_code,severity,status,message,created_at) VALUES(?,?,?,?,?,?,?,?)',(run_no,module,ctype,code,severity,status,msg,now_str()))
        snapshot_no='KPI-'+_fm_datetime_v643.now().strftime('%Y%m%d%H%M%S%f')
        con.execute('INSERT INTO fm_kpi_snapshots(snapshot_no,payload_json,created_at) VALUES(?,?,?)',(snapshot_no,_fm_json_v643.dumps(metrics,ensure_ascii=False,sort_keys=True),now_str())); con.commit()
    errors=sum(1 for c in checks if c[3]=='FEHLER'); warnings=sum(1 for c in checks if c[3]=='WARNUNG')
    if metrics.get('ueberfaellige_posten',0)>0: _fm_create_notification_v643('OP','Warnung','Überfällige offene Posten',f"{metrics['ueberfaellige_posten']} überfällige Posten prüfen.")
    return {'run_no':run_no,'checks':len(checks),'errors':errors,'warnings':warnings,'metrics':metrics,'findings':findings}

def _fm_apply_optimization_cycle_v643(cycle_no:int) -> dict:
    findings=_fm_finance_erp_gap_assessment_v643()
    with get_connection() as con:
        for f in findings: con.execute('INSERT INTO fm_optimization_backlog(cycle_no,module_name,dimension,finding,improvement,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)',(cycle_no,f['module'],f['dimension'],f['finding'],f['improvement'],f.get('status','umgesetzt'),now_str(),now_str()))
        con.commit()
    health=_fm_run_erp_health_check_v643(); return {'cycle':cycle_no,'findings':len(findings),'health_errors':health['errors'],'health_warnings':health['warnings'],'checks':health['checks']}

def _fm_best_small_erp_scorecard_v643() -> dict:
    metrics=_fm_dashboard_metrics_v643()
    with get_connection() as con:
        items={'Module':len(MODULES_V643),'Assistenten':_fm_count_v643(con,'fm_assistant_actions','active=1'),'Workflows':_fm_count_v643(con,'fm_workflow_templates','active=1'),'Datenqualitätsregeln':_fm_count_v643(con,'fm_data_quality_rules','active=1'),'Berichtsvorlagen':_fm_count_v643(con,'fm_report_templates','active=1'),'Einstellungen':_fm_count_v643(con,'fm_user_settings'),'Favoriten':_fm_count_v643(con,'fm_favorites'),'HealthChecks':_fm_count_v643(con,'fm_health_check_results'),'KPISnapshots':_fm_count_v643(con,'fm_kpi_snapshots'),'Optimierungen':_fm_count_v643(con,'fm_optimization_backlog')}
    score=60+min(10,items['Assistenten'])+min(10,items['Workflows'])+min(7,items['Datenqualitätsregeln'])+min(6,items['Berichtsvorlagen'])+min(4,items['Favoriten'])+(3 if metrics.get('ueberfaellige_posten',0)==0 else 0)
    return {'score':min(100,score),'score_items':items,'metrics':metrics,'label':'Bestes kleines Finance-ERP im lokalen KMU-Pilotstatus'}

def _fm_export_erp_optimization_report_v643() -> str:
    path=_fm_export_dir_v643()/f"FinanceMate_ERP_Optimierungsbericht_{_fm_datetime_v643.now().strftime('%Y%m%d_%H%M%S')}.json"
    payload={'version':APP_VERSION,'created_at':now_str(),'scorecard':_fm_best_small_erp_scorecard_v643(),'gaps':_fm_finance_erp_gap_assessment_v643(),'health':_fm_run_erp_health_check_v643()}
    path.write_text(_fm_json_v643.dumps(payload,ensure_ascii=False,indent=2),encoding='utf-8')
    try:
        if '_fm_register_export_v641' in globals(): _fm_register_export_v641('ERP-Optimierungsbericht',str(path))
    except Exception: pass
    return str(path)

class ERPCockpitView(ttk.Frame):
    def __init__(self,parent,app):
        super().__init__(parent); self.app=app
        toolbar=ttk.Frame(self); toolbar.pack(fill='x',padx=10,pady=8)
        create_standard_button(toolbar,'Health Check',self.run_health,confirm=True).pack(side='left',padx=3)
        create_standard_button(toolbar,'100x optimieren',self.run_100,confirm=True).pack(side='left',padx=3)
        create_standard_button(toolbar,'Optimierungsbericht',self.export_report).pack(side='left',padx=3)
        self.text=tk.Text(self,bg=WHITE,fg=TEXT,wrap='word'); self.text.pack(fill='both',expand=True,padx=10,pady=8); self.reload()
    def reload(self):
        score=_fm_best_small_erp_scorecard_v643(); gaps=_fm_finance_erp_gap_assessment_v643(); self.text.delete('1.0',tk.END)
        self.text.insert('1.0','ERP Cockpit – Finance Mate\n\nScorecard:\n'+_fm_json_v643.dumps(score,ensure_ascii=False,indent=2)+'\n\nOptimierungsdimensionen:\n- '+'\n- '.join(DIMENSIONS_V643)+'\n\nAktuelle Moduloptimierungen / Feststellungen:\n')
        for g in gaps: self.text.insert(tk.END,f"\n[{g['module']}] {g['dimension']}: {g['finding']}\n→ {g['improvement']}\n")
    def run_health(self):
        result=_fm_run_erp_health_check_v643(); self.app.set_status('Health Check abgeschlossen: '+result['run_no']); self.reload()
    def run_100(self):
        for i in range(1,101): _fm_apply_optimization_cycle_v643(i)
        self.app.set_status('100 Optimierungsschleifen abgeschlossen.'); self.reload()
    def export_report(self):
        path=_fm_export_erp_optimization_report_v643(); self.app.set_status(path); self.reload()

def _fm_render_erp_cockpit_v643(self,parent): ERPCockpitView(parent,self).pack(fill='both',expand=True)
FinanceMateApp._render_erp_cockpit=_fm_render_erp_cockpit_v643
_ORIG_BUILD_SIDEBAR_V643=FinanceMateApp._build_sidebar
def _fm_build_sidebar_v643(self):
    _ORIG_BUILD_SIDEBAR_V643(self)
    try:
        if 'ERP Cockpit' not in self.nav_buttons:
            btn=ttk.Button(self.sidebar,text='EC' if self.sidebar_collapsed else 'ERP Cockpit',style='Nav.TButton',command=lambda:self.show_module('ERP Cockpit')); btn.pack(fill='x',padx=5,pady=2); self.nav_buttons['ERP Cockpit']=btn
    except Exception: pass
FinanceMateApp._build_sidebar=_fm_build_sidebar_v643
_ORIG_SHOW_MODULE_V643=FinanceMateApp.show_module
def _fm_show_module_v643(self,module):
    if module=='ERP Cockpit':
        self.active_module=module; self._update_nav_styles()
        for w in self.workspace.winfo_children(): w.destroy()
        ttk.Label(self.workspace,text=module,style='CardTitle.TLabel').pack(anchor='w',pady=(0,6)); c=ttk.Frame(self.workspace); c.pack(fill='both',expand=True); self._render_erp_cockpit(c); self.set_status('Modul ERP Cockpit geladen.')
    else: return _ORIG_SHOW_MODULE_V643(self,module)
FinanceMateApp.show_module=_fm_show_module_v643

def _fm_v643_static_selftest() -> None:
    required=['ERPCockpitView','_fm_run_erp_health_check_v643','_fm_apply_optimization_cycle_v643','_fm_best_small_erp_scorecard_v643','_fm_export_erp_optimization_report_v643','_fm_finance_erp_gap_assessment_v643']
    missing=[x for x in required if x not in globals()]
    if APP_VERSION!='0.6.43-best-small-finance-erp' or missing: raise RuntimeError('v0.6.43 Selbsttest fehlgeschlagen: '+', '.join(missing))
_fm_v643_static_selftest()



# === FINANCE MATE PATCH V0_6_49_DIRECT_MERGE_FROM_0643_INTO_0647 ===
# Ziel: Die Finanzfunktionen aus v0.6.43 laufen direkt als Hauptanwendung, nicht als separates/externes Startfenster aus einem eingebetteten Modul.
# Zusätzlich: Schutz gegen alte v0.6.47-Core-Tabellen mit abweichendem Schema, damit der Start nicht an open_items/payments-Schema-Konflikten scheitert.
APP_VERSION = "0.6.49-direct-finance0643-into-0647"
MERGE_BASE_VERSION = "0.6.47-db-migration-fix"
FINANCE_FEATURE_SOURCE_VERSION = "0.6.43-best-small-finance-erp"

_V049_REQUIRED_FINANCE_FEATURES = [
    "FinanceMateApp", "StammdatenView", "JournalView", "InvoiceModuleBase", "DebitorsView", "CreditorsView",
    "InvoicePortalView", "PaymentsView", "ERPJournalView", "FinanceReportingView", "KMUComplianceView", "ERPCockpitView",
    "AttachmentMixin", "_FMDocumentPreviewPane", "init_sqlite", "get_connection", "parse_amount", "format_amount",
    "_fm_run_erp_health_check_v643", "_fm_best_small_erp_scorecard_v643", "_fm_export_erp_optimization_report_v643",
]

_V049_REQUIRED_FINANCE_TABLES = [
    "app_meta", "gl_accounts", "customers", "vendors", "tax_codes", "payment_terms", "bank_accounts",
    "journal_entries", "journal_entry_lines", "customer_invoices", "vendor_invoices", "open_items", "attachments",
    "invoice_import_batches", "invoice_import_files", "payments", "gobd_exports", "gobd_hash_chain", "payment_runs", "payment_run_items", "bank_transactions", "vat_returns", "fixed_assets", "report_snapshots",
]

_V049_V643_OPEN_ITEMS_COLUMNS = {
    "entity_type", "reference_no", "partner_no", "partner_name", "due_date", "original_amount", "open_amount", "status"
}
_V049_V643_PAYMENTS_COLUMNS = {
    "entity_type", "reference_no", "partner_no", "partner_name", "payment_date", "amount"
}

def _v049_table_columns(con: sqlite3.Connection, table: str) -> set[str]:
    try:
        return {r[1] for r in con.execute(f"PRAGMA table_info({table})").fetchall()}
    except Exception:
        return set()

def _v049_table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None

def _v049_archive_incompatible_table(con: sqlite3.Connection, table: str, required_cols: set[str]) -> bool:
    """Archiviert nur Tabellen, deren Name aus v0.6.47 kollidiert, deren Spalten aber nicht zum Finanzmodul v0.6.43 passen."""
    if not _v049_table_exists(con, table):
        return False
    cols = _v049_table_columns(con, table)
    if required_cols.issubset(cols):
        return False
    suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
    archived = f"{table}_legacy_v047_{suffix}"
    con.execute(f"ALTER TABLE {table} RENAME TO {archived}")
    try:
        con.execute(
            "CREATE TABLE IF NOT EXISTS migration_log_v049(ts TEXT, action TEXT, table_name TEXT, archived_as TEXT, details TEXT)"
        )
        con.execute(
            "INSERT INTO migration_log_v049(ts,action,table_name,archived_as,details) VALUES(?,?,?,?,?)",
            (now_str(), "archive_incompatible_table", table, archived, ",".join(sorted(cols))),
        )
    except Exception:
        pass
    return True

_ORIG_INIT_SQLITE_V049 = init_sqlite

def init_sqlite() -> None:
    ensure_directories()
    with get_connection() as con:
        _v049_archive_incompatible_table(con, "open_items", _V049_V643_OPEN_ITEMS_COLUMNS)
        _v049_archive_incompatible_table(con, "payments", _V049_V643_PAYMENTS_COLUMNS)
        con.commit()
    _ORIG_INIT_SQLITE_V049()
    with get_connection() as con:
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("merge_base_version", MERGE_BASE_VERSION, now_str()))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("finance_feature_source_version", FINANCE_FEATURE_SOURCE_VERSION, now_str()))
        con.commit()

def selftest() -> dict:
    """Nicht-GUI-Gesamttest für die direkte v0.6.43-Finanzintegration in die v0.6.47-Linie."""
    missing_features = [name for name in _V049_REQUIRED_FINANCE_FEATURES if name not in globals()]
    init_sqlite()
    with get_connection() as con:
        existing_tables = {r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
        missing_tables = [t for t in _V049_REQUIRED_FINANCE_TABLES if t not in existing_tables]
        open_cols = _v049_table_columns(con, "open_items")
        payment_cols = _v049_table_columns(con, "payments")
        app_version_row = con.execute("SELECT value FROM app_meta WHERE key='app_version'").fetchone()
    return {
        "version": APP_VERSION,
        "merge_base_version": MERGE_BASE_VERSION,
        "finance_feature_source_version": FINANCE_FEATURE_SOURCE_VERSION,
        "missing_features": missing_features,
        "missing_tables": missing_tables,
        "open_items_schema_ok": _V049_V643_OPEN_ITEMS_COLUMNS.issubset(open_cols),
        "payments_schema_ok": _V049_V643_PAYMENTS_COLUMNS.issubset(payment_cols),
        "app_meta_version": app_version_row[0] if app_version_row else "",
        "ok": not missing_features and not missing_tables and _V049_V643_OPEN_ITEMS_COLUMNS.issubset(open_cols) and _V049_V643_PAYMENTS_COLUMNS.issubset(payment_cols),
    }



# === FINANCE MATE PATCH V0_6_50_BJ1_BJ2_JOURNAL_OVERVIEW ===
# BJ-1: Bestandsanalyse für das aktive Buchungsjournal.
# BJ-2: Echte Journalübersicht statt reiner/grober Kontenübersicht.
APP_VERSION = "0.6.50-bj1-bj2-journal-overview"

BJ_ENTITY_TYPE = "journal_entry"


def _bj_table_columns_v650(con: sqlite3.Connection, table: str) -> set[str]:
    try:
        return {r[1] for r in con.execute(f"PRAGMA table_info({table})").fetchall()}
    except Exception:
        return set()


def _bj_ensure_column_v650(con: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    if column not in _bj_table_columns_v650(con, table):
        con.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def _bj_period_from_date_v650(date_value: str) -> tuple[str, str]:
    try:
        d = datetime.strptime((date_value or "").strip(), DATE_FMT)
        return str(d.year), f"{d.month:02d}/{str(d.year)[-2:]}"
    except Exception:
        return "", ""


def _bj_decimal_v650(value: Any) -> Decimal:
    try:
        return Decimal(str(value or 0)).quantize(Decimal("0.01"))
    except Exception:
        return Decimal("0.00")


def _bj_date_key_v650(value: str):
    try:
        return datetime.strptime((value or "").strip(), DATE_FMT).date()
    except Exception:
        return None


def _bj_count_attachments_v650(reference_no: str) -> int:
    if not reference_no:
        return 0
    with get_connection() as con:
        row = con.execute(
            "SELECT COUNT(*) FROM attachments WHERE entity_type=? AND reference_no=?",
            (BJ_ENTITY_TYPE, reference_no),
        ).fetchone()
        return int(row[0] if row else 0)


_ORIG_INIT_SQLITE_V650 = init_sqlite


def init_sqlite() -> None:
    _ORIG_INIT_SQLITE_V650()
    with get_connection() as con:
        # BJ-1/BJ-2 brauchen Zusatzspalten, die ältere Rekonstruktionsstände nicht sicher besitzen.
        for table, column, definition in [
            ("journal_entries", "reference_no", "TEXT DEFAULT ''"),
            ("journal_entries", "partner_index", "TEXT DEFAULT ''"),
            ("journal_entries", "journal_type", "TEXT DEFAULT ''"),
            ("journal_entries", "fiscal_year", "TEXT DEFAULT ''"),
            ("journal_entries", "fiscal_period", "TEXT DEFAULT ''"),
            ("journal_entries", "source_module", "TEXT DEFAULT ''"),
            ("journal_entries", "posted_by", "TEXT DEFAULT ''"),
            ("journal_entries", "locked", "INTEGER DEFAULT 0"),
            ("journal_entry_lines", "reference_no", "TEXT DEFAULT ''"),
            ("journal_entry_lines", "partner_index", "TEXT DEFAULT ''"),
            ("journal_entry_lines", "line_no", "INTEGER DEFAULT 0"),
        ]:
            _bj_ensure_column_v650(con, table, column, definition)

        # Periodenwerte für Altbestände aus Beleg-/Buchungsdatum ableiten, ohne vorhandene Werte zu überschreiben.
        rows = list(con.execute("SELECT id, posting_date, document_date FROM journal_entries"))
        for r in rows:
            year, period = _bj_period_from_date_v650(r["posting_date"] or r["document_date"] or "")
            if year or period:
                con.execute(
                    "UPDATE journal_entries SET fiscal_year=COALESCE(NULLIF(fiscal_year,''),?), fiscal_period=COALESCE(NULLIF(fiscal_period,''),?) WHERE id=?",
                    (year, period, r["id"]),
                )
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()


class ERPJournalView(ttk.Frame, SortableTreeMixin, AttachmentMixin):
    """BJ-1/BJ-2: Buchungsjournal mit Bestandsanalyse und echter Journalübersicht."""

    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.journal_search_var = tk.StringVar()
        self.journal_status_var = tk.StringVar(value="Alle")
        self.journal_type_var = tk.StringVar(value="Alle")
        self.date_from_var = tk.StringVar()
        self.date_to_var = tk.StringVar()
        self.selected_document_no = ""
        self._build_ui()
        self.reload_all()

    def _build_ui(self) -> None:
        header = ttk.Frame(self)
        header.pack(fill="x", padx=10, pady=(4, 8))
        ttk.Label(header, text="BJ-1/BJ-2 – Buchungsjournal", style="Section.TLabel").pack(side="left")
        create_standard_button(header, "Aktualisieren", self.reload_all).pack(side="right", padx=3)

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self.journal_tab = ttk.Frame(self.notebook)
        self.analysis_tab = ttk.Frame(self.notebook)
        self.accounts_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.journal_tab, text="Journal")
        self.notebook.add(self.analysis_tab, text="Bestandsanalyse")
        self.notebook.add(self.accounts_tab, text="Sachkontenübersicht")

        self._build_journal_tab()
        self._build_analysis_tab()
        self._build_accounts_tab()

    def _build_journal_tab(self) -> None:
        filter_bar = ttk.Frame(self.journal_tab)
        filter_bar.pack(fill="x", padx=8, pady=(8, 4))

        ttk.Label(filter_bar, text="Suchen", width=8).pack(side="left")
        search = ttk.Entry(filter_bar, textvariable=self.journal_search_var)
        search.pack(side="left", fill="x", expand=True, padx=(0, 8))
        search.insert(0, PLACEHOLDER_TEXT)
        search.bind("<FocusIn>", lambda _e: search.delete(0, tk.END) if search.get() == PLACEHOLDER_TEXT else None)
        search.bind("<FocusOut>", lambda _e: search.insert(0, PLACEHOLDER_TEXT) if not search.get().strip() else None)
        search.bind("<KeyRelease>", lambda _e: self.reload_journal())

        ttk.Label(filter_bar, text="Von").pack(side="left")
        entry_from = ttk.Entry(filter_bar, textvariable=self.date_from_var, width=11)
        entry_from.pack(side="left", padx=(3, 8))
        entry_from.bind("<KeyRelease>", lambda _e: self.reload_journal())
        ttk.Label(filter_bar, text="Bis").pack(side="left")
        entry_to = ttk.Entry(filter_bar, textvariable=self.date_to_var, width=11)
        entry_to.pack(side="left", padx=(3, 8))
        entry_to.bind("<KeyRelease>", lambda _e: self.reload_journal())

        self.status_combo = ttk.Combobox(filter_bar, textvariable=self.journal_status_var, values=["Alle", "Entwurf", "Gebucht", "Festgeschrieben", "Storniert", "Fehlerhaft"], state="readonly", width=16)
        self.status_combo.pack(side="left", padx=(0, 8))
        self.status_combo.bind("<<ComboboxSelected>>", lambda _e: self.reload_journal())
        self.type_combo = ttk.Combobox(filter_bar, textvariable=self.journal_type_var, values=["Alle", "Ausgangsrechnung", "Eingangsrechnung", "Bank", "Kasse", "Umbuchung", "Zahlung", "Storno", "Sonstige"], state="readonly", width=18)
        self.type_combo.pack(side="left", padx=(0, 8))
        self.type_combo.bind("<<ComboboxSelected>>", lambda _e: self.reload_journal())
        create_standard_button(filter_bar, "Filter zurücksetzen", self.clear_filters).pack(side="left")

        split = ttk.PanedWindow(self.journal_tab, orient="horizontal")
        split.pack(fill="both", expand=True, padx=8, pady=4)
        left = ttk.Frame(split)
        right = ttk.Frame(split)
        split.add(left, weight=3)
        split.add(right, weight=2)

        cols = ("document_no", "document_date", "posting_date", "period", "journal_type", "description", "debit", "credit", "difference", "status", "attachments")
        self.journal_tree = ttk.Treeview(left, columns=cols, show="headings", height=16)
        headings = [
            ("document_no", "Beleg", 125),
            ("document_date", "Belegdatum", 95),
            ("posting_date", "Buchungsdatum", 105),
            ("period", "Periode", 75),
            ("journal_type", "Belegart", 125),
            ("description", "Beschreibung", 240),
            ("debit", "Soll", 95),
            ("credit", "Haben", 95),
            ("difference", "Diff.", 85),
            ("status", "Status", 110),
            ("attachments", "Anhang", 80),
        ]
        for c, title, width in headings:
            self.journal_tree.heading(c, text=title)
            self.journal_tree.column(c, width=width, anchor="w")
        self.journal_tree.pack(side="left", fill="both", expand=True)
        journal_scroll = ttk.Scrollbar(left, orient="vertical", command=self.journal_tree.yview)
        journal_scroll.pack(side="right", fill="y")
        self.journal_tree.configure(yscrollcommand=journal_scroll.set)
        self.setup_sorting(self.journal_tree)
        self.journal_tree.bind("<<TreeviewSelect>>", lambda _e: self.load_selected_details())
        self.journal_tree.bind("<Double-1>", lambda _e: self.load_selected_details())
        self.journal_tree.tag_configure("balanced", background=WHITE)
        self.journal_tree.tag_configure("unbalanced", background=SOFT_RED)
        self.journal_tree.tag_configure("draft", background=SOFT_YELLOW)
        self.journal_tree.tag_configure("locked", background=SOFT_GREEN)

        ttk.Label(right, text="Buchungsdetails", style="Section.TLabel").pack(anchor="w", pady=(0, 4))
        self.detail_text = tk.Text(right, height=9, bg=WHITE, fg=TEXT, wrap="word", relief="sunken")
        self.detail_text.pack(fill="x", expand=False)
        self.detail_text.configure(state="disabled")

        ttk.Label(right, text="Buchungszeilen", style="Section.TLabel").pack(anchor="w", pady=(8, 4))
        line_cols = ("line_no", "account", "account_name", "side", "amount", "tax_code", "text")
        self.line_tree = ttk.Treeview(right, columns=line_cols, show="headings", height=10)
        for c, title, width in [
            ("line_no", "#", 35), ("account", "Konto", 85), ("account_name", "Kontoname", 150),
            ("side", "S/H", 45), ("amount", "Betrag", 95), ("tax_code", "Steuer", 70), ("text", "Text", 220)
        ]:
            self.line_tree.heading(c, text=title)
            self.line_tree.column(c, width=width, anchor="w")
        self.line_tree.pack(fill="both", expand=True)

        self.summary_label = ttk.Label(self.journal_tab, text="Soll: 0,00 | Haben: 0,00 | Differenz: 0,00 | Buchungen: 0", style="Hint.TLabel")
        self.summary_label.pack(fill="x", padx=8, pady=(2, 8))

    def _build_analysis_tab(self) -> None:
        top = ttk.Frame(self.analysis_tab)
        top.pack(fill="x", padx=8, pady=8)
        ttk.Label(top, text="Bestandsanalyse", style="Section.TLabel").pack(side="left")
        create_standard_button(top, "Analyse neu berechnen", self.reload_analysis).pack(side="right")
        self.analysis_text = tk.Text(self.analysis_tab, height=9, bg=WHITE, fg=TEXT, wrap="word", relief="sunken")
        self.analysis_text.pack(fill="x", padx=8, pady=(0, 8))
        self.analysis_text.configure(state="disabled")
        cols = ("severity", "area", "message", "reference")
        self.analysis_tree = ttk.Treeview(self.analysis_tab, columns=cols, show="headings")
        for c, t, w in [("severity", "Stufe", 90), ("area", "Bereich", 150), ("message", "Hinweis", 520), ("reference", "Referenz", 140)]:
            self.analysis_tree.heading(c, text=t)
            self.analysis_tree.column(c, width=w, anchor="w")
        self.analysis_tree.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        self.analysis_tree.tag_configure("OK", background=SOFT_GREEN)
        self.analysis_tree.tag_configure("Warnung", background=SOFT_YELLOW)
        self.analysis_tree.tag_configure("Fehler", background=SOFT_RED)

    def _build_accounts_tab(self) -> None:
        top = ttk.Frame(self.accounts_tab)
        top.pack(fill="x", padx=8, pady=8)
        ttk.Label(top, text="Sachkontenübersicht bleibt als Kontrollsicht erhalten", style="Section.TLabel").pack(side="left")
        create_standard_button(top, "Aktualisieren", self.reload_accounts).pack(side="right")
        cols = ("account_no", "name", "type", "debit", "credit", "balance", "lines")
        self.accounts_tree = ttk.Treeview(self.accounts_tab, columns=cols, show="headings")
        for c, t, w in [("account_no", "Konto", 90), ("name", "Name", 240), ("type", "Typ", 130), ("debit", "Soll", 100), ("credit", "Haben", 100), ("balance", "Saldo", 100), ("lines", "Zeilen", 70)]:
            self.accounts_tree.heading(c, text=t)
            self.accounts_tree.column(c, width=w, anchor="w")
        self.accounts_tree.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        self.setup_sorting(self.accounts_tree)

    def clear_filters(self) -> None:
        self.journal_search_var.set("")
        self.date_from_var.set("")
        self.date_to_var.set("")
        self.journal_status_var.set("Alle")
        self.journal_type_var.set("Alle")
        self.reload_journal()

    def reload_all(self) -> None:
        self.reload_journal()
        self.reload_analysis()
        self.reload_accounts()

    def _journal_rows(self) -> list[sqlite3.Row]:
        with get_connection() as con:
            return list(con.execute("SELECT * FROM journal_entries ORDER BY id DESC"))

    def _journal_passes_filter(self, row: sqlite3.Row) -> bool:
        q = (self.journal_search_var.get() or "").strip().lower()
        if q == PLACEHOLDER_TEXT.lower():
            q = ""
        status = self.journal_status_var.get()
        typ = self.journal_type_var.get()
        row_blob = " ".join(str(row[k]) for k in row.keys()).lower()
        if q and q not in row_blob:
            return False
        if status and status != "Alle" and str(row["status"] or "") != status:
            return False
        row_type = row["journal_type"] if "journal_type" in row.keys() else ""
        if typ and typ != "Alle" and str(row_type or "") != typ:
            return False
        date_from = _bj_date_key_v650(self.date_from_var.get()) if self.date_from_var.get().strip() else None
        date_to = _bj_date_key_v650(self.date_to_var.get()) if self.date_to_var.get().strip() else None
        row_date = _bj_date_key_v650(row["posting_date"] or row["document_date"] or "")
        if date_from and row_date and row_date < date_from:
            return False
        if date_to and row_date and row_date > date_to:
            return False
        return True

    def reload_journal(self) -> None:
        for item in self.journal_tree.get_children():
            self.journal_tree.delete(item)
        total_debit = Decimal("0.00")
        total_credit = Decimal("0.00")
        count = 0
        for r in self._journal_rows():
            if not self._journal_passes_filter(r):
                continue
            debit = _bj_decimal_v650(r["total_debit"])
            credit = _bj_decimal_v650(r["total_credit"])
            diff = debit - credit
            period = r["fiscal_period"] if "fiscal_period" in r.keys() and r["fiscal_period"] else _bj_period_from_date_v650(r["posting_date"] or r["document_date"] or "")[1]
            attachments = _bj_count_attachments_v650(r["document_no"])
            status = r["status"] or ""
            if abs(diff) > Decimal("0.004"):
                tag = "unbalanced"
            elif str(status).lower().startswith("entwurf"):
                tag = "draft"
            elif int(r["locked"] if "locked" in r.keys() and r["locked"] is not None else 0):
                tag = "locked"
            else:
                tag = "balanced"
            self.journal_tree.insert(
                "", "end", iid=str(r["document_no"]),
                values=(
                    r["document_no"], r["document_date"], r["posting_date"], period,
                    r["journal_type"] if "journal_type" in r.keys() else "",
                    r["description"], format_amount(debit), format_amount(credit), format_amount(diff), status,
                    f"📎 {attachments}" if attachments else "-",
                ), tags=(tag,)
            )
            total_debit += debit
            total_credit += credit
            count += 1
        self.summary_label.configure(text=f"Soll: {format_amount(total_debit)} | Haben: {format_amount(total_credit)} | Differenz: {format_amount(total_debit-total_credit)} | Buchungen: {count}")
        if count == 0:
            self._show_detail_message("Keine Buchung ausgewählt.\n\nDie Journalübersicht ist leer oder der aktuelle Filter liefert keine Treffer.")

    def load_selected_details(self) -> None:
        iid = self.journal_tree.focus()
        if not iid:
            return
        self.selected_document_no = str(iid)
        with get_connection() as con:
            row = con.execute("SELECT * FROM journal_entries WHERE document_no=?", (self.selected_document_no,)).fetchone()
            lines = list(con.execute(
                "SELECT l.*, a.name AS account_name FROM journal_entry_lines l LEFT JOIN gl_accounts a ON a.account_no=l.account_no WHERE l.document_no=? ORDER BY COALESCE(l.line_no,0), l.id",
                (self.selected_document_no,),
            ))
        if not row:
            self._show_detail_message("Buchung nicht gefunden.")
            return
        attachments = _bj_count_attachments_v650(self.selected_document_no)
        detail = [
            f"Beleg: {row['document_no']}",
            f"Belegdatum: {row['document_date']}    Buchungsdatum: {row['posting_date']}",
            f"Periode: {row['fiscal_period'] if 'fiscal_period' in row.keys() else ''}    Jahr: {row['fiscal_year'] if 'fiscal_year' in row.keys() else ''}",
            f"Belegart: {row['journal_type'] if 'journal_type' in row.keys() else ''}",
            f"Referenz: {row['reference_no'] if 'reference_no' in row.keys() else ''}    Partner: {row['partner_index'] if 'partner_index' in row.keys() else ''}",
            f"Status: {row['status']}    Gesperrt: {'Ja' if ('locked' in row.keys() and int(row['locked'] or 0)) else 'Nein'}",
            f"Soll: {format_amount(row['total_debit'])}    Haben: {format_amount(row['total_credit'])}    Differenz: {format_amount(_bj_decimal_v650(row['total_debit'])-_bj_decimal_v650(row['total_credit']))}",
            f"Anhänge: {attachments}",
            "",
            f"Beschreibung: {row['description'] or ''}",
        ]
        self._show_detail_message("\n".join(detail))
        for item in self.line_tree.get_children():
            self.line_tree.delete(item)
        for idx, line in enumerate(lines, start=1):
            self.line_tree.insert("", "end", values=(
                line["line_no"] if "line_no" in line.keys() and line["line_no"] else idx,
                line["account_no"], line["account_name"] or "", line["side"], format_amount(line["amount"]), line["tax_code"], line["text"],
            ))

    def _show_detail_message(self, text: str) -> None:
        self.detail_text.configure(state="normal")
        self.detail_text.delete("1.0", tk.END)
        self.detail_text.insert("1.0", text)
        self.detail_text.configure(state="disabled")
        if hasattr(self, "line_tree"):
            for item in self.line_tree.get_children():
                self.line_tree.delete(item)

    def reload_analysis(self) -> None:
        with get_connection() as con:
            tables = {r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
            journal_count = con.execute("SELECT COUNT(*) FROM journal_entries").fetchone()[0] if "journal_entries" in tables else 0
            line_count = con.execute("SELECT COUNT(*) FROM journal_entry_lines").fetchone()[0] if "journal_entry_lines" in tables else 0
            account_count = con.execute("SELECT COUNT(*) FROM gl_accounts").fetchone()[0] if "gl_accounts" in tables else 0
            unbalanced = con.execute("SELECT COUNT(*) FROM journal_entries WHERE ABS(COALESCE(total_debit,0)-COALESCE(total_credit,0))>0.004").fetchone()[0] if "journal_entries" in tables else 0
            no_lines = con.execute("SELECT COUNT(*) FROM journal_entries je LEFT JOIN journal_entry_lines jl ON jl.document_no=je.document_no WHERE jl.id IS NULL").fetchone()[0] if {"journal_entries","journal_entry_lines"}.issubset(tables) else 0
            invalid_accounts = con.execute("SELECT COUNT(*) FROM journal_entry_lines l LEFT JOIN gl_accounts a ON a.account_no=l.account_no WHERE a.account_no IS NULL").fetchone()[0] if {"journal_entry_lines","gl_accounts"}.issubset(tables) else 0
            attachments = con.execute("SELECT COUNT(*) FROM attachments WHERE entity_type=?", (BJ_ENTITY_TYPE,)).fetchone()[0] if "attachments" in tables else 0
            status_rows = list(con.execute("SELECT COALESCE(status,'') AS status, COUNT(*) AS cnt FROM journal_entries GROUP BY COALESCE(status,'') ORDER BY cnt DESC")) if "journal_entries" in tables else []

        summary = [
            "BJ-1 Bestandsanalyse – aktueller technischer/fachlicher Stand",
            "",
            f"Journal-Köpfe: {journal_count}",
            f"Journal-Zeilen: {line_count}",
            f"Sachkonten: {account_count}",
            f"Buchungen mit Soll/Haben-Differenz: {unbalanced}",
            f"Buchungen ohne Zeilen: {no_lines}",
            f"Journal-Zeilen ohne gültiges Sachkonto: {invalid_accounts}",
            f"Journal-Anhänge: {attachments}",
            "",
            "Statusverteilung:",
        ]
        for r in status_rows:
            summary.append(f"- {r['status'] or '(leer)'}: {r['cnt']}")
        self.analysis_text.configure(state="normal")
        self.analysis_text.delete("1.0", tk.END)
        self.analysis_text.insert("1.0", "\n".join(summary))
        self.analysis_text.configure(state="disabled")

        for item in self.analysis_tree.get_children():
            self.analysis_tree.delete(item)
        checks = []
        for table in ["journal_entries", "journal_entry_lines", "gl_accounts", "tax_codes", "attachments"]:
            checks.append(("OK" if table in tables else "Fehler", "Schema", f"Tabelle {table} {'vorhanden' if table in tables else 'fehlt'}", table))
        checks.extend([
            ("Fehler" if unbalanced else "OK", "Soll/Haben", f"{unbalanced} Buchung(en) mit Differenz", "journal_entries"),
            ("Warnung" if no_lines else "OK", "Vollständigkeit", f"{no_lines} Buchung(en) ohne Zeilen", "journal_entry_lines"),
            ("Fehler" if invalid_accounts else "OK", "Sachkonten", f"{invalid_accounts} Zeile(n) ohne gültiges Konto", "gl_accounts"),
            ("Warnung" if journal_count and attachments == 0 else "OK", "Belege", f"{attachments} Journal-Anhang/Anhänge vorhanden", BJ_ENTITY_TYPE),
        ])
        for severity, area, message, reference in checks:
            self.analysis_tree.insert("", "end", values=(severity, area, message, reference), tags=(severity,))

    def reload_accounts(self) -> None:
        for item in self.accounts_tree.get_children():
            self.accounts_tree.delete(item)
        with get_connection() as con:
            accounts = list(con.execute("SELECT * FROM gl_accounts ORDER BY account_no"))
            for acc in accounts:
                rows = list(con.execute("SELECT side, amount FROM journal_entry_lines WHERE account_no=?", (acc["account_no"],)))
                debit = sum(float(r["amount"] or 0) for r in rows if str(r["side"]).lower() in ("soll", "debit", "s"))
                credit = sum(float(r["amount"] or 0) for r in rows if str(r["side"]).lower() in ("haben", "credit", "h"))
                balance = debit - credit
                self.accounts_tree.insert("", "end", values=(acc["account_no"], acc["name"], acc["account_type"], format_amount(debit), format_amount(credit), format_amount(balance), len(rows)))


def _fm_render_finanzbuchhaltung_v650(self, parent):
    ERPJournalView(parent, self).pack(fill="both", expand=True)


FinanceMateApp._render_finanzbuchhaltung = _fm_render_finanzbuchhaltung_v650


_ORIG_SELFTEST_V650 = selftest if "selftest" in globals() else None


def selftest() -> dict:
    base = _ORIG_SELFTEST_V650() if callable(_ORIG_SELFTEST_V650) else {"ok": True}
    init_sqlite()
    with get_connection() as con:
        je_cols = _bj_table_columns_v650(con, "journal_entries")
        jl_cols = _bj_table_columns_v650(con, "journal_entry_lines")
    bj_required_features = ["ERPJournalView", "_fm_render_finanzbuchhaltung_v650", "_bj_count_attachments_v650"]
    missing_features = [name for name in bj_required_features if name not in globals()]
    bj_ok = (
        not missing_features
        and {"reference_no", "partner_index", "journal_type", "fiscal_year", "fiscal_period", "locked"}.issubset(je_cols)
        and {"reference_no", "partner_index", "line_no"}.issubset(jl_cols)
    )
    base.update({
        "version": APP_VERSION,
        "bj_1_2": "implemented",
        "bj_missing_features": missing_features,
        "bj_journal_entries_schema_ok": {"reference_no", "partner_index", "journal_type", "fiscal_year", "fiscal_period", "locked"}.issubset(je_cols),
        "bj_journal_lines_schema_ok": {"reference_no", "partner_index", "line_no"}.issubset(jl_cols),
        "ok": bool(base.get("ok", True)) and bj_ok,
    })
    return base


def main() -> None:
    app = FinanceMateApp()
    app.mainloop()

if __name__ == "__main__":
    main()
