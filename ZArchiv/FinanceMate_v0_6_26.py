
# -*- coding: utf-8 -*-
"""
Finance Mate – rekonstruierte lokale Desktop-App
Stand: v0.6.16-rekonstruiert

Rekonstruktionsbasis: Projektprotokoll FinanceMate – Rekonstruktionsmaster V4 / Patch-für-Patch-Systemreferenz.
Technik: Python, tkinter/ttk, SQLite, lokale Dateipersistenz.

Hinweis: Diese Datei ist als konsolidierte, startfähige Rekonstruktion aufgebaut. Sie bildet die im
Projektprotokoll beschriebene Fachsemantik in einer bereinigten Ein-Datei-Struktur ab.
"""

from __future__ import annotations

import os
import sys
import shutil
import sqlite3
import mimetypes
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from dataclasses import dataclass
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

APP_NAME = "Finance Mate"
APP_VERSION = "0.6.16-rekonstruiert"
STANDARD_DESIGN_NAME = "Reliefgrau"
DATE_FMT = "%d.%m.%Y"

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
ATTACHMENTS_DIR = BASE_DIR / "attachments"
IMPORTS_DIR = BASE_DIR / "imports"
DB_PATH = DATA_DIR / "finance_mate.sqlite3"

# Farben / Design
BG = "#eef1f4"
HEADER = "#dde3ea"
FOOTER = "#d8dee6"
LINE = "#b8c0ca"
TEXT = "#1f2933"
TEXT2 = "#5d6672"
WHITE = "#ffffff"
CARD_BG = "#f9fafb"
CARD_BORDER = "#cbd3dc"
BUTTON_GREEN = "#cfead1"
BUTTON_GREEN_ACTIVE = "#b9dfbd"
STANDARD_BUTTON_BG = "#d9dde3"
STANDARD_BUTTON_ACTIVE = "#cbd1d8"
STANDARD_BUTTON_BORDER = "#aeb7c2"
PLACEHOLDER_TEXT = "Suchen…"
PLACEHOLDER_COLOR = "#8a929c"
SOFT_GREEN = "#dff3e2"
SOFT_YELLOW = "#fff4c2"
SOFT_ORANGE = "#ffe1bf"
SOFT_RED = "#ffd4d4"

STATUS_OPEN = "Offen"
STATUS_PARTIAL = "Teilweise bezahlt"
STATUS_PAID = "Bezahlt"
STATUS_OVERDUE = "Überfällig"
STATUS_STORNO = "Storno"
PAYMENT_STATUS_VALUES = [STATUS_OPEN, STATUS_PARTIAL, STATUS_PAID, STATUS_OVERDUE, STATUS_STORNO]
TAX_SCOPE_VALUES = ["Inland", "EU", "Drittland"]
TRADE_COUNTRIES = ["Deutschland", "Österreich", "Schweiz", "Frankreich", "Italien", "Spanien", "Niederlande", "Belgien", "Polen", "Sonstige"]
SUPPORTED_ARCHIVE_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".txt", ".csv", ".doc", ".docx", ".xls", ".xlsx"}


def ensure_directories() -> None:
    for p in (DATA_DIR, ATTACHMENTS_DIR, IMPORTS_DIR):
        p.mkdir(parents=True, exist_ok=True)


def get_connection() -> sqlite3.Connection:
    ensure_directories()
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys = ON")
    return con


def table_has_column(con: sqlite3.Connection, table: str, column: str) -> bool:
    rows = con.execute(f"PRAGMA table_info({table})").fetchall()
    return any(r[1] == column for r in rows)


def ensure_column(con: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    if not table_has_column(con, table, column):
        con.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def init_sqlite() -> None:
    ensure_directories()
    with get_connection() as con:
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS app_meta (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS gl_accounts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                account_no TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                account_type TEXT DEFAULT '',
                tax_code TEXT DEFAULT '',
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS customers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                customer_no TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                country TEXT DEFAULT 'Deutschland',
                street TEXT DEFAULT '', zip TEXT DEFAULT '', city TEXT DEFAULT '',
                email TEXT DEFAULT '', phone TEXT DEFAULT '', tax_id TEXT DEFAULT '', vat_id TEXT DEFAULT '',
                iban TEXT DEFAULT '', bic TEXT DEFAULT '', bank_name TEXT DEFAULT '',
                payment_term TEXT DEFAULT '', active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS vendors (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                vendor_no TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                country TEXT DEFAULT 'Deutschland',
                street TEXT DEFAULT '', zip TEXT DEFAULT '', city TEXT DEFAULT '',
                email TEXT DEFAULT '', phone TEXT DEFAULT '', tax_id TEXT DEFAULT '', vat_id TEXT DEFAULT '',
                iban TEXT DEFAULT '', bic TEXT DEFAULT '', bank_name TEXT DEFAULT '',
                payment_term TEXT DEFAULT '', active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS tax_codes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                rate REAL DEFAULT 0,
                tax_scope TEXT DEFAULT 'Inland',
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS payment_terms (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                days INTEGER DEFAULT 0,
                discount_days INTEGER DEFAULT 0,
                discount_percent REAL DEFAULT 0,
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS bank_accounts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                iban TEXT DEFAULT '', bic TEXT DEFAULT '', bank_name TEXT DEFAULT '',
                active INTEGER DEFAULT 1,
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS journal_entries (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_no TEXT UNIQUE NOT NULL,
                document_date TEXT NOT NULL,
                posting_date TEXT NOT NULL,
                description TEXT DEFAULT '',
                total_debit REAL DEFAULT 0,
                total_credit REAL DEFAULT 0,
                status TEXT DEFAULT 'Gebucht',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS journal_entry_lines (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_no TEXT NOT NULL,
                account_no TEXT NOT NULL,
                side TEXT NOT NULL,
                amount REAL NOT NULL,
                tax_code TEXT DEFAULT '',
                text TEXT DEFAULT '',
                FOREIGN KEY(document_no) REFERENCES journal_entries(document_no) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS customer_invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_no TEXT UNIQUE NOT NULL,
                customer_no TEXT NOT NULL,
                customer_name TEXT NOT NULL,
                customer_address TEXT DEFAULT '',
                invoice_date TEXT NOT NULL,
                due_date TEXT NOT NULL,
                payment_term TEXT DEFAULT '',
                tax_code TEXT DEFAULT '',
                net_amount REAL DEFAULT 0,
                tax_amount REAL DEFAULT 0,
                gross_amount REAL DEFAULT 0,
                open_amount REAL DEFAULT 0,
                status TEXT DEFAULT 'Offen',
                linked_journal_no TEXT DEFAULT '',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS vendor_invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_no TEXT UNIQUE NOT NULL,
                vendor_no TEXT NOT NULL,
                vendor_name TEXT NOT NULL,
                vendor_address TEXT DEFAULT '',
                invoice_date TEXT NOT NULL,
                due_date TEXT NOT NULL,
                payment_term TEXT DEFAULT '',
                tax_code TEXT DEFAULT '',
                net_amount REAL DEFAULT 0,
                tax_amount REAL DEFAULT 0,
                gross_amount REAL DEFAULT 0,
                open_amount REAL DEFAULT 0,
                status TEXT DEFAULT 'Offen',
                linked_journal_no TEXT DEFAULT '',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS open_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                entity_type TEXT NOT NULL,
                reference_no TEXT NOT NULL,
                partner_no TEXT NOT NULL,
                partner_name TEXT NOT NULL,
                due_date TEXT NOT NULL,
                original_amount REAL NOT NULL,
                open_amount REAL NOT NULL,
                status TEXT NOT NULL,
                linked_journal_no TEXT DEFAULT '',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS attachments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                entity_type TEXT NOT NULL,
                reference_no TEXT NOT NULL,
                file_name TEXT NOT NULL,
                file_path TEXT NOT NULL,
                added_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS invoice_import_batches (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                hist_no TEXT UNIQUE NOT NULL,
                title TEXT NOT NULL,
                vendor_no TEXT DEFAULT '',
                vendor_name TEXT DEFAULT '',
                status TEXT DEFAULT 'Importiert',
                marked INTEGER DEFAULT 0,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS invoice_import_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                hist_no TEXT NOT NULL,
                file_name TEXT NOT NULL,
                file_path TEXT NOT NULL,
                mime_type TEXT DEFAULT '',
                created_at TEXT NOT NULL,
                FOREIGN KEY(hist_no) REFERENCES invoice_import_batches(hist_no) ON DELETE CASCADE
            );
            """
        )
        # defensive schema compatibility
        for table, col, definition in [
            ("customer_invoices", "customer_address", "TEXT DEFAULT ''"),
            ("customer_invoices", "linked_journal_no", "TEXT DEFAULT ''"),
            ("vendor_invoices", "vendor_address", "TEXT DEFAULT ''"),
            ("vendor_invoices", "linked_journal_no", "TEXT DEFAULT ''"),
            ("open_items", "linked_journal_no", "TEXT DEFAULT ''"),
            ("tax_codes", "tax_scope", "TEXT DEFAULT 'Inland'"),
        ]:
            ensure_column(con, table, col, definition)
        now = now_str()
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now))
        seed_defaults(con)
        con.commit()


def seed_defaults(con: sqlite3.Connection) -> None:
    now = now_str()
    defaults = [
        ("gl_accounts", "account_no", "1000", {"name": "Kasse", "account_type": "Aktiv"}),
        ("gl_accounts", "account_no", "1200", {"name": "Bank", "account_type": "Aktiv"}),
        ("gl_accounts", "account_no", "1400", {"name": "Forderungen", "account_type": "Aktiv"}),
        ("gl_accounts", "account_no", "1600", {"name": "Verbindlichkeiten", "account_type": "Passiv"}),
        ("gl_accounts", "account_no", "8400", {"name": "Erlöse", "account_type": "Ertrag"}),
        ("gl_accounts", "account_no", "3400", {"name": "Wareneingang", "account_type": "Aufwand"}),
    ]
    for table, keycol, keyval, vals in defaults:
        exists = con.execute(f"SELECT 1 FROM {table} WHERE {keycol}=?", (keyval,)).fetchone()
        if not exists:
            cols = [keycol] + list(vals.keys()) + ["created_at", "updated_at"]
            params = [keyval] + list(vals.values()) + [now, now]
            con.execute(f"INSERT INTO {table}({','.join(cols)}) VALUES({','.join(['?']*len(cols))})", params)
    for code, name, rate in [("V19", "Vorsteuer/Umsatzsteuer 19%", 19), ("V7", "Vorsteuer/Umsatzsteuer 7%", 7), ("O0", "Ohne Steuer", 0)]:
        if not con.execute("SELECT 1 FROM tax_codes WHERE code=?", (code,)).fetchone():
            con.execute("INSERT INTO tax_codes(code,name,rate,tax_scope,created_at,updated_at) VALUES(?,?,?,?,?,?)", (code, name, rate, "Inland", now, now))
    for code, name, days in [("SOFORT", "Sofort fällig", 0), ("14T", "14 Tage netto", 14), ("30T", "30 Tage netto", 30)]:
        if not con.execute("SELECT 1 FROM payment_terms WHERE code=?", (code,)).fetchone():
            con.execute("INSERT INTO payment_terms(code,name,days,created_at,updated_at) VALUES(?,?,?,?,?)", (code, name, days, now, now))


def now_str() -> str:
    return datetime.now().strftime("%d.%m.%Y %H:%M:%S")


def today_str() -> str:
    return datetime.now().strftime(DATE_FMT)


def validate_date(value: str) -> bool:
    try:
        datetime.strptime(value.strip(), DATE_FMT)
        return True
    except Exception:
        return False


def parse_amount(value: Any) -> Decimal:
    if value is None or value == "":
        return Decimal("0.00")
    s = str(value).strip().replace("€", "").replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    else:
        s = s.replace(",", ".")
    try:
        return Decimal(s).quantize(Decimal("0.01"))
    except InvalidOperation:
        raise ValueError(f"Ungültiger Betrag: {value}")


def format_amount(value: Any) -> str:
    try:
        d = Decimal(str(value)).quantize(Decimal("0.01"))
    except Exception:
        d = Decimal("0.00")
    return f"{d:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def calc_due_date(date_str: str, payment_term_code: str) -> str:
    base = datetime.strptime(date_str, DATE_FMT)
    days = 0
    with get_connection() as con:
        row = con.execute("SELECT days FROM payment_terms WHERE code=?", (payment_term_code,)).fetchone()
        if row:
            days = int(row[0] or 0)
    return (base + timedelta(days=days)).strftime(DATE_FMT)


def compute_status_from_open_amount(open_amount: Any, due_date: str) -> str:
    amount = parse_amount(open_amount)
    if amount <= 0:
        return STATUS_PAID
    if validate_date(due_date) and datetime.strptime(due_date, DATE_FMT).date() < datetime.now().date():
        return STATUS_OVERDUE
    return STATUS_OPEN


def urgency_bucket(due_date: str, open_amount: Any) -> str:
    if parse_amount(open_amount) <= 0:
        return "paid"
    if not validate_date(due_date):
        return "open"
    d = datetime.strptime(due_date, DATE_FMT).date()
    delta = (d - datetime.now().date()).days
    if delta < 0:
        return "overdue"
    if delta <= 7:
        return "soon"
    return "open"


def normalize_tax_code(code: str) -> str:
    return (code or "").strip().upper()


def yes_no_to_int(v: Any) -> int:
    return 1 if str(v).lower() in ("1", "ja", "true", "aktiv", "yes") else 0


def int_to_yes_no(v: Any) -> str:
    return "Ja" if int(v or 0) else "Nein"


def generate_number(prefix: str, meta_key: str, width: int = 5) -> str:
    with get_connection() as con:
        row = con.execute("SELECT value FROM app_meta WHERE key=?", (meta_key,)).fetchone()
        current = int(row[0]) if row else 0
        nxt = current + 1
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", (meta_key, str(nxt), now_str()))
        con.commit()
    return f"{prefix}{nxt:0{width}d}"


def next_hist_no() -> str:
    return generate_number("HIST-", "counter_hist", 6)


def clean_document_title(path: str) -> str:
    return Path(path).stem.replace("_", " ").replace("-", " ").strip() or Path(path).name


def iter_supported_files_from_folder(folder: str) -> Iterable[str]:
    for root, _dirs, files in os.walk(folder):
        for f in files:
            p = Path(root) / f
            if p.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
                yield str(p)


def collect_preview_text(path: str, limit: int = 8000) -> str:
    p = Path(path)
    if not p.exists():
        return "Datei nicht gefunden."
    if p.suffix.lower() in {".txt", ".csv"}:
        try:
            return p.read_text(encoding="utf-8", errors="replace")[:limit]
        except Exception as exc:
            return f"Textvorschau nicht möglich: {exc}"
    return f"Dokumentkarte\n\nDatei: {p.name}\nTyp: {mimetypes.guess_type(str(p))[0] or p.suffix}\nPfad: {p}\n\nFür PDF-/Office-Dateien wird eine Dateikarte angezeigt. Öffnen über Doppelklick oder Kontextmenü."


def open_path(path: str) -> None:
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as exc:
        messagebox.showerror("Öffnen nicht möglich", str(exc))


def apply_search_placeholder(entry: tk.Entry, placeholder: str = PLACEHOLDER_TEXT) -> None:
    def on_focus_in(_event=None):
        if entry.get() == placeholder and entry.cget("fg") == PLACEHOLDER_COLOR:
            entry.delete(0, tk.END)
            entry.config(fg=TEXT)
    def on_focus_out(_event=None):
        if not entry.get().strip():
            entry.insert(0, placeholder)
            entry.config(fg=PLACEHOLDER_COLOR)
    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)
    on_focus_out()


def create_standard_button(parent, text: str, command=None, confirm: bool = False, **kw):
    style = "Confirm.TButton" if confirm else "StandardGray.TButton"
    return ttk.Button(parent, text=text, command=command, style=style, **kw)


def configure_tree_tags(tree: ttk.Treeview) -> None:
    tree.tag_configure("paid", background=SOFT_GREEN)
    tree.tag_configure("soon", background=SOFT_YELLOW)
    tree.tag_configure("overdue", background=SOFT_RED)
    tree.tag_configure("open", background=WHITE)


def load_ui_preference(key: str, default: str = "") -> str:
    with get_connection() as con:
        row = con.execute("SELECT value FROM app_meta WHERE key=?", (f"ui:{key}",)).fetchone()
        return row[0] if row else default


def save_ui_preference(key: str, value: str) -> None:
    with get_connection() as con:
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", (f"ui:{key}", value, now_str()))
        con.commit()


class ScrollableFrame(ttk.Frame):
    def __init__(self, parent, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.canvas = tk.Canvas(self, bg=BG, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview, style="Vertical.TScrollbar")
        self.content = ttk.Frame(self.canvas)
        self.content_id = self.canvas.create_window((0, 0), window=self.content, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        self.content.bind("<Configure>", self._on_content_configure)
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfigure(self.content_id, width=e.width))
        self._bind_mousewheel(self.canvas)
        self._bind_descendant_wheel(self.content)

    def _on_content_configure(self, _event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _bind_mousewheel(self, widget):
        widget.bind("<Enter>", lambda e: self.canvas.bind_all("<MouseWheel>", self._on_mousewheel), add="+")
        widget.bind("<Leave>", lambda e: self.canvas.unbind_all("<MouseWheel>"), add="+")

    def _bind_descendant_wheel(self, widget):
        # Wheel always scrolls frame; comboboxes explicitly do not change values by wheel.
        def bind_child(child):
            if isinstance(child, ttk.Combobox):
                child.bind("<MouseWheel>", lambda e: "break")
            else:
                child.bind("<MouseWheel>", self._on_mousewheel, add="+")
            for sub in child.winfo_children():
                bind_child(sub)
        self.after(100, lambda: bind_child(widget))

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        return "break"


class SortableTreeMixin:
    def _coerce_sort_value(self, value: str):
        if value is None:
            return ""
        s = str(value)
        try:
            return float(s.replace(".", "").replace(",", "."))
        except Exception:
            pass
        if validate_date(s):
            return datetime.strptime(s, DATE_FMT)
        return s.lower()

    def setup_sorting(self, tree: ttk.Treeview) -> None:
        for col in tree["columns"]:
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

    def sort_treeview(self, tree: ttk.Treeview, col: str, reverse: bool) -> None:
        data = [(self._coerce_sort_value(tree.set(k, col)), k) for k in tree.get_children("")]
        data.sort(reverse=reverse)
        for idx, (_val, k) in enumerate(data):
            tree.move(k, "", idx)
        tree.heading(col, command=lambda: self.sort_treeview(tree, col, not reverse))


class AttachmentMixin:
    def add_attachment_paths(self, entity_type: str, reference_no: str, paths: Sequence[str]) -> None:
        if not reference_no:
            messagebox.showwarning("Anhang", "Bitte zuerst einen Datensatz speichern oder auswählen.")
            return
        target_dir = ATTACHMENTS_DIR / entity_type / reference_no
        target_dir.mkdir(parents=True, exist_ok=True)
        with get_connection() as con:
            for p in paths:
                src = Path(p)
                if not src.exists():
                    continue
                dest = target_dir / src.name
                if dest.exists():
                    dest = target_dir / f"{src.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{src.suffix}"
                shutil.copy2(src, dest)
                con.execute(
                    "INSERT INTO attachments(entity_type,reference_no,file_name,file_path,added_at) VALUES(?,?,?,?,?)",
                    (entity_type, reference_no, dest.name, str(dest), now_str()),
                )
            con.commit()

    def get_attachment_count(self, entity_type: str, reference_no: str) -> int:
        with get_connection() as con:
            row = con.execute("SELECT COUNT(*) FROM attachments WHERE entity_type=? AND reference_no=?", (entity_type, reference_no)).fetchone()
            return int(row[0] if row else 0)

    def get_attachments(self, entity_type: str, reference_no: str) -> List[sqlite3.Row]:
        with get_connection() as con:
            return list(con.execute("SELECT * FROM attachments WHERE entity_type=? AND reference_no=? ORDER BY added_at DESC", (entity_type, reference_no)).fetchall())

    def open_attachment_popup(self, entity_type: str, reference_no: str) -> None:
        pop = tk.Toplevel()
        pop.title(f"Anhänge – {reference_no}")
        pop.geometry("900x520")
        pop.configure(bg=BG)
        ttk.Label(pop, text=f"Anhänge zu {entity_type} / {reference_no}", style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
        cols = ("file_name", "file_path", "added_at")
        tree = ttk.Treeview(pop, columns=cols, show="headings")
        for col, w, txt in [("file_name", 220, "Dateiname"), ("file_path", 480, "Pfad"), ("added_at", 150, "Zeitstempel")]:
            tree.heading(col, text=txt)
            tree.column(col, width=w, anchor="w")
        tree.pack(fill="both", expand=True, padx=12, pady=6)
        for r in self.get_attachments(entity_type, reference_no):
            tree.insert("", "end", values=(r["file_name"], r["file_path"], r["added_at"]))
        tree.bind("<Double-1>", lambda e: open_path(tree.set(tree.focus(), "file_path")) if tree.focus() else None)
        bar = ttk.Frame(pop)
        bar.pack(fill="x", padx=12, pady=10)
        create_standard_button(bar, "Datei hinzufügen", lambda: self._attachment_popup_add(entity_type, reference_no, tree), confirm=True).pack(side="left")
        create_standard_button(bar, "Schließen", pop.destroy).pack(side="right")

    def _attachment_popup_add(self, entity_type: str, reference_no: str, tree: ttk.Treeview):
        paths = filedialog.askopenfilenames(title="Anhänge auswählen")
        if paths:
            self.add_attachment_paths(entity_type, reference_no, paths)
            for i in tree.get_children(): tree.delete(i)
            for r in self.get_attachments(entity_type, reference_no):
                tree.insert("", "end", values=(r["file_name"], r["file_path"], r["added_at"]))


class TwoBlock(ttk.Frame):
    """Zentrale Zweiblock-Engine: freie Größenänderung ohne starre Wand."""
    def __init__(self, parent, left_title: str, right_title: str):
        super().__init__(parent)
        self.pane = ttk.PanedWindow(self, orient="horizontal")
        self.pane.pack(fill="both", expand=True)
        self.left = ttk.Frame(self.pane, style="Card.TFrame")
        self.right = ttk.Frame(self.pane, style="Card.TFrame")
        self.pane.add(self.left, weight=3)
        self.pane.add(self.right, weight=2)
        ttk.Label(self.left, text=left_title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))
        ttk.Label(self.right, text=right_title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))


def build_two_block_shell(parent, left_title: str, right_title: str) -> TwoBlock:
    tb = TwoBlock(parent, left_title, right_title)
    tb.pack(fill="both", expand=True)
    return tb


class StammdatenView(ttk.Frame, SortableTreeMixin):
    CONFIG = {
        "Sachkonten": {
            "table": "gl_accounts", "key": "account_no", "label": "Sachkonto",
            "fields": [("account_no", "Konto-Nr."), ("name", "Name"), ("account_type", "Kontotyp"), ("tax_code", "Steuerkennz."), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("account_no", 100, "Konto"), ("name", 220, "Name"), ("account_type", 120, "Typ"), ("tax_code", 100, "Steuer"), ("active", 70, "Aktiv")],
            "prefix": "SK"
        },
        "Debitoren": {
            "table": "customers", "key": "customer_no", "label": "Debitor",
            "fields": [("customer_no", "Debitor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"), ("email", "E-Mail"), ("phone", "Telefon"), ("tax_id", "Steuer-Nr."), ("vat_id", "USt-IdNr."), ("iban", "IBAN"), ("bic", "BIC"), ("bank_name", "Bank"), ("payment_term", "Zahlungsbed."), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("customer_no", 110, "Nr."), ("name", 220, "Name"), ("city", 140, "Ort"), ("payment_term", 110, "ZB"), ("active", 70, "Aktiv")],
            "prefix": "D"
        },
        "Kreditoren": {
            "table": "vendors", "key": "vendor_no", "label": "Kreditor",
            "fields": [("vendor_no", "Kreditor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"), ("email", "E-Mail"), ("phone", "Telefon"), ("tax_id", "Steuer-Nr."), ("vat_id", "USt-IdNr."), ("iban", "IBAN"), ("bic", "BIC"), ("bank_name", "Bank"), ("payment_term", "Zahlungsbed."), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("vendor_no", 110, "Nr."), ("name", 220, "Name"), ("city", 140, "Ort"), ("payment_term", 110, "ZB"), ("active", 70, "Aktiv")],
            "prefix": "K"
        },
        "Steuerkennzeichen": {
            "table": "tax_codes", "key": "code", "label": "Steuerkennzeichen",
            "fields": [("code", "Code"), ("name", "Name"), ("rate", "Satz %"), ("tax_scope", "Gültigkeit"), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("code", 90, "Code"), ("name", 240, "Name"), ("rate", 90, "%"), ("tax_scope", 120, "Gültigkeit"), ("active", 70, "Aktiv")],
            "prefix": "T"
        },
        "Zahlungsbedingungen": {
            "table": "payment_terms", "key": "code", "label": "Zahlungsbedingung",
            "fields": [("code", "Code"), ("name", "Name"), ("days", "Tage netto"), ("discount_days", "Skontotage"), ("discount_percent", "Skonto %"), ("active", "Aktiv")],
            "columns": [("id", 50, "ID"), ("code", 90, "Code"), ("name", 240, "Name"), ("days", 90, "Tage"), ("discount_percent", 100, "Skonto %"), ("active", 70, "Aktiv")],
            "prefix": "ZB"
        },
    }

    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.vars: Dict[str, Dict[str, tk.StringVar]] = {}
        self.trees: Dict[str, ttk.Treeview] = {}
        self.search_vars: Dict[str, tk.StringVar] = {}
        self.selected_ids: Dict[str, Optional[int]] = {}
        self._build_ui()
        self.load_all_tables()

    def _build_ui(self):
        top = ttk.Frame(self)
        top.pack(fill="x", pady=(0, 6))
        ttk.Label(top, text="Stammdaten", style="CardTitle.TLabel").pack(side="left")
        create_standard_button(top, "Daten-Refresh", self.load_all_tables).pack(side="right")
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)
        for title in self.CONFIG:
            self._build_tab(title)
        self.nb.bind("<<NotebookTabChanged>>", lambda e: self.load_all_tables())

    def _build_tab(self, title: str):
        cfg = self.CONFIG[title]
        frame = ttk.Frame(self.nb)
        self.nb.add(frame, text=title)
        shell = build_two_block_shell(frame, f"{title} – Liste", f"{title} – Formular")
        # left
        left_body = ttk.Frame(shell.left)
        left_body.pack(fill="both", expand=True, padx=10, pady=8)
        search = tk.Entry(left_body, relief="sunken", bg=WHITE, fg=TEXT)
        search.pack(fill="x", pady=(0, 6))
        apply_search_placeholder(search)
        search.bind("<KeyRelease>", lambda e, t=title, s=search: self.load_table_data(t, s.get()))
        cols = [c[0] for c in cfg["columns"]]
        tree = ttk.Treeview(left_body, columns=cols, show="headings", selectmode="browse")
        for col, w, txt in cfg["columns"]:
            tree.heading(col, text=txt)
            tree.column(col, width=w, anchor="w")
        tree.pack(fill="both", expand=True)
        tree.bind("<<TreeviewSelect>>", lambda e, t=title: self.load_selected_record(t))
        self.setup_sorting(tree)
        self.trees[title] = tree
        self.selected_ids[title] = None
        # right
        sf = ScrollableFrame(shell.right)
        sf.pack(fill="both", expand=True, padx=10, pady=8)
        self.vars[title] = {}
        for field, label in cfg["fields"]:
            row = ttk.Frame(sf.content)
            row.pack(fill="x", pady=3)
            ttk.Label(row, text=label, width=18).pack(side="left")
            var = tk.StringVar(value="Ja" if field == "active" else "")
            self.vars[title][field] = var
            if field == "country":
                cb = ttk.Combobox(row, textvariable=var, values=TRADE_COUNTRIES, state="readonly")
                cb.pack(side="left", fill="x", expand=True)
                cb.bind("<MouseWheel>", lambda e: "break")
            elif field == "tax_scope":
                cb = ttk.Combobox(row, textvariable=var, values=TAX_SCOPE_VALUES, state="readonly")
                cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
            elif field == "payment_term":
                cb = ttk.Combobox(row, textvariable=var, values=self._load_payment_term_codes(), state="readonly")
                cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
            elif field == "active":
                cb = ttk.Combobox(row, textvariable=var, values=["Ja", "Nein"], state="readonly")
                cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
            else:
                ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)
        btns = ttk.Frame(sf.content)
        btns.pack(fill="x", pady=10)
        create_standard_button(btns, "Neu / Felder leeren", lambda t=title: self.clear_form(t)).pack(side="left", padx=3)
        create_standard_button(btns, "Speichern", lambda t=title: self.save_record(t), confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Löschen", lambda t=title: self.delete_record(t)).pack(side="left", padx=3)

    def _load_payment_term_codes(self):
        with get_connection() as con:
            return [r[0] for r in con.execute("SELECT code FROM payment_terms ORDER BY code")]

    def load_all_tables(self):
        for t in self.CONFIG:
            self.load_table_data(t)

    def load_table_data(self, title: str, search: str = ""):
        if search == PLACEHOLDER_TEXT: search = ""
        cfg = self.CONFIG[title]
        tree = self.trees[title]
        for i in tree.get_children(): tree.delete(i)
        cols = [c[0] for c in cfg["columns"]]
        sql = f"SELECT {','.join(cols)} FROM {cfg['table']}"
        params: List[Any] = []
        if search.strip():
            where_cols = [c for c in cols if c != "id"]
            sql += " WHERE " + " OR ".join([f"CAST({c} AS TEXT) LIKE ?" for c in where_cols])
            params = [f"%{search.strip()}%"] * len(where_cols)
        sql += " ORDER BY id DESC"
        with get_connection() as con:
            for r in con.execute(sql, params):
                vals = [int_to_yes_no(r[c]) if c == "active" else r[c] for c in cols]
                tree.insert("", "end", values=vals)

    def load_selected_record(self, title: str):
        tree = self.trees[title]
        iid = tree.focus()
        if not iid: return
        item = tree.item(iid)
        rec_id = item["values"][0]
        self.selected_ids[title] = int(rec_id)
        cfg = self.CONFIG[title]
        with get_connection() as con:
            row = con.execute(f"SELECT * FROM {cfg['table']} WHERE id=?", (rec_id,)).fetchone()
        if not row: return
        for field, var in self.vars[title].items():
            if field == "active": var.set(int_to_yes_no(row[field]))
            else: var.set(str(row[field] or ""))

    def clear_form(self, title: str):
        self.selected_ids[title] = None
        cfg = self.CONFIG[title]
        for field, var in self.vars[title].items():
            if field == "active": var.set("Ja")
            elif field == "country": var.set("Deutschland")
            elif field == cfg["key"]: var.set(generate_number(cfg["prefix"], f"counter_{cfg['table']}", 5))
            else: var.set("")

    def save_record(self, title: str):
        cfg = self.CONFIG[title]
        vals = {field: var.get().strip() for field, var in self.vars[title].items()}
        if not vals.get(cfg["key"]):
            vals[cfg["key"]] = generate_number(cfg["prefix"], f"counter_{cfg['table']}", 5)
        if not vals.get("name") and "name" in vals:
            messagebox.showwarning("Pflichtfeld", "Bitte einen Namen eingeben.")
            return
        if "active" in vals: vals["active"] = yes_no_to_int(vals["active"])
        now = now_str()
        with get_connection() as con:
            if self.selected_ids[title]:
                cols = list(vals.keys()) + ["updated_at"]
                params = list(vals.values()) + [now, self.selected_ids[title]]
                con.execute(f"UPDATE {cfg['table']} SET " + ",".join([f"{c}=?" for c in cols]) + " WHERE id=?", params)
            else:
                vals["created_at"] = now; vals["updated_at"] = now
                cols = list(vals.keys())
                con.execute(f"INSERT INTO {cfg['table']}({','.join(cols)}) VALUES({','.join(['?']*len(cols))})", list(vals.values()))
                self.selected_ids[title] = con.execute("SELECT last_insert_rowid()").fetchone()[0]
            con.commit()
        self.app.set_status(f"{title}: Datensatz gespeichert.")
        self.load_table_data(title)

    def delete_record(self, title: str):
        rec_id = self.selected_ids[title]
        if not rec_id:
            messagebox.showinfo("Löschen", "Bitte zuerst einen Datensatz auswählen.")
            return
        cfg = self.CONFIG[title]
        name = self.vars[title].get("name", tk.StringVar(value="")).get() or self.vars[title].get("code", tk.StringVar(value="")).get()
        if not messagebox.askyesno("Unwiderruflich löschen", f"[{rec_id}] [{name}] unwiderruflich löschen?"):
            return
        with get_connection() as con:
            con.execute(f"DELETE FROM {cfg['table']} WHERE id=?", (rec_id,)); con.commit()
        self.clear_form(title)
        self.load_table_data(title)
        self.app.set_status(f"{title}: Datensatz gelöscht.")


class JournalView(ttk.Frame, SortableTreeMixin, AttachmentMixin):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.pending_lines: List[Dict[str, Any]] = []
        self.pending_attachments: List[str] = []
        self.selected_doc = ""
        self._build_ui()
        self.reload_journal_entries()

    def _build_ui(self):
        shell = build_two_block_shell(self, "Buchung erfassen", "Historie / Kontrolle")
        left = ScrollableFrame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        self.doc_no = tk.StringVar(value=generate_number("B", "counter_journal", 6))
        self.doc_date = tk.StringVar(value=today_str()); self.posting_date = tk.StringVar(value=today_str()); self.desc = tk.StringVar()
        for label, var in [("Beleg-Nr.", self.doc_no), ("Belegdatum", self.doc_date), ("Buchungsdatum", self.posting_date), ("Beschreibung", self.desc)]:
            self._labeled_entry(left.content, label, var)
        ttk.Separator(left.content).pack(fill="x", pady=8)
        self.line_account = tk.StringVar(); self.line_side = tk.StringVar(value="Soll"); self.line_amount = tk.StringVar(); self.line_tax = tk.StringVar(); self.line_text = tk.StringVar()
        self._labeled_combo(left.content, "Konto", self.line_account, self._load_accounts())
        self._labeled_combo(left.content, "Soll/Haben", self.line_side, ["Soll", "Haben"])
        self._labeled_entry(left.content, "Betrag", self.line_amount)
        self._labeled_combo(left.content, "Steuerkennz.", self.line_tax, self._load_tax_codes())
        self._labeled_entry(left.content, "Text", self.line_text)
        btns = ttk.Frame(left.content); btns.pack(fill="x", pady=6)
        create_standard_button(btns, "Zeile hinzufügen", self.add_line_item, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Zeile entfernen", self.remove_selected_line).pack(side="left", padx=3)
        cols = ("account", "side", "amount", "tax", "text")
        self.lines_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=7)
        for c, t, w in [("account", "Konto", 110), ("side", "S/H", 70), ("amount", "Betrag", 100), ("tax", "Steuer", 90), ("text", "Text", 240)]:
            self.lines_tree.heading(c, text=t); self.lines_tree.column(c, width=w, anchor="w")
        self.lines_tree.pack(fill="x", pady=6)
        self.totals = ttk.Label(left.content, text="Soll 0,00 | Haben 0,00 | Differenz 0,00", style="Hint.TLabel")
        self.totals.pack(anchor="w", pady=5)
        btns2 = ttk.Frame(left.content); btns2.pack(fill="x", pady=10)
        create_standard_button(btns2, "Anhang vormerken", self.add_pending_attachments).pack(side="left", padx=3)
        create_standard_button(btns2, "Buchung speichern", self.save_journal_entry, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns2, "Neu", self.clear_journal_form).pack(side="left", padx=3)
        # right history
        body = ttk.Frame(shell.right); body.pack(fill="both", expand=True, padx=10, pady=8)
        cols = ("document_no", "posting_date", "description", "debit", "credit", "att")
        self.hist_tree = ttk.Treeview(body, columns=cols, show="headings")
        for c, t, w in [("document_no", "Beleg", 110), ("posting_date", "Datum", 90), ("description", "Beschreibung", 240), ("debit", "Soll", 90), ("credit", "Haben", 90), ("att", "Anh.", 55)]:
            self.hist_tree.heading(c, text=t); self.hist_tree.column(c, width=w, anchor="w")
        self.hist_tree.pack(fill="both", expand=True)
        self.hist_tree.bind("<Double-1>", self._on_history_click)
        create_standard_button(body, "Anhang zu ausgewähltem Beleg", self.add_attachment_to_selected_history).pack(anchor="e", pady=8)
        self.setup_sorting(self.hist_tree)

    def _labeled_entry(self, parent, label, var):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)

    def _labeled_combo(self, parent, label, var, values):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        cb = ttk.Combobox(row, textvariable=var, values=values)
        cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")

    def _load_accounts(self):
        with get_connection() as con: return [f"{r['account_no']} {r['name']}" for r in con.execute("SELECT account_no,name FROM gl_accounts WHERE active=1 ORDER BY account_no")]
    def _load_tax_codes(self):
        with get_connection() as con: return [r[0] for r in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]

    def add_line_item(self):
        try: amount = parse_amount(self.line_amount.get())
        except ValueError as exc: messagebox.showwarning("Betrag", str(exc)); return
        acc = self.line_account.get().split(" ")[0]
        if not acc or amount <= 0: messagebox.showwarning("Zeile", "Bitte Konto und positiven Betrag erfassen."); return
        self.pending_lines.append({"account": acc, "side": self.line_side.get(), "amount": float(amount), "tax": self.line_tax.get(), "text": self.line_text.get()})
        self._refresh_lines_tree(); self.clear_line_inputs()

    def remove_selected_line(self):
        iid = self.lines_tree.focus()
        if iid:
            idx = self.lines_tree.index(iid); self.pending_lines.pop(idx); self._refresh_lines_tree()

    def clear_line_inputs(self):
        self.line_account.set(""); self.line_side.set("Soll"); self.line_amount.set(""); self.line_tax.set(""); self.line_text.set("")

    def _refresh_lines_tree(self):
        for i in self.lines_tree.get_children(): self.lines_tree.delete(i)
        debit = Decimal("0"); credit = Decimal("0")
        for ln in self.pending_lines:
            self.lines_tree.insert("", "end", values=(ln["account"], ln["side"], format_amount(ln["amount"]), ln["tax"], ln["text"]))
            if ln["side"] == "Soll": debit += parse_amount(ln["amount"])
            else: credit += parse_amount(ln["amount"])
        diff = debit - credit
        self.totals.config(text=f"Soll {format_amount(debit)} | Haben {format_amount(credit)} | Differenz {format_amount(diff)}")

    def add_pending_attachments(self):
        paths = filedialog.askopenfilenames(title="Beleganhänge vormerken")
        if paths:
            self.pending_attachments.extend(paths)
            self.app.set_status(f"{len(paths)} Anhang/Anhänge vorgemerkt.")

    def save_journal_entry(self):
        if not validate_date(self.doc_date.get()) or not validate_date(self.posting_date.get()):
            messagebox.showwarning("Datum", "Datumsformat bitte TT.MM.JJJJ verwenden."); return
        if not self.pending_lines: messagebox.showwarning("Buchung", "Bitte mindestens eine Buchungszeile erfassen."); return
        debit = sum(parse_amount(ln["amount"]) for ln in self.pending_lines if ln["side"] == "Soll")
        credit = sum(parse_amount(ln["amount"]) for ln in self.pending_lines if ln["side"] == "Haben")
        if debit != credit:
            messagebox.showwarning("Nicht ausgeglichen", "Soll und Haben müssen identisch sein."); return
        with get_connection() as con:
            con.execute("INSERT INTO journal_entries(document_no,document_date,posting_date,description,total_debit,total_credit,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)",
                        (self.doc_no.get(), self.doc_date.get(), self.posting_date.get(), self.desc.get(), float(debit), float(credit), "Gebucht", now_str(), now_str()))
            for ln in self.pending_lines:
                con.execute("INSERT INTO journal_entry_lines(document_no,account_no,side,amount,tax_code,text) VALUES(?,?,?,?,?,?)", (self.doc_no.get(), ln["account"], ln["side"], ln["amount"], ln["tax"], ln["text"]))
            con.commit()
        if self.pending_attachments:
            self.add_attachment_paths("journal", self.doc_no.get(), self.pending_attachments)
        self.app.set_status(f"Buchung {self.doc_no.get()} gespeichert.")
        self.clear_journal_form(); self.reload_journal_entries()

    def clear_journal_form(self):
        self.doc_no.set(generate_number("B", "counter_journal", 6)); self.doc_date.set(today_str()); self.posting_date.set(today_str()); self.desc.set("")
        self.pending_lines.clear(); self.pending_attachments.clear(); self._refresh_lines_tree(); self.clear_line_inputs()

    def reload_journal_entries(self):
        for i in self.hist_tree.get_children(): self.hist_tree.delete(i)
        with get_connection() as con:
            for r in con.execute("SELECT * FROM journal_entries ORDER BY id DESC"):
                att = self.get_attachment_count("journal", r["document_no"])
                self.hist_tree.insert("", "end", values=(r["document_no"], r["posting_date"], r["description"], format_amount(r["total_debit"]), format_amount(r["total_credit"]), f"📎 {att}" if att else ""))

    def _on_history_click(self, _event=None):
        iid = self.hist_tree.focus()
        if iid:
            doc = self.hist_tree.set(iid, "document_no")
            self.open_attachment_popup("journal", doc)

    def add_attachment_to_selected_history(self):
        iid = self.hist_tree.focus()
        if not iid: messagebox.showinfo("Anhang", "Bitte Buchung auswählen."); return
        doc = self.hist_tree.set(iid, "document_no")
        paths = filedialog.askopenfilenames(title="Anhang auswählen")
        if paths:
            self.add_attachment_paths("journal", doc, paths); self.reload_journal_entries()


class InvoiceModuleBase(ttk.Frame, SortableTreeMixin, AttachmentMixin):
    entity_label = "Rechnung"
    partner_table = "customers"
    partner_no_col = "customer_no"
    invoice_table = "customer_invoices"
    entity_type = "customer_invoice"
    prefix = "AR"

    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.pending_attachments: List[str] = []
        self.selected_invoice = ""
        self._load_reference_data()
        self._build_ui()
        self.reload_invoices(); self.reload_open_items()

    def _load_reference_data(self):
        with get_connection() as con:
            self.partners = [f"{r[self.partner_no_col]} {r['name']}" for r in con.execute(f"SELECT {self.partner_no_col},name FROM {self.partner_table} WHERE active=1 ORDER BY name")]
            self.tax_codes = [r[0] for r in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]
            self.payment_terms = [r[0] for r in con.execute("SELECT code FROM payment_terms WHERE active=1 ORDER BY code")]

    def _build_ui(self):
        shell = build_two_block_shell(self, f"{self.entity_label} erfassen", "Offene Posten / Kontrolle")
        left = ScrollableFrame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
        self.partner = tk.StringVar(); self.invoice_date = tk.StringVar(value=today_str()); self.due_date = tk.StringVar(value=today_str())
        self.payment_term = tk.StringVar(); self.tax_code = tk.StringVar(value="V19"); self.net_amount = tk.StringVar(); self.tax_amount = tk.StringVar(); self.gross_amount = tk.StringVar()
        self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
        self._labeled_combo(left.content, "Partner", self.partner, self.partners)
        self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
        self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
        self._labeled_entry(left.content, "Fälligkeit", self.due_date)
        self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
        self._labeled_entry(left.content, "Netto", self.net_amount)
        self._labeled_entry(left.content, "Steuer", self.tax_amount)
        self._labeled_entry(left.content, "Brutto", self.gross_amount)
        self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
        self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
        self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
        btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
        create_standard_button(btns, "Anhang vormerken", self.add_pending_attachments).pack(side="left", padx=3)
        create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
        ttk.Label(left.content, text="Erfasste Rechnungen", style="Section.TLabel").pack(anchor="w", pady=(10, 4))
        cols = ("invoice_no", "partner", "date", "due", "gross", "open", "status", "att")
        self.inv_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=9)
        for c, t, w in [("invoice_no", "Rechnung", 110), ("partner", "Partner", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anh.", 55)]:
            self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
        self.inv_tree.pack(fill="both", expand=True)
        self.inv_tree.bind("<Double-1>", lambda e: self.add_attachment_to_selected_invoice())
        configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
        # right
        right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
        cols2 = ("reference_no", "partner", "due", "original", "open", "status")
        self.op_tree = ttk.Treeview(right, columns=cols2, show="headings")
        for c, t, w in [("reference_no", "Referenz", 120), ("partner", "Partner", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120)]:
            self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
        self.op_tree.pack(fill="both", expand=True)
        configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
        create_standard_button(right, "Anhang zu Rechnung", self.add_attachment_to_selected_invoice).pack(anchor="e", pady=8)

    def _generate_invoice_no(self): return generate_number(self.prefix, f"counter_{self.invoice_table}", 6)
    def _labeled_entry(self, parent, label, var):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3); ttk.Label(row, text=label, width=16).pack(side="left"); ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)
    def _labeled_combo(self, parent, label, var, values):
        row = ttk.Frame(parent); row.pack(fill="x", pady=3); ttk.Label(row, text=label, width=16).pack(side="left")
        cb = ttk.Combobox(row, textvariable=var, values=values); cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")

    def apply_payment_term(self):
        if validate_date(self.invoice_date.get()) and self.payment_term.get():
            try: self.due_date.set(calc_due_date(self.invoice_date.get(), self.payment_term.get()))
            except Exception: pass

    def _calculate_amounts(self):
        try: net = parse_amount(self.net_amount.get())
        except Exception: return
        rate = Decimal("0")
        with get_connection() as con:
            row = con.execute("SELECT rate FROM tax_codes WHERE code=?", (self.tax_code.get(),)).fetchone()
            if row: rate = Decimal(str(row[0] or 0))
        tax = (net * rate / Decimal("100")).quantize(Decimal("0.01")); gross = net + tax
        self.tax_amount.set(format_amount(tax)); self.gross_amount.set(format_amount(gross))

    def _partner_parts(self):
        no = self.partner.get().split(" ")[0] if self.partner.get() else ""
        with get_connection() as con:
            row = con.execute(f"SELECT * FROM {self.partner_table} WHERE {self.partner_no_col}=?", (no,)).fetchone()
        if row:
            address = f"{row['street']}, {row['zip']} {row['city']}, {row['country']}".strip(', ')
            return no, row["name"], address
        return no, self.partner.get(), ""

    def add_pending_attachments(self):
        paths = filedialog.askopenfilenames(title="Beleganhänge vormerken")
        if paths: self.pending_attachments.extend(paths); self.app.set_status(f"{len(paths)} Anhang/Anhänge vorgemerkt.")

    def save_invoice(self):
        if not validate_date(self.invoice_date.get()) or not validate_date(self.due_date.get()): messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden."); return
        no, name, address = self._partner_parts()
        if not no: messagebox.showwarning("Partner", "Bitte Partner auswählen."); return
        gross = parse_amount(self.gross_amount.get()); net = parse_amount(self.net_amount.get()); tax = parse_amount(self.tax_amount.get())
        status = compute_status_from_open_amount(gross, self.due_date.get())
        address_col = "customer_address" if self.invoice_table == "customer_invoices" else "vendor_address"
        partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
        partner_no_col = "customer_no" if self.invoice_table == "customer_invoices" else "vendor_no"
        with get_connection() as con:
            con.execute(f"INSERT INTO {self.invoice_table}(invoice_no,{partner_no_col},{partner_name_col},{address_col},invoice_date,due_date,payment_term,tax_code,net_amount,tax_amount,gross_amount,open_amount,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                        (self.invoice_no.get(), no, name, address, self.invoice_date.get(), self.due_date.get(), self.payment_term.get(), self.tax_code.get(), float(net), float(tax), float(gross), float(gross), status, now_str(), now_str()))
            con.execute("INSERT INTO open_items(entity_type,reference_no,partner_no,partner_name,due_date,original_amount,open_amount,status,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                        (self.entity_type, self.invoice_no.get(), no, name, self.due_date.get(), float(gross), float(gross), status, now_str(), now_str()))
            con.commit()
        if self.pending_attachments: self.add_attachment_paths(self.entity_type, self.invoice_no.get(), self.pending_attachments)
        self.app.set_status(f"{self.entity_label} {self.invoice_no.get()} gespeichert.")
        self.clear_form(); self.reload_invoices(); self.reload_open_items()

    def clear_form(self):
        self.invoice_no.set(self._generate_invoice_no()); self.partner.set(""); self.invoice_date.set(today_str()); self.due_date.set(today_str()); self.payment_term.set(""); self.tax_code.set("V19"); self.net_amount.set(""); self.tax_amount.set(""); self.gross_amount.set(""); self.pending_attachments.clear()

    def reload_invoices(self):
        for i in self.inv_tree.get_children(): self.inv_tree.delete(i)
        partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
        with get_connection() as con:
            for r in con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"):
                att = self.get_attachment_count(self.entity_type, r["invoice_no"])
                tag = urgency_bucket(r["due_date"], r["open_amount"])
                self.inv_tree.insert("", "end", values=(r["invoice_no"], r[partner_name_col], r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else ""), tags=(tag,))

    def reload_open_items(self):
        for i in self.op_tree.get_children(): self.op_tree.delete(i)
        with get_connection() as con:
            for r in con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date", (self.entity_type,)):
                tag = urgency_bucket(r["due_date"], r["open_amount"])
                self.op_tree.insert("", "end", values=(r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"]), tags=(tag,))

    def add_attachment_to_selected_invoice(self):
        iid = self.inv_tree.focus()
        if not iid: messagebox.showinfo("Anhang", "Bitte Rechnung auswählen."); return
        inv = self.inv_tree.set(iid, "invoice_no")
        paths = filedialog.askopenfilenames(title="Anhang auswählen")
        if paths: self.add_attachment_paths(self.entity_type, inv, paths); self.reload_invoices()


class DebitorsView(InvoiceModuleBase):
    entity_label = "Ausgangsrechnung"
    partner_table = "customers"; partner_no_col = "customer_no"; invoice_table = "customer_invoices"; entity_type = "customer_invoice"; prefix = "AR"


class CreditorsView(InvoiceModuleBase):
    entity_label = "Eingangsrechnung"
    partner_table = "vendors"; partner_no_col = "vendor_no"; invoice_table = "vendor_invoices"; entity_type = "vendor_invoice"; prefix = "ER"

    def _build_ui(self):
        outer = ttk.PanedWindow(self, orient="vertical"); outer.pack(fill="both", expand=True)
        invoice_area = ttk.Frame(outer); stack_area = ttk.Frame(outer)
        outer.add(invoice_area, weight=3); outer.add(stack_area, weight=2)
        # temporarily build base UI into invoice_area by monkey parent composition
        old_parent = self.master
        shell = build_two_block_shell(invoice_area, "Eingangsrechnung erfassen", "OP / Freigabe")
        left = ScrollableFrame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
        self.partner = tk.StringVar(); self.invoice_date = tk.StringVar(value=today_str()); self.due_date = tk.StringVar(value=today_str())
        self.payment_term = tk.StringVar(); self.tax_code = tk.StringVar(value="V19"); self.net_amount = tk.StringVar(); self.tax_amount = tk.StringVar(); self.gross_amount = tk.StringVar()
        self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
        prow = ttk.Frame(left.content); prow.pack(fill="x", pady=3)
        ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
        cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners); cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)
        self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
        self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
        self._labeled_entry(left.content, "Fälligkeit", self.due_date)
        self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
        self._labeled_entry(left.content, "Netto", self.net_amount); self._labeled_entry(left.content, "Steuer", self.tax_amount); self._labeled_entry(left.content, "Brutto", self.gross_amount)
        self.payment_term.trace_add("write", lambda *_: self.apply_payment_term()); self.net_amount.trace_add("write", lambda *_: self._calculate_amounts()); self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
        btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
        create_standard_button(btns, "Anhang vormerken", self.add_pending_attachments).pack(side="left", padx=3)
        create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
        ttk.Label(left.content, text="Eingangsrechnungen", style="Section.TLabel").pack(anchor="w", pady=(10, 4))
        cols = ("invoice_no", "partner", "date", "due", "gross", "open", "status", "att")
        self.inv_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=7)
        for c, t, w in [("invoice_no", "Rechnung", 110), ("partner", "Kreditor", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anh.", 55)]: self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
        self.inv_tree.pack(fill="both", expand=True); self.inv_tree.bind("<Double-1>", lambda e: self.add_attachment_to_selected_invoice()); configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
        right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
        cols2 = ("reference_no", "partner", "due", "original", "open", "status")
        self.op_tree = ttk.Treeview(right, columns=cols2, show="headings")
        for c, t, w in [("reference_no", "Referenz", 120), ("partner", "Kreditor", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120)]: self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
        self.op_tree.pack(fill="both", expand=True); configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
        create_standard_button(right, "Anhang zu Rechnung", self.add_attachment_to_selected_invoice).pack(anchor="e", pady=8)
        self._build_stack_ui(stack_area)

    def _open_vendor_quick_popup(self):
        pop = tk.Toplevel(self); pop.title("Quick-Kreditor"); pop.geometry("520x360"); pop.configure(bg=BG)
        vars = {k: tk.StringVar() for k in ["vendor_no", "name", "country", "street", "zip", "city", "payment_term"]}
        vars["vendor_no"].set(generate_number("K", "counter_vendors", 5)); vars["country"].set("Deutschland")
        for key, label in [("vendor_no", "Kreditor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"), ("payment_term", "Zahlungsbed.")]:
            row = ttk.Frame(pop); row.pack(fill="x", padx=12, pady=4); ttk.Label(row, text=label, width=16).pack(side="left"); ttk.Entry(row, textvariable=vars[key]).pack(side="left", fill="x", expand=True)
        def save():
            if not vars["name"].get().strip(): messagebox.showwarning("Pflichtfeld", "Name fehlt."); return
            with get_connection() as con:
                con.execute("INSERT INTO vendors(vendor_no,name,country,street,zip,city,payment_term,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)", (vars["vendor_no"].get(), vars["name"].get(), vars["country"].get(), vars["street"].get(), vars["zip"].get(), vars["city"].get(), vars["payment_term"].get(), 1, now_str(), now_str())); con.commit()
            self._load_reference_data(); self.partner.set(f"{vars['vendor_no'].get()} {vars['name'].get()}"); pop.destroy()
        create_standard_button(pop, "Speichern", save, confirm=True).pack(pady=12)

    def _build_stack_ui(self, parent):
        shell = build_two_block_shell(parent, "Dokumentenstapel", "Vorschau")
        left = ttk.Frame(shell.left); left.pack(fill="both", expand=True, padx=10, pady=8)
        bar = ttk.Frame(left); bar.pack(fill="x", pady=(0, 6))
        create_standard_button(bar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=2)
        create_standard_button(bar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=2)
        create_standard_button(bar, "Alle markieren", self._toggle_all_marks).pack(side="left", padx=2)
        create_standard_button(bar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=2)
        create_standard_button(bar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=2)
        cols = ("mark", "hist_no", "title", "status", "files", "created")
        self.stack_tree = ttk.Treeview(left, columns=cols, show="headings", selectmode="extended")
        for c, t, w in [("mark", "☐", 45), ("hist_no", "Hist.-Nr.", 120), ("title", "Titel", 260), ("status", "Status", 100), ("files", "Dateien", 70), ("created", "Importiert", 150)]: self.stack_tree.heading(c, text=t); self.stack_tree.column(c, width=w, anchor="w")
        self.stack_tree.pack(fill="both", expand=True); self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select); self.stack_tree.bind("<Double-1>", lambda e: self._open_selected_stack_file())
        right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
        self.preview = tk.Text(right, wrap="word", bg=WHITE, fg=TEXT, relief="sunken")
        self.preview.pack(fill="both", expand=True)
        self._load_import_batches()

    def import_creditor_files(self):
        paths = filedialog.askopenfilenames(title="Dokumente importieren", filetypes=[("Unterstützte Dateien", "*.pdf *.png *.jpg *.jpeg *.bmp *.gif *.txt *.csv *.doc *.docx *.xls *.xlsx"), ("Alle Dateien", "*.*")])
        if paths: self._import_archive_paths(paths)

    def import_creditor_folder(self):
        folder = filedialog.askdirectory(title="Importordner auswählen")
        if folder: self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

    def _import_archive_paths(self, paths: Sequence[str]):
        imported = 0
        with get_connection() as con:
            for src in paths:
                p = Path(src)
                if not p.exists() or p.suffix.lower() not in SUPPORTED_ARCHIVE_EXTENSIONS: continue
                hist = next_hist_no(); target_dir = IMPORTS_DIR / hist; target_dir.mkdir(parents=True, exist_ok=True)
                dest = target_dir / p.name; shutil.copy2(p, dest)
                title = clean_document_title(str(p))
                con.execute("INSERT INTO invoice_import_batches(hist_no,title,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?)", (hist, title, "Importiert", 0, now_str(), now_str()))
                con.execute("INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)", (hist, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()))
                imported += 1
            con.commit()
        self._load_import_batches(); self.app.set_status(f"{imported} Dokument(e) importiert.")

    def _load_import_batches(self):
        if not hasattr(self, "stack_tree"): return
        for i in self.stack_tree.get_children(): self.stack_tree.delete(i)
        with get_connection() as con:
            rows = con.execute("SELECT b.*, COUNT(f.id) files FROM invoice_import_batches b LEFT JOIN invoice_import_files f ON f.hist_no=b.hist_no GROUP BY b.hist_no ORDER BY b.id DESC").fetchall()
            for r in rows:
                self.stack_tree.insert("", "end", values=("☑" if r["marked"] else "☐", r["hist_no"], r["title"], r["status"], r["files"], r["created_at"]))

    def _on_stack_select(self, event=None):
        iid = self.stack_tree.focus()
        if not iid: return
        col = self.stack_tree.identify_column(event.x) if event else ""
        hist = self.stack_tree.set(iid, "hist_no")
        if col == "#1":
            with get_connection() as con:
                row = con.execute("SELECT marked FROM invoice_import_batches WHERE hist_no=?", (hist,)).fetchone(); new = 0 if row and row[0] else 1
                con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=? WHERE hist_no=?", (new, now_str(), hist)); con.commit()
            self._load_import_batches(); return
        self._refresh_stack_preview(hist)

    def _refresh_stack_preview(self, hist_no: str):
        with get_connection() as con: row = con.execute("SELECT * FROM invoice_import_files WHERE hist_no=? ORDER BY id LIMIT 1", (hist_no,)).fetchone()
        self.preview.delete("1.0", tk.END)
        if row: self.preview.insert("1.0", collect_preview_text(row["file_path"]))
        else: self.preview.insert("1.0", "Keine Datei im Stapel.")

    def _open_selected_stack_file(self):
        iid = self.stack_tree.focus()
        if not iid: return
        hist = self.stack_tree.set(iid, "hist_no")
        with get_connection() as con: row = con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id LIMIT 1", (hist,)).fetchone()
        if row: open_path(row["file_path"])

    def _toggle_all_marks(self):
        with get_connection() as con:
            any_unmarked = con.execute("SELECT 1 FROM invoice_import_batches WHERE marked=0 LIMIT 1").fetchone()
            con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=?", (1 if any_unmarked else 0, now_str())); con.commit()
        self._load_import_batches()

    def _delete_marked_batches(self):
        if not messagebox.askyesno("Bulk-Löschung", "Alle markierten Stapel unwiderruflich löschen?"): return
        with get_connection() as con:
            rows = con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1").fetchall()
            for r in rows:
                shutil.rmtree(IMPORTS_DIR / r["hist_no"], ignore_errors=True)
                con.execute("DELETE FROM invoice_import_batches WHERE hist_no=?", (r["hist_no"],))
            con.commit()
        self._load_import_batches(); self.preview.delete("1.0", tk.END)

    def merge_selected_batches(self):
        with get_connection() as con:
            rows = con.execute("SELECT hist_no,title FROM invoice_import_batches WHERE marked=1 ORDER BY id").fetchall()
            if len(rows) < 2: messagebox.showinfo("Zusammenführen", "Bitte mindestens zwei Stapel markieren."); return
            new_hist = next_hist_no(); target_dir = IMPORTS_DIR / new_hist; target_dir.mkdir(parents=True, exist_ok=True)
            title = " + ".join([r["title"] for r in rows])[:240]
            con.execute("INSERT INTO invoice_import_batches(hist_no,title,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?)", (new_hist, title, "Zusammengeführt", 0, now_str(), now_str()))
            for r in rows:
                for f in con.execute("SELECT * FROM invoice_import_files WHERE hist_no=?", (r["hist_no"],)).fetchall():
                    src = Path(f["file_path"]); dest = target_dir / f"{r['hist_no']}_{src.name}"
                    if src.exists(): shutil.copy2(src, dest)
                    con.execute("INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)", (new_hist, dest.name, str(dest), f["mime_type"], now_str()))
                con.execute("UPDATE invoice_import_batches SET status='Zusammengeführt in ' || ?, marked=0, updated_at=? WHERE hist_no=?", (new_hist, now_str(), r["hist_no"]))
            con.commit()
        self._load_import_batches(); self.app.set_status(f"Stapel {new_hist} erzeugt.")


class FinanceMateApp(tk.Tk):
    def __init__(self):
        super().__init__()
        init_sqlite()
        self.title(f"{APP_NAME} {APP_VERSION}")
        self.geometry("1420x860")
        self.minsize(1100, 680)
        self.configure(bg=BG)
        self.sidebar_collapsed = load_ui_preference("sidebar_collapsed", "0") == "1"
        self.active_module = "Dashboard"
        self.nav_buttons: Dict[str, ttk.Button] = {}
        self._configure_ttk()
        self._build_layout()
        self.show_module("Dashboard")

    def _configure_ttk(self):
        style = ttk.Style(self)
        try: style.theme_use("clam")
        except Exception: pass
        style.configure("TFrame", background=BG)
        style.configure("Card.TFrame", background=CARD_BG, relief="raised", borderwidth=1)
        style.configure("TLabel", background=BG, foreground=TEXT)
        style.configure("CardTitle.TLabel", background=BG, foreground=TEXT, font=("Segoe UI", 15, "bold"))
        style.configure("Section.TLabel", background=CARD_BG, foreground=TEXT, font=("Segoe UI", 10, "bold"))
        style.configure("Hint.TLabel", background=BG, foreground=TEXT2)
        style.configure("StandardGray.TButton", background=STANDARD_BUTTON_BG, foreground=TEXT, bordercolor=STANDARD_BUTTON_BORDER, relief="raised", padding=(8, 4))
        style.map("StandardGray.TButton", background=[("active", STANDARD_BUTTON_ACTIVE)])
        style.configure("Confirm.TButton", background=BUTTON_GREEN, foreground=TEXT, relief="raised", padding=(8, 4))
        style.map("Confirm.TButton", background=[("active", BUTTON_GREEN_ACTIVE)])
        style.configure("Nav.TButton", background=STANDARD_BUTTON_BG, anchor="w", padding=(8, 7), relief="raised")
        style.configure("NavActive.TButton", background="#c5d3e6", anchor="w", padding=(8, 7), relief="sunken")
        style.configure("Treeview.Heading", background=STANDARD_BUTTON_BG, relief="raised", foreground=TEXT)
        style.configure("Treeview", rowheight=24, fieldbackground=WHITE, background=WHITE)
        style.configure("TNotebook", background=BG)
        style.configure("TNotebook.Tab", background=STANDARD_BUTTON_BG, padding=(10, 5))
        style.map("TNotebook.Tab", background=[("selected", CARD_BG)])
        style.configure("Vertical.TScrollbar", background=STANDARD_BUTTON_BG, troughcolor="#edf0f3", arrowsize=14, width=15)
        style.configure("TCombobox", fieldbackground=WHITE, background=STANDARD_BUTTON_BG)

    def _build_layout(self):
        self.header = ttk.Frame(self, style="Card.TFrame"); self.header.pack(fill="x", padx=6, pady=(6, 3))
        ttk.Label(self.header, text=APP_NAME, style="CardTitle.TLabel").pack(side="left", padx=12, pady=7)
        ttk.Label(self.header, text=f"{STANDARD_DESIGN_NAME} · lokale SQLite-Persistenz · {APP_VERSION}", style="Hint.TLabel").pack(side="left", padx=10)
        create_standard_button(self.header, "Datenbankordner", lambda: open_path(str(DATA_DIR))).pack(side="right", padx=8)
        body = ttk.Frame(self); body.pack(fill="both", expand=True, padx=6, pady=3)
        self.sidebar = ttk.Frame(body, style="Card.TFrame"); self.sidebar.pack(side="left", fill="y", padx=(0, 5))
        self.workspace = ttk.Frame(body); self.workspace.pack(side="left", fill="both", expand=True)
        self._build_sidebar()
        self.footer = ttk.Frame(self, style="Card.TFrame"); self.footer.pack(fill="x", padx=6, pady=(3, 6))
        self.status_var = tk.StringVar(value="Bereit.")
        ttk.Label(self.footer, textvariable=self.status_var, style="Hint.TLabel").pack(side="left", padx=12, pady=5)
        ttk.Label(self.footer, text="Mini-Widgets: Änderung vorschlagen · [i] Hilfe", style="Hint.TLabel").pack(side="right", padx=12)

    def _build_sidebar(self):
        for w in self.sidebar.winfo_children(): w.destroy()
        width = 42 if self.sidebar_collapsed else 210
        self.sidebar.configure(width=width)
        self.sidebar.pack_propagate(False)
        arrow = "»" if self.sidebar_collapsed else "«"
        create_standard_button(self.sidebar, arrow, self._toggle_sidebar).pack(fill="x", padx=5, pady=5)
        modules = [("Dashboard", "DB"), ("Stammdaten", "ST"), ("Finanzbuchhaltung", "FB"), ("Debitoren", "DE"), ("Kreditoren", "KR"), ("Zahlungen", "ZA"), ("Reporting", "RE"), ("Audit", "AU"), ("Einstellungen", "EI")]
        self.nav_buttons.clear()
        for name, abbr in modules:
            text = abbr if self.sidebar_collapsed else name
            btn = ttk.Button(self.sidebar, text=text, style="Nav.TButton", command=lambda n=name: self.show_module(n))
            btn.pack(fill="x", padx=5, pady=2)
            self.nav_buttons[name] = btn

    def _toggle_sidebar(self):
        self.sidebar_collapsed = not self.sidebar_collapsed
        save_ui_preference("sidebar_collapsed", "1" if self.sidebar_collapsed else "0")
        self._build_sidebar(); self._update_nav_styles()

    def _update_nav_styles(self):
        for name, btn in self.nav_buttons.items(): btn.configure(style="NavActive.TButton" if name == self.active_module else "Nav.TButton")

    def set_status(self, msg: str):
        self.status_var.set(msg)

    def show_module(self, module: str):
        self.active_module = module; self._update_nav_styles()
        for w in self.workspace.winfo_children(): w.destroy()
        title = ttk.Label(self.workspace, text=module, style="CardTitle.TLabel"); title.pack(anchor="w", pady=(0, 6))
        container = ttk.Frame(self.workspace); container.pack(fill="both", expand=True)
        render = {
            "Dashboard": self._render_dashboard,
            "Stammdaten": self._render_stammdaten,
            "Finanzbuchhaltung": self._render_finanzbuchhaltung,
            "Debitoren": self._render_debitoren,
            "Kreditoren": self._render_kreditoren,
            "Zahlungen": self._render_zahlungen,
            "Reporting": self._render_reporting,
            "Audit": self._render_audit,
            "Einstellungen": self._render_einstellungen,
        }.get(module, self._render_dashboard)
        render(container); self.set_status(f"Modul {module} geladen.")

    def _card(self, parent, title, body):
        f = ttk.Frame(parent, style="Card.TFrame"); f.pack(fill="both", expand=True, padx=6, pady=6)
        ttk.Label(f, text=title, style="Section.TLabel").pack(anchor="w", padx=12, pady=(10, 4))
        ttk.Label(f, text=body, style="Hint.TLabel", wraplength=520).pack(anchor="nw", padx=12, pady=(0, 12))
        return f

    def _render_dashboard(self, parent):
        grid = ttk.Frame(parent); grid.pack(fill="both", expand=True)
        for i in range(2): grid.columnconfigure(i, weight=1); grid.rowconfigure(i, weight=1)
        for idx, (title, text) in enumerate([
            ("Systemstatus", "Alle Kernmodule sind über die Sidebar erreichbar. SQLite wird lokal initialisiert."),
            ("Offene Posten", "Debitoren- und Kreditoren-OPs werden nach Fälligkeit farblich gekennzeichnet."),
            ("Dokumentenstapel", "Kreditorenimporte mit Hist.-Nr., Markierung, Bulk-Löschung, Merge und Vorschau."),
            ("Nächste Ausbaustufe", "Zahlungen, Reporting und Audit sind als funktionsfähige Vorbereitungsflächen angelegt."),
        ]):
            f = ttk.Frame(grid, style="Card.TFrame"); f.grid(row=idx//2, column=idx%2, sticky="nsew", padx=6, pady=6)
            ttk.Label(f, text=title, style="Section.TLabel").pack(anchor="w", padx=12, pady=10)
            ttk.Label(f, text=text, style="Hint.TLabel", wraplength=520).pack(anchor="nw", padx=12, pady=4)

    def _render_stammdaten(self, parent): StammdatenView(parent, self).pack(fill="both", expand=True)
    def _render_finanzbuchhaltung(self, parent): JournalView(parent, self).pack(fill="both", expand=True)
    def _render_debitoren(self, parent): DebitorsView(parent, self).pack(fill="both", expand=True)
    def _render_kreditoren(self, parent): CreditorsView(parent, self).pack(fill="both", expand=True)

    def _render_zahlungen(self, parent):
        shell = build_two_block_shell(parent, "Zahlungserfassung / OP-Ausgleich", "Vorbereitung Zahlungsläufe")
        ttk.Label(shell.left, text="Teilzahlungen, Vollausgleich und OP-Abgleich sind im Datenmodell vorbereitet.\nDie OP-Tabellen aus Debitoren/Kreditoren dienen als Grundlage für Block 7.", style="Hint.TLabel", wraplength=650).pack(anchor="nw", padx=12, pady=12)
        ttk.Label(shell.right, text="Workflow-Platzhalter für spätere Zahlungsläufe, Zuordnungsvorschläge und Statuswechsel.", style="Hint.TLabel", wraplength=420).pack(anchor="nw", padx=12, pady=12)

    def _render_reporting(self, parent):
        shell = build_two_block_shell(parent, "Standardberichte", "Kennzahlen")
        with get_connection() as con:
            counts = {"Debitorenrechnungen": con.execute("SELECT COUNT(*) FROM customer_invoices").fetchone()[0], "Kreditorenrechnungen": con.execute("SELECT COUNT(*) FROM vendor_invoices").fetchone()[0], "Offene Posten": con.execute("SELECT COUNT(*) FROM open_items WHERE open_amount>0").fetchone()[0], "Buchungen": con.execute("SELECT COUNT(*) FROM journal_entries").fetchone()[0]}
        for k, v in counts.items(): ttk.Label(shell.left, text=f"{k}: {v}", style="Hint.TLabel").pack(anchor="w", padx=12, pady=4)
        ttk.Label(shell.right, text="Periodenfilter, Druck und Exportpfade sind als Zielarchitektur vorgesehen.", style="Hint.TLabel", wraplength=420).pack(anchor="nw", padx=12, pady=12)

    def _render_audit(self, parent):
        self._card(parent, "Audit / Kontrollpfade", "Änderungshistorien, Kontrollpfade, Berechtigungsoptionen und Compliance-Prüfungen sind als eigene Ausbauschicht vorbereitet. Revisionsrelevante Zeitstempel werden bereits in den Kerntabellen geführt.")

    def _render_einstellungen(self, parent):
        shell = build_two_block_shell(parent, "Systemparameter", "UI-Präferenzen")
        ttk.Label(shell.left, text=f"Datenbank: {DB_PATH}\nAnhänge: {ATTACHMENTS_DIR}\nImporte: {IMPORTS_DIR}\nVersion: {APP_VERSION}", style="Hint.TLabel", wraplength=700).pack(anchor="nw", padx=12, pady=12)
        create_standard_button(shell.left, "SQLite initialisieren / prüfen", lambda: (init_sqlite(), self.set_status("SQLite geprüft.")), confirm=True).pack(anchor="w", padx=12, pady=6)
        ttk.Label(shell.right, text=f"Sidebar kollabiert: {'Ja' if self.sidebar_collapsed else 'Nein'}\nDesignstandard: {STANDARD_DESIGN_NAME}", style="Hint.TLabel").pack(anchor="nw", padx=12, pady=12)


def maximize_window(root: tk.Tk) -> None:
    try: root.state("zoomed")
    except Exception: pass



# === FINANCE MATE PATCH V0_6_17 ===
APP_VERSION = "0.6.17-relief-docflow"

# --- Benutzerbezogene UI-Präferenzen / Splitterpersistenz ---
_FM_UI_USER = os.environ.get("FINANCEMATE_UI_USER") or os.environ.get("USERNAME") or os.environ.get("USER") or "default"


def _fm_pref_key(key: str) -> str:
    return f"ui:{_FM_UI_USER}:{key}"


def load_ui_preference(key: str, default: str = "") -> str:
    with get_connection() as con:
        row = con.execute("SELECT value FROM app_meta WHERE key=?", (_fm_pref_key(key),)).fetchone()
        return row[0] if row else default


def save_ui_preference(key: str, value: str) -> None:
    with get_connection() as con:
        con.execute(
            "INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)",
            (_fm_pref_key(key), str(value), now_str()),
        )
        con.commit()


def _fm_attach_paned_persistence(paned, key: str, default_ratio: float = 0.60) -> None:
    if getattr(paned, "_fm_pref_attached", False):
        return
    paned._fm_pref_attached = True

    def _restore():
        try:
            paned.update_idletasks()
            saved = load_ui_preference(f"pane:{key}", "")
            total = max(240, paned.winfo_width())
            if saved:
                pos = int(float(saved))
            else:
                pos = int(total * default_ratio)
            pos = max(80, min(total - 80, pos))
            try:
                paned.sashpos(0, pos)
            except Exception:
                try:
                    paned.tk.call(paned._w, "sashpos", 0, pos)
                except Exception:
                    pass
        except Exception:
            pass

    def _save(_event=None):
        try:
            try:
                pos = int(paned.sashpos(0))
            except Exception:
                pos = int(paned.tk.call(paned._w, "sashpos", 0))
            save_ui_preference(f"pane:{key}", str(pos))
        except Exception:
            pass

    paned.bind("<ButtonRelease-1>", _save, add="+")
    paned.bind("<Configure>", lambda _e: paned.after(40, _restore), add="+")
    paned.after(120, _restore)


_orig_TwoBlock_init = TwoBlock.__init__

def _fm_twoblock_init(self, parent, left_title: str, right_title: str):
    _orig_TwoBlock_init(self, parent, left_title, right_title)
    self._fm_pane_key = f"{parent.__class__.__name__}:{left_title}|{right_title}"
    _fm_attach_paned_persistence(self.pane, self._fm_pane_key, default_ratio=0.60)


TwoBlock.__init__ = _fm_twoblock_init

# --- Globale Dokumentlogik ---
def _fm_refresh_attachment_tree(tree: ttk.Treeview, rows) -> None:
    for item in tree.get_children():
        tree.delete(item)
    for row in rows:
        tree.insert("", "end", iid=str(row["id"]), values=(row["file_name"], row["file_path"], row["added_at"]))


def _fm_remove_attachment_record(attachment_id: int) -> None:
    with get_connection() as con:
        row = con.execute("SELECT file_path FROM attachments WHERE id=?", (attachment_id,)).fetchone()
        if not row:
            return
        file_path = Path(row[0])
        con.execute("DELETE FROM attachments WHERE id=?", (attachment_id,))
        con.commit()
    try:
        if file_path.exists():
            file_path.unlink()
    except Exception:
        pass


def _fm_replace_attachment_record(attachment_id: int, new_source: str) -> None:
    src_path = Path(new_source)
    if not src_path.exists():
        return
    with get_connection() as con:
        row = con.execute("SELECT entity_type, reference_no, file_path FROM attachments WHERE id=?", (attachment_id,)).fetchone()
        if not row:
            return
        entity_type = row["entity_type"]
        reference_no = row["reference_no"]
        old_file = Path(row["file_path"])
        target_dir = ATTACHMENTS_DIR / entity_type / reference_no
        target_dir.mkdir(parents=True, exist_ok=True)
        dest = target_dir / src_path.name
        if dest.exists():
            dest = target_dir / f"{src_path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{src_path.suffix}"
        shutil.copy2(src_path, dest)
        con.execute(
            "UPDATE attachments SET file_name=?, file_path=?, added_at=? WHERE id=?",
            (dest.name, str(dest), now_str(), attachment_id),
        )
        con.commit()
    try:
        if old_file.exists() and old_file.resolve() != dest.resolve():
            old_file.unlink()
    except Exception:
        pass


def _fm_pending_popup(self, title: str = "Dokumente") -> None:
    pop = tk.Toplevel(self)
    pop.title(title)
    pop.geometry("860x500")
    pop.configure(bg=BG)
    ttk.Label(pop, text=title, style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
    cols = ("file_name", "file_path")
    tree = ttk.Treeview(pop, columns=cols, show="headings")
    for col, w, txt in [("file_name", 260, "Dateiname"), ("file_path", 540, "Pfad")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=12, pady=8)

    def _refresh():
        for item in tree.get_children():
            tree.delete(item)
        for idx, file_path in enumerate(getattr(self, "pending_attachments", [])):
            q = Path(file_path)
            tree.insert("", "end", iid=str(idx), values=(q.name, str(q)))

    def _open_selected():
        iid = tree.focus()
        if iid:
            fp = tree.set(iid, "file_path")
            if fp:
                open_path(fp)

    def _delete_selected():
        iid = tree.focus()
        if iid:
            idx = int(iid)
            if 0 <= idx < len(self.pending_attachments):
                self.pending_attachments.pop(idx)
            _refresh()

    def _add_more():
        paths = list(filedialog.askopenfilenames(title="Dokument anhängen"))
        if paths:
            self.pending_attachments.extend(paths)
            _refresh()
            try:
                self.app.set_status(f"{len(self.pending_attachments)} Dokument(e) vorgemerkt.")
            except Exception:
                pass

    bar = ttk.Frame(pop)
    bar.pack(fill="x", padx=12, pady=10)
    create_standard_button(bar, "Öffnen", _open_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Löschen", _delete_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Weitere hinzufügen", _add_more, confirm=True).pack(side="left", padx=3)
    create_standard_button(bar, "Schließen", pop.destroy).pack(side="right", padx=3)
    tree.bind("<Double-1>", lambda _e: _open_selected())
    _refresh()


def _fm_manage_attachment_request(self, entity_type: str, reference_no: str, refresh_callback=None) -> None:
    if not reference_no:
        if not getattr(self, "pending_attachments", []):
            paths = list(filedialog.askopenfilenames(title="Dokument anhängen"))
            if paths:
                self.pending_attachments.extend(paths)
                try:
                    self.app.set_status(f"{len(paths)} Dokument(e) vorgemerkt.")
                except Exception:
                    pass
        else:
            _fm_pending_popup(self, "Vorgemerkte Dokumente")
        if callable(refresh_callback):
            refresh_callback()
        return

    count = self.get_attachment_count(entity_type, reference_no)
    if count <= 0:
        paths = list(filedialog.askopenfilenames(title="Dokument anhängen"))
        if paths:
            self.add_attachment_paths(entity_type, reference_no, paths)
            if callable(refresh_callback):
                refresh_callback()
    else:
        self.open_attachment_popup(entity_type, reference_no, refresh_callback=refresh_callback)


def _fm_open_attachment_popup(self, entity_type: str, reference_no: str, refresh_callback=None) -> None:
    pop = tk.Toplevel()
    pop.title(f"Dokumente – {reference_no}")
    pop.geometry("980x560")
    pop.configure(bg=BG)
    ttk.Label(pop, text=f"Dokumente zu {reference_no}", style="Section.TLabel").pack(anchor="w", padx=12, pady=8)
    cols = ("file_name", "file_path", "added_at")
    tree = ttk.Treeview(pop, columns=cols, show="headings")
    for col, w, txt in [("file_name", 220, "Dateiname"), ("file_path", 560, "Pfad"), ("added_at", 150, "Hinzugefügt")]:
        tree.heading(col, text=txt)
        tree.column(col, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=12, pady=8)

    def _reload():
        _fm_refresh_attachment_tree(tree, self.get_attachments(entity_type, reference_no))
        if callable(refresh_callback):
            refresh_callback()

    def _selected_id():
        iid = tree.focus()
        return int(iid) if iid else None

    def _open_selected():
        iid = tree.focus()
        if iid:
            fp = tree.set(iid, "file_path")
            if fp:
                open_path(fp)

    def _add_more():
        paths = list(filedialog.askopenfilenames(title="Weitere Dokumente hinzufügen"))
        if paths:
            self.add_attachment_paths(entity_type, reference_no, paths)
            _reload()

    def _delete_selected():
        attachment_id = _selected_id()
        if attachment_id is None:
            return
        if not messagebox.askyesno("Dokument löschen", "Ausgewähltes Dokument unwiderruflich löschen?", parent=pop):
            return
        _fm_remove_attachment_record(attachment_id)
        _reload()
        if not tree.get_children():
            pop.destroy()

    def _replace_selected():
        attachment_id = _selected_id()
        if attachment_id is None:
            return
        path = filedialog.askopenfilename(title="Dokument ersetzen")
        if path:
            _fm_replace_attachment_record(attachment_id, path)
            _reload()

    bar = ttk.Frame(pop)
    bar.pack(fill="x", padx=12, pady=10)
    create_standard_button(bar, "Öffnen", _open_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Ändern", _replace_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Löschen", _delete_selected).pack(side="left", padx=3)
    create_standard_button(bar, "Dokument hinzufügen", _add_more, confirm=True).pack(side="left", padx=3)
    create_standard_button(bar, "Schließen", pop.destroy).pack(side="right", padx=3)
    tree.bind("<Double-1>", lambda _e: _open_selected())
    _reload()


AttachmentMixin.open_attachment_popup = _fm_open_attachment_popup
AttachmentMixin.manage_attachment_request = _fm_manage_attachment_request

# --- Gemeinsame Rechnungslogik ---
def _fm_invoice_tree_click(self, event, tree_attr: str = "inv_tree"):
    tree = getattr(self, tree_attr, None)
    if tree is None:
        return None
    region = tree.identify("region", event.x, event.y)
    col = tree.identify_column(event.x)
    row = tree.identify_row(event.y)
    if region == "cell" and row and col == f"#{len(tree['columns'])}":
        inv_no = tree.item(row, "values")[0]
        self.manage_attachment_request(self.entity_type, inv_no, refresh_callback=self.reload_invoices)
        return "break"
    return None


def _fm_op_tree_click(self, event, tree_attr: str = "op_tree"):
    tree = getattr(self, tree_attr, None)
    if tree is None:
        return None
    region = tree.identify("region", event.x, event.y)
    col = tree.identify_column(event.x)
    row = tree.identify_row(event.y)
    if region == "cell" and row and col == f"#{len(tree['columns'])}":
        ref_no = tree.item(row, "values")[0]
        self.manage_attachment_request(self.entity_type, ref_no, refresh_callback=self.reload_open_items)
        return "break"
    return None


def _fm_invoice_add_pending(self):
    self.manage_attachment_request(self.entity_type, "", refresh_callback=None)


def _fm_invoice_open_selected(self):
    iid = self.inv_tree.focus()
    if not iid:
        messagebox.showinfo("Dokument", "Bitte Rechnung auswählen.")
        return
    inv = self.inv_tree.set(iid, "invoice_no")
    self.manage_attachment_request(self.entity_type, inv, refresh_callback=self.reload_invoices)


def _fm_invoice_reload_invoices(self):
    trees = []
    for attr in ('inv_tree',):
        tree = getattr(self, attr, None)
        if tree is not None:
            trees.append(tree)
    for tree in getattr(self, 'invoice_tabs_trees', []):
        if tree not in trees:
            trees.append(tree)
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    with get_connection() as con:
        rows = list(con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"))
    for tree in trees:
        for i in tree.get_children():
            tree.delete(i)
    for r in rows:
        att = self.get_attachment_count(self.entity_type, r["invoice_no"])
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        values = (r["invoice_no"], r[partner_name_col], r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else "Dokument anhängen")
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))


def _fm_invoice_reload_open_items(self):
    trees = []
    for attr in ('op_tree',):
        tree = getattr(self, attr, None)
        if tree is not None:
            trees.append(tree)
    for tree in getattr(self, 'open_item_tabs_trees', []):
        if tree not in trees:
            trees.append(tree)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date", (self.entity_type,)))
    for tree in trees:
        for i in tree.get_children():
            tree.delete(i)
    for r in rows:
        if parse_amount(r['open_amount']) <= 0:
            continue
        att = self.get_attachment_count(self.entity_type, r['reference_no'])
        tag = urgency_bucket(r['due_date'], r['open_amount'])
        values = (r['reference_no'], r['partner_name'], r['due_date'], format_amount(r['original_amount']), format_amount(r['open_amount']), r['status'], f"📎 {att}" if att else "Dokument anhängen")
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))


InvoiceModuleBase.add_pending_attachments = _fm_invoice_add_pending
InvoiceModuleBase.add_attachment_to_selected_invoice = _fm_invoice_open_selected
InvoiceModuleBase.reload_invoices = _fm_invoice_reload_invoices
InvoiceModuleBase.reload_open_items = _fm_invoice_reload_open_items

_orig_invoice_save = InvoiceModuleBase.save_invoice


def _fm_invoice_save_capture(self):
    self._last_saved_invoice_no = self.invoice_no.get().strip()
    selected_hist = getattr(self, 'current_hist_no', None)
    selected_partner_display = self.partner.get().strip()
    if self.entity_type == 'vendor_invoice' and not (selected_hist or getattr(self, 'pending_attachments', [])):
        messagebox.showwarning('Dokument', 'Eingangsrechnungen können ohne Dokument nicht erfasst werden.')
        return
    _orig_invoice_save(self)
    invoice_no = self._last_saved_invoice_no or ''
    if self.entity_type == 'vendor_invoice' and selected_hist and invoice_no:
        with get_connection() as con:
            files = list(con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id", (selected_hist,)))
            partner_no = selected_partner_display.split(' ')[0] if selected_partner_display else ''
            partner_name = ' '.join(selected_partner_display.split(' ')[1:]) if ' ' in selected_partner_display else selected_partner_display
            con.execute("UPDATE invoice_import_batches SET status='Erfasst', vendor_no=?, vendor_name=?, updated_at=? WHERE hist_no=?", (partner_no, partner_name, now_str(), selected_hist))
            con.commit()
        if files:
            self.add_attachment_paths(self.entity_type, invoice_no, [r[0] for r in files])
        self.current_hist_no = None
        try:
            self.stack_tree.selection_remove(self.stack_tree.selection())
        except Exception:
            pass
        if hasattr(self, '_clear_stack_preview'):
            self._clear_stack_preview()
        if hasattr(self, '_load_import_batches'):
            self._load_import_batches()
        self.reload_invoices()
        self.reload_open_items()


InvoiceModuleBase.save_invoice = _fm_invoice_save_capture

_orig_clear_form = InvoiceModuleBase.clear_form

def _fm_invoice_clear_form(self):
    _orig_clear_form(self)
    if hasattr(self, '_update_vendor_info'):
        try:
            self._update_vendor_info()
        except Exception:
            pass
    if hasattr(self, '_clear_stack_preview'):
        try:
            self._clear_stack_preview()
        except Exception:
            pass


InvoiceModuleBase.clear_form = _fm_invoice_clear_form

# --- Debitoren: nur explizit gewünschte Änderungen ---

def _walk_children(widget):
    for child in widget.winfo_children():
        yield child
        yield from _walk_children(child)


_orig_debitors_build_ui = InvoiceModuleBase._build_ui

def _fm_debitors_build_ui(self):
    _orig_debitors_build_ui(self)
    for ch in _walk_children(self):
        if isinstance(ch, ttk.Button):
            try:
                if ch.cget('text') == 'Anhang vormerken':
                    ch.configure(text='Dokument anhängen', command=self.add_pending_attachments)
            except Exception:
                pass
    try:
        self.inv_tree.bind('<ButtonRelease-1>', lambda e: _fm_invoice_tree_click(self, e), add='+')
    except Exception:
        pass
    try:
        self.op_tree.configure(columns=("reference_no", "partner", "due", "original", "open", "status", "att"))
        for c, t, w in [("reference_no", "Referenz", 120), ("partner", "Partner", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
            self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor='w')
        self.op_tree.bind('<ButtonRelease-1>', lambda e: _fm_op_tree_click(self, e), add='+')
    except Exception:
        pass
    self.reload_invoices()
    self.reload_open_items()


DebitorsView._build_ui = _fm_debitors_build_ui

# --- Kreditoren: 4 Fensterblöcke beibehalten, nur gewünschte Änderungen ---

def _fm_creditors_update_vendor_info(self):
    body = getattr(self, 'vendor_info_body', None)
    if body is None:
        return
    vendor_no = self.partner.get().split(' ')[0] if self.partner.get() else ''
    content = 'Kein Kreditor ausgewählt.'
    if vendor_no:
        with get_connection() as con:
            row = con.execute("SELECT * FROM vendors WHERE vendor_no=?", (vendor_no,)).fetchone()
        if row:
            content = "\n".join([
                f"Kreditor-Nr.: {row['vendor_no']}",
                f"Name: {row['name']}",
                f"Land: {row['country']}",
                f"Adresse: {row['street']}, {row['zip']} {row['city']}",
                f"E-Mail: {row['email']}",
                f"Telefon: {row['phone']}",
                f"Steuer-Nr.: {row['tax_id']}",
                f"USt-IdNr.: {row['vat_id']}",
                f"IBAN/BIC: {row['iban']} / {row['bic']}",
                f"Bank: {row['bank_name']}",
                f"Zahlungsbedingung: {row['payment_term']}",
                f"Aktiv: {int_to_yes_no(row['active'])}",
            ])
    body.configure(state='normal')
    body.delete('1.0', tk.END)
    body.insert('1.0', content)
    body.configure(state='disabled')


def _fm_creditors_clear_stack_preview(self):
    if hasattr(self, 'preview_title_label'):
        self.preview_title_label.configure(text='Kein Dokument ausgewählt')
    if hasattr(self, 'preview_image_label'):
        self.preview_image_label.configure(text='Keine Vorschau verfügbar', image='')
        try:
            self.preview_image_label.image = None
        except Exception:
            pass


def _fm_creditors_build_ui(self):
    outer = ttk.PanedWindow(self, orient='vertical')
    outer.pack(fill='both', expand=True)
    _fm_attach_paned_persistence(outer, 'CreditorsView:outer_vertical', default_ratio=0.62)
    invoice_area = ttk.Frame(outer)
    stack_area = ttk.Frame(outer)
    outer.add(invoice_area, weight=3)
    outer.add(stack_area, weight=2)

    shell = build_two_block_shell(invoice_area, 'Eingangsrechnung erfassen', 'Eingangsrechnungen')
    left = ScrollableFrame(shell.left)
    left.pack(fill='both', expand=True, padx=10, pady=8)
    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.partner = tk.StringVar(); self.invoice_date = tk.StringVar(value=today_str()); self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar(); self.tax_code = tk.StringVar(value='V19'); self.net_amount = tk.StringVar(); self.tax_amount = tk.StringVar(); self.gross_amount = tk.StringVar()
    self.current_hist_no = None
    self.pending_attachments = []

    self._labeled_entry(left.content, 'Rechnungs-Nr.', self.invoice_no)
    prow = ttk.Frame(left.content); prow.pack(fill='x', pady=3)
    ttk.Label(prow, text='Kreditor', width=16).pack(side='left')
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side='left', fill='x', expand=True)
    self.partner_cb.bind('<MouseWheel>', lambda e: 'break')
    self.partner_cb.bind('<<ComboboxSelected>>', lambda _e: self._update_vendor_info())
    create_standard_button(prow, '+', self._open_vendor_quick_popup).pack(side='left', padx=3)
    self._labeled_entry(left.content, 'Rechnungsdatum', self.invoice_date)
    self._labeled_combo(left.content, 'Zahlungsbed.', self.payment_term, self.payment_terms)
    self._labeled_entry(left.content, 'Fälligkeit', self.due_date)
    self._labeled_combo(left.content, 'Steuerkennz.', self.tax_code, self.tax_codes)
    self._labeled_entry(left.content, 'Netto', self.net_amount)
    self._labeled_entry(left.content, 'Steuer', self.tax_amount)
    self._labeled_entry(left.content, 'Brutto', self.gross_amount)
    self.payment_term.trace_add('write', lambda *_: self.apply_payment_term())
    self.net_amount.trace_add('write', lambda *_: self._calculate_amounts())
    self.tax_code.trace_add('write', lambda *_: self._calculate_amounts())

    btns = ttk.Frame(left.content); btns.pack(fill='x', pady=8)
    create_standard_button(btns, 'Dokument anhängen', self.add_pending_attachments).pack(side='left', padx=3)
    create_standard_button(btns, 'Speichern', self.save_invoice, confirm=True).pack(side='left', padx=3)
    create_standard_button(btns, 'Neu', self.clear_form).pack(side='left', padx=3)

    info_frame = ttk.LabelFrame(left.content, text='Stammdaten des Kreditors')
    info_frame.pack(fill='x', pady=(8, 4))
    self.vendor_info_body = tk.Text(info_frame, height=8, wrap='word', bg=WHITE, fg=TEXT, relief='flat')
    self.vendor_info_body.pack(fill='both', expand=True, padx=6, pady=6)
    self.vendor_info_body.configure(state='disabled')

    right = ttk.Frame(shell.right); right.pack(fill='both', expand=True, padx=10, pady=8)
    nb = ttk.Notebook(right)
    nb.pack(fill='both', expand=True)
    tab_inv = ttk.Frame(nb)
    tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text='Eingangsrechnungen')
    nb.add(tab_op, text='Offene Posten')

    inv_cols = ('invoice_no', 'partner', 'date', 'due', 'gross', 'open', 'status', 'att')
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show='headings', height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [('invoice_no', 'Rechnung', 110), ('partner', 'Kreditor', 180), ('date', 'Datum', 90), ('due', 'Fällig', 90), ('gross', 'Brutto', 90), ('open', 'Offen', 90), ('status', 'Status', 110), ('att', 'Anhang', 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor='w')
    self.inv_tree.pack(fill='both', expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind('<ButtonRelease-1>', lambda e: _fm_invoice_tree_click(self, e), add='+')

    op_cols = ('reference_no', 'partner', 'due', 'original', 'open', 'status', 'att')
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show='headings', height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [('reference_no', 'Referenz', 120), ('partner', 'Kreditor', 180), ('due', 'Fällig', 90), ('original', 'Original', 90), ('open', 'Offen', 90), ('status', 'Status', 120), ('att', 'Anhang', 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor='w')
    self.op_tree.pack(fill='both', expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind('<ButtonRelease-1>', lambda e: _fm_op_tree_click(self, e), add='+')

    self._build_stack_ui(stack_area)
    self.partner.trace_add('write', lambda *_: self._update_vendor_info())
    self._update_vendor_info()
    self.reload_invoices(); self.reload_open_items(); self._load_import_batches(); self._clear_stack_preview()


def _fm_creditors_build_stack_ui(self, parent):
    split = build_two_block_shell(parent, 'Dokumentenstapel', 'Vorschau')
    left = ttk.Frame(split.left); left.pack(fill='both', expand=True, padx=10, pady=8)
    right = ttk.Frame(split.right); right.pack(fill='both', expand=True, padx=10, pady=8)
    toolbar = ttk.Frame(left); toolbar.pack(fill='x', pady=(0, 6))
    create_standard_button(toolbar, 'Dateien importieren', self.import_creditor_files, confirm=True).pack(side='left', padx=3)
    create_standard_button(toolbar, 'Ordner importieren', self.import_creditor_folder).pack(side='left', padx=3)
    create_standard_button(toolbar, 'Markierte löschen', self._delete_marked_batches).pack(side='left', padx=3)
    create_standard_button(toolbar, 'Markierte zusammenführen', self.merge_selected_batches).pack(side='left', padx=3)

    cols = ('hist_no', 'title', 'status', 'mark')
    self.stack_tree = ttk.Treeview(left, columns=cols, show='headings', height=10)
    for c, t, w in [('hist_no', 'Hist.-Nr.', 110), ('title', 'Dokument', 260), ('status', 'Status', 110), ('mark', '☑', 40)]:
        self.stack_tree.heading(c, text=t)
        self.stack_tree.column(c, width=w, anchor='w' if c != 'mark' else 'center')
    self.stack_tree.pack(fill='both', expand=True)
    self.stack_tree.bind('<ButtonRelease-1>', self._on_stack_select, add='+')
    self.stack_tree.bind('<Double-1>', lambda _e: self._open_selected_stack_file())

    self.preview_title_label = ttk.Label(right, text='Kein Dokument ausgewählt', style='Hint.TLabel')
    self.preview_title_label.pack(anchor='w', pady=(0, 6))
    self.preview_image_label = ttk.Label(right, text='Keine Vorschau verfügbar', anchor='nw', justify='left')
    self.preview_image_label.pack(fill='both', expand=True)


def _fm_creditors_import_archive_paths(self, paths):
    valid = []
    for raw in paths:
        fp = Path(raw)
        if fp.exists() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo('Dokumentenstapel', 'Keine unterstützten Dateien ausgewählt.')
        return
    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute("INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)", (hist_no, clean_document_title(fp.name), '', '', 'Importiert', 0, now_str(), now_str()))
            con.execute("INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)", (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or '', now_str()))
        con.commit()
    self._load_import_batches()
    self.app.set_status(f"{len(valid)} Dokument(e) in den Dokumentenstapel importiert.")


def _fm_creditors_import_files(self):
    paths = filedialog.askopenfilenames(title='Dateien für Dokumentenstapel importieren')
    if paths:
        self._import_archive_paths(paths)


def _fm_creditors_import_folder(self):
    folder = filedialog.askdirectory(title='Ordner für Dokumentenstapel importieren')
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))


def _fm_creditors_load_import_batches(self):
    for i in self.stack_tree.get_children():
        self.stack_tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_batches ORDER BY id DESC"))
    for r in rows:
        self.stack_tree.insert('', 'end', iid=r['hist_no'], values=(r['hist_no'], r['title'], r['status'], '☑' if int(r['marked'] or 0) else '☐'))


def _fm_creditors_on_stack_select(self, event=None):
    row_id = self.stack_tree.identify_row(event.y) if event is not None else None
    col = self.stack_tree.identify_column(event.x) if event is not None else None
    if row_id and col == '#4':
        with get_connection() as con:
            row = con.execute("SELECT marked FROM invoice_import_batches WHERE hist_no=?", (row_id,)).fetchone()
            new_val = 0 if row and int(row[0] or 0) else 1
            con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=? WHERE hist_no=?", (new_val, now_str(), row_id))
            con.commit()
        self._load_import_batches()
        return 'break'
    if not row_id:
        self.current_hist_no = None
        self._clear_stack_preview()
        try:
            self.stack_tree.selection_remove(self.stack_tree.selection())
        except Exception:
            pass
        return None
    if self.current_hist_no == row_id:
        self.current_hist_no = None
        try:
            self.stack_tree.selection_remove(self.stack_tree.selection())
        except Exception:
            pass
        self._clear_stack_preview()
        return 'break'
    self.current_hist_no = row_id
    try:
        self.stack_tree.selection_set(row_id)
    except Exception:
        pass
    self._refresh_stack_preview()
    return None


def _fm_creditors_refresh_stack_preview(self):
    if not getattr(self, 'current_hist_no', None):
        self._clear_stack_preview()
        return
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_files WHERE hist_no=? ORDER BY id", (self.current_hist_no,)))
    if not rows:
        self._clear_stack_preview()
        return
    first = rows[0]
    title = clean_document_title(first['file_name'])
    if len(rows) > 1:
        title += f" (+ {len(rows)-1} weitere)"
    self.preview_title_label.configure(text=title)
    self.preview_image_label.configure(text=collect_preview_text(first['file_path']), image='')
    try:
        self.preview_image_label.image = None
    except Exception:
        pass


def _fm_creditors_open_selected_stack_file(self):
    if not getattr(self, 'current_hist_no', None):
        return
    with get_connection() as con:
        row = con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id LIMIT 1", (self.current_hist_no,)).fetchone()
    if row:
        open_path(row[0])


def _fm_creditors_delete_marked(self):
    with get_connection() as con:
        hist_nos = [r['hist_no'] for r in con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1 ORDER BY id")]
    if not hist_nos:
        messagebox.showinfo('Dokumentenstapel', 'Keine markierten Dokumente vorhanden.')
        return
    if not messagebox.askyesno('Dokumentenstapel', f"{len(hist_nos)} markierte Dokument(e) löschen?", parent=self):
        return
    placeholders = ','.join(['?'] * len(hist_nos))
    with get_connection() as con:
        files = list(con.execute(f"SELECT file_path FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos))
        con.execute(f"DELETE FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos)
        con.execute(f"DELETE FROM invoice_import_batches WHERE hist_no IN ({placeholders})", hist_nos)
        con.commit()
    for r in files:
        try:
            fp = Path(r['file_path'])
            if fp.exists():
                fp.unlink()
        except Exception:
            pass
    if getattr(self, 'current_hist_no', None) in hist_nos:
        self.current_hist_no = None
        self._clear_stack_preview()
    self._load_import_batches()


def _fm_creditors_merge_selected(self):
    with get_connection() as con:
        hist_nos = [r['hist_no'] for r in con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1 ORDER BY id")]
        if len(hist_nos) < 2:
            messagebox.showinfo('Dokumentenstapel', 'Bitte mindestens zwei markierte Dokumente wählen.')
            return
        target = hist_nos[0]
        for hist_no in hist_nos[1:]:
            con.execute("UPDATE invoice_import_files SET hist_no=? WHERE hist_no=?", (target, hist_no))
            con.execute("DELETE FROM invoice_import_batches WHERE hist_no=?", (hist_no,))
        con.execute("UPDATE invoice_import_batches SET marked=0, updated_at=? WHERE hist_no=?", (now_str(), target))
        con.commit()
    self.current_hist_no = target
    self._load_import_batches()
    self._refresh_stack_preview()


CreditorsView._build_ui = _fm_creditors_build_ui
CreditorsView._build_stack_ui = _fm_creditors_build_stack_ui
CreditorsView._update_vendor_info = _fm_creditors_update_vendor_info
CreditorsView._clear_stack_preview = _fm_creditors_clear_stack_preview
CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths
CreditorsView.import_creditor_files = _fm_creditors_import_files
CreditorsView.import_creditor_folder = _fm_creditors_import_folder
CreditorsView._load_import_batches = _fm_creditors_load_import_batches
CreditorsView._on_stack_select = _fm_creditors_on_stack_select
CreditorsView._refresh_stack_preview = _fm_creditors_refresh_stack_preview
CreditorsView._open_selected_stack_file = _fm_creditors_open_selected_stack_file
CreditorsView._delete_marked_batches = _fm_creditors_delete_marked
CreditorsView.merge_selected_batches = _fm_creditors_merge_selected



# === FINANCE MATE V0.6.19 - CLEAN PATCH FROM STABLE V0.6.18 ===
# Grundlage: stabile v0.6.18. Ziel: keine leeren/grauen Kreditoren-Module mehr, nur gewünschte Weiterentwicklung.
APP_VERSION = "0.6.19-clean-from-v0.6.18"

TRADE_COUNTRIES = [
    "Deutschland", "Vereinigte Staaten", "China", "Japan", "Vereinigtes Königreich", "Frankreich", "Italien", "Niederlande", "Belgien", "Schweiz",
    "Österreich", "Spanien", "Polen", "Tschechien", "Ungarn", "Slowakei", "Schweden", "Dänemark", "Finnland", "Norwegen",
    "Irland", "Portugal", "Griechenland", "Türkei", "Rumänien", "Bulgarien", "Slowenien", "Kroatien", "Serbien", "Ukraine",
    "Kanada", "Mexiko", "Brasilien", "Argentinien", "Chile", "Kolumbien", "Peru", "Indien", "Südkorea", "Taiwan",
    "Singapur", "Malaysia", "Thailand", "Indonesien", "Vietnam", "Philippinen", "Australien", "Neuseeland", "Vereinigte Arabische Emirate", "Südafrika"
]
INVOICE_KIND_VALUES = ["Rechnung", "Gutschrift"]
FIBU_VISIBLE_NAME = "Buchungsjournal"


def _v619_status(open_amount: Any) -> str:
    try:
        return STATUS_PAID if parse_amount(open_amount) <= 0 else STATUS_OPEN
    except Exception:
        return STATUS_OPEN


def compute_status_from_open_amount(open_amount: Any, due_date: str) -> str:
    return _v619_status(open_amount)


def urgency_bucket(due_date: str, open_amount: Any) -> str:
    try:
        amount = parse_amount(open_amount)
    except Exception:
        amount = Decimal("0.00")
    if amount <= 0:
        return "paid"
    if not validate_date(due_date):
        return "open"
    delta = (datetime.strptime(due_date, DATE_FMT).date() - datetime.now().date()).days
    if delta <= 0:
        return "overdue"
    if delta <= 5:
        return "soon"
    return "open"


def configure_tree_tags(tree: ttk.Treeview) -> None:
    tree.tag_configure("paid", background=SOFT_GREEN)
    tree.tag_configure("soon", background=SOFT_ORANGE)
    tree.tag_configure("overdue", background=SOFT_RED)
    tree.tag_configure("open", background=WHITE)


_orig_init_sqlite_v619_clean = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v619_clean()
    with get_connection() as con:
        for table, column, definition in [
            ("customer_invoices", "invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("customer_invoices", "description", "TEXT DEFAULT ''"),
            ("vendor_invoices", "invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("vendor_invoices", "description", "TEXT DEFAULT ''"),
            ("open_items", "invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("open_items", "description", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, table, column, definition)
        # bestehende Doppler in offenen Posten bereinigen
        con.execute("DELETE FROM open_items WHERE id NOT IN (SELECT MIN(id) FROM open_items GROUP BY entity_type, reference_no)")
        # Status für vorhandene Rechnungen und OPs fachlich korrigieren
        for table in ("customer_invoices", "vendor_invoices"):
            for row in con.execute(f"SELECT id, open_amount FROM {table}").fetchall():
                con.execute(f"UPDATE {table} SET status=?, updated_at=? WHERE id=?", (_v619_status(row["open_amount"]), now_str(), row["id"]))
        for row in con.execute("SELECT id, open_amount FROM open_items").fetchall():
            con.execute("UPDATE open_items SET status=?, updated_at=? WHERE id=?", (_v619_status(row["open_amount"]), now_str(), row["id"]))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()


# ---------- kleine stabile UI-Helfer ----------
def _v619_labeled_text(parent, label: str, height: int = 2) -> tk.Text:
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=3)
    ttk.Label(row, text=label, width=16).pack(side="left", anchor="n")
    txt = tk.Text(row, height=height, wrap="word", bg=WHITE, fg=TEXT, relief="sunken")
    txt.pack(side="left", fill="x", expand=True)
    return txt


def _v619_text_get(widget: Optional[tk.Text]) -> str:
    if widget is None:
        return ""
    return widget.get("1.0", "end-1c").strip()


def _v619_text_set(widget: Optional[tk.Text], value: str = "") -> None:
    if widget is None:
        return
    widget.delete("1.0", tk.END)
    if value:
        widget.insert("1.0", value)


def _v619_unique_trees(*trees):
    result = []
    seen = set()
    for tree in trees:
        if tree is None:
            continue
        if id(tree) in seen:
            continue
        result.append(tree)
        seen.add(id(tree))
    return result


def _v619_add_attachment_column_click(self, event, tree, ref_getter, refresh_callback):
    row = tree.identify_row(event.y)
    col = tree.identify_column(event.x)
    if row and col == f"#{len(tree['columns'])}":
        self.manage_attachment_request(self.entity_type, ref_getter(row), refresh_callback=refresh_callback)
        return "break"
    return None


# ---------- Referenzdaten stabil ----------
def _v619_load_reference_data(self):
    with get_connection() as con:
        self.partners = [f"{r[self.partner_no_col]} {r['name']}" for r in con.execute(f"SELECT {self.partner_no_col},name FROM {self.partner_table} WHERE active=1 ORDER BY name")]
        self.tax_codes = [r[0] for r in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]
        self.payment_terms = [r[0] for r in con.execute("SELECT code FROM payment_terms WHERE active=1 ORDER BY code")]

InvoiceModuleBase._load_reference_data = _v619_load_reference_data


# ---------- Speichern/Leeren/Reload für Debitoren und Kreditoren ----------
def _v619_save_invoice(self):
    if not validate_date(self.invoice_date.get()) or not validate_date(self.due_date.get()):
        messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.")
        return
    no, name, address = self._partner_parts()
    if not no:
        messagebox.showwarning("Partner", "Bitte Partner auswählen.")
        return
    selected_hist = getattr(self, "current_hist_no", None)
    if self.entity_type == "vendor_invoice" and not (selected_hist or getattr(self, "pending_attachments", [])):
        messagebox.showwarning("Dokument", "Eingangsrechnungen können ohne Dokument nicht erfasst werden.")
        return
    gross = parse_amount(self.gross_amount.get())
    net = parse_amount(self.net_amount.get())
    tax = parse_amount(self.tax_amount.get())
    status = _v619_status(gross)
    invoice_no = self.invoice_no.get().strip()
    invoice_kind = self.invoice_kind.get().strip() if hasattr(self, "invoice_kind") else "Rechnung"
    description = _v619_text_get(getattr(self, "description_widget", None))
    address_col = "customer_address" if self.invoice_table == "customer_invoices" else "vendor_address"
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    partner_no_col = "customer_no" if self.invoice_table == "customer_invoices" else "vendor_no"
    stack_files = []
    with get_connection() as con:
        con.execute(
            f"INSERT INTO {self.invoice_table}(invoice_no,{partner_no_col},{partner_name_col},{address_col},invoice_date,due_date,payment_term,tax_code,net_amount,tax_amount,gross_amount,open_amount,status,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (invoice_no, no, name, address, self.invoice_date.get(), self.due_date.get(), self.payment_term.get(), self.tax_code.get(), float(net), float(tax), float(gross), float(gross), status, invoice_kind, description, now_str(), now_str()),
        )
        con.execute("DELETE FROM open_items WHERE entity_type=? AND reference_no=?", (self.entity_type, invoice_no))
        con.execute(
            "INSERT INTO open_items(entity_type,reference_no,partner_no,partner_name,due_date,original_amount,open_amount,status,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
            (self.entity_type, invoice_no, no, name, self.due_date.get(), float(gross), float(gross), status, invoice_kind, description, now_str(), now_str()),
        )
        if self.entity_type == "vendor_invoice" and selected_hist:
            stack_files = list(con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id", (selected_hist,)))
            con.execute("UPDATE invoice_import_batches SET status='Erfasst', vendor_no=?, vendor_name=?, updated_at=? WHERE hist_no=?", (no, name, now_str(), selected_hist))
        con.commit()
    if stack_files:
        self.add_attachment_paths(self.entity_type, invoice_no, [r["file_path"] for r in stack_files])
    if getattr(self, "pending_attachments", []):
        self.add_attachment_paths(self.entity_type, invoice_no, self.pending_attachments)
    self.app.set_status(f"{self.entity_label} {invoice_no} gespeichert.")
    self.clear_form()
    self.reload_invoices()
    self.reload_open_items()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    if hasattr(self, "_clear_stack_preview"):
        self._clear_stack_preview()

InvoiceModuleBase.save_invoice = _v619_save_invoice


def _v619_clear_form(self):
    self.invoice_no.set(self._generate_invoice_no())
    if hasattr(self, "invoice_kind"):
        self.invoice_kind.set("Rechnung")
    self.partner.set("")
    self.invoice_date.set(today_str())
    self.due_date.set(today_str())
    self.payment_term.set("")
    self.tax_code.set("V19")
    self.net_amount.set("")
    self.tax_amount.set("")
    self.gross_amount.set("")
    _v619_text_set(getattr(self, "description_widget", None), "")
    self.pending_attachments.clear()
    if hasattr(self, "current_hist_no"):
        self.current_hist_no = None
    if hasattr(self, "_update_vendor_info"):
        try:
            self._update_vendor_info()
        except Exception:
            pass
    if hasattr(self, "_clear_stack_preview"):
        try:
            self._clear_stack_preview()
        except Exception:
            pass

InvoiceModuleBase.clear_form = _v619_clear_form


def _v619_reload_invoices(self):
    trees = _v619_unique_trees(getattr(self, "inv_tree", None), *getattr(self, "invoice_tabs_trees", []))
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    with get_connection() as con:
        rows = list(con.execute(f"SELECT * FROM {self.invoice_table} ORDER BY id DESC"))
    for tree in trees:
        for item in tree.get_children():
            tree.delete(item)
    for r in rows:
        att = self.get_attachment_count(self.entity_type, r["invoice_no"])
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        values = (
            r["invoice_no"],
            r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung",
            r[partner_name_col],
            r["description"] if "description" in r.keys() else "",
            r["invoice_date"],
            r["due_date"],
            format_amount(r["gross_amount"]),
            format_amount(r["open_amount"]),
            _v619_status(r["open_amount"]),
            f"📎 {att}" if att else "Dokument anhängen",
        )
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))

InvoiceModuleBase.reload_invoices = _v619_reload_invoices


def _v619_reload_open_items(self):
    trees = _v619_unique_trees(getattr(self, "op_tree", None), *getattr(self, "open_item_tabs_trees", []))
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date, id", (self.entity_type,)))
    for tree in trees:
        for item in tree.get_children():
            tree.delete(item)
    for r in rows:
        if parse_amount(r["open_amount"]) <= 0:
            continue
        att = self.get_attachment_count(self.entity_type, r["reference_no"])
        tag = urgency_bucket(r["due_date"], r["open_amount"])
        values = (
            r["reference_no"],
            r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung",
            r["partner_name"],
            r["description"] if "description" in r.keys() else "",
            r["due_date"],
            format_amount(r["original_amount"]),
            format_amount(r["open_amount"]),
            _v619_status(r["open_amount"]),
            f"📎 {att}" if att else "Dokument anhängen",
        )
        for tree in trees:
            tree.insert("", "end", values=values, tags=(tag,))

InvoiceModuleBase.reload_open_items = _v619_reload_open_items


# ---------- Debitoren UI stabil erweitern ----------
def _v619_debitors_build_ui(self):
    shell = build_two_block_shell(self, f"{self.entity_label} erfassen", "Offene Posten / Kontrolle")
    left = ScrollableFrame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()
    self._labeled_combo(left.content, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
    self._labeled_combo(left.content, "Partner", self.partner, self.partners)
    self.description_widget = _v619_labeled_text(left.content, "Rechnungsbeschreibung", height=2)
    self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(left.content, "Fälligkeit", self.due_date)
    self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(left.content, "Netto", self.net_amount)
    self._labeled_entry(left.content, "Steuer", self.tax_amount)
    self._labeled_entry(left.content, "Brutto", self.gross_amount)
    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
    btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
    ttk.Label(left.content, text="Erfasste Rechnungen", style="Section.TLabel").pack(anchor="w", pady=(10, 4))
    cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(left.content, columns=cols, show="headings", height=9)
    for c, t, w in [("invoice_no", "Rechnung", 110), ("kind", "Art", 90), ("partner", "Partner", 170), ("description", "Beschreibung", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")
    right = ttk.Frame(shell.right); right.pack(fill="both", expand=True, padx=10, pady=8)
    cols2 = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(right, columns=cols2, show="headings")
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Partner", 170), ("description", "Beschreibung", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

DebitorsView._build_ui = _v619_debitors_build_ui


# ---------- Kreditoren UI vollständig aus v0.6.18 wiederhergestellt und erweitert ----------
def _v619_creditors_build_ui(self):
    # Fehler im Kreditoren-Modul dürfen nicht mehr still in einem grauen Hintergrund enden.
    outer = ttk.PanedWindow(self, orient="vertical")
    outer.pack(fill="both", expand=True)
    try:
        _fm_attach_paned_persistence(outer, "CreditorsView:outer_vertical", default_ratio=0.62)
    except Exception:
        pass
    invoice_area = ttk.Frame(outer)
    stack_area = ttk.Frame(outer)
    outer.add(invoice_area, weight=3)
    outer.add(stack_area, weight=2)

    shell = build_two_block_shell(invoice_area, "Eingangsrechnung erfassen", "Eingangsrechnungen")
    left = ScrollableFrame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()
    self.current_hist_no = None
    self.pending_attachments = []

    self._labeled_combo(left.content, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(left.content, "Rechnungs-Nr.", self.invoice_no)
    prow = ttk.Frame(left.content); prow.pack(fill="x", pady=3)
    ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side="left", fill="x", expand=True)
    self.partner_cb.bind("<MouseWheel>", lambda e: "break")
    self.partner_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_vendor_info())
    create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)
    self.description_widget = _v619_labeled_text(left.content, "Rechnungsbeschreibung", height=2)
    self._labeled_entry(left.content, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(left.content, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(left.content, "Fälligkeit", self.due_date)
    self._labeled_combo(left.content, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(left.content, "Netto", self.net_amount)
    self._labeled_entry(left.content, "Steuer", self.tax_amount)
    self._labeled_entry(left.content, "Brutto", self.gross_amount)
    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())
    btns = ttk.Frame(left.content); btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)
    info_frame = ttk.LabelFrame(left.content, text="Stammdaten des Kreditors")
    info_frame.pack(fill="x", pady=(8, 4))
    self.vendor_info_body = tk.Text(info_frame, height=8, wrap="word", bg=WHITE, fg=TEXT, relief="flat")
    self.vendor_info_body.pack(fill="both", expand=True, padx=6, pady=6)
    self.vendor_info_body.configure(state="disabled")

    right = ttk.Frame(shell.right)
    right.pack(fill="both", expand=True, padx=10, pady=8)
    nb = ttk.Notebook(right)
    nb.pack(fill="both", expand=True)
    tab_inv = ttk.Frame(nb)
    tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text="Eingangsrechnungen")
    nb.add(tab_op, text="Offene Posten")

    inv_cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show="headings", height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [("invoice_no", "Rechnung", 110), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 180), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 110), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")

    op_cols = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show="headings", height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 180), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

    # Dokumentenstapel/Vorschau exakt als eigener Bereich beibehalten.
    self._build_stack_ui(stack_area)
    self.partner.trace_add("write", lambda *_: self._update_vendor_info())
    self._update_vendor_info()
    self.reload_invoices()
    self.reload_open_items()
    self._load_import_batches()
    self._clear_stack_preview()

CreditorsView._build_ui = _v619_creditors_build_ui


# ---------- Kreditoren: Stapel erfasste ausblenden, Löschen stabilisieren ----------
def _v619_load_import_batches(self):
    if not hasattr(self, "stack_tree"):
        return
    for i in self.stack_tree.get_children():
        self.stack_tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_batches WHERE status <> 'Erfasst' ORDER BY id DESC"))
    for r in rows:
        self.stack_tree.insert("", "end", iid=r["hist_no"], values=(r["hist_no"], r["title"], r["status"], "☑" if int(r["marked"] or 0) else "☐"))

CreditorsView._load_import_batches = _v619_load_import_batches


def _v619_delete_marked(self):
    with get_connection() as con:
        hist_nos = [r["hist_no"] for r in con.execute("SELECT hist_no FROM invoice_import_batches WHERE marked=1 AND status <> 'Erfasst' ORDER BY id")]
    if not hist_nos:
        messagebox.showinfo("Dokumentenstapel", "Keine markierten Dokumente vorhanden.")
        return
    if not messagebox.askyesno("Dokumentenstapel", f"{len(hist_nos)} markierte Dokument(e) löschen?", parent=self):
        return
    placeholders = ",".join(["?"] * len(hist_nos))
    with get_connection() as con:
        files = list(con.execute(f"SELECT file_path FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos))
        con.execute(f"DELETE FROM invoice_import_files WHERE hist_no IN ({placeholders})", hist_nos)
        con.execute(f"DELETE FROM invoice_import_batches WHERE hist_no IN ({placeholders})", hist_nos)
        con.commit()
    for r in files:
        try:
            fp = Path(r["file_path"])
            if fp.exists():
                fp.unlink()
        except Exception:
            pass
    if getattr(self, "current_hist_no", None) in hist_nos:
        self.current_hist_no = None
        self._clear_stack_preview()
    self._load_import_batches()

CreditorsView._delete_marked_batches = _v619_delete_marked


# ---------- Quick-Kreditor erweitert ----------
def _v619_open_vendor_quick_popup(self):
    pop = tk.Toplevel(self)
    pop.title("Quick-Kreditor")
    pop.geometry("620x650")
    pop.configure(bg=BG)
    fields = ["vendor_no", "name", "country", "street", "zip", "city", "email", "phone", "tax_id", "vat_id", "iban", "bic", "bank_name", "payment_term", "active"]
    vars = {k: tk.StringVar() for k in fields}
    vars["vendor_no"].set(generate_number("K", "counter_vendors", 5))
    vars["country"].set("Deutschland")
    vars["active"].set("Ja")
    labels = [
        ("vendor_no", "Kreditor-Nr."), ("name", "Name"), ("country", "Land"), ("street", "Straße"), ("zip", "PLZ"), ("city", "Ort"),
        ("email", "E-Mail"), ("phone", "Telefon"), ("tax_id", "Steuer-Nr."), ("vat_id", "USt-IdNr."), ("iban", "IBAN"), ("bic", "BIC"),
        ("bank_name", "Bank"), ("payment_term", "Zahlungsbed."), ("active", "Aktiv")
    ]
    for key, label in labels:
        row = ttk.Frame(pop); row.pack(fill="x", padx=12, pady=4)
        ttk.Label(row, text=label, width=16).pack(side="left")
        if key == "country":
            cb = ttk.Combobox(row, textvariable=vars[key], values=TRADE_COUNTRIES, state="readonly")
            cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        elif key == "payment_term":
            cb = ttk.Combobox(row, textvariable=vars[key], values=self.payment_terms, state="readonly")
            cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        elif key == "active":
            cb = ttk.Combobox(row, textvariable=vars[key], values=["Ja", "Nein"], state="readonly")
            cb.pack(side="left", fill="x", expand=True); cb.bind("<MouseWheel>", lambda e: "break")
        else:
            ttk.Entry(row, textvariable=vars[key]).pack(side="left", fill="x", expand=True)

    def save():
        if not vars["name"].get().strip():
            messagebox.showwarning("Pflichtfeld", "Name fehlt.", parent=pop)
            return
        with get_connection() as con:
            con.execute(
                "INSERT INTO vendors(vendor_no,name,country,street,zip,city,email,phone,tax_id,vat_id,iban,bic,bank_name,payment_term,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (vars["vendor_no"].get().strip(), vars["name"].get().strip(), vars["country"].get().strip(), vars["street"].get().strip(), vars["zip"].get().strip(), vars["city"].get().strip(), vars["email"].get().strip(), vars["phone"].get().strip(), vars["tax_id"].get().strip(), vars["vat_id"].get().strip(), vars["iban"].get().strip(), vars["bic"].get().strip(), vars["bank_name"].get().strip(), vars["payment_term"].get().strip(), yes_no_to_int(vars["active"].get()), now_str(), now_str())
            )
            con.commit()
        self._load_reference_data()
        self.partner.set(f"{vars['vendor_no'].get().strip()} {vars['name'].get().strip()}")
        try:
            self._update_vendor_info()
        except Exception:
            pass
        pop.destroy()

    create_standard_button(pop, "Speichern", save, confirm=True).pack(pady=12)

CreditorsView._open_vendor_quick_popup = _v619_open_vendor_quick_popup


# ---------- Navigation: Finanzbuchhaltung sichtbar umbenennen ----------
def _v619_build_sidebar(self):
    for w in self.sidebar.winfo_children():
        w.destroy()
    width = 42 if self.sidebar_collapsed else 210
    self.sidebar.configure(width=width)
    self.sidebar.pack_propagate(False)
    arrow = "»" if self.sidebar_collapsed else "«"
    create_standard_button(self.sidebar, arrow, self._toggle_sidebar).pack(fill="x", padx=5, pady=5)
    modules = [("Dashboard", "DB"), ("Stammdaten", "ST"), (FIBU_VISIBLE_NAME, "BJ"), ("Debitoren", "DE"), ("Kreditoren", "KR"), ("Zahlungen", "ZA"), ("Reporting", "RE"), ("Audit", "AU"), ("Einstellungen", "EI")]
    self.nav_buttons.clear()
    for name, abbr in modules:
        text = abbr if self.sidebar_collapsed else name
        btn = ttk.Button(self.sidebar, text=text, style="Nav.TButton", command=lambda n=name: self.show_module(n))
        btn.pack(fill="x", padx=5, pady=2)
        self.nav_buttons[name] = btn

FinanceMateApp._build_sidebar = _v619_build_sidebar


def _v619_show_module(self, module: str):
    module_key = FIBU_VISIBLE_NAME if module == "Finanzbuchhaltung" else module
    self.active_module = module_key
    self._update_nav_styles()
    for w in self.workspace.winfo_children():
        w.destroy()
    ttk.Label(self.workspace, text=module_key, style="CardTitle.TLabel").pack(anchor="w", pady=(0, 6))
    container = ttk.Frame(self.workspace)
    container.pack(fill="both", expand=True)
    render = {
        "Dashboard": self._render_dashboard,
        "Stammdaten": self._render_stammdaten,
        FIBU_VISIBLE_NAME: self._render_finanzbuchhaltung,
        "Debitoren": self._render_debitoren,
        "Kreditoren": self._render_kreditoren,
        "Zahlungen": self._render_zahlungen,
        "Reporting": self._render_reporting,
        "Audit": self._render_audit,
        "Einstellungen": self._render_einstellungen,
    }.get(module_key, self._render_dashboard)
    render(container)
    self.set_status(f"Modul {module_key} geladen.")

FinanceMateApp.show_module = _v619_show_module



# === FINANCE MATE PATCH V0_6_20_RECHNUNGSPORTAL_PREVIEW_APPROVAL ===
APP_VERSION = "0.6.20-rechnungsportal-preview-approval"
APPROVAL_THRESHOLDS = [
    (1, Decimal("1000.00")),
    (2, Decimal("5000.00")),
    (3, Decimal("25000.00")),
    (4, None),
]
APPROVAL_STATUS_REVIEW = "Sachliche Prüfung ausstehend"
APPROVAL_STATUS_APPROVAL1 = "1. Freigabe ausstehend"
APPROVAL_STATUS_APPROVAL2 = "2. Freigabe ausstehend"
APPROVAL_STATUS_RELEASED = "Freigegeben"
APPROVAL_STATUS_VALUES = [APPROVAL_STATUS_REVIEW, APPROVAL_STATUS_APPROVAL1, APPROVAL_STATUS_APPROVAL2, APPROVAL_STATUS_RELEASED]


def _fm_current_user() -> str:
    return os.environ.get("FINANCEMATE_USER") or os.environ.get("USERNAME") or os.environ.get("USER") or "default"


def _fm_current_approval_level() -> int:
    # Übergangslösung bis zur echten Benutzerverwaltung: technisch Ebene 4, damit Funktionen testbar sind.
    try:
        return int(os.environ.get("FINANCEMATE_APPROVAL_LEVEL", "4"))
    except Exception:
        return 4


def _fm_invoice_direction(entity_type: str) -> str:
    return "Eingangsrechnung" if entity_type == "vendor_invoice" else "Ausgangsrechnung"


def _fm_invoice_table(entity_type: str) -> str:
    return "vendor_invoices" if entity_type == "vendor_invoice" else "customer_invoices"


def _fm_partner_cols(entity_type: str):
    if entity_type == "vendor_invoice":
        return "vendor_no", "vendor_name", "vendor_address"
    return "customer_no", "customer_name", "customer_address"


def _fm_calc_approval_level(amount) -> int:
    try:
        amt = abs(parse_amount(amount))
    except Exception:
        amt = Decimal("0.00")
    for level, threshold in APPROVAL_THRESHOLDS:
        if threshold is None or amt <= threshold:
            return level
    return 4


def _fm_status_from_open(open_amount: Any) -> str:
    try:
        return STATUS_PAID if parse_amount(open_amount) <= 0 else STATUS_OPEN
    except Exception:
        return STATUS_OPEN


def compute_status_from_open_amount(open_amount: Any, due_date: str) -> str:
    return _fm_status_from_open(open_amount)


def urgency_bucket(due_date: str, open_amount: Any) -> str:
    try:
        amount = parse_amount(open_amount)
    except Exception:
        amount = Decimal("0.00")
    if amount <= 0:
        return "paid"
    if not validate_date(due_date):
        return "open"
    delta = (datetime.strptime(due_date, DATE_FMT).date() - datetime.now().date()).days
    if delta <= 0:
        return "overdue"
    if delta <= 5:
        return "soon"
    return "open"


def configure_tree_tags(tree: ttk.Treeview) -> None:
    tree.tag_configure("paid", background=SOFT_GREEN)
    tree.tag_configure("soon", background=SOFT_ORANGE)
    tree.tag_configure("overdue", background=SOFT_RED)
    tree.tag_configure("open", background=WHITE)


_orig_init_sqlite_v620 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v620()
    with get_connection() as con:
        for table in ("customer_invoices", "vendor_invoices"):
            for col, definition in [
                ("invoice_kind", "TEXT DEFAULT 'Rechnung'"),
                ("description", "TEXT DEFAULT ''"),
                ("approval_status", f"TEXT DEFAULT '{APPROVAL_STATUS_REVIEW}'"),
                ("released", "INTEGER DEFAULT 0"),
                ("released_at", "TEXT DEFAULT ''"),
                ("reviewed_by", "TEXT DEFAULT ''"),
                ("reviewed_at", "TEXT DEFAULT ''"),
                ("approval_1_by", "TEXT DEFAULT ''"),
                ("approval_1_at", "TEXT DEFAULT ''"),
                ("approval_2_by", "TEXT DEFAULT ''"),
                ("approval_2_at", "TEXT DEFAULT ''"),
                ("approval_level", "INTEGER DEFAULT 1"),
                ("last_checked_by", "TEXT DEFAULT ''"),
                ("last_checked_at", "TEXT DEFAULT ''"),
            ]:
                ensure_column(con, table, col, definition)
        for col, definition in [
            ("invoice_kind", "TEXT DEFAULT 'Rechnung'"),
            ("description", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, "open_items", col, definition)
        # Einmalige Rollout-Bereinigung aller bestehenden Rechnungsdatensätze.
        cleanup_done = con.execute("SELECT value FROM app_meta WHERE key='v0_6_20_invoice_cleanup_done'").fetchone()
        if not cleanup_done:
            old_attachment_paths = [r["file_path"] for r in con.execute("SELECT file_path FROM attachments WHERE entity_type IN ('customer_invoice','vendor_invoice')")]
            con.execute("DELETE FROM open_items")
            con.execute("DELETE FROM attachments WHERE entity_type IN ('customer_invoice','vendor_invoice')")
            con.execute("DELETE FROM customer_invoices")
            con.execute("DELETE FROM vendor_invoices")
            con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("v0_6_20_invoice_cleanup_done", "1", now_str()))
            for path in old_attachment_paths:
                try:
                    fp = Path(path)
                    if fp.exists():
                        fp.unlink()
                except Exception:
                    pass
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()


# ---------- Mehrseiten-Bildvorschau mit Pan/Zoom und Grenzen ----------
class _FMDocumentPreviewPane(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.canvas = tk.Canvas(self, bg=WHITE, highlightthickness=0)
        self.canvas.pack(fill="both", expand=True)
        self.base_image = None
        self.tk_image = None
        self.fit_zoom = 1.0
        self.zoom = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.drag_start = None
        self.canvas.bind("<Configure>", lambda e: self._render(reset=False))
        self.canvas.bind("<ButtonPress-1>", self._on_press)
        self.canvas.bind("<B1-Motion>", self._on_drag)
        self.canvas.bind("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind("<Button-4>", lambda e: self._wheel(120))
        self.canvas.bind("<Button-5>", lambda e: self._wheel(-120))
        self.clear()

    def clear(self):
        self.base_image = None
        self.tk_image = None
        self.canvas.delete("all")
        self.canvas.create_text(20, 20, anchor="nw", text="Keine Vorschau verfügbar", fill=TEXT2, font=("Segoe UI", 10))

    def set_image(self, pil_image):
        self.base_image = pil_image.convert("RGBA") if pil_image is not None else None
        self.zoom = 1.0
        self.offset_x = 0
        self.offset_y = 0
        if self.base_image is None:
            self.clear()
        else:
            self._render(reset=True)

    def _on_press(self, event):
        self.drag_start = (event.x, event.y)

    def _on_drag(self, event):
        if self.base_image is None or self.drag_start is None:
            return
        dx = event.x - self.drag_start[0]
        dy = event.y - self.drag_start[1]
        self.drag_start = (event.x, event.y)
        self.offset_x += dx
        self.offset_y += dy
        self._clamp_offsets()
        self._draw()

    def _on_mousewheel(self, event):
        self._wheel(event.delta)

    def _wheel(self, delta):
        if self.base_image is None:
            return
        factor = 1.12 if delta > 0 else 0.88
        old_zoom = self.zoom
        self.zoom = max(self.fit_zoom, min(self.zoom * factor, self.fit_zoom * 8.0))
        if abs(self.zoom - old_zoom) > 0.0001:
            self._clamp_offsets()
            self._render(reset=False)

    def _fit_zoom(self):
        if self.base_image is None:
            return 1.0
        cw = max(50, self.canvas.winfo_width())
        ch = max(50, self.canvas.winfo_height())
        iw, ih = self.base_image.size
        return min(cw / iw, ch / ih, 1.0)

    def _render(self, reset=False):
        if self.base_image is None:
            return
        try:
            from PIL import ImageTk
            self.fit_zoom = self._fit_zoom()
            if reset or self.zoom < self.fit_zoom:
                self.zoom = self.fit_zoom
                self.offset_x = 0
                self.offset_y = 0
            self._clamp_offsets()
            iw, ih = self.base_image.size
            size = (max(1, int(iw * self.zoom)), max(1, int(ih * self.zoom)))
            img = self.base_image.resize(size)
            self.tk_image = ImageTk.PhotoImage(img)
            self._draw()
        except Exception as exc:
            self.clear()
            self.canvas.create_text(20, 48, anchor="nw", text=f"Vorschaufehler: {exc}", fill=TEXT2, font=("Segoe UI", 10))

    def _clamp_offsets(self):
        if self.base_image is None:
            return
        cw = max(50, self.canvas.winfo_width())
        ch = max(50, self.canvas.winfo_height())
        rw = int(self.base_image.size[0] * self.zoom)
        rh = int(self.base_image.size[1] * self.zoom)
        # Mindestens 20% bzw. 80px müssen sichtbar bleiben.
        min_visible_x = min(max(80, cw * 0.2), cw)
        min_visible_y = min(max(80, ch * 0.2), ch)
        max_x = max(0, (rw + cw) / 2 - min_visible_x)
        max_y = max(0, (rh + ch) / 2 - min_visible_y)
        self.offset_x = max(-max_x, min(max_x, self.offset_x))
        self.offset_y = max(-max_y, min(max_y, self.offset_y))

    def _draw(self):
        self.canvas.delete("all")
        if self.tk_image is None:
            return
        x = self.canvas.winfo_width() // 2 + int(self.offset_x)
        y = self.canvas.winfo_height() // 2 + int(self.offset_y)
        self.canvas.create_image(x, y, image=self.tk_image)


def _fm_preview_font(size=14, bold=False):
    try:
        from PIL import ImageFont
        return ImageFont.truetype("DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf", size)
    except Exception:
        from PIL import ImageFont
        return ImageFont.load_default()


def _fm_text_pages_to_image(title: str, lines: list[str], page_lines: int = 42):
    from PIL import Image, ImageDraw
    pages = [lines[i:i+page_lines] for i in range(0, max(1, len(lines)), page_lines)] or [["(Keine Inhalte gefunden)"]]
    width = 1200
    page_height = 1500
    gap = 28
    img = Image.new("RGBA", (width, len(pages) * page_height + (len(pages)-1) * gap), "#f2f4f7")
    draw = ImageDraw.Draw(img)
    title_font = _fm_preview_font(22, True)
    body_font = _fm_preview_font(15, False)
    for pidx, page in enumerate(pages):
        y0 = pidx * (page_height + gap)
        draw.rectangle((0, y0, width, y0 + page_height), fill="white", outline="#cbd3dc")
        draw.text((38, y0 + 30), f"{title} – Seite {pidx+1}/{len(pages)}", fill=TEXT, font=title_font)
        y = y0 + 82
        for line in page:
            draw.text((38, y), str(line)[:150], fill=TEXT, font=body_font)
            y += 31
    return img


def _fm_combine_pages(images):
    from PIL import Image, ImageDraw
    if not images:
        return _fm_text_pages_to_image("Vorschau", ["Keine Vorschau verfügbar"])
    width = max(im.width for im in images)
    gap = 28
    total_height = sum(im.height for im in images) + gap * (len(images)-1)
    out = Image.new("RGBA", (width, total_height), "#f2f4f7")
    y = 0
    for im in images:
        rgba = im.convert("RGBA")
        x = (width - rgba.width) // 2
        out.paste(rgba, (x, y))
        y += rgba.height + gap
    return out


def _fm_generate_document_preview(path: str):
    from PIL import Image
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(p))
            pages = []
            for idx in range(len(doc)):
                page = doc.load_page(idx)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.65, 1.65), alpha=False)
                mode = "RGB" if pix.n < 4 else "RGBA"
                pages.append(Image.frombytes(mode, [pix.width, pix.height], pix.samples))
            doc.close()
            return _fm_combine_pages(pages)
        except Exception as exc:
            return _fm_text_pages_to_image(f"PDF-Vorschau – {p.name}", [f"PDF konnte nicht gerendert werden: {exc}"])
    if ext in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}:
        return Image.open(p)
    if ext in {".txt", ".csv"}:
        return _fm_text_pages_to_image(f"Text-Vorschau – {p.name}", p.read_text(encoding="utf-8", errors="replace").splitlines())
    if ext in {".xlsx", ".xls"}:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"Arbeitsblatt: {ws.title}")
                for row in ws.iter_rows(min_row=1, max_row=120, max_col=10, values_only=True):
                    values = ["" if v is None else str(v) for v in row]
                    if any(v.strip() for v in values):
                        lines.append(" | ".join(values).rstrip())
            wb.close()
            return _fm_text_pages_to_image(f"Excel-Vorschau – {p.name}", lines or ["(Keine Inhalte gefunden)"])
        except Exception as exc:
            return _fm_text_pages_to_image(f"Excel-Vorschau – {p.name}", [f"Excel konnte nicht gelesen werden: {exc}"])
    if ext in {".docx", ".doc"}:
        try:
            import docx
            doc = docx.Document(str(p))
            lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    vals = [cell.text.strip() for cell in row.cells]
                    if any(vals):
                        lines.append(" | ".join(vals))
            return _fm_text_pages_to_image(f"Word-Vorschau – {p.name}", lines or ["(Keine Inhalte gefunden)"])
        except Exception as exc:
            return _fm_text_pages_to_image(f"Word-Vorschau – {p.name}", [f"Word konnte nicht gelesen werden: {exc}"])
    return _fm_text_pages_to_image(f"Dokument-Vorschau – {p.name}", [collect_preview_text(str(p), 8000)])


def _fm_creditors_build_stack_ui_v620(self, parent):
    split = build_two_block_shell(parent, "Dokumentenstapel", "Vorschau")
    left = ttk.Frame(split.left); left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(split.right); right.pack(fill="both", expand=True, padx=10, pady=8)
    toolbar = ttk.Frame(left); toolbar.pack(fill="x", pady=(0, 6))
    create_standard_button(toolbar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=3)
    create_standard_button(toolbar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=3)
    cols = ("hist_no", "title", "status", "mark")
    self.stack_tree = ttk.Treeview(left, columns=cols, show="headings", height=10)
    for c, t, w in [("hist_no", "Hist.-Nr.", 110), ("title", "Dokument", 260), ("status", "Status", 110), ("mark", "☑", 45)]:
        self.stack_tree.heading(c, text=t, command=self._toggle_all_marks if c == "mark" and hasattr(self, "_toggle_all_marks") else None)
        self.stack_tree.column(c, width=w, anchor="w" if c != "mark" else "center")
    self.stack_tree.pack(fill="both", expand=True)
    self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select, add="+")
    self.stack_tree.bind("<Double-1>", lambda _e: self._open_selected_stack_file())
    self.preview_title_label = ttk.Label(right, text="Kein Dokument ausgewählt", style="Hint.TLabel")
    self.preview_title_label.pack(anchor="w", pady=(0, 6))
    self.preview_pane = _FMDocumentPreviewPane(right)
    self.preview_pane.pack(fill="both", expand=True)

CreditorsView._build_stack_ui = _fm_creditors_build_stack_ui_v620


def _fm_creditors_clear_stack_preview_v620(self):
    if hasattr(self, "preview_title_label"):
        self.preview_title_label.configure(text="Kein Dokument ausgewählt")
    if hasattr(self, "preview_pane"):
        self.preview_pane.clear()
    if hasattr(self, "preview_image_label"):
        self.preview_image_label.configure(text="Keine Vorschau verfügbar", image="")

CreditorsView._clear_stack_preview = _fm_creditors_clear_stack_preview_v620


def _fm_creditors_refresh_stack_preview_v620(self):
    if not getattr(self, "current_hist_no", None):
        self._clear_stack_preview(); return
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_files WHERE hist_no=? ORDER BY id", (self.current_hist_no,)))
    if not rows:
        self._clear_stack_preview(); return
    title = clean_document_title(rows[0]["file_name"])
    if len(rows) > 1:
        title += f" (+ {len(rows)-1} weitere)"
    self.preview_title_label.configure(text=title)
    try:
        images = [_fm_generate_document_preview(r["file_path"]) for r in rows]
        self.preview_pane.set_image(_fm_combine_pages(images))
    except Exception as exc:
        self.preview_pane.set_image(_fm_text_pages_to_image("Vorschaufehler", [str(exc)]))

CreditorsView._refresh_stack_preview = _fm_creditors_refresh_stack_preview_v620


# ---------- Freigabe- und Rechnungslogik ----------
def _fm_validate_positive_invoice_amounts(self) -> bool:
    for label, var in [("Netto", self.net_amount), ("Steuer", self.tax_amount), ("Brutto", self.gross_amount)]:
        try:
            if parse_amount(var.get()) < 0:
                messagebox.showwarning("Negativbetrag", f"{label} darf nicht negativ sein. Bitte Belegart 'Gutschrift' verwenden.")
                return False
        except Exception as exc:
            messagebox.showwarning("Betrag", str(exc))
            return False
    return True


def _fm_description_get(self):
    w = getattr(self, "description_widget", None)
    return w.get("1.0", "end-1c").strip() if w is not None else ""


def _fm_create_open_item_for_invoice(con, entity_type: str, invoice_row):
    if entity_type == "vendor_invoice":
        partner_no, partner_name = invoice_row["vendor_no"], invoice_row["vendor_name"]
    else:
        partner_no, partner_name = invoice_row["customer_no"], invoice_row["customer_name"]
    con.execute("DELETE FROM open_items WHERE entity_type=? AND reference_no=?", (entity_type, invoice_row["invoice_no"]))
    con.execute(
        "INSERT INTO open_items(entity_type,reference_no,partner_no,partner_name,due_date,original_amount,open_amount,status,linked_journal_no,invoice_kind,description,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (entity_type, invoice_row["invoice_no"], partner_no, partner_name, invoice_row["due_date"], float(invoice_row["gross_amount"] or 0), float(invoice_row["open_amount"] or invoice_row["gross_amount"] or 0), _fm_status_from_open(invoice_row["open_amount"] or invoice_row["gross_amount"] or 0), invoice_row["linked_journal_no"] if "linked_journal_no" in invoice_row.keys() else "", invoice_row["invoice_kind"] if "invoice_kind" in invoice_row.keys() else "Rechnung", invoice_row["description"] if "description" in invoice_row.keys() else "", now_str(), now_str())
    )


def _fm_save_invoice_v620(self):
    if not _fm_validate_positive_invoice_amounts(self):
        return
    if not validate_date(self.invoice_date.get()) or not validate_date(self.due_date.get()):
        messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden."); return
    no, name, address = self._partner_parts()
    if not no:
        messagebox.showwarning("Partner", "Bitte Partner auswählen."); return
    selected_hist = getattr(self, "current_hist_no", None)
    if self.entity_type == "vendor_invoice" and not (selected_hist or getattr(self, "pending_attachments", [])):
        messagebox.showwarning("Dokument", "Eingangsrechnungen können ohne Dokument nicht erfasst werden."); return
    gross = parse_amount(self.gross_amount.get()); net = parse_amount(self.net_amount.get()); tax = parse_amount(self.tax_amount.get())
    invoice_no = self.invoice_no.get().strip()
    invoice_kind = self.invoice_kind.get().strip() if hasattr(self, "invoice_kind") else "Rechnung"
    description = _fm_description_get(self)
    address_col = "customer_address" if self.invoice_table == "customer_invoices" else "vendor_address"
    partner_name_col = "customer_name" if self.invoice_table == "customer_invoices" else "vendor_name"
    partner_no_col = "customer_no" if self.invoice_table == "customer_invoices" else "vendor_no"
    approval_status = APPROVAL_STATUS_REVIEW
    approval_level = _fm_calc_approval_level(gross)
    stack_files = []
    with get_connection() as con:
        con.execute(
            f"INSERT INTO {self.invoice_table}(invoice_no,{partner_no_col},{partner_name_col},{address_col},invoice_date,due_date,payment_term,tax_code,net_amount,tax_amount,gross_amount,open_amount,status,invoice_kind,description,approval_status,released,approval_level,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (invoice_no, no, name, address, self.invoice_date.get(), self.due_date.get(), self.payment_term.get(), self.tax_code.get(), float(net), float(tax), float(gross), float(gross), STATUS_OPEN, invoice_kind, description, approval_status, 0, approval_level, now_str(), now_str())
        )
        # Keine offenen Posten vor Freigabe.
        con.execute("DELETE FROM open_items WHERE entity_type=? AND reference_no=?", (self.entity_type, invoice_no))
        if self.entity_type == "vendor_invoice" and selected_hist:
            stack_files = list(con.execute("SELECT file_path FROM invoice_import_files WHERE hist_no=? ORDER BY id", (selected_hist,)))
            con.execute("UPDATE invoice_import_batches SET status='Erfasst', vendor_no=?, vendor_name=?, updated_at=? WHERE hist_no=?", (no, name, now_str(), selected_hist))
        con.commit()
    if stack_files:
        self.add_attachment_paths(self.entity_type, invoice_no, [r["file_path"] for r in stack_files])
    if getattr(self, "pending_attachments", []):
        self.add_attachment_paths(self.entity_type, invoice_no, self.pending_attachments)
    self.app.set_status(f"{self.entity_label} {invoice_no} gespeichert – Freigabe ausstehend.")
    self.clear_form(); self.reload_invoices(); self.reload_open_items()
    if hasattr(self, "_load_import_batches"): self._load_import_batches()
    if hasattr(self, "_clear_stack_preview"): self._clear_stack_preview()

InvoiceModuleBase.save_invoice = _fm_save_invoice_v620


def _fm_reload_open_items_v620(self):
    trees = _v619_unique_trees(getattr(self, "op_tree", None), *getattr(self, "open_item_tabs_trees", [])) if "_v619_unique_trees" in globals() else [getattr(self, "op_tree", None)]
    trees = [t for t in trees if t is not None]
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM open_items WHERE entity_type=? ORDER BY due_date, id", (self.entity_type,)))
    for tree in trees:
        for i in tree.get_children(): tree.delete(i)
        for r in rows:
            if parse_amount(r["open_amount"]) <= 0: continue
            att = self.get_attachment_count(self.entity_type, r["reference_no"])
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            cols = tree["columns"]
            if "kind" in cols:
                vals = (r["reference_no"], r["invoice_kind"] if "invoice_kind" in r.keys() else "Rechnung", r["partner_name"], r["description"] if "description" in r.keys() else "", r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else "Dokument anhängen")
            else:
                vals = (r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"], f"📎 {att}" if att else "Dokument anhängen")
            tree.insert("", "end", values=vals, tags=(tag,))

InvoiceModuleBase.reload_open_items = _fm_reload_open_items_v620


# ---------- Rechnungsportal ----------
class InvoicePortalView(ttk.Frame, SortableTreeMixin):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.archive_filter = tk.StringVar(value="Alle")
        self.approval_filter = tk.StringVar(value="Alle")
        self._build_ui()
        self.reload_all()

    def _build_ui(self):
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)
        self.archive_tab = ttk.Frame(self.nb)
        self.approval_tab = ttk.Frame(self.nb)
        self.nb.add(self.archive_tab, text="Rechnungsarchiv")
        self.nb.add(self.approval_tab, text="Rechnungsfreigabe")
        self._build_archive()
        self._build_approval()

    def _invoice_tree(self, parent):
        cols = ("direction", "invoice_no", "kind", "partner", "date", "due", "gross", "open", "approval_status", "last_checked_by", "released")
        tree = ttk.Treeview(parent, columns=cols, show="headings")
        for c, t, w in [
            ("direction", "Richtung", 125), ("invoice_no", "Rechnung", 120), ("kind", "Art", 95), ("partner", "Partner", 190),
            ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 95), ("open", "Offen", 95),
            ("approval_status", "Freigabestatus", 190), ("last_checked_by", "Letzter Prüfender", 150), ("released", "Freigegeben", 90)
        ]:
            tree.heading(c, text=t); tree.column(c, width=w, anchor="w")
        tree.pack(fill="both", expand=True, padx=10, pady=8)
        configure_tree_tags(tree); self.setup_sorting(tree)
        return tree

    def _build_archive(self):
        top = ttk.Frame(self.archive_tab); top.pack(fill="x", padx=10, pady=8)
        ttk.Label(top, text="Filter", width=8).pack(side="left")
        cb = ttk.Combobox(top, textvariable=self.archive_filter, values=["Alle", "Eingangsrechnungen", "Ausgangsrechnungen"], state="readonly", width=24)
        cb.pack(side="left", padx=4); cb.bind("<<ComboboxSelected>>", lambda _e: self.reload_archive()); cb.bind("<MouseWheel>", lambda e: "break")
        create_standard_button(top, "Aktualisieren", self.reload_archive).pack(side="left", padx=4)
        create_standard_button(top, "Bearbeiten", self.edit_selected_archive).pack(side="left", padx=4)
        self.archive_tree = self._invoice_tree(self.archive_tab)

    def _build_approval(self):
        top = ttk.Frame(self.approval_tab); top.pack(fill="x", padx=10, pady=8)
        ttk.Label(top, text="Filter", width=8).pack(side="left")
        cb = ttk.Combobox(top, textvariable=self.approval_filter, values=["Alle", "Eingangsrechnungen", "Ausgangsrechnungen"], state="readonly", width=24)
        cb.pack(side="left", padx=4); cb.bind("<<ComboboxSelected>>", lambda _e: self.reload_approval()); cb.bind("<MouseWheel>", lambda e: "break")
        create_standard_button(top, "Aktualisieren", self.reload_approval).pack(side="left", padx=4)
        create_standard_button(top, "Bearbeiten", self.edit_selected_approval).pack(side="left", padx=4)
        create_standard_button(top, "Sachlich prüfen", lambda: self.perform_action("review")).pack(side="left", padx=4)
        create_standard_button(top, "1. Freigabe", lambda: self.perform_action("approval1"), confirm=True).pack(side="left", padx=4)
        create_standard_button(top, "2. Freigabe", lambda: self.perform_action("approval2"), confirm=True).pack(side="left", padx=4)
        self.approval_tree = self._invoice_tree(self.approval_tab)

    def _filter_sql(self, flt):
        if flt == "Eingangsrechnungen": return ["vendor_invoice"]
        if flt == "Ausgangsrechnungen": return ["customer_invoice"]
        return ["vendor_invoice", "customer_invoice"]

    def _load_rows(self, flt, only_unreleased=False):
        rows = []
        with get_connection() as con:
            if "vendor_invoice" in self._filter_sql(flt):
                q = "SELECT 'vendor_invoice' AS entity_type, invoice_no, invoice_kind, vendor_name AS partner, invoice_date, due_date, gross_amount, open_amount, approval_status, last_checked_by, released FROM vendor_invoices"
                if only_unreleased: q += " WHERE COALESCE(released,0)=0"
                rows.extend(con.execute(q).fetchall())
            if "customer_invoice" in self._filter_sql(flt):
                q = "SELECT 'customer_invoice' AS entity_type, invoice_no, invoice_kind, customer_name AS partner, invoice_date, due_date, gross_amount, open_amount, approval_status, last_checked_by, released FROM customer_invoices"
                if only_unreleased: q += " WHERE COALESCE(released,0)=0"
                rows.extend(con.execute(q).fetchall())
        return sorted(rows, key=lambda r: (r["invoice_date"], r["invoice_no"]), reverse=True)

    def _fill_tree(self, tree, rows):
        for i in tree.get_children(): tree.delete(i)
        for r in rows:
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            values = (_fm_invoice_direction(r["entity_type"]), r["invoice_no"], r["invoice_kind"] or "Rechnung", r["partner"], r["invoice_date"], r["due_date"], format_amount(r["gross_amount"]), format_amount(r["open_amount"]), r["approval_status"] or APPROVAL_STATUS_REVIEW, r["last_checked_by"] or "", "Ja" if int(r["released"] or 0) else "Nein")
            tree.insert("", "end", iid=f"{r['entity_type']}|{r['invoice_no']}", values=values, tags=(tag,))

    def reload_archive(self): self._fill_tree(self.archive_tree, self._load_rows(self.archive_filter.get(), False))
    def reload_approval(self): self._fill_tree(self.approval_tree, self._load_rows(self.approval_filter.get(), True))
    def reload_all(self): self.reload_archive(); self.reload_approval()

    def _selected(self, tree):
        iid = tree.focus()
        if not iid or "|" not in iid:
            messagebox.showinfo("Rechnungsportal", "Bitte Rechnung auswählen."); return None
        entity_type, invoice_no = iid.split("|", 1)
        return entity_type, invoice_no

    def edit_selected_archive(self):
        sel = self._selected(self.archive_tree)
        if sel: self.open_edit_popup(*sel)

    def edit_selected_approval(self):
        sel = self._selected(self.approval_tree)
        if sel: self.open_edit_popup(*sel)

    def open_edit_popup(self, entity_type, invoice_no):
        table = _fm_invoice_table(entity_type)
        no_col, name_col, addr_col = _fm_partner_cols(entity_type)
        with get_connection() as con:
            r = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
            payment_terms = [x[0] for x in con.execute("SELECT code FROM payment_terms WHERE active=1 ORDER BY code")]
            tax_codes = [x[0] for x in con.execute("SELECT code FROM tax_codes WHERE active=1 ORDER BY code")]
        if not r: return
        if int(r["released"] or 0):
            messagebox.showinfo("Bearbeitung gesperrt", "Diese Rechnung ist freigegeben und kann nicht mehr geändert werden."); return
        pop = tk.Toplevel(self); pop.title(f"Rechnung bearbeiten – {invoice_no}"); pop.geometry("620x620"); pop.configure(bg=BG)
        vars = {
            "invoice_kind": tk.StringVar(value=r["invoice_kind"] or "Rechnung"), "invoice_no": tk.StringVar(value=r["invoice_no"]),
            "invoice_date": tk.StringVar(value=r["invoice_date"]), "due_date": tk.StringVar(value=r["due_date"]),
            "payment_term": tk.StringVar(value=r["payment_term"] or ""), "tax_code": tk.StringVar(value=r["tax_code"] or ""),
            "net_amount": tk.StringVar(value=format_amount(r["net_amount"])), "tax_amount": tk.StringVar(value=format_amount(r["tax_amount"])),
            "gross_amount": tk.StringVar(value=format_amount(r["gross_amount"])),
        }
        def row(label, widget):
            fr = ttk.Frame(pop); fr.pack(fill="x", padx=12, pady=4); ttk.Label(fr, text=label, width=18).pack(side="left"); widget(fr).pack(side="left", fill="x", expand=True)
        row("Belegart", lambda fr: ttk.Combobox(fr, textvariable=vars["invoice_kind"], values=INVOICE_KIND_VALUES, state="readonly"))
        row("Rechnungs-Nr.", lambda fr: ttk.Entry(fr, textvariable=vars["invoice_no"]))
        row("Rechnungsdatum", lambda fr: ttk.Entry(fr, textvariable=vars["invoice_date"]))
        row("Zahlungsbed.", lambda fr: ttk.Combobox(fr, textvariable=vars["payment_term"], values=payment_terms, state="readonly"))
        row("Fälligkeit", lambda fr: ttk.Entry(fr, textvariable=vars["due_date"]))
        row("Steuerkennz.", lambda fr: ttk.Combobox(fr, textvariable=vars["tax_code"], values=tax_codes, state="readonly"))
        row("Netto", lambda fr: ttk.Entry(fr, textvariable=vars["net_amount"]))
        row("Steuer", lambda fr: ttk.Entry(fr, textvariable=vars["tax_amount"]))
        row("Brutto", lambda fr: ttk.Entry(fr, textvariable=vars["gross_amount"]))
        txtrow = ttk.Frame(pop); txtrow.pack(fill="both", expand=True, padx=12, pady=4); ttk.Label(txtrow, text="Beschreibung", width=18).pack(side="left", anchor="n")
        desc = tk.Text(txtrow, height=4, wrap="word", bg=WHITE, fg=TEXT); desc.pack(side="left", fill="both", expand=True); desc.insert("1.0", r["description"] if "description" in r.keys() else "")
        def save():
            try:
                net, tax, gross = parse_amount(vars["net_amount"].get()), parse_amount(vars["tax_amount"].get()), parse_amount(vars["gross_amount"].get())
                if net < 0 or tax < 0 or gross < 0:
                    messagebox.showwarning("Negativbetrag", "Rechnungen/Gutschriften werden mit positivem Betrag erfasst.", parent=pop); return
            except Exception as exc:
                messagebox.showwarning("Betrag", str(exc), parent=pop); return
            if not validate_date(vars["invoice_date"].get()) or not validate_date(vars["due_date"].get()):
                messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.", parent=pop); return
            with get_connection() as con:
                con.execute(f"UPDATE {table} SET invoice_no=?, invoice_kind=?, invoice_date=?, due_date=?, payment_term=?, tax_code=?, net_amount=?, tax_amount=?, gross_amount=?, open_amount=?, description=?, approval_level=?, updated_at=? WHERE invoice_no=?",
                            (vars["invoice_no"].get().strip(), vars["invoice_kind"].get(), vars["invoice_date"].get(), vars["due_date"].get(), vars["payment_term"].get(), vars["tax_code"].get(), float(net), float(tax), float(gross), float(gross), desc.get("1.0", "end-1c").strip(), _fm_calc_approval_level(gross), now_str(), invoice_no))
                con.commit()
            pop.destroy(); self.reload_all()
        create_standard_button(pop, "Speichern", save, confirm=True).pack(pady=12)

    def perform_action(self, action):
        sel = self._selected(self.approval_tree)
        if not sel: return
        entity_type, invoice_no = sel
        table = _fm_invoice_table(entity_type)
        user = _fm_current_user()
        level = _fm_current_approval_level()
        with get_connection() as con:
            r = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
            if not r or int(r["released"] or 0): return
            if entity_type == "customer_invoice":
                if action != "review":
                    messagebox.showinfo("Ausgangsrechnung", "Ausgangsrechnungen benötigen nur die sachliche Prüfung durch Ebene 3 oder 4."); return
                if level < 3:
                    messagebox.showwarning("Berechtigung", "Sachliche Prüfung für Ausgangsrechnungen benötigt Freigabeebene 3 oder 4."); return
                con.execute(f"UPDATE {table} SET reviewed_by=?, reviewed_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, released=1, released_at=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_RELEASED, now_str(), now_str(), invoice_no))
                rr = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
                _fm_create_open_item_for_invoice(con, entity_type, rr)
            else:
                if action == "review":
                    con.execute(f"UPDATE {table} SET reviewed_by=?, reviewed_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_APPROVAL1, now_str(), invoice_no))
                elif action == "approval1":
                    if not r["reviewed_by"]:
                        messagebox.showwarning("Freigabe", "Sachliche Prüfung ist noch ausstehend."); return
                    con.execute(f"UPDATE {table} SET approval_1_by=?, approval_1_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_APPROVAL2, now_str(), invoice_no))
                elif action == "approval2":
                    if not r["approval_1_by"]:
                        messagebox.showwarning("Freigabe", "1. Freigabe ist noch ausstehend."); return
                    if r["approval_1_by"] == user:
                        messagebox.showwarning("Freigabe", "1. und 2. Freigabe müssen von unterschiedlichen Benutzern erfolgen."); return
                    con.execute(f"UPDATE {table} SET approval_2_by=?, approval_2_at=?, last_checked_by=?, last_checked_at=?, approval_status=?, released=1, released_at=?, updated_at=? WHERE invoice_no=?", (user, now_str(), user, now_str(), APPROVAL_STATUS_RELEASED, now_str(), now_str(), invoice_no))
                    rr = con.execute(f"SELECT * FROM {table} WHERE invoice_no=?", (invoice_no,)).fetchone()
                    _fm_create_open_item_for_invoice(con, entity_type, rr)
            con.commit()
        self.reload_all()


def _render_rechnungsportal(self, parent):
    InvoicePortalView(parent, self).pack(fill="both", expand=True)

FinanceMateApp._render_rechnungsportal = _render_rechnungsportal


def _fm_build_sidebar_v620(self):
    for w in self.sidebar.winfo_children(): w.destroy()
    width = 42 if self.sidebar_collapsed else 210
    self.sidebar.configure(width=width); self.sidebar.pack_propagate(False)
    arrow = "»" if self.sidebar_collapsed else "«"
    create_standard_button(self.sidebar, arrow, self._toggle_sidebar).pack(fill="x", padx=5, pady=5)
    modules = [("Dashboard", "DB"), ("Stammdaten", "ST"), ("Buchungsjournal", "BJ"), ("Debitoren", "DE"), ("Kreditoren", "KR"), ("Rechnungsportal", "RP"), ("Zahlungen", "ZA"), ("Reporting", "RE"), ("Audit", "AU"), ("Einstellungen", "EI")]
    self.nav_buttons.clear()
    for name, abbr in modules:
        text = abbr if self.sidebar_collapsed else name
        btn = ttk.Button(self.sidebar, text=text, style="Nav.TButton", command=lambda n=name: self.show_module(n))
        btn.pack(fill="x", padx=5, pady=2); self.nav_buttons[name] = btn

FinanceMateApp._build_sidebar = _fm_build_sidebar_v620


def _fm_show_module_v620(self, module: str):
    module_key = "Buchungsjournal" if module == "Finanzbuchhaltung" else module
    self.active_module = module_key; self._update_nav_styles()
    for w in self.workspace.winfo_children(): w.destroy()
    ttk.Label(self.workspace, text=module_key, style="CardTitle.TLabel").pack(anchor="w", pady=(0, 6))
    container = ttk.Frame(self.workspace); container.pack(fill="both", expand=True)
    render = {
        "Dashboard": self._render_dashboard, "Stammdaten": self._render_stammdaten, "Buchungsjournal": self._render_finanzbuchhaltung,
        "Debitoren": self._render_debitoren, "Kreditoren": self._render_kreditoren, "Rechnungsportal": self._render_rechnungsportal,
        "Zahlungen": self._render_zahlungen, "Reporting": self._render_reporting, "Audit": self._render_audit, "Einstellungen": self._render_einstellungen,
    }.get(module_key, self._render_dashboard)
    render(container); self.set_status(f"Modul {module_key} geladen.")

FinanceMateApp.show_module = _fm_show_module_v620



# === FINANCE MATE PATCH V0_6_21_CREDITORS_HARD_RESTORE_SCALING ===
APP_VERSION = "0.6.21-creditors-hard-restore-scaling"
BLOCK7_PREPARATION_STATUS = "vorbereitet - nach erfolgreicher Kreditorenstabilisierung fachlich fortsetzen"

# Windows-DPI möglichst früh auf 100%-Logik normalisieren. Fehler werden bewusst ignoriert,
# damit Linux/macOS und ältere Windows-Versionen weiter starten.
try:
    if sys.platform.startswith("win"):
        import ctypes
        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(1)  # system-DPI-aware, vermeidet Windows-Bitmapskalierung
        except Exception:
            try:
                ctypes.windll.user32.SetProcessDPIAware()
            except Exception:
                pass
except Exception:
    pass


def _fm_apply_1920x1080_scaling(root: tk.Tk) -> None:
    """Robuste Gesamtlogik: Basis 1920x1080 bei 100%, kleinere Anzeigen werden proportional verdichtet."""
    try:
        sw = max(1, root.winfo_screenwidth())
        sh = max(1, root.winfo_screenheight())
        scale = min(sw / 1920.0, sh / 1080.0)
        scale = max(0.72, min(1.15, scale))
        # Tk 100% entspricht i. d. R. 96/72 = 1.3333 Pixel je Punkt.
        root.tk.call("tk", "scaling", 96.0 / 72.0 * scale)
        w = min(1920, max(1100, int(1920 * scale)))
        h = min(1080, max(680, int(1080 * scale)))
        root.geometry(f"{w}x{h}")
        root.minsize(min(1100, w), min(680, h))
    except Exception:
        pass


_orig_app_init_v621 = FinanceMateApp.__init__

def _fm_app_init_v621(self):
    _orig_app_init_v621(self)
    _fm_apply_1920x1080_scaling(self)

FinanceMateApp.__init__ = _fm_app_init_v621


def _fm_creditors_build_ui_v621(self):
    """Kreditoren-Hard-Restore ohne TwoBlock-Abhängigkeit: erzeugt sichtbar alle vier Bereiche."""
    self.current_hist_no = None
    self.pending_attachments = []

    root = ttk.PanedWindow(self, orient="vertical")
    root.pack(fill="both", expand=True, padx=0, pady=0)
    top = ttk.Frame(root)
    bottom = ttk.Frame(root)
    root.add(top, weight=3)
    root.add(bottom, weight=2)

    top_split = ttk.PanedWindow(top, orient="horizontal")
    top_split.pack(fill="both", expand=True)
    form_card = ttk.LabelFrame(top_split, text="Eingangsrechnung erfassen")
    list_card = ttk.LabelFrame(top_split, text="Eingangsrechnungen")
    top_split.add(form_card, weight=2)
    top_split.add(list_card, weight=3)

    form_scroll = ScrollableFrame(form_card)
    form_scroll.pack(fill="both", expand=True, padx=10, pady=8)
    f = form_scroll.content

    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()

    self._labeled_combo(f, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(f, "Rechnungs-Nr.", self.invoice_no)

    prow = ttk.Frame(f); prow.pack(fill="x", pady=3)
    ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side="left", fill="x", expand=True)
    self.partner_cb.bind("<MouseWheel>", lambda e: "break")
    self.partner_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_vendor_info())
    create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)

    self.description_widget = _v619_labeled_text(f, "Rechnungsbeschreibung", height=2) if "_v619_labeled_text" in globals() else tk.Text(f, height=2)
    if not getattr(self.description_widget, "winfo_manager", lambda: "")():
        self.description_widget.pack(fill="x", pady=3)

    self._labeled_entry(f, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(f, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(f, "Fälligkeit", self.due_date)
    self._labeled_combo(f, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(f, "Netto", self.net_amount)
    self._labeled_entry(f, "Steuer", self.tax_amount)
    self._labeled_entry(f, "Brutto", self.gross_amount)

    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())

    btns = ttk.Frame(f); btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)

    info = ttk.LabelFrame(f, text="Stammdaten des Kreditors")
    info.pack(fill="both", expand=False, pady=(8, 4))
    self.vendor_info_body = tk.Text(info, height=8, wrap="word", bg=WHITE, fg=TEXT, relief="flat")
    self.vendor_info_body.pack(fill="both", expand=True, padx=6, pady=6)
    self.vendor_info_body.configure(state="disabled")

    nb = ttk.Notebook(list_card)
    nb.pack(fill="both", expand=True, padx=10, pady=8)
    tab_inv = ttk.Frame(nb); tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text="Eingangsrechnungen")
    nb.add(tab_op, text="Offene Posten")

    inv_cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show="headings", height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [("invoice_no", "Rechnung", 115), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t); self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree); self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")

    op_cols = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show="headings", height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t); self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree); self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

    # Unterer Bereich: Dokumentenstapel + Bildvorschau.
    self._build_stack_ui(bottom)

    self.partner.trace_add("write", lambda *_: self._update_vendor_info())
    self._update_vendor_info()
    self.reload_invoices()
    self.reload_open_items()
    self._load_import_batches()
    self._clear_stack_preview()

CreditorsView._build_ui = _fm_creditors_build_ui_v621


def _fm_render_kreditoren_v621(self, parent):
    try:
        view = CreditorsView(parent, self)
        view.pack(fill="both", expand=True)
        # Harte Mindestprüfung direkt nach Aufbau: zentrale Widgets müssen existieren.
        required = ["inv_tree", "op_tree", "stack_tree"]
        missing = [name for name in required if not hasattr(view, name)]
        if missing:
            raise RuntimeError("Kreditoren-Aufbau unvollständig, fehlende Widgets: " + ", ".join(missing))
    except Exception as exc:
        # Kein stiller grauer Hintergrund mehr: Fehler sichtbar ausgeben.
        for w in parent.winfo_children():
            w.destroy()
        err = tk.Text(parent, height=12, bg="#fff4c2", fg=TEXT, wrap="word")
        err.pack(fill="both", expand=True, padx=12, pady=12)
        err.insert("1.0", "Kreditoren-Modul konnte nicht vollständig aufgebaut werden.\n\n" + str(exc))
        err.configure(state="disabled")
        raise

FinanceMateApp._render_kreditoren = _fm_render_kreditoren_v621


# Stammdaten-Doppeltitel: Im Modul Stammdaten bleibt nur die interne Modulüberschrift sichtbar.
def _fm_show_module_v621(self, module: str):
    module_key = "Buchungsjournal" if module == "Finanzbuchhaltung" else module
    self.active_module = module_key
    self._update_nav_styles()
    for w in self.workspace.winfo_children():
        w.destroy()
    if module_key != "Stammdaten":
        ttk.Label(self.workspace, text=module_key, style="CardTitle.TLabel").pack(anchor="w", pady=(0, 6))
    container = ttk.Frame(self.workspace)
    container.pack(fill="both", expand=True)
    render = {
        "Dashboard": self._render_dashboard,
        "Stammdaten": self._render_stammdaten,
        "Buchungsjournal": self._render_finanzbuchhaltung,
        "Debitoren": self._render_debitoren,
        "Kreditoren": self._render_kreditoren,
        "Rechnungsportal": self._render_rechnungsportal,
        "Zahlungen": self._render_zahlungen,
        "Reporting": self._render_reporting,
        "Audit": self._render_audit,
        "Einstellungen": self._render_einstellungen,
    }.get(module_key, self._render_dashboard)
    render(container)
    self.set_status(f"Modul {module_key} geladen.")

FinanceMateApp.show_module = _fm_show_module_v621


import inspect

def _fm_creditors_static_selftest_v621() -> None:
    # Statische Auslieferungsprüfung: Der aktive Kreditoren-Build darf nicht nur Container anlegen.
    try:
        src = inspect.getsource(CreditorsView._build_ui)
    except Exception:
        src = ""
    required_tokens = [
        "Eingangsrechnung erfassen", "Eingangsrechnungen", "Offene Posten",
        "Dokumentenstapel", "self.inv_tree", "self.op_tree", "self._build_stack_ui(bottom)"
    ]
    missing = [t for t in required_tokens if t not in src]
    if missing:
        raise RuntimeError("Kreditoren-Selbsttest fehlgeschlagen: " + ", ".join(missing))

_fm_creditors_static_selftest_v621()



# === FINANCE MATE PATCH V0_6_22_CREDITORS_COMMAND_FIX_SEARCH_FULLSCREEN ===
APP_VERSION = "0.6.22-creditors-command-fix-search-fullscreen"

# Kernkorrektur für Python/Tk 3.14: ttk.Treeview.heading darf command=None NICHT erhalten.
# Wenn kein Command existiert, wird die Option command vollständig weggelassen.
def _fm_heading_safe(tree: ttk.Treeview, column: str, text: str, command=None) -> None:
    if callable(command):
        tree.heading(column, text=text, command=command)
    else:
        tree.heading(column, text=text)


def _fm_creditors_toggle_all_marks_v622(self):
    try:
        with get_connection() as con:
            rows = list(con.execute("SELECT hist_no, marked FROM invoice_import_batches WHERE status <> 'Erfasst' ORDER BY id"))
            if not rows:
                return
            all_marked = all(int(r["marked"] or 0) for r in rows)
            new_val = 0 if all_marked else 1
            con.execute("UPDATE invoice_import_batches SET marked=?, updated_at=? WHERE status <> 'Erfasst'", (new_val, now_str()))
            con.commit()
        self._load_import_batches()
    except Exception as exc:
        messagebox.showerror("Dokumentenstapel", str(exc), parent=self)

CreditorsView._toggle_all_marks = _fm_creditors_toggle_all_marks_v622


def _fm_creditors_build_stack_ui_v622(self, parent):
    split = build_two_block_shell(parent, "Dokumentenstapel", "Vorschau")
    left = ttk.Frame(split.left); left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(split.right); right.pack(fill="both", expand=True, padx=10, pady=8)

    toolbar = ttk.Frame(left); toolbar.pack(fill="x", pady=(0, 6))
    create_standard_button(toolbar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=3)
    create_standard_button(toolbar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=3)

    # Suchleiste Dokumentenstapel
    self.stack_search_var = tk.StringVar()
    search_row = ttk.Frame(left); search_row.pack(fill="x", pady=(0, 6))
    ttk.Label(search_row, text="Suchen", width=10).pack(side="left")
    stack_search = ttk.Entry(search_row, textvariable=self.stack_search_var)
    stack_search.pack(side="left", fill="x", expand=True)
    stack_search.insert(0, PLACEHOLDER_TEXT)
    stack_search.bind("<FocusIn>", lambda e: (stack_search.delete(0, tk.END) if stack_search.get() == PLACEHOLDER_TEXT else None))
    stack_search.bind("<FocusOut>", lambda e: (stack_search.insert(0, PLACEHOLDER_TEXT) if not stack_search.get().strip() else None))
    stack_search.bind("<KeyRelease>", lambda _e: self._load_import_batches())

    cols = ("hist_no", "title", "status", "mark")
    self.stack_tree = ttk.Treeview(left, columns=cols, show="headings", height=10)
    for c, t, w in [("hist_no", "Hist.-Nr.", 110), ("title", "Dokument", 260), ("status", "Status", 110), ("mark", "☑", 45)]:
        _fm_heading_safe(self.stack_tree, c, t, self._toggle_all_marks if c == "mark" else None)
        self.stack_tree.column(c, width=w, anchor="w" if c != "mark" else "center")
    self.stack_tree.pack(fill="both", expand=True)
    self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select, add="+")
    self.stack_tree.bind("<Double-1>", lambda _e: self._open_selected_stack_file())

    self.preview_title_label = ttk.Label(right, text="Kein Dokument ausgewählt", style="Hint.TLabel")
    self.preview_title_label.pack(anchor="w", pady=(0, 6))
    self.preview_pane = _FMDocumentPreviewPane(right)
    self.preview_pane.pack(fill="both", expand=True)

CreditorsView._build_stack_ui = _fm_creditors_build_stack_ui_v622


def _fm_creditors_load_import_batches_v622(self):
    if not hasattr(self, "stack_tree"):
        return
    query = ""
    try:
        query = self.stack_search_var.get().strip().lower()
        if query == PLACEHOLDER_TEXT.lower():
            query = ""
    except Exception:
        query = ""
    for i in self.stack_tree.get_children():
        self.stack_tree.delete(i)
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM invoice_import_batches WHERE status <> 'Erfasst' ORDER BY id DESC"))
    for r in rows:
        text_blob = f"{r['hist_no']} {r['title']} {r['status']}".lower()
        if query and query not in text_blob:
            continue
        self.stack_tree.insert("", "end", iid=r["hist_no"], values=(r["hist_no"], r["title"], r["status"], "☑" if int(r["marked"] or 0) else "☐"))

CreditorsView._load_import_batches = _fm_creditors_load_import_batches_v622


# ---------- generische Suchleisten Debitoren/Kreditoren ----------
def _fm_tree_apply_filter(tree: ttk.Treeview, query: str) -> None:
    q = (query or "").strip().lower()
    if not q or q == PLACEHOLDER_TEXT.lower():
        return
    for iid in list(tree.get_children()):
        values = tree.item(iid, "values")
        if q not in " ".join(str(v) for v in values).lower():
            tree.delete(iid)


def _fm_install_tree_search(owner, tree_attr: str, reload_callback, var_attr: str, label: str = "Suchen") -> None:
    tree = getattr(owner, tree_attr, None)
    if tree is None or getattr(owner, f"_{var_attr}_installed", False):
        return
    setattr(owner, f"_{var_attr}_installed", True)
    var = tk.StringVar()
    setattr(owner, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=0, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=0, pady=(0, 4))
    ttk.Label(row, text=label, width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    entry.insert(0, PLACEHOLDER_TEXT)
    entry.bind("<FocusIn>", lambda e: (entry.delete(0, tk.END) if entry.get() == PLACEHOLDER_TEXT else None))
    entry.bind("<FocusOut>", lambda e: (entry.insert(0, PLACEHOLDER_TEXT) if not entry.get().strip() else None))
    entry.bind("<KeyRelease>", lambda _e: reload_callback())


_orig_invoice_reload_invoices_v622 = InvoiceModuleBase.reload_invoices
_orig_invoice_reload_open_items_v622 = InvoiceModuleBase.reload_open_items

def _fm_reload_invoices_v622(self):
    _orig_invoice_reload_invoices_v622(self)
    q = getattr(getattr(self, "invoice_search_var", None), "get", lambda: "")()
    if hasattr(self, "inv_tree"):
        _fm_tree_apply_filter(self.inv_tree, q)


def _fm_reload_open_items_v622(self):
    _orig_invoice_reload_open_items_v622(self)
    q = getattr(getattr(self, "open_items_search_var", None), "get", lambda: "")()
    if hasattr(self, "op_tree"):
        _fm_tree_apply_filter(self.op_tree, q)

InvoiceModuleBase.reload_invoices = _fm_reload_invoices_v622
InvoiceModuleBase.reload_open_items = _fm_reload_open_items_v622

_orig_debitors_build_ui_v622 = DebitorsView._build_ui

def _fm_debitors_build_ui_v622(self):
    _orig_debitors_build_ui_v622(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices(); self.reload_open_items()

DebitorsView._build_ui = _fm_debitors_build_ui_v622

_orig_creditors_build_ui_v622 = CreditorsView._build_ui

def _fm_creditors_build_ui_v622(self):
    _orig_creditors_build_ui_v622(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices(); self.reload_open_items(); self._load_import_batches()

CreditorsView._build_ui = _fm_creditors_build_ui_v622


# ---------- Suchleisten Rechnungsportal ----------
def _fm_install_portal_search(portal, tree_attr: str, reload_callback, var_attr: str) -> None:
    tree = getattr(portal, tree_attr, None)
    if tree is None or getattr(portal, f"_{var_attr}_installed", False):
        return
    setattr(portal, f"_{var_attr}_installed", True)
    var = tk.StringVar()
    setattr(portal, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=10, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=10, pady=(0, 4))
    ttk.Label(row, text="Suchen", width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    entry.insert(0, PLACEHOLDER_TEXT)
    entry.bind("<FocusIn>", lambda e: (entry.delete(0, tk.END) if entry.get() == PLACEHOLDER_TEXT else None))
    entry.bind("<FocusOut>", lambda e: (entry.insert(0, PLACEHOLDER_TEXT) if not entry.get().strip() else None))
    entry.bind("<KeyRelease>", lambda _e: reload_callback())

if 'InvoicePortalView' in globals():
    _orig_portal_init_v622 = InvoicePortalView.__init__
    def _fm_portal_init_v622(self, parent, app):
        _orig_portal_init_v622(self, parent, app)
        _fm_install_portal_search(self, "archive_tree", self.reload_archive, "archive_search_var")
        _fm_install_portal_search(self, "approval_tree", self.reload_approval, "approval_search_var")
        self.reload_all()
    InvoicePortalView.__init__ = _fm_portal_init_v622

    _orig_portal_reload_archive_v622 = InvoicePortalView.reload_archive
    _orig_portal_reload_approval_v622 = InvoicePortalView.reload_approval
    def _fm_portal_reload_archive_v622(self):
        _orig_portal_reload_archive_v622(self)
        _fm_tree_apply_filter(self.archive_tree, getattr(getattr(self, "archive_search_var", None), "get", lambda: "")())
    def _fm_portal_reload_approval_v622(self):
        _orig_portal_reload_approval_v622(self)
        _fm_tree_apply_filter(self.approval_tree, getattr(getattr(self, "approval_search_var", None), "get", lambda: "")())
    InvoicePortalView.reload_archive = _fm_portal_reload_archive_v622
    InvoicePortalView.reload_approval = _fm_portal_reload_approval_v622


# ---------- Fenstermodus Vollbild + stabilere 100%-Darstellungsannahme ----------
_orig_app_init_v622 = FinanceMateApp.__init__

def _fm_app_init_v622(self):
    _orig_app_init_v622(self)
    try:
        # 100%-Referenz in Tk: 96 dpi / 72 pt.
        self.tk.call("tk", "scaling", 96.0 / 72.0)
    except Exception:
        pass
    try:
        self.state("zoomed")  # Fenstermodus Vollbild, keine echte borderless Fullscreen-Ansicht.
    except Exception:
        try:
            self.attributes("-zoomed", True)
        except Exception:
            pass
    try:
        self.update_idletasks()
    except Exception:
        pass

FinanceMateApp.__init__ = _fm_app_init_v622


# ---------- Kreditoren-Selbsttest ohne GUI-Display ----------
def _fm_creditors_static_selftest_v622() -> None:
    import inspect as _inspect
    stack_src = _inspect.getsource(CreditorsView._build_stack_ui)
    build_src = _inspect.getsource(CreditorsView._build_ui)
    must_stack = ["_fm_heading_safe", "self.stack_tree", "Dokumentenstapel", "Vorschau"]
    must_build = ["_orig_creditors_build_ui_v622", "invoice_search_var", "open_items_search_var"]
    missing = [x for x in must_stack if x not in stack_src] + [x for x in must_build if x not in build_src]
    if missing:
        raise RuntimeError("Kreditoren-Selbsttest v0.6.22 fehlgeschlagen: " + ", ".join(missing))

_fm_creditors_static_selftest_v622()

# === FINANCE MATE PATCH V0_6_23_CREDITORS_SPLITLINES_FRAMELESS ===
APP_VERSION = "0.6.23-creditors-splitlines-frameless"


def _fm_pack_section_header(parent, title: str):
    """Optisch neutrale Bereichsüberschrift ohne grauen LabelFrame-Rahmen."""
    ttk.Label(parent, text=title, style="Section.TLabel").pack(anchor="w", padx=10, pady=(8, 4))


def _fm_make_plain_card(parent):
    # Kein LabelFrame, kein grauer Rahmen um den Block.
    frame = ttk.Frame(parent)
    return frame


def _fm_creditors_build_left_stack_area_v623(self, parent):
    """Nur linker unterer Block: Dokumentenstapel ohne Vorschau; Vorschau sitzt rechts unten separat."""
    _fm_pack_section_header(parent, "Dokumentenstapel")
    body = ttk.Frame(parent)
    body.pack(fill="both", expand=True, padx=10, pady=(0, 8))

    toolbar = ttk.Frame(body)
    toolbar.pack(fill="x", pady=(0, 6))
    create_standard_button(toolbar, "Dateien importieren", self.import_creditor_files, confirm=True).pack(side="left", padx=3)
    create_standard_button(toolbar, "Ordner importieren", self.import_creditor_folder).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte löschen", self._delete_marked_batches).pack(side="left", padx=3)
    create_standard_button(toolbar, "Markierte zusammenführen", self.merge_selected_batches).pack(side="left", padx=3)

    self.stack_search_var = tk.StringVar()
    search_row = ttk.Frame(body)
    search_row.pack(fill="x", pady=(0, 6))
    ttk.Label(search_row, text="Suchen", width=10).pack(side="left")
    stack_search = ttk.Entry(search_row, textvariable=self.stack_search_var)
    stack_search.pack(side="left", fill="x", expand=True)
    stack_search.insert(0, PLACEHOLDER_TEXT)
    stack_search.bind("<FocusIn>", lambda e: (stack_search.delete(0, tk.END) if stack_search.get() == PLACEHOLDER_TEXT else None))
    stack_search.bind("<FocusOut>", lambda e: (stack_search.insert(0, PLACEHOLDER_TEXT) if not stack_search.get().strip() else None))
    stack_search.bind("<KeyRelease>", lambda _e: self._load_import_batches())

    cols = ("hist_no", "title", "status", "mark")
    self.stack_tree = ttk.Treeview(body, columns=cols, show="headings", height=10)
    for c, t, w in [("hist_no", "Hist.-Nr.", 110), ("title", "Dokument", 260), ("status", "Status", 110), ("mark", "☑", 45)]:
        if c == "mark" and hasattr(self, "_toggle_all_marks"):
            self.stack_tree.heading(c, text=t, command=self._toggle_all_marks)
        else:
            self.stack_tree.heading(c, text=t)
        self.stack_tree.column(c, width=w, anchor="w" if c != "mark" else "center")
    self.stack_tree.pack(fill="both", expand=True)
    self.stack_tree.bind("<ButtonRelease-1>", self._on_stack_select, add="+")
    self.stack_tree.bind("<Double-1>", lambda _e: self._open_selected_stack_file())


def _fm_creditors_build_right_preview_area_v623(self, parent):
    """Nur rechter unterer Block: Vorschau separat, damit die unteren Blöcke unabhängig skalierbar sind."""
    _fm_pack_section_header(parent, "Vorschau")
    body = ttk.Frame(parent)
    body.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    self.preview_title_label = ttk.Label(body, text="Kein Dokument ausgewählt", style="Hint.TLabel")
    self.preview_title_label.pack(anchor="w", pady=(0, 6))
    self.preview_pane = _FMDocumentPreviewPane(body)
    self.preview_pane.pack(fill="both", expand=True)


def _fm_creditors_build_ui_v623(self):
    """Vier-Block-Kreditorenlayout mit zweigeteilter horizontaler Verschiebelinie.

    Struktur:
    - outer_horizontal: links/rechts Breite
    - left_vertical: oben Eingangsrechnung erfassen, unten Dokumentenstapel
    - right_vertical: oben Eingangsrechnungen/Offene Posten, unten Vorschau

    Dadurch sind die unteren Blöcke wieder getrennt vergrößerbar.
    """
    self.current_hist_no = None
    self.pending_attachments = []

    outer_horizontal = ttk.PanedWindow(self, orient="horizontal")
    outer_horizontal.pack(fill="both", expand=True, padx=0, pady=0)
    left_col = ttk.Frame(outer_horizontal)
    right_col = ttk.Frame(outer_horizontal)
    outer_horizontal.add(left_col, weight=2)
    outer_horizontal.add(right_col, weight=3)

    left_vertical = ttk.PanedWindow(left_col, orient="vertical")
    right_vertical = ttk.PanedWindow(right_col, orient="vertical")
    left_vertical.pack(fill="both", expand=True)
    right_vertical.pack(fill="both", expand=True)

    form_card = _fm_make_plain_card(left_vertical)
    stack_card = _fm_make_plain_card(left_vertical)
    list_card = _fm_make_plain_card(right_vertical)
    preview_card = _fm_make_plain_card(right_vertical)

    left_vertical.add(form_card, weight=3)
    left_vertical.add(stack_card, weight=2)
    right_vertical.add(list_card, weight=3)
    right_vertical.add(preview_card, weight=2)

    # Linker oberer Block: Eingangsrechnung erfassen, ohne grauen Rahmen.
    _fm_pack_section_header(form_card, "Eingangsrechnung erfassen")
    form_scroll = ScrollableFrame(form_card)
    form_scroll.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    f = form_scroll.content

    self.invoice_no = tk.StringVar(value=self._generate_invoice_no())
    self.invoice_kind = tk.StringVar(value="Rechnung")
    self.partner = tk.StringVar()
    self.invoice_date = tk.StringVar(value=today_str())
    self.due_date = tk.StringVar(value=today_str())
    self.payment_term = tk.StringVar()
    self.tax_code = tk.StringVar(value="V19")
    self.net_amount = tk.StringVar()
    self.tax_amount = tk.StringVar()
    self.gross_amount = tk.StringVar()

    self._labeled_combo(f, "Belegart", self.invoice_kind, INVOICE_KIND_VALUES)
    self._labeled_entry(f, "Rechnungs-Nr.", self.invoice_no)

    prow = ttk.Frame(f)
    prow.pack(fill="x", pady=3)
    ttk.Label(prow, text="Kreditor", width=16).pack(side="left")
    self.partner_cb = ttk.Combobox(prow, textvariable=self.partner, values=self.partners)
    self.partner_cb.pack(side="left", fill="x", expand=True)
    self.partner_cb.bind("<MouseWheel>", lambda e: "break")
    self.partner_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_vendor_info())
    create_standard_button(prow, "+", self._open_vendor_quick_popup).pack(side="left", padx=3)

    self.description_widget = _v619_labeled_text(f, "Rechnungsbeschreibung", height=2) if "_v619_labeled_text" in globals() else tk.Text(f, height=2)
    if not getattr(self.description_widget, "winfo_manager", lambda: "")():
        self.description_widget.pack(fill="x", pady=3)

    self._labeled_entry(f, "Rechnungsdatum", self.invoice_date)
    self._labeled_combo(f, "Zahlungsbed.", self.payment_term, self.payment_terms)
    self._labeled_entry(f, "Fälligkeit", self.due_date)
    self._labeled_combo(f, "Steuerkennz.", self.tax_code, self.tax_codes)
    self._labeled_entry(f, "Netto", self.net_amount)
    self._labeled_entry(f, "Steuer", self.tax_amount)
    self._labeled_entry(f, "Brutto", self.gross_amount)

    self.payment_term.trace_add("write", lambda *_: self.apply_payment_term())
    self.net_amount.trace_add("write", lambda *_: self._calculate_amounts())
    self.tax_code.trace_add("write", lambda *_: self._calculate_amounts())

    btns = ttk.Frame(f)
    btns.pack(fill="x", pady=8)
    create_standard_button(btns, "Dokument anhängen", self.add_pending_attachments).pack(side="left", padx=3)
    create_standard_button(btns, "Speichern", self.save_invoice, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Neu", self.clear_form).pack(side="left", padx=3)

    info = ttk.Frame(f)
    info.pack(fill="both", expand=False, pady=(8, 4))
    ttk.Label(info, text="Stammdaten des Kreditors", style="Section.TLabel").pack(anchor="w", pady=(0, 4))
    self.vendor_info_body = tk.Text(info, height=8, wrap="word", bg=WHITE, fg=TEXT, relief="flat")
    self.vendor_info_body.pack(fill="both", expand=True, padx=0, pady=0)
    self.vendor_info_body.configure(state="disabled")

    # Rechter oberer Block: Listen ohne grauen Rahmen.
    _fm_pack_section_header(list_card, "Eingangsrechnungen")
    nb = ttk.Notebook(list_card)
    nb.pack(fill="both", expand=True, padx=10, pady=(0, 8))
    tab_inv = ttk.Frame(nb)
    tab_op = ttk.Frame(nb)
    nb.add(tab_inv, text="Eingangsrechnungen")
    nb.add(tab_op, text="Offene Posten")

    inv_cols = ("invoice_no", "kind", "partner", "description", "date", "due", "gross", "open", "status", "att")
    self.inv_tree = ttk.Treeview(tab_inv, columns=inv_cols, show="headings", height=12)
    self.invoice_tabs_trees = [self.inv_tree]
    for c, t, w in [("invoice_no", "Rechnung", 115), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("date", "Datum", 90), ("due", "Fällig", 90), ("gross", "Brutto", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.inv_tree.heading(c, text=t)
        self.inv_tree.column(c, width=w, anchor="w")
    self.inv_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.inv_tree)
    self.setup_sorting(self.inv_tree)
    self.inv_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.inv_tree, lambda row: self.inv_tree.set(row, "invoice_no"), self.reload_invoices), add="+")

    op_cols = ("reference_no", "kind", "partner", "description", "due", "original", "open", "status", "att")
    self.op_tree = ttk.Treeview(tab_op, columns=op_cols, show="headings", height=12)
    self.open_item_tabs_trees = [self.op_tree]
    for c, t, w in [("reference_no", "Referenz", 120), ("kind", "Art", 90), ("partner", "Kreditor", 170), ("description", "Beschreibung", 190), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120), ("att", "Anhang", 145)]:
        self.op_tree.heading(c, text=t)
        self.op_tree.column(c, width=w, anchor="w")
    self.op_tree.pack(fill="both", expand=True)
    configure_tree_tags(self.op_tree)
    self.setup_sorting(self.op_tree)
    self.op_tree.bind("<ButtonRelease-1>", lambda e: _v619_add_attachment_column_click(self, e, self.op_tree, lambda row: self.op_tree.set(row, "reference_no"), self.reload_open_items), add="+")

    # Untere Blöcke getrennt, jeweils eigener vertikaler Sash links/rechts.
    _fm_creditors_build_left_stack_area_v623(self, stack_card)
    _fm_creditors_build_right_preview_area_v623(self, preview_card)

    self.partner.trace_add("write", lambda *_: self._update_vendor_info())
    self._update_vendor_info()

    # Suchleisten aus v0.6.22 weiter nutzen, falls vorhanden.
    if "_fm_install_tree_search" in globals():
        _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
        _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")

    self.reload_invoices()
    self.reload_open_items()
    self._load_import_batches()
    self._clear_stack_preview()

CreditorsView._build_ui = _fm_creditors_build_ui_v623


def _fm_render_kreditoren_v623(self, parent):
    try:
        view = CreditorsView(parent, self)
        view.pack(fill="both", expand=True)
        required = ["inv_tree", "op_tree", "stack_tree", "preview_pane", "vendor_info_body"]
        missing = [name for name in required if not hasattr(view, name)]
        if missing:
            raise RuntimeError("Kreditoren-Aufbau unvollständig, fehlende Widgets: " + ", ".join(missing))
    except Exception as exc:
        for w in parent.winfo_children():
            w.destroy()
        err = tk.Text(parent, height=12, bg="#fff4c2", fg=TEXT, wrap="word")
        err.pack(fill="both", expand=True, padx=12, pady=12)
        err.insert("1.0", "Kreditoren-Modul konnte nicht vollständig aufgebaut werden.\n\n" + str(exc))
        err.configure(state="disabled")
        raise

FinanceMateApp._render_kreditoren = _fm_render_kreditoren_v623



def _fm_creditors_static_selftest_v623() -> None:
    import inspect as _inspect
    src_build = _inspect.getsource(CreditorsView._build_ui)
    src_stack = _inspect.getsource(_fm_creditors_build_left_stack_area_v623)
    src_preview = _inspect.getsource(_fm_creditors_build_right_preview_area_v623)
    src_combined = src_build + "\n" + src_stack + "\n" + src_preview
    forbidden = ["ttk.LabelFrame", "form_card = ttk.LabelFrame", "list_card = ttk.LabelFrame", "style=\"Card.TFrame\""]
    required = [
        "outer_horizontal", "left_vertical", "right_vertical",
        "_fm_creditors_build_left_stack_area_v623", "_fm_creditors_build_right_preview_area_v623",
        "self.inv_tree", "self.op_tree", "self.stack_tree", "self.preview_pane"
    ]
    missing = [x for x in required if x not in src_combined]
    found_forbidden = [x for x in forbidden if x in src_combined]
    if missing or found_forbidden:
        raise RuntimeError("Kreditoren-Selbsttest v0.6.23 fehlgeschlagen. Fehlend: " + ", ".join(missing) + "; verboten: " + ", ".join(found_forbidden))

_fm_creditors_static_selftest_v623()



# === FINANCE MATE PATCH V0_6_24_PAYMENTS_BLOCK7_LABELWRAP ===
APP_VERSION = "0.6.24-payments-block7-labelwrap"

# ---------- Debitoren/Kreditoren: Label-Umbruch Rechnungsbeschreibung ----------
_orig_v619_labeled_text_v624 = _v619_labeled_text if "_v619_labeled_text" in globals() else None

def _v619_labeled_text(parent, label: str, height: int = 2) -> tk.Text:
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=3)
    display_label = "Rechnungs-\nbeschreibung" if label == "Rechnungsbeschreibung" else label
    ttk.Label(row, text=display_label, width=16, justify="left").pack(side="left", anchor="n")
    txt = tk.Text(row, height=height, wrap="word", bg=WHITE, fg=TEXT, relief="sunken")
    txt.pack(side="left", fill="x", expand=True)
    return txt

# ---------- Block 7 Zahlungen: Persistenz ----------
def _fm_init_payments_schema_v624() -> None:
    ensure_directories()
    with get_connection() as con:
        con.executescript("""
        CREATE TABLE IF NOT EXISTS payments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entity_type TEXT NOT NULL,
            reference_no TEXT NOT NULL,
            partner_no TEXT NOT NULL,
            partner_name TEXT NOT NULL,
            payment_date TEXT NOT NULL,
            amount REAL NOT NULL,
            payment_method TEXT DEFAULT '',
            bank_account TEXT DEFAULT '',
            booking_text TEXT DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );
        """)
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()

_orig_init_sqlite_v624 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v624()
    _fm_init_payments_schema_v624()


def _fm_payment_partner_label(entity_type: str) -> str:
    return "Kreditor" if entity_type == "vendor_invoice" else "Debitor"


def _fm_invoice_table_for_entity_v624(entity_type: str) -> str:
    return "vendor_invoices" if entity_type == "vendor_invoice" else "customer_invoices"


def _fm_payment_status_from_amount_v624(amount) -> str:
    try:
        return STATUS_PAID if parse_amount(amount) <= 0 else STATUS_OPEN
    except Exception:
        return STATUS_OPEN


class PaymentsView(ttk.Frame, SortableTreeMixin):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.selected_open_item_id = None
        self.selected_entity_type = ""
        self.selected_reference_no = ""
        self.search_var = tk.StringVar()
        self.history_search_var = tk.StringVar()
        self.payment_date = tk.StringVar(value=today_str())
        self.amount = tk.StringVar()
        self.payment_method = tk.StringVar(value="Überweisung")
        self.bank_account = tk.StringVar()
        self.booking_text = tk.StringVar()
        self.full_clear = tk.IntVar(value=1)
        self._build_ui()
        self.reload_open_items()
        self.reload_history()

    def _build_ui(self):
        shell = build_two_block_shell(self, "Zahlungserfassung / OP-Ausgleich", "Zahlungshistorie")

        left = ttk.Frame(shell.left)
        left.pack(fill="both", expand=True, padx=10, pady=8)
        form = ttk.Frame(left)
        form.pack(fill="x", pady=(0, 8))

        self._row_entry(form, "Zahlungsdatum", self.payment_date)
        self._row_entry(form, "Zahlbetrag", self.amount)
        self._row_combo(form, "Zahlungsart", self.payment_method, ["Überweisung", "Lastschrift", "Bar", "Karte", "Sonstige"])
        self._row_entry(form, "Bank/Konto", self.bank_account)
        self._row_entry(form, "Buchungstext", self.booking_text)
        ttk.Checkbutton(form, text="OP vollständig ausgleichen", variable=self.full_clear).pack(anchor="w", pady=3)

        btns = ttk.Frame(form)
        btns.pack(fill="x", pady=6)
        create_standard_button(btns, "Zahlung buchen", self.book_payment, confirm=True).pack(side="left", padx=3)
        create_standard_button(btns, "Formular leeren", self.clear_form).pack(side="left", padx=3)
        create_standard_button(btns, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)

        search_row = ttk.Frame(left)
        search_row.pack(fill="x", pady=(0, 4))
        ttk.Label(search_row, text="Suchen", width=10).pack(side="left")
        search_entry = ttk.Entry(search_row, textvariable=self.search_var)
        search_entry.pack(side="left", fill="x", expand=True)
        search_entry.insert(0, PLACEHOLDER_TEXT)
        search_entry.bind("<FocusIn>", lambda e: (search_entry.delete(0, tk.END) if search_entry.get() == PLACEHOLDER_TEXT else None))
        search_entry.bind("<FocusOut>", lambda e: (search_entry.insert(0, PLACEHOLDER_TEXT) if not search_entry.get().strip() else None))
        search_entry.bind("<KeyRelease>", lambda _e: self.reload_open_items())

        ttk.Label(left, text="Offene Posten", style="Section.TLabel").pack(anchor="w", pady=(6, 4))
        cols = ("id", "type", "reference", "partner", "due", "original", "open", "status")
        self.op_tree = ttk.Treeview(left, columns=cols, show="headings", height=13)
        headings = [("id", "ID", 55), ("type", "Art", 90), ("reference", "Referenz", 120), ("partner", "Partner", 210), ("due", "Fällig", 90), ("original", "Original", 90), ("open", "Offen", 90), ("status", "Status", 120)]
        for c, t, w in headings:
            self.op_tree.heading(c, text=t)
            self.op_tree.column(c, width=w, anchor="w")
        self.op_tree.pack(fill="both", expand=True)
        configure_tree_tags(self.op_tree)
        self.setup_sorting(self.op_tree)
        self.op_tree.bind("<<TreeviewSelect>>", self.on_open_item_select)

        right = ttk.Frame(shell.right)
        right.pack(fill="both", expand=True, padx=10, pady=8)
        hist_search_row = ttk.Frame(right)
        hist_search_row.pack(fill="x", pady=(0, 4))
        ttk.Label(hist_search_row, text="Suchen", width=10).pack(side="left")
        hist_search = ttk.Entry(hist_search_row, textvariable=self.history_search_var)
        hist_search.pack(side="left", fill="x", expand=True)
        hist_search.insert(0, PLACEHOLDER_TEXT)
        hist_search.bind("<FocusIn>", lambda e: (hist_search.delete(0, tk.END) if hist_search.get() == PLACEHOLDER_TEXT else None))
        hist_search.bind("<FocusOut>", lambda e: (hist_search.insert(0, PLACEHOLDER_TEXT) if not hist_search.get().strip() else None))
        hist_search.bind("<KeyRelease>", lambda _e: self.reload_history())

        ttk.Label(right, text="Gebuchte Zahlungen", style="Section.TLabel").pack(anchor="w", pady=(6, 4))
        hcols = ("id", "date", "type", "reference", "partner", "amount", "method", "text")
        self.history_tree = ttk.Treeview(right, columns=hcols, show="headings", height=16)
        for c, t, w in [("id", "ID", 55), ("date", "Datum", 90), ("type", "Art", 90), ("reference", "Referenz", 120), ("partner", "Partner", 190), ("amount", "Betrag", 90), ("method", "Zahlungsart", 110), ("text", "Text", 220)]:
            self.history_tree.heading(c, text=t)
            self.history_tree.column(c, width=w, anchor="w")
        self.history_tree.pack(fill="both", expand=True)
        self.setup_sorting(self.history_tree)

    def _row_entry(self, parent, label, var):
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)

    def _row_combo(self, parent, label, var, values):
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=3)
        ttk.Label(row, text=label, width=16).pack(side="left")
        cb = ttk.Combobox(row, textvariable=var, values=values, state="readonly")
        cb.pack(side="left", fill="x", expand=True)
        cb.bind("<MouseWheel>", lambda e: "break")

    def _query_text(self, var):
        q = (var.get() or "").strip().lower()
        return "" if q == PLACEHOLDER_TEXT.lower() else q

    def reload_open_items(self):
        for i in self.op_tree.get_children():
            self.op_tree.delete(i)
        q = self._query_text(self.search_var)
        with get_connection() as con:
            rows = list(con.execute("SELECT * FROM open_items WHERE open_amount > 0 ORDER BY due_date, id"))
        for r in rows:
            art = _fm_payment_partner_label(r["entity_type"])
            values = (r["id"], art, r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"])
            blob = " ".join(str(v) for v in values).lower()
            if q and q not in blob:
                continue
            tag = urgency_bucket(r["due_date"], r["open_amount"])
            self.op_tree.insert("", "end", iid=str(r["id"]), values=values, tags=(tag,))

    def reload_history(self):
        for i in self.history_tree.get_children():
            self.history_tree.delete(i)
        q = self._query_text(self.history_search_var)
        with get_connection() as con:
            rows = list(con.execute("SELECT * FROM payments ORDER BY id DESC"))
        for r in rows:
            art = _fm_payment_partner_label(r["entity_type"])
            values = (r["id"], r["payment_date"], art, r["reference_no"], r["partner_name"], format_amount(r["amount"]), r["payment_method"], r["booking_text"])
            blob = " ".join(str(v) for v in values).lower()
            if q and q not in blob:
                continue
            self.history_tree.insert("", "end", values=values)

    def on_open_item_select(self, _event=None):
        iid = self.op_tree.focus()
        if not iid:
            return
        self.selected_open_item_id = int(iid)
        with get_connection() as con:
            row = con.execute("SELECT * FROM open_items WHERE id=?", (self.selected_open_item_id,)).fetchone()
        if not row:
            return
        self.selected_entity_type = row["entity_type"]
        self.selected_reference_no = row["reference_no"]
        self.amount.set(format_amount(row["open_amount"]))
        if not self.booking_text.get().strip():
            self.booking_text.set(f"Zahlung {row['reference_no']} {row['partner_name']}")

    def clear_form(self):
        self.selected_open_item_id = None
        self.selected_entity_type = ""
        self.selected_reference_no = ""
        self.payment_date.set(today_str())
        self.amount.set("")
        self.payment_method.set("Überweisung")
        self.bank_account.set("")
        self.booking_text.set("")
        self.full_clear.set(1)
        try:
            self.op_tree.selection_remove(self.op_tree.selection())
        except Exception:
            pass

    def refresh_all(self):
        self.reload_open_items()
        self.reload_history()
        try:
            self.app.set_status("Zahlungen aktualisiert.")
        except Exception:
            pass

    def book_payment(self):
        if not self.selected_open_item_id:
            messagebox.showwarning("Zahlung", "Bitte zuerst einen offenen Posten auswählen.", parent=self)
            return
        if not validate_date(self.payment_date.get()):
            messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.", parent=self)
            return
        try:
            amount = parse_amount(self.amount.get())
        except Exception as exc:
            messagebox.showwarning("Betrag", str(exc), parent=self)
            return
        if amount <= 0:
            messagebox.showwarning("Betrag", "Der Zahlbetrag muss größer als 0 sein.", parent=self)
            return
        with get_connection() as con:
            op = con.execute("SELECT * FROM open_items WHERE id=?", (self.selected_open_item_id,)).fetchone()
            if not op:
                messagebox.showwarning("Zahlung", "Der offene Posten wurde nicht gefunden.", parent=self)
                return
            open_amount = parse_amount(op["open_amount"])
            pay_amount = open_amount if int(self.full_clear.get() or 0) else amount
            if pay_amount > open_amount:
                messagebox.showwarning("Betrag", "Der Zahlbetrag darf den offenen Betrag nicht überschreiten.", parent=self)
                return
            new_open = (open_amount - pay_amount).quantize(Decimal("0.01"))
            new_status = _fm_payment_status_from_amount_v624(new_open)
            con.execute(
                "INSERT INTO payments(entity_type,reference_no,partner_no,partner_name,payment_date,amount,payment_method,bank_account,booking_text,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                (op["entity_type"], op["reference_no"], op["partner_no"], op["partner_name"], self.payment_date.get(), float(pay_amount), self.payment_method.get(), self.bank_account.get(), self.booking_text.get(), now_str(), now_str())
            )
            con.execute("UPDATE open_items SET open_amount=?, status=?, updated_at=? WHERE id=?", (float(new_open), new_status, now_str(), self.selected_open_item_id))
            table = _fm_invoice_table_for_entity_v624(op["entity_type"])
            con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), new_status, now_str(), op["reference_no"]))
            con.commit()
        try:
            self.app.set_status(f"Zahlung zu {self.selected_reference_no} gebucht.")
        except Exception:
            pass
        self.clear_form()
        self.refresh_all()


def _fm_render_zahlungen_v624(self, parent):
    PaymentsView(parent, self).pack(fill="both", expand=True)

FinanceMateApp._render_zahlungen = _fm_render_zahlungen_v624


def _fm_block7_static_selftest_v624() -> None:
    # Dateibasierte Prüfung, robust auch beim dynamischen Import über importlib.
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "class PaymentsView",
        "def book_payment",
        "CREATE TABLE IF NOT EXISTS payments",
        "SELECT * FROM open_items WHERE open_amount > 0",
        "UPDATE open_items SET open_amount",
        "UPDATE {table} SET open_amount",
        "PaymentsView(parent, self).pack",
    ]
    missing = [x for x in required if x not in src]
    if ("Rechnungs-" not in src) or ("beschreibung" not in src):
        missing.append("Rechnungs-beschreibung Labelumbruch")
    if missing:
        raise RuntimeError("Block-7-/Label-Selbsttest v0.6.24 fehlgeschlagen: " + ", ".join(missing))

_fm_block7_static_selftest_v624()



# === FINANCE MATE PATCH V0_6_25_BLOCK7A_STABILISIERUNG ===
APP_VERSION = "0.6.25-block7a-stabilisierung"

# ---------- 7A.1 Rechnungsbeschreibung: blockgebundenes 2-Zeilen-Textfeld ----------
_orig_v619_labeled_text_v625 = _v619_labeled_text if "_v619_labeled_text" in globals() else None

def _v619_labeled_text(parent, label: str, height: int = 2) -> tk.Text:
    """Gemeinsames Rechnungsbeschreibung-Feld für Debitoren/Kreditoren.

    Das Textfeld bleibt strikt im Formularblock, nutzt Wortumbruch und ist optisch auf zwei Zeilen
    ausgelegt. Damit läuft weder das Label noch das Textfeld in Nachbarblöcke hinein.
    """
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=3)
    row.columnconfigure(1, weight=1)
    display_label = "Rechnungs-\nbeschreibung" if label == "Rechnungsbeschreibung" else label
    ttk.Label(row, text=display_label, width=16, justify="left").grid(row=0, column=0, sticky="nw", padx=(0, 4))
    txt = tk.Text(row, height=height, width=1, wrap="word", bg=WHITE, fg=TEXT, relief="sunken", undo=True)
    txt.grid(row=0, column=1, sticky="ew")
    return txt

# ---------- 7A.2 Suchleisten: einheitlich, pro Tabelle statt nur einmal je Modul ----------
def _fm_entry_placeholder_bind(entry: ttk.Entry, placeholder: str = PLACEHOLDER_TEXT, callback=None) -> None:
    entry.insert(0, placeholder)
    def _focus_in(_e=None):
        if entry.get() == placeholder:
            entry.delete(0, tk.END)
    def _focus_out(_e=None):
        if not entry.get().strip():
            entry.insert(0, placeholder)
        if callable(callback):
            callback()
    entry.bind("<FocusIn>", _focus_in)
    entry.bind("<FocusOut>", _focus_out)
    entry.bind("<KeyRelease>", lambda _e: callback() if callable(callback) else None)

def _fm_install_tree_search(owner, tree_attr: str, reload_callback, var_attr: str, label: str = "Suchen") -> None:
    tree = getattr(owner, tree_attr, None)
    flag = f"_fm_search_installed_{tree_attr}"
    if tree is None or getattr(owner, flag, False):
        return
    setattr(owner, flag, True)
    var = getattr(owner, var_attr, None)
    if var is None or not hasattr(var, "get"):
        var = tk.StringVar()
        setattr(owner, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=0, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=0, pady=(0, 4))
    ttk.Label(row, text=label, width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, reload_callback)

def _fm_install_portal_search(portal, tree_attr: str, reload_callback, var_attr: str) -> None:
    tree = getattr(portal, tree_attr, None)
    flag = f"_fm_portal_search_installed_{tree_attr}"
    if tree is None or getattr(portal, flag, False):
        return
    setattr(portal, flag, True)
    var = getattr(portal, var_attr, None)
    if var is None or not hasattr(var, "get"):
        var = tk.StringVar()
        setattr(portal, var_attr, var)
    parent = tree.master
    row = ttk.Frame(parent)
    try:
        row.pack(fill="x", padx=10, pady=(0, 4), before=tree)
    except Exception:
        row.pack(fill="x", padx=10, pady=(0, 4))
    ttk.Label(row, text="Suchen", width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, reload_callback)

# Debitoren: vorhandene Features beibehalten, Suchleisten verlässlich für beide Tabellen.
_orig_debitors_build_ui_v625 = DebitorsView._build_ui

def _fm_debitors_build_ui_v625(self):
    _orig_debitors_build_ui_v625(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices()
    self.reload_open_items()

DebitorsView._build_ui = _fm_debitors_build_ui_v625

# Kreditoren: v0.6.23-Vierblocklayout bleibt erhalten, Suchleisten werden wieder ergänzt.
_orig_creditors_build_ui_v625 = CreditorsView._build_ui

def _fm_creditors_build_ui_v625(self):
    _orig_creditors_build_ui_v625(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    _fm_install_tree_search(self, "stack_tree", self._load_import_batches, "stack_search_var")
    self.reload_invoices()
    self.reload_open_items()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()

CreditorsView._build_ui = _fm_creditors_build_ui_v625

# Rechnungsportal: Archiv- und Freigabetabelle bekommen je eine eigene Suchleiste.
if "InvoicePortalView" in globals():
    _orig_portal_init_v625 = InvoicePortalView.__init__
    def _fm_portal_init_v625(self, parent, app):
        _orig_portal_init_v625(self, parent, app)
        _fm_install_portal_search(self, "archive_tree", self.reload_archive, "archive_search_var")
        _fm_install_portal_search(self, "approval_tree", self.reload_approval, "approval_search_var")
        try:
            self.reload_all()
        except Exception:
            pass
    InvoicePortalView.__init__ = _fm_portal_init_v625

# ---------- 7A.3 Dokumentenstapel: Mehrfachimport stabilisieren ----------
def _fm_creditors_import_archive_paths_v625(self, paths):
    valid = []
    seen = set()
    for raw in paths or []:
        fp = Path(raw)
        key = str(fp.resolve()) if fp.exists() else str(fp)
        if key in seen:
            continue
        seen.add(key)
        if fp.exists() and fp.is_file() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo("Dokumentenstapel", "Keine unterstützten Dateien ausgewählt.", parent=self)
        return
    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    imported = 0
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute(
                "INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
                (hist_no, clean_document_title(fp.name), "", "", "Importiert", 0, now_str(), now_str()),
            )
            con.execute(
                "INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)",
                (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()),
            )
            imported += 1
        con.commit()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    try:
        self.app.set_status(f"{imported} Dokument(e) in den Dokumentenstapel importiert.")
    except Exception:
        pass

def _fm_creditors_import_files_v625(self):
    paths = filedialog.askopenfilenames(title="Dateien für Dokumentenstapel importieren")
    if paths:
        self._import_archive_paths(list(paths))

def _fm_creditors_import_folder_v625(self):
    folder = filedialog.askdirectory(title="Ordner für Dokumentenstapel importieren")
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v625
CreditorsView.import_creditor_files = _fm_creditors_import_files_v625
CreditorsView.import_creditor_folder = _fm_creditors_import_folder_v625

# ---------- 7A.4 Vorschau-Zoom: Standard höher, Maximalzoom deutlich höher ----------
def _fm_preview_set_image_v625(self, pil_image):
    self.base_image = pil_image.convert("RGBA") if pil_image is not None else None
    self.zoom = 1.0
    self.offset_x = 0
    self.offset_y = 0
    if self.base_image is None:
        self.clear()
    else:
        self._render(reset=True)

def _fm_preview_render_v625(self, reset=False):
    if self.base_image is None:
        return
    try:
        from PIL import ImageTk
        self.fit_zoom = self._fit_zoom()
        if reset:
            # Etwas größer als reine Einpassen-Ansicht, aber weiterhin vollständig navigierbar.
            self.zoom = max(self.fit_zoom * 1.18, self.fit_zoom)
            self.offset_x = 0
            self.offset_y = 0
        elif self.zoom < self.fit_zoom:
            self.zoom = self.fit_zoom
        self._clamp_offsets()
        iw, ih = self.base_image.size
        size = (max(1, int(iw * self.zoom)), max(1, int(ih * self.zoom)))
        img = self.base_image.resize(size)
        self.tk_image = ImageTk.PhotoImage(img)
        self._draw()
    except Exception as exc:
        self.clear()
        self.canvas.create_text(20, 48, anchor="nw", text=f"Vorschaufehler: {exc}", fill=TEXT2, font=("Segoe UI", 10))

def _fm_preview_wheel_v625(self, delta):
    if self.base_image is None:
        return
    factor = 1.15 if delta > 0 else 0.87
    old_zoom = self.zoom
    max_zoom = max(self.fit_zoom * 20.0, 20.0)
    self.zoom = max(self.fit_zoom, min(self.zoom * factor, max_zoom))
    if abs(self.zoom - old_zoom) > 0.0001:
        self._clamp_offsets()
        self._render(reset=False)

_FMDocumentPreviewPane.set_image = _fm_preview_set_image_v625
_FMDocumentPreviewPane._render = _fm_preview_render_v625
_FMDocumentPreviewPane._wheel = _fm_preview_wheel_v625

# ---------- 7A.5 Startverhalten: Fenstermodus Vollbild + frühere Skalierungsannahme ----------
def _fm_apply_start_scaling_v625(root: tk.Tk) -> None:
    try:
        root.tk.call("tk", "scaling", 96.0 / 72.0)
    except Exception:
        pass
    try:
        sw = max(1100, root.winfo_screenwidth())
        sh = max(680, root.winfo_screenheight())
        root.geometry(f"{min(sw, 1920)}x{min(sh, 1080)}")
        root.minsize(1100, 680)
    except Exception:
        pass

def _fm_maximize_window_v625(root: tk.Tk) -> None:
    try:
        root.state("zoomed")
    except Exception:
        try:
            root.attributes("-zoomed", True)
        except Exception:
            pass

# Frischer App-Start: Skalierung vor Layoutaufbau, danach maximierter Fenstermodus.
def _fm_app_init_v625(self):
    tk.Tk.__init__(self)
    _fm_apply_start_scaling_v625(self)
    init_sqlite()
    self.title(f"{APP_NAME} {APP_VERSION}")
    self.geometry("1420x860")
    self.minsize(1100, 680)
    self.configure(bg=BG)
    self.sidebar_collapsed = load_ui_preference("sidebar_collapsed", "0") == "1"
    self.active_module = "Dashboard"
    self.nav_buttons: Dict[str, ttk.Button] = {}
    self._configure_ttk()
    self._build_layout()
    _fm_maximize_window_v625(self)
    self.show_module("Dashboard")
    try:
        self.after(250, lambda: _fm_maximize_window_v625(self))
    except Exception:
        pass

FinanceMateApp.__init__ = _fm_app_init_v625

# ---------- 7A-Selbsttest ----------
def _fm_block7a_static_selftest_v625() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required_src = [
        "APP_VERSION = \"0.6.25-block7a-stabilisierung\"",
        "def _fm_creditors_import_archive_paths_v625",
        "self.zoom = max(self.fit_zoom * 1.18",
        "self.fit_zoom * 20.0",
        "def _fm_debitors_build_ui_v625",
        "def _fm_creditors_build_ui_v625",
        "def _fm_app_init_v625",
        "Rechnungs-\\nbeschreibung",
    ]
    missing = [x for x in required_src if x not in src]
    try:
        csrc = _inspect.getsource(CreditorsView._build_ui)
    except Exception:
        csrc = ""
    for token in ["_orig_creditors_build_ui_v625", "stack_search_var", "invoice_search_var", "open_items_search_var"]:
        if token not in csrc:
            missing.append(token)
    if missing:
        raise RuntimeError("Block-7A-Selbsttest v0.6.25 fehlgeschlagen: " + ", ".join(missing))

_fm_block7a_static_selftest_v625()



# === FINANCE MATE PATCH V0_6_26_BLOCK7B_PAYMENT_TABS_IMPORT_PREVIEW ===
APP_VERSION = "0.6.26-block7b-payment-tabs-import-preview"

# ---------- 7A-Nachkorrektur: nur eine Suchleiste im Dokumentenstapel ----------
# Der Dokumentenstapel besitzt bereits im v0.6.23-Layout eine eigene Suchleiste. Die zusätzliche
# generische Suchleiste aus v0.6.25 wird hier bewusst nicht mehr installiert.
_orig_creditors_build_ui_v626 = _orig_creditors_build_ui_v625 if "_orig_creditors_build_ui_v625" in globals() else CreditorsView._build_ui

def _fm_creditors_build_ui_v626(self):
    _orig_creditors_build_ui_v626(self)
    _fm_install_tree_search(self, "inv_tree", self.reload_invoices, "invoice_search_var")
    _fm_install_tree_search(self, "op_tree", self.reload_open_items, "open_items_search_var")
    self.reload_invoices()
    self.reload_open_items()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()

CreditorsView._build_ui = _fm_creditors_build_ui_v626

# ---------- 7A-Nachkorrektur: Mehrfachauswahl für Dateien importieren robust aus Tk lesen ----------
def _fm_normalize_dialog_paths_v626(widget, raw_paths) -> list[str]:
    """Tk kann je nach Plattform Tuple, Tcl-Liste oder String liefern; hier wird alles normalisiert."""
    if not raw_paths:
        return []
    if isinstance(raw_paths, (list, tuple)):
        return [str(p) for p in raw_paths if str(p).strip()]
    try:
        return [str(p) for p in widget.tk.splitlist(raw_paths) if str(p).strip()]
    except Exception:
        return [str(raw_paths)] if str(raw_paths).strip() else []

def _fm_creditors_import_archive_paths_v626(self, paths):
    valid: list[Path] = []
    seen = set()
    for raw in _fm_normalize_dialog_paths_v626(self, paths):
        fp = Path(raw)
        try:
            key = str(fp.resolve()) if fp.exists() else str(fp)
        except Exception:
            key = str(fp)
        if key in seen:
            continue
        seen.add(key)
        if fp.exists() and fp.is_file() and fp.suffix.lower() in SUPPORTED_ARCHIVE_EXTENSIONS:
            valid.append(fp)
    if not valid:
        messagebox.showinfo("Dokumentenstapel", "Keine unterstützten Dateien ausgewählt.", parent=self)
        return
    IMPORTS_DIR.mkdir(parents=True, exist_ok=True)
    imported = 0
    with get_connection() as con:
        for fp in valid:
            hist_no = next_hist_no()
            hist_dir = IMPORTS_DIR / hist_no
            hist_dir.mkdir(parents=True, exist_ok=True)
            dest = hist_dir / fp.name
            if dest.exists():
                dest = hist_dir / f"{fp.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{fp.suffix}"
            shutil.copy2(fp, dest)
            con.execute(
                "INSERT INTO invoice_import_batches(hist_no,title,vendor_no,vendor_name,status,marked,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
                (hist_no, clean_document_title(fp.name), "", "", "Importiert", 0, now_str(), now_str()),
            )
            con.execute(
                "INSERT INTO invoice_import_files(hist_no,file_name,file_path,mime_type,created_at) VALUES(?,?,?,?,?)",
                (hist_no, dest.name, str(dest), mimetypes.guess_type(str(dest))[0] or "", now_str()),
            )
            imported += 1
        con.commit()
    if hasattr(self, "_load_import_batches"):
        self._load_import_batches()
    try:
        self.app.set_status(f"{imported} Dokument(e) in den Dokumentenstapel importiert.")
    except Exception:
        pass

def _fm_creditors_import_files_v626(self):
    # Direkter Tk-Aufruf mit -multiple 1 ist robuster als die Wrapperfunktion auf einzelnen Plattformen.
    try:
        raw = self.tk.call("tk_getOpenFile", "-multiple", "1", "-title", "Dateien für Dokumentenstapel importieren")
        paths = _fm_normalize_dialog_paths_v626(self, raw)
    except Exception:
        paths = list(filedialog.askopenfilenames(title="Dateien für Dokumentenstapel importieren", parent=self))
    if paths:
        self._import_archive_paths(paths)

def _fm_creditors_import_folder_v626(self):
    folder = filedialog.askdirectory(title="Ordner für Dokumentenstapel importieren", parent=self)
    if folder:
        self._import_archive_paths(list(iter_supported_files_from_folder(folder)))

CreditorsView._import_archive_paths = _fm_creditors_import_archive_paths_v626
CreditorsView.import_creditor_files = _fm_creditors_import_files_v626
CreditorsView.import_creditor_folder = _fm_creditors_import_folder_v626

# ---------- 7A-Nachkorrektur: Vorschau-Initialzoom = obere 50% des ersten Vorschaubilds ----------
_orig_fm_combine_pages_v626 = _fm_combine_pages
_orig_fm_text_pages_to_image_v626 = _fm_text_pages_to_image
_orig_fm_generate_document_preview_v626 = _fm_generate_document_preview

def _fm_text_pages_to_image(title: str, lines: list[str], page_lines: int = 42):
    img = _orig_fm_text_pages_to_image_v626(title, lines, page_lines)
    try:
        img._fm_first_page_height = 1500
    except Exception:
        pass
    return img

def _fm_combine_pages(images):
    out = _orig_fm_combine_pages_v626(images)
    try:
        out._fm_first_page_height = images[0].height if images else out.height
    except Exception:
        pass
    return out

def _fm_generate_document_preview(path: str):
    img = _orig_fm_generate_document_preview_v626(path)
    try:
        if not hasattr(img, "_fm_first_page_height"):
            img._fm_first_page_height = img.height
    except Exception:
        pass
    return img

def _fm_preview_render_v626(self, reset=False):
    if self.base_image is None:
        return
    try:
        from PIL import ImageTk
        self.fit_zoom = self._fit_zoom()
        cw = max(50, self.canvas.winfo_width())
        ch = max(50, self.canvas.winfo_height())
        first_h = float(getattr(self.base_image, "_fm_first_page_height", self.base_image.size[1]) or self.base_image.size[1])
        if reset:
            # Ziel: In der Vorschau ist initial die obere Hälfte des ersten Vorschaubilds sichtbar.
            target_zoom = ch / max(1.0, first_h * 0.5)
            self.zoom = max(self.fit_zoom, target_zoom)
            rw = int(self.base_image.size[0] * self.zoom)
            rh = int(self.base_image.size[1] * self.zoom)
            self.offset_x = 0
            # Bildoberkante an Canvas-Oberkante ausrichten, damit die oberen 50% sichtbar sind.
            self.offset_y = max(0, (rh - ch) / 2)
        elif self.zoom < self.fit_zoom:
            self.zoom = self.fit_zoom
        self._clamp_offsets()
        iw, ih = self.base_image.size
        size = (max(1, int(iw * self.zoom)), max(1, int(ih * self.zoom)))
        img = self.base_image.resize(size)
        self.tk_image = ImageTk.PhotoImage(img)
        self._draw()
    except Exception as exc:
        self.clear()
        self.canvas.create_text(20, 48, anchor="nw", text=f"Vorschaufehler: {exc}", fill=TEXT2, font=("Segoe UI", 10))

def _fm_preview_wheel_v626(self, delta):
    if self.base_image is None:
        return
    factor = 1.15 if delta > 0 else 0.87
    old_zoom = self.zoom
    max_zoom = max(self.fit_zoom * 20.0, 20.0)
    self.zoom = max(self.fit_zoom, min(self.zoom * factor, max_zoom))
    if abs(self.zoom - old_zoom) > 0.0001:
        self._clamp_offsets()
        self._render(reset=False)

_FMDocumentPreviewPane._render = _fm_preview_render_v626
_FMDocumentPreviewPane._wheel = _fm_preview_wheel_v626

# ---------- Block 7B: Zahlungsmodul fachlich in Ein-/Ausgänge + Bankbewegungen strukturieren ----------
def _fm_init_block7b_schema_v626() -> None:
    with get_connection() as con:
        # Bankkonten fachlich für spätere Schnittstellen / Kontierung vorbereiten.
        for col, definition in [
            ("gl_account_no", "TEXT DEFAULT ''"),
            ("incoming_clearing_account", "TEXT DEFAULT ''"),
            ("outgoing_clearing_account", "TEXT DEFAULT ''"),
            ("online_banking_enabled", "INTEGER DEFAULT 0"),
            ("interface_profile", "TEXT DEFAULT ''"),
        ]:
            ensure_column(con, "bank_accounts", col, definition)
        con.executescript("""
CREATE TABLE IF NOT EXISTS bank_transactions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    bank_account_id INTEGER DEFAULT 0,
    booking_date TEXT DEFAULT '',
    value_date TEXT DEFAULT '',
    amount REAL DEFAULT 0,
    currency TEXT DEFAULT 'EUR',
    direction TEXT DEFAULT '',
    partner_name TEXT DEFAULT '',
    partner_iban TEXT DEFAULT '',
    purpose_text TEXT DEFAULT '',
    end_to_end_id TEXT DEFAULT '',
    bank_reference TEXT DEFAULT '',
    imported_from TEXT DEFAULT '',
    import_batch_id TEXT DEFAULT '',
    match_status TEXT DEFAULT 'Ungeklärt',
    matched_entity_type TEXT DEFAULT '',
    matched_reference_no TEXT DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS payment_matches (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    payment_id INTEGER DEFAULT 0,
    bank_transaction_id INTEGER DEFAULT 0,
    entity_type TEXT NOT NULL,
    reference_no TEXT NOT NULL,
    matched_amount REAL NOT NULL,
    difference_amount REAL DEFAULT 0,
    match_type TEXT DEFAULT 'Manuell',
    created_by TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS posting_rules (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    rule_name TEXT NOT NULL,
    direction TEXT DEFAULT '',
    text_contains TEXT DEFAULT '',
    debit_account TEXT DEFAULT '',
    credit_account TEXT DEFAULT '',
    clearing_strategy TEXT DEFAULT '',
    active INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
""")
        # Basis-Regeln nur anlegen, wenn noch keine Regeln vorhanden sind.
        if not con.execute("SELECT 1 FROM posting_rules LIMIT 1").fetchone():
            now = now_str()
            con.execute("INSERT INTO posting_rules(rule_name,direction,text_contains,debit_account,credit_account,clearing_strategy,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)", ("Zahlungseingang Debitor", "Eingang", "", "1200", "1400", "OP-Ausgleich Debitor", 1, now, now))
            con.execute("INSERT INTO posting_rules(rule_name,direction,text_contains,debit_account,credit_account,clearing_strategy,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)", ("Zahlungsausgang Kreditor", "Ausgang", "", "1600", "1200", "OP-Ausgleich Kreditor", 1, now, now))
        con.execute("INSERT OR REPLACE INTO app_meta(key,value,updated_at) VALUES(?,?,?)", ("app_version", APP_VERSION, now_str()))
        con.commit()

_orig_init_sqlite_v626 = init_sqlite

def init_sqlite() -> None:
    _orig_init_sqlite_v626()
    _fm_init_block7b_schema_v626()

# Hilfsfunktionen für Zahlungsoberfläche
def _fm_payment_direction_label(entity_type: str) -> str:
    return "Zahlungsausgang" if entity_type == "vendor_invoice" else "Zahlungseingang"

def _fm_payment_direction_from_amount(amount) -> str:
    try:
        return "Eingang" if parse_amount(amount) >= 0 else "Ausgang"
    except Exception:
        return ""

def _fm_payment_query_text(var) -> str:
    try:
        q = var.get().strip().lower()
        return "" if q == PLACEHOLDER_TEXT.lower() else q
    except Exception:
        return ""

def _fm_payments_make_search(parent, var, callback):
    row = ttk.Frame(parent)
    row.pack(fill="x", pady=(0, 4))
    ttk.Label(row, text="Suchen", width=10).pack(side="left")
    entry = ttk.Entry(row, textvariable=var)
    entry.pack(side="left", fill="x", expand=True)
    _fm_entry_placeholder_bind(entry, PLACEHOLDER_TEXT, callback)
    return entry

# Neue Block-7B-UI: Reiter Zahlungseingänge / Zahlungsausgänge / Bankbewegungen.
def _fm_payments_build_ui_v626(self):
    self.in_search_var = tk.StringVar()
    self.out_search_var = tk.StringVar()
    self.in_history_search_var = tk.StringVar()
    self.out_history_search_var = tk.StringVar()
    self.bank_search_var = tk.StringVar()

    nb = ttk.Notebook(self)
    nb.pack(fill="both", expand=True)
    self.payments_notebook = nb

    tab_in = ttk.Frame(nb)
    tab_out = ttk.Frame(nb)
    tab_bank = ttk.Frame(nb)
    nb.add(tab_in, text="Zahlungseingänge")
    nb.add(tab_out, text="Zahlungsausgänge")
    nb.add(tab_bank, text="Bankbewegungen")

    self._build_direction_tab(tab_in, "customer_invoice")
    self._build_direction_tab(tab_out, "vendor_invoice")
    self._build_bank_transactions_tab(tab_bank)


def _fm_payments_build_direction_tab(self, parent, entity_type: str):
    is_in = entity_type == "customer_invoice"
    shell = build_two_block_shell(
        parent,
        "Zahlungseingang / Debitoren-OP" if is_in else "Zahlungsausgang / Kreditoren-OP",
        "Zahlungshistorie",
    )
    left = ttk.Frame(shell.left)
    left.pack(fill="both", expand=True, padx=10, pady=8)
    right = ttk.Frame(shell.right)
    right.pack(fill="both", expand=True, padx=10, pady=8)

    form = ttk.Frame(left)
    form.pack(fill="x", pady=(0, 8))
    self._row_entry(form, "Zahlungsdatum", self.payment_date)
    self._row_entry(form, "Zahlbetrag", self.amount)
    self._row_combo(form, "Zahlungsart", self.payment_method, ["Überweisung", "Lastschrift", "Bar", "Karte", "Sonstige"])
    self._row_entry(form, "Bank/Konto", self.bank_account)
    self._row_entry(form, "Buchungstext", self.booking_text)
    ttk.Checkbutton(form, text="OP vollständig ausgleichen", variable=self.full_clear).pack(anchor="w", pady=3)
    btns = ttk.Frame(form)
    btns.pack(fill="x", pady=6)
    create_standard_button(btns, "Zahlung buchen", self.book_payment, confirm=True).pack(side="left", padx=3)
    create_standard_button(btns, "Formular leeren", self.clear_form).pack(side="left", padx=3)
    create_standard_button(btns, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)

    search_var = self.in_search_var if is_in else self.out_search_var
    _fm_payments_make_search(left, search_var, self.reload_open_items)
    cols = ("id", "type", "reference", "partner", "due", "original", "open", "status")
    tree = ttk.Treeview(left, columns=cols, show="headings", height=12)
    for c, t, w in [
        ("id", "ID", 55), ("type", "Vorgang", 130), ("reference", "Referenz", 120),
        ("partner", "Partner", 190), ("due", "Fällig", 90), ("original", "Original", 95),
        ("open", "Offen", 95), ("status", "Status", 120),
    ]:
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w")
    tree.pack(fill="both", expand=True)
    configure_tree_tags(tree)
    self.setup_sorting(tree)
    tree.bind("<<TreeviewSelect>>", lambda _e, tr=tree: self._select_open_item_from_tree(tr))

    hist_search_var = self.in_history_search_var if is_in else self.out_history_search_var
    _fm_payments_make_search(right, hist_search_var, self.reload_history)
    hcols = ("date", "direction", "reference", "partner", "amount", "method", "bank", "text")
    hist = ttk.Treeview(right, columns=hcols, show="headings")
    for c, t, w in [
        ("date", "Datum", 90), ("direction", "Richtung", 115), ("reference", "Referenz", 120),
        ("partner", "Partner", 180), ("amount", "Betrag", 90), ("method", "Art", 100),
        ("bank", "Bank/Konto", 130), ("text", "Text", 180),
    ]:
        hist.heading(c, text=t)
        hist.column(c, width=w, anchor="w")
    hist.pack(fill="both", expand=True)
    self.setup_sorting(hist)

    if is_in:
        self.incoming_tree = tree
        self.incoming_history_tree = hist
    else:
        self.outgoing_tree = tree
        self.outgoing_history_tree = hist


def _fm_payments_build_bank_transactions_tab(self, parent):
    top = ttk.Frame(parent)
    top.pack(fill="x", padx=10, pady=8)
    create_standard_button(top, "Bankimport vorbereiten", lambda: messagebox.showinfo("Bankbewegungen", "CSV/camt.053/EBICS-Struktur ist vorbereitet. Der echte Import folgt in Block 7D.", parent=self)).pack(side="left", padx=3)
    create_standard_button(top, "Daten-Refresh", self.refresh_all).pack(side="left", padx=3)
    _fm_payments_make_search(parent, self.bank_search_var, self.reload_bank_transactions)
    cols = ("id", "date", "value", "direction", "amount", "partner", "purpose", "match")
    tree = ttk.Treeview(parent, columns=cols, show="headings")
    for c, t, w in [
        ("id", "ID", 55), ("date", "Buchung", 90), ("value", "Valuta", 90), ("direction", "Richtung", 90),
        ("amount", "Betrag", 90), ("partner", "Partner", 180), ("purpose", "Verwendungszweck", 320), ("match", "Status", 120),
    ]:
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w")
    tree.pack(fill="both", expand=True, padx=10, pady=(0, 10))
    self.bank_tx_tree = tree
    self.setup_sorting(tree)


def _fm_payments_select_open_item_from_tree(self, tree):
    iid = tree.focus()
    if not iid:
        return
    try:
        op_id = int(tree.set(iid, "id"))
    except Exception:
        return
    with get_connection() as con:
        row = con.execute("SELECT * FROM open_items WHERE id=?", (op_id,)).fetchone()
    if not row:
        return
    self.selected_open_item_id = op_id
    self.selected_entity_type = row["entity_type"]
    self.selected_reference_no = row["reference_no"]
    self.amount.set(format_amount(row["open_amount"]))
    if not self.booking_text.get().strip():
        self.booking_text.set(f"{_fm_payment_direction_label(row['entity_type'])} {row['reference_no']}")


def _fm_payments_reload_open_items_v626(self):
    mapping = [
        (getattr(self, "incoming_tree", None), "customer_invoice", getattr(self, "in_search_var", None)),
        (getattr(self, "outgoing_tree", None), "vendor_invoice", getattr(self, "out_search_var", None)),
    ]
    with get_connection() as con:
        all_rows = list(con.execute("SELECT * FROM open_items WHERE open_amount > 0 ORDER BY due_date, id"))
    for tree, entity_type, var in mapping:
        if tree is None:
            continue
        for i in tree.get_children():
            tree.delete(i)
        q = _fm_payment_query_text(var)
        for r in all_rows:
            if r["entity_type"] != entity_type:
                continue
            values = (r["id"], _fm_payment_direction_label(r["entity_type"]), r["reference_no"], r["partner_name"], r["due_date"], format_amount(r["original_amount"]), format_amount(r["open_amount"]), r["status"])
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values, tags=(urgency_bucket(r["due_date"], r["open_amount"]),))


def _fm_payments_reload_history_v626(self):
    mapping = [
        (getattr(self, "incoming_history_tree", None), "customer_invoice", getattr(self, "in_history_search_var", None)),
        (getattr(self, "outgoing_history_tree", None), "vendor_invoice", getattr(self, "out_history_search_var", None)),
    ]
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM payments ORDER BY id DESC"))
    for tree, entity_type, var in mapping:
        if tree is None:
            continue
        for i in tree.get_children():
            tree.delete(i)
        q = _fm_payment_query_text(var)
        for r in rows:
            if r["entity_type"] != entity_type:
                continue
            values = (r["payment_date"], _fm_payment_direction_label(r["entity_type"]), r["reference_no"], r["partner_name"], format_amount(r["amount"]), r["payment_method"], r["bank_account"], r["booking_text"])
            if q and q not in " ".join(str(v) for v in values).lower():
                continue
            tree.insert("", "end", values=values)


def _fm_payments_reload_bank_transactions_v626(self):
    tree = getattr(self, "bank_tx_tree", None)
    if tree is None:
        return
    for i in tree.get_children():
        tree.delete(i)
    q = _fm_payment_query_text(getattr(self, "bank_search_var", None))
    with get_connection() as con:
        rows = list(con.execute("SELECT * FROM bank_transactions ORDER BY id DESC"))
    for r in rows:
        values = (r["id"], r["booking_date"], r["value_date"], r["direction"], format_amount(r["amount"]), r["partner_name"], r["purpose_text"], r["match_status"])
        if q and q not in " ".join(str(v) for v in values).lower():
            continue
        tree.insert("", "end", values=values)


def _fm_payments_book_payment_v626(self):
    if not self.selected_open_item_id:
        messagebox.showwarning("Zahlung", "Bitte zuerst einen offenen Posten auswählen.", parent=self)
        return
    if not validate_date(self.payment_date.get()):
        messagebox.showwarning("Datum", "Datumsformat TT.MM.JJJJ verwenden.", parent=self)
        return
    try:
        amount = parse_amount(self.amount.get())
    except Exception as exc:
        messagebox.showwarning("Betrag", str(exc), parent=self)
        return
    if amount <= 0:
        messagebox.showwarning("Betrag", "Der Zahlbetrag muss größer als 0 sein.", parent=self)
        return
    with get_connection() as con:
        op = con.execute("SELECT * FROM open_items WHERE id=?", (self.selected_open_item_id,)).fetchone()
        if not op:
            messagebox.showwarning("Zahlung", "Der offene Posten wurde nicht gefunden.", parent=self)
            return
        open_amount = parse_amount(op["open_amount"])
        pay_amount = open_amount if int(self.full_clear.get() or 0) else amount
        if pay_amount > open_amount:
            messagebox.showwarning("Betrag", "Der Zahlbetrag darf den offenen Betrag nicht überschreiten.", parent=self)
            return
        new_open = (open_amount - pay_amount).quantize(Decimal("0.01"))
        new_status = _fm_payment_status_from_amount_v624(new_open)
        cur = con.execute(
            "INSERT INTO payments(entity_type,reference_no,partner_no,partner_name,payment_date,amount,payment_method,bank_account,booking_text,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
            (op["entity_type"], op["reference_no"], op["partner_no"], op["partner_name"], self.payment_date.get(), float(pay_amount), self.payment_method.get(), self.bank_account.get(), self.booking_text.get(), now_str(), now_str()),
        )
        payment_id = cur.lastrowid
        con.execute("UPDATE open_items SET open_amount=?, status=?, updated_at=? WHERE id=?", (float(new_open), new_status, now_str(), self.selected_open_item_id))
        table = _fm_invoice_table_for_entity_v624(op["entity_type"])
        con.execute(f"UPDATE {table} SET open_amount=?, status=?, updated_at=? WHERE invoice_no=?", (float(new_open), new_status, now_str(), op["reference_no"]))
        con.execute(
            "INSERT INTO payment_matches(payment_id,bank_transaction_id,entity_type,reference_no,matched_amount,difference_amount,match_type,created_by,created_at) VALUES(?,?,?,?,?,?,?,?,?)",
            (payment_id, 0, op["entity_type"], op["reference_no"], float(pay_amount), float(new_open), "Manuell", _fm_current_user() if "_fm_current_user" in globals() else "", now_str()),
        )
        con.commit()
    try:
        self.app.set_status(f"Zahlung zu {self.selected_reference_no} gebucht.")
    except Exception:
        pass
    self.clear_form()
    self.refresh_all()


def _fm_payments_clear_form_v626(self):
    self.selected_open_item_id = None
    self.selected_entity_type = ""
    self.selected_reference_no = ""
    self.payment_date.set(today_str())
    self.amount.set("")
    self.payment_method.set("Überweisung")
    self.bank_account.set("")
    self.booking_text.set("")
    self.full_clear.set(1)
    for tree in [getattr(self, "incoming_tree", None), getattr(self, "outgoing_tree", None)]:
        try:
            tree.selection_remove(tree.selection())
        except Exception:
            pass


def _fm_payments_refresh_all_v626(self):
    self.reload_open_items()
    self.reload_history()
    self.reload_bank_transactions()
    try:
        self.app.set_status("Zahlungen aktualisiert.")
    except Exception:
        pass

PaymentsView._build_ui = _fm_payments_build_ui_v626
PaymentsView._build_direction_tab = _fm_payments_build_direction_tab
PaymentsView._build_bank_transactions_tab = _fm_payments_build_bank_transactions_tab
PaymentsView._select_open_item_from_tree = _fm_payments_select_open_item_from_tree
PaymentsView.reload_open_items = _fm_payments_reload_open_items_v626
PaymentsView.reload_history = _fm_payments_reload_history_v626
PaymentsView.reload_bank_transactions = _fm_payments_reload_bank_transactions_v626
PaymentsView.book_payment = _fm_payments_book_payment_v626
PaymentsView.clear_form = _fm_payments_clear_form_v626
PaymentsView.refresh_all = _fm_payments_refresh_all_v626

# ---------- v0.6.26 Selbsttest ----------
def _fm_block7b_static_selftest_v626() -> None:
    import inspect as _inspect
    src = Path(__file__).read_text(encoding="utf-8", errors="replace")
    required = [
        "APP_VERSION = \"0.6.26-block7b-payment-tabs-import-preview\"",
        "tk_getOpenFile",
        "-multiple",
        "self.fit_zoom * 20.0",
        "target_zoom = ch / max(1.0, first_h * 0.5)",
        "bank_transactions",
        "payment_matches",
        "posting_rules",
        "Zahlungseingänge",
        "Zahlungsausgänge",
        "Bankbewegungen",
    ]
    missing = [x for x in required if x not in src]
    csrc = _inspect.getsource(CreditorsView._build_ui)
    if "stack_tree" in csrc or "stack_search_var" in csrc:
        missing.append("Dokumentenstapel generische Zusatzsuchleiste nicht entfernt")
    psrc = _inspect.getsource(PaymentsView._build_ui)
    for token in ["ttk.Notebook", "Zahlungseingänge", "Zahlungsausgänge", "Bankbewegungen"]:
        if token not in psrc:
            missing.append(token)
    if missing:
        raise RuntimeError("Block-7B-Selbsttest v0.6.26 fehlgeschlagen: " + ", ".join(missing))

_fm_block7b_static_selftest_v626()


def main() -> None:
    app = FinanceMateApp()
    app.mainloop()

if __name__ == "__main__":
    main()
